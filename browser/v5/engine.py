"""
Form Filler Engine V5 - Universal Form Filler

Multi-layer architecture:
1. DETECTION: HTML → Probe → Vision
2. RESOLUTION: Profile → Learned → AI → Human  
3. INPUT: Type/Select/Click based on field type
4. VALIDATION: Read back + Error check + Vision verify
5. LEARNING: Save successful fills

Modes:
- PRE_FLIGHT: Analyze form, generate readiness report (no actual fill)
- INTERACTIVE: Fill with human confirmation for unknowns
- AUTONOMOUS: Fill everything, skip unknowns
"""

import json
import re
import time
import requests  # For Ollama API
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass, field as dataclass_field
from enum import Enum

from playwright.sync_api import Page, ElementHandle

from .browser_manager import BrowserManager, BrowserMode
from .form_logger import FormLogger

# ═══════════════════════════════════════════════════════════════════════════
# PATHS
# ═══════════════════════════════════════════════════════════════════════════

V5_DIR = Path(__file__).parent
BROWSER_DIR = V5_DIR.parent
DATA_DIR = V5_DIR / "data"
DATA_DIR.mkdir(exist_ok=True)

PROFILE_PATH = BROWSER_DIR / "profiles" / "anton_tpm.json"
# Use shared learned database with V3.5
LEARNED_DB_PATH = BROWSER_DIR / "learned_database.json"
KNOWLEDGE_BASE_PATH = BROWSER_DIR / "knowledge_base.json"
FIELD_PATTERNS_PATH = DATA_DIR / "field_patterns.json"


# ═══════════════════════════════════════════════════════════════════════════
# ENUMS
# ═══════════════════════════════════════════════════════════════════════════

class FillMode(Enum):
    """Form filling modes"""
    PRE_FLIGHT = "pre_flight"      # Analyze only, no fill
    INTERACTIVE = "interactive"    # Human confirms unknowns
    AUTONOMOUS = "autonomous"      # Fill all, skip unknowns


class FieldType(Enum):
    """Detected field types"""
    TEXT = "text"
    EMAIL = "email"
    PHONE = "phone"
    TEXTAREA = "textarea"
    SELECT = "select"              # Native <select>
    AUTOCOMPLETE = "autocomplete"  # React Select, combobox
    CHECKBOX = "checkbox"
    RADIO = "radio"
    FILE = "file"
    DATE = "date"
    HIDDEN = "hidden"
    UNKNOWN = "unknown"


class DetectionMethod(Enum):
    """How field type was detected"""
    HTML = "html"           # From HTML tag/type
    ARIA = "aria"           # From ARIA attributes
    PROBE = "probe"         # From click behavior
    VISION = "vision"       # From screenshot AI
    PATTERN = "pattern"     # From known patterns
    DEFAULT = "default"     # Fallback


class AnswerSource(Enum):
    """Where the answer came from"""
    PROFILE = "profile"
    LEARNED = "learned"
    AI = "ai"
    HUMAN = "human"
    DEFAULT = "default"
    NONE = "none"


class FillStatus(Enum):
    """Field fill status"""
    READY = "ready"           # Has answer, ready to fill
    FILLED = "filled"         # Successfully filled
    VERIFIED = "verified"     # Filled and verified
    NEEDS_INPUT = "needs_input"   # No answer found
    SKIPPED = "skipped"       # Intentionally skipped
    ERROR = "error"           # Fill failed


# ═══════════════════════════════════════════════════════════════════════════
# DATA CLASSES
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class FormField:
    """Detected form field with all metadata"""
    selector: str
    element_id: str
    name: str
    label: str
    
    # Detection
    field_type: FieldType
    detection_method: DetectionMethod
    html_tag: str
    input_type: str
    
    # Options (for select/autocomplete)
    options: List[str] = dataclass_field(default_factory=list)
    
    # State
    required: bool = False
    visible: bool = True
    current_value: str = ""
    
    # Answer
    answer: str = ""
    answer_source: AnswerSource = AnswerSource.NONE
    confidence: float = 0.0
    
    # Fill status
    status: FillStatus = FillStatus.NEEDS_INPUT
    error_message: str = ""
    
    # Profile mapping
    profile_key: str = ""

    # DOM constraints (for format adaptation)
    placeholder: str = ""
    maxlength: int = 0
    pattern: str = ""       # HTML5 pattern attribute


@dataclass
class FillReport:
    """Report of form filling attempt"""
    url: str
    title: str
    ats_type: str
    
    total_fields: int = 0
    ready_fields: int = 0
    filled_fields: int = 0
    verified_fields: int = 0
    needs_input: int = 0
    skipped: int = 0
    errors: int = 0
    
    fields: List[FormField] = dataclass_field(default_factory=list)
    
    def summary(self) -> str:
        """Generate text summary."""
        lines = [
            f"═══════════════════════════════════════════════════════════",
            f"FORM FILL REPORT",
            f"═══════════════════════════════════════════════════════════",
            f"URL: {self.url[:70]}...",
            f"Title: {self.title[:50]}",
            f"ATS: {self.ats_type}",
            f"",
            f"📊 SUMMARY:",
            f"   Total fields: {self.total_fields}",
            f"   ✅ Ready/Filled: {self.filled_fields}",
            f"   ✓ Verified: {self.verified_fields}",
            f"   ⚠️ Needs input: {self.needs_input}",
            f"   ⏭️ Skipped: {self.skipped}",
            f"   ❌ Errors: {self.errors}",
        ]
        
        if self.needs_input > 0:
            lines.append(f"\n⚠️ FIELDS NEEDING INPUT:")
            for f in self.fields:
                if f.status == FillStatus.NEEDS_INPUT:
                    lines.append(f"   • {f.label[:50]}")
        
        if self.errors > 0:
            lines.append(f"\n❌ ERRORS:")
            for f in self.fields:
                if f.status == FillStatus.ERROR:
                    lines.append(f"   • {f.label[:40]}: {f.error_message}")
        
        lines.append("═══════════════════════════════════════════════════════════")
        return "\n".join(lines)

    def detailed_report(self) -> str:
        """Generate detailed field-by-field text report."""
        lines = [
            "╔══════════════════════════════════════════════════════════════════════════════════════════════╗",
            f"║  FORM FILL REPORT — {self.title[:60]:<60}        ║",
            f"║  ATS: {self.ats_type:<15}  URL: {self.url[:55]:<55} ║",
            "╠══════════════════════════════════════════════════════════════════════════════════════════════╣",
        ]

        # Source stats
        source_counts = {}
        for f in self.fields:
            src = f.answer_source.value if f.answer_source else "none"
            source_counts[src] = source_counts.get(src, 0) + 1

        # Group fields by category
        status_icons = {
            FillStatus.VERIFIED: "✅",
            FillStatus.FILLED: "📝",
            FillStatus.ERROR: "❌",
            FillStatus.SKIPPED: "⏭️",
            FillStatus.NEEDS_INPUT: "❓",
            FillStatus.READY: "🔵",
        }

        for i, f in enumerate(self.fields, 1):
            icon = status_icons.get(f.status, "?")
            src = f.answer_source.value if f.answer_source else "-"
            answer = (f.answer or "")
            label = (f.label or "").split(" [")[0]  # Remove [name=...] suffix

            # Truncate for display
            label_display = label[:55]
            answer_display = answer[:50]

            lines.append(
                f"║ {i:>2}. {icon} {f.field_type.value:12} │ {label_display:<55} │ {answer_display:<50} │ {src:<8} ║"
            )

        lines.append("╠══════════════════════════════════════════════════════════════════════════════════════════════╣")

        # Summary
        lines.append(f"║  📊 RESULTS: {self.verified_fields} verified, {self.filled_fields} filled, {self.errors} errors, {self.skipped} skipped / {self.total_fields} total")

        # Source breakdown
        src_parts = [f"{src}: {cnt}" for src, cnt in sorted(source_counts.items(), key=lambda x: -x[1])]
        lines.append(f"║  📦 SOURCES: {', '.join(src_parts)}")

        # Errors
        error_fields = [f for f in self.fields if f.status == FillStatus.ERROR]
        if error_fields:
            lines.append(f"║  ❌ ERRORS:")
            for f in error_fields:
                lines.append(f"║     • {f.label[:50]}: {f.error_message}")

        # Needs input
        needs = [f for f in self.fields if f.status == FillStatus.NEEDS_INPUT]
        if needs:
            lines.append(f"║  ⚠️ NEEDS INPUT:")
            for f in needs:
                lines.append(f"║     • {f.label[:60]}")

        lines.append("╚══════════════════════════════════════════════════════════════════════════════════════════════╝")
        return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════════
# PROFILE
# ═══════════════════════════════════════════════════════════════════════════

class Profile:
    """User profile data manager."""
    
    # Label patterns → profile keys
    LABEL_MAPPINGS = {
        # Personal info
        "first name": "personal.first_name",
        "last name": "personal.last_name",
        "email": "personal.email",
        "phone": "personal.phone",
        "city": "personal.city",
        "beverly hills": "personal.city",  # Common city placeholder
        "new york": "personal.city",
        "los angeles": "personal.city",
        "san francisco": "personal.city",
        "state": "personal.state",
        "country": "personal.country",
        "zip": "personal.zip_code",
        "postal": "personal.zip_code",
        "90210": "personal.zip_code",  # Common ZIP placeholder
        "12345": "personal.zip_code",
        "00000": "personal.zip_code",
        "street": "personal.street_address",
        "address": "personal.street_address",
        "address line": "personal.street_address",
        "main st": "personal.street_address",  # Common placeholder pattern
        "your address": "personal.street_address",
        "home address": "personal.street_address",
        "linkedin": "links.linkedin",
        "github": "links.github",
        # Education dates — MUST be before generic start/end patterns
        # because "Education Start Month" contains "start month" which would
        # match work_experience patterns if checked first
        "education start month": "education.0.start_month",
        "education start year": "education.0.start_year",
        "education end month": "education.0.end_month",
        "education end year": "education.0.end_year",
        "school": "education.0.school",
        "university": "education.0.school",
        "degree": "education.0.degree",
        "discipline": "education.0.discipline",
        "your major": "education.0.discipline",
        "field of study": "education.0.discipline",
        # Work experience (generic date patterns — checked after education-specific)
        "company name": "work_experience.0.company",
        "employer": "work_experience.0.company",
        "job title": "work_experience.0.title",
        "title": "work_experience.0.title",
        "start date month": "work_experience.0.start_month",
        "start month": "work_experience.0.start_month",
        "start date year": "work_experience.0.start_year",
        "start year": "work_experience.0.start_year",
        "end date month": "work_experience.0.end_month",
        "end month": "work_experience.0.end_month",
        "end date year": "work_experience.0.end_year",
        "end year": "work_experience.0.end_year",
        # Common questions
        "how did you hear": "common_answers.how_heard",
        "full name": "personal.full_name",
        "your name": "personal.full_name",
        "legal name": "personal.full_name",
        "preferred name": "personal.first_name",
        "preferred first name": "personal.first_name",
        "candidate name": "personal.full_name",
        "name of company": "work_experience.0.company",
        "name of your current company": "work_experience.0.company",
        "current company": "work_experience.0.company",
        "current employer": "work_experience.0.company",
        "location": "personal.location",
        "current location": "personal.location",
        "where are you located": "personal.location",
        "website": "links.linkedin",
        "portfolio": "links.linkedin",
        "personal website": "links.linkedin",
    }
    
    # Yes/No question patterns
    YES_NO_PATTERNS = {
        "18 years": "Yes",
        "authorized to work": "Yes",
        "legally authorized": "Yes",
        "eligible to work": "Yes",
        "require sponsorship": "No",
        "visa sponsorship": "No",
        "government official": "No",
        "close relative of a government": "No",
        "conflict of interest": "No",
        "connected to": "No",
        "financial interest": "No",
        "referred to this position by": "No",
        "senior leader": "No",
        "previously employed": "No",
        "previously been employed": "No",
        "ever worked for": "No",
        "currently or have you ever worked": "No",
        "worked for": "No",
        "former employee": "No",
        "confirm": "Yes",
        "acknowledge": "Yes",
        "confirm receipt": "Confirmed",
        # IMPORTANT: More specific patterns MUST come before generic ones
        # "non-compete" must be before "agree" because questions like
        # "Do you work under any agreement, such as a non-compete agreement..."
        # would match "agree" first and answer "Yes" incorrectly
        "non-compete": "No",
        "non-disclosure": "Yes",
        "agree": "Yes",
        "i understand": "Yes",
        "current role": "Yes",
        "currently work here": "Yes",
        "currently work": "Yes",
        "i currently work": "Yes",
        # Additional from V35
        "willing to relocate": "Yes",
        "background check": "Yes",
        "drug test": "Yes",
    }
    
    # Text field defaults (for common questions)
    TEXT_DEFAULTS = {
        "years of experience": "15",
        "years experience": "15",
        "how many years": "15",
        "how did you hear": "LinkedIn",
        "how did you find": "LinkedIn",
        "where did you hear": "LinkedIn",
        "how were you referred": "LinkedIn",
        "referred by": "LinkedIn",
        "referral source": "LinkedIn",
        "source of application": "LinkedIn",
        "salary": "150000",
        "desired salary": "150000",
        "expected salary": "150000",
        "compensation": "150000",
        "salary expectation": "150000",
        "annual salary": "150000",
        "notice period": "2 weeks",
        "when can you start": "2 weeks",
        "earliest start": "2 weeks",
        "availability": "2 weeks",
        "start date": "2 weeks",
        "website": "https://linkedin.com/in/antonkondakov",
        "portfolio": "https://linkedin.com/in/antonkondakov",
        "personal website": "https://linkedin.com/in/antonkondakov",
        "additional information": "I am excited about this opportunity and believe my 15+ years of experience in technical program management align well with this role.",
    }
    
    # Demographic defaults
    DEMOGRAPHIC_DEFAULTS = {
        "gender": "Decline to self-identify",
        "race": "Decline to self-identify",
        "ethnicity": "Decline to self-identify",
        "hispanic": "Decline to self-identify",
        "veteran": "I am not a protected veteran",
        "disability": "I do not want to answer",
    }
    
    def __init__(self, path: Path = PROFILE_PATH):
        self.data = {}
        if path.exists():
            with open(path) as f:
                self.data = json.load(f)
    
    def get(self, key: str) -> str:
        """Get value by dot-notation key."""
        parts = key.split(".")
        val = self.data
        for p in parts:
            if val is None:
                return ""
            if p.isdigit():
                idx = int(p)
                val = val[idx] if isinstance(val, list) and idx < len(val) else None
            else:
                val = val.get(p) if isinstance(val, dict) else None
        return str(val) if val else ""

    # Alias for FormSchemaDB compatibility
    get_by_dotpath = get

    def find_by_label(self, label: str) -> Tuple[Optional[str], Optional[str]]:
        """Find profile value by label text. Returns (value, profile_key)."""
        ll = label.lower()
        for pattern, key in self.LABEL_MAPPINGS.items():
            if re.search(r'\b' + re.escape(pattern) + r'\b', ll):
                val = self.get(key)
                if val:
                    return val, key
        return None, None
    
    def find_yes_no(self, label: str) -> Optional[str]:
        """Find Yes/No answer for common questions."""
        ll = label.lower()
        for pattern, answer in self.YES_NO_PATTERNS.items():
            if pattern in ll:
                return answer
        return None
    
    def find_demographic(self, label: str) -> Optional[str]:
        """Find demographic answer — profile first, then defaults."""
        ll = label.lower()
        # Try profile demographics first
        demo = self.data.get("demographics", {})
        if demo:
            demo_field_map = {
                "gender": "gender",
                "hispanic": "hispanic_latino",
                "latino": "hispanic_latino",
                "race": "race_ethnicity",
                "ethnicity": "race_ethnicity",
                "veteran": "veteran_status",
                "disability": "disability_status",
            }
            for pattern, field_key in demo_field_map.items():
                if pattern in ll:
                    val = demo.get(field_key, "")
                    if val:
                        return val
        # Fallback to hardcoded defaults
        for pattern, answer in self.DEMOGRAPHIC_DEFAULTS.items():
            if pattern in ll:
                return answer
        return None
    
    def find_text_default(self, label: str) -> Optional[str]:
        """Find text default for common questions like years of experience."""
        ll = label.lower()
        for pattern, answer in self.TEXT_DEFAULTS.items():
            if pattern in ll:
                return answer
        return None
    
    def get_context(self) -> str:
        """Get summary context for AI."""
        p = self.data.get("personal", {})
        w = self.data.get("work_experience", [{}])[0]
        return f"Name: {p.get('first_name', '')} {p.get('last_name', '')}\nLocation: {p.get('location', '')}\nRole: {w.get('title', '')} at {w.get('company', '')}"
    
    def get_files_for_role(self, job_title: str) -> Tuple[Optional[Path], Optional[Path]]:
        """
        Get CV and Cover Letter paths based on job title.
        
        Args:
            job_title: Job title from the application (e.g. "Senior Product Manager, Finance")
            
        Returns:
            Tuple of (cv_path, cover_letter_path) - either can be None if not found
        """
        files_config = self.data.get("files", {})
        base_path = Path(files_config.get("base_path", ""))
        by_role = files_config.get("by_role", {})
        default_role = files_config.get("default_role", "TPM")
        
        if not base_path.exists():
            return None, None
        
        # Determine role from job title
        job_title_lower = job_title.lower()
        detected_role = None
        
        # Role detection patterns (order matters - more specific first)
        role_patterns = [
            ("TPM", ["technical program manager", "tpm"]),
            ("Product Manager", ["product manager"]),
            ("Product Owner", ["product owner"]),
            ("Project Manager", ["project manager"]),
            ("Scrum Master", ["scrum master", "agile coach"]),
            ("Delivery Lead", ["delivery lead", "delivery manager"]),
        ]
        
        for role_name, patterns in role_patterns:
            for pattern in patterns:
                if pattern in job_title_lower:
                    detected_role = role_name
                    break
            if detected_role:
                break
        
        # Use default if no match
        if not detected_role:
            detected_role = default_role
        
        # Get files for detected role
        role_files = by_role.get(detected_role, by_role.get(default_role, {}))
        
        cv_filename = role_files.get("cv")
        cover_letter_filename = role_files.get("cover_letter")
        
        cv_path = base_path / cv_filename if cv_filename else None
        cover_letter_path = base_path / cover_letter_filename if cover_letter_filename else None
        
        # Verify files exist
        if cv_path and not cv_path.exists():
            print(f"   ⚠️ CV not found: {cv_path}")
            cv_path = None
        if cover_letter_path and not cover_letter_path.exists():
            print(f"   ⚠️ Cover letter not found: {cover_letter_path}")
            cover_letter_path = None
        
        if cv_path:
            print(f"   📄 Selected CV for \'{detected_role}\': {cv_path.name}")
        
        return cv_path, cover_letter_path


# ═══════════════════════════════════════════════════════════════════════════
# LEARNED DATABASE
# ═══════════════════════════════════════════════════════════════════════════

class LearnedDB:
    """Database of learned field answers."""
    
    def __init__(self, path: Path = LEARNED_DB_PATH):
        self.path = path
        self.data = self._load()
    
    def _load(self) -> dict:
        if self.path.exists():
            with open(self.path) as f:
                data = json.load(f)
            # Support V3.5 format (field_answers) and V5 format (answers)
            if "field_answers" in data and "answers" not in data:
                # Convert V3.5 format to V5
                return {
                    "answers": data.get("field_answers", {}),
                    "dropdown_choices": data.get("dropdown_choices", {})
                }
            return data
        return {"answers": {}, "dropdown_choices": {}}
    
    def save(self):
        with open(self.path, "w") as f:
            json.dump(self.data, f, indent=2, ensure_ascii=False)
    
    def _normalize_key(self, label: str) -> str:
        key = label.lower().strip()
        key = re.sub(r'[*?!:\-_()\"\']+', ' ', key)
        key = re.sub(r'\s+', ' ', key).strip()
        return key[:100]
    
    def find(self, label: str, is_dropdown: bool = False) -> Optional[str]:
        """Find saved answer for field label."""
        key = self._normalize_key(label)
        store = self.data["dropdown_choices" if is_dropdown else "answers"]
        
        if key in store:
            return store[key]
        
        # Partial match
        for k, v in store.items():
            if k in key or key in k:
                return v
        
        return None
    
    def save_answer(self, label: str, answer: str, is_dropdown: bool = False):
        """Save answer for future use."""
        key = self._normalize_key(label)
        store = "dropdown_choices" if is_dropdown else "answers"
        self.data[store][key] = answer
        self.save()
        print(f"   💾 Learned: '{label[:30]}' → '{answer[:25]}'")


# ═══════════════════════════════════════════════════════════════════════════
# FORM SCHEMA DB - Remembers form field structures per ATS type
# ═══════════════════════════════════════════════════════════════════════════

FORM_SCHEMAS_PATH = BROWSER_DIR / "form_schemas.json"

class FormSchemaDB:
    """Database of form schemas per ATS type.

    After each successful fill, saves field mappings:
      field_id → {profile_key, source, field_type, label, options, answer_hint}

    Before next fill on same ATS, loads schema and uses it as "step 0" in cascade:
      - Known profile_key fields → resolve directly from profile
      - Known dropdown choices → apply immediately
      - Known field types → skip detection
      - Tracks success_count per field → higher count = higher confidence

    Schema structure:
    {
      "Greenhouse": {
        "last_updated": "2026-02-14T...",
        "fill_count": 5,
        "fields": {
          "first_name": {
            "label": "First Name*",
            "field_type": "text",
            "profile_key": "personal.first_name",
            "source": "profile",
            "answer_hint": "",
            "options": [],
            "success_count": 5,
            "last_success": "2026-02-14T..."
          },
          "question_11097818007": {
            "label": "Are you 18 years of age or older?*",
            "field_type": "autocomplete",
            "profile_key": "",
            "source": "default",
            "answer_hint": "Yes",
            "options": ["Yes", "No"],
            "success_count": 3,
            "last_success": "2026-02-14T..."
          }
        },
        "repeatable_field_ids": {
          "education": ["school--{N}", "degree--{N}", "discipline--{N}",
                        "start-month--{N}", "start-year--{N}"],
          "work_experience": ["company-name-{N}", "title-{N}"]
        }
      }
    }
    """

    def __init__(self, path: Path = FORM_SCHEMAS_PATH):
        self.path = path
        self.data = self._load()

    def _load(self) -> dict:
        if self.path.exists():
            try:
                with open(self.path) as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def save(self):
        with open(self.path, "w") as f:
            json.dump(self.data, f, indent=2, ensure_ascii=False)

    def get_schema(self, ats_type: str) -> Optional[dict]:
        """Get form schema for an ATS type."""
        return self.data.get(ats_type)

    def get_field_mapping(self, ats_type: str, field_id: str) -> Optional[dict]:
        """Get known mapping for a specific field in an ATS schema."""
        schema = self.data.get(ats_type)
        if not schema:
            return None
        return schema.get("fields", {}).get(field_id)

    def resolve_from_schema(self, ats_type: str, field: 'FormField',
                            profile: 'Profile') -> Optional[Tuple[str, 'AnswerSource', float]]:
        """Try to resolve a field answer using saved schema.

        Returns (answer, source, confidence) or None.
        Uses the saved profile_key or answer_hint from previous successful fills.
        """
        mapping = self.get_field_mapping(ats_type, field.element_id or "")
        if not mapping:
            # Try by normalized label match
            schema = self.data.get(ats_type, {})
            fields = schema.get("fields", {})
            label_lower = (field.label or "").lower().strip()
            for fid, fmap in fields.items():
                saved_label = (fmap.get("label", "") or "").lower().strip()
                if saved_label and saved_label == label_lower:
                    mapping = fmap
                    break
            if not mapping:
                return None

        success_count = mapping.get("success_count", 0)
        if success_count < 1:
            return None

        # Confidence based on success count (min 0.92, grows with count)
        confidence = min(0.98, 0.92 + success_count * 0.01)

        # 1. If profile_key is saved, resolve from profile directly
        profile_key = mapping.get("profile_key", "")
        if profile_key:
            val = profile.get_by_dotpath(profile_key)
            if val:
                return (val, AnswerSource.PROFILE, confidence)

        # 2. If source was "default" or "learned" and answer_hint is saved
        answer_hint = mapping.get("answer_hint", "")
        source_str = mapping.get("source", "")
        if answer_hint and source_str in ("default", "learned", "ai"):
            src = AnswerSource.LEARNED if source_str == "learned" else AnswerSource.DEFAULT
            return (answer_hint, src, confidence)

        return None

    def save_schema_from_fill(self, ats_type: str, fields: List['FormField'],
                               url: str = ""):
        """Save/update schema from a successful form fill.

        Called after each fill — merges new field data into existing schema.
        Only saves fields that were successfully filled (FILLED or VERIFIED).
        """
        from datetime import datetime as _dt
        now = _dt.now().isoformat()

        if ats_type not in self.data:
            self.data[ats_type] = {
                "last_updated": now,
                "fill_count": 0,
                "fields": {},
                "repeatable_field_ids": {},
            }

        schema = self.data[ats_type]
        schema["last_updated"] = now
        schema["fill_count"] = schema.get("fill_count", 0) + 1

        saved_fields = schema.get("fields", {})
        repeatable_ids = schema.get("repeatable_field_ids", {})

        for field in fields:
            # Only learn from successful fills
            if field.status not in (FillStatus.FILLED, FillStatus.VERIFIED):
                continue
            if not field.answer:
                continue

            # Determine the field key — use element_id (most stable)
            fid = field.element_id or field.name or ""
            if not fid:
                continue

            # Skip repeatable section fields with numeric index — normalize them
            # e.g., "school--0" → "school--{N}", "company-name-1" → "company-name-{N}"
            normalized_fid, section_name = self._normalize_repeatable_id(fid)
            if section_name:
                # Track repeatable field patterns per section
                if section_name not in repeatable_ids:
                    repeatable_ids[section_name] = []
                if normalized_fid not in repeatable_ids[section_name]:
                    repeatable_ids[section_name].append(normalized_fid)
                # Don't save individual repeatable entries (school--0, school--1, etc.)
                # — they're handled by REPEATABLE_SECTIONS config
                continue

            # Build/update field mapping
            existing = saved_fields.get(fid, {})
            success_count = existing.get("success_count", 0) + 1

            # Determine answer_hint — save static answers, skip dynamic ones
            answer_hint = ""
            source_str = field.answer_source.value if field.answer_source else "none"
            if source_str in ("default", "learned"):
                answer_hint = field.answer
            elif source_str == "ai" and field.field_type in (FieldType.SELECT, FieldType.AUTOCOMPLETE):
                # For dropdowns, AI choices are stable — save them
                answer_hint = field.answer

            saved_fields[fid] = {
                "label": (field.label or "")[:100],
                "field_type": field.field_type.value if field.field_type else "unknown",
                "profile_key": field.profile_key or existing.get("profile_key", ""),
                "source": source_str,
                "answer_hint": answer_hint or existing.get("answer_hint", ""),
                "options": field.options[:20] if field.options else existing.get("options", []),
                "success_count": success_count,
                "last_success": now,
            }

        schema["fields"] = saved_fields
        schema["repeatable_field_ids"] = repeatable_ids
        self.save()

        print(f"   📐 Schema saved: {ats_type} ({len(saved_fields)} fields, fill #{schema['fill_count']})")

    def _normalize_repeatable_id(self, field_id: str) -> Tuple[str, str]:
        """Normalize repeatable field IDs: 'school--0' → ('school--{N}', 'education').

        Returns (normalized_id, section_name) or (field_id, '') if not repeatable.
        """
        # Patterns for known repeatable sections
        education_patterns = [
            (r'(school|degree|discipline|start-month|start-year|end-month|end-year)--(\d+)$', 'education'),
            (r'(education_school_name|education_degree|education_discipline)_(\d+)$', 'education'),
        ]
        work_patterns = [
            (r'(company-name|title|start-date-month|start-date-year|end-date-month|end-date-year)-(\d+)$', 'work_experience'),
            (r'(current-role)-(\d+)_\d+$', 'work_experience'),
        ]

        for patterns, section in [(education_patterns, 'education'), (work_patterns, 'work_experience')]:
            for pattern, _ in patterns:
                m = re.match(pattern, field_id)
                if m:
                    base = m.group(1)
                    # Reconstruct with {N}
                    normalized = field_id[:m.start(2)] + '{N}' + field_id[m.end(2):]
                    return (normalized, section)

        return (field_id, '')

    def get_stats(self) -> dict:
        """Get schema statistics for display."""
        stats = {}
        for ats, schema in self.data.items():
            stats[ats] = {
                "fill_count": schema.get("fill_count", 0),
                "field_count": len(schema.get("fields", {})),
                "last_updated": schema.get("last_updated", ""),
                "repeatable_sections": list(schema.get("repeatable_field_ids", {}).keys()),
            }
        return stats


# ═══════════════════════════════════════════════════════════════════════════
# KNOWLEDGE BASE - Experience snippets + common answers for AI context
# ═══════════════════════════════════════════════════════════════════════════

class KnowledgeBase:
    """Knowledge base with experience snippets and common answers."""

    def __init__(self, path: Path = KNOWLEDGE_BASE_PATH):
        self.path = path
        try:
            self.data = json.loads(path.read_text())
        except Exception:
            self.data = {}
        self.snippets = self.data.get("experience_snippets", {})
        self.common_answers = self.data.get("common_answers", {})
        self.skills = self.data.get("skills", {})
        self.achievements = self.data.get("achievements", [])
        print(f"   📚 KnowledgeBase loaded: {len(self.snippets)} snippets, {len(self.common_answers)} common answers")

    def find_relevant_snippets(self, question: str) -> List[str]:
        """Find relevant experience snippets for a question."""
        question_lower = question.lower()
        found = []

        for keyword, snippet in self.snippets.items():
            # Check if keyword appears in question
            if keyword.lower() in question_lower:
                found.append(snippet)
            # Also check individual words from multi-word keywords
            elif len(keyword.split()) > 1:
                words = keyword.lower().split()
                if any(w in question_lower for w in words if len(w) > 3):
                    found.append(snippet)

        return found[:3]  # Max 3 snippets

    def get_context_for_question(self, question: str) -> str:
        """Get formatted context with relevant snippets for AI prompt."""
        snippets = self.find_relevant_snippets(question)
        if snippets:
            return "Relevant experience:\n" + "\n".join(f"- {s}" for s in snippets)
        return ""

    def find_common_answer(self, question: str) -> Optional[str]:
        """Find pre-written answer for common questions (salary, why interested, etc.)."""
        question_lower = question.lower()

        for answer_key, answer_data in self.common_answers.items():
            keywords = answer_data.get("keywords", [])
            for kw in keywords:
                if kw.lower() in question_lower:
                    return answer_data.get("answer")

        return None


# ═══════════════════════════════════════════════════════════════════════════
# OLLAMA HELPER - Local LLM for custom questions
# ═══════════════════════════════════════════════════════════════════════════

class OllamaHelper:
    """Local LLM using llava:7b for generating answers to custom questions."""
    
    OLLAMA_URL = "http://localhost:11434/api/generate"
    MODEL = "llava:7b"
    
    PROMPT_TEMPLATE = """Answer this job application question based on the candidate profile.

CANDIDATE:
{profile_context}

QUESTION: {question}
{options_text}

RULES:
1. "How many years of experience" -> answer with a number like "15+" 
2. Questions about specific software (NetSuite, SAP, Salesforce, Oracle):
   - Answer "No" unless that EXACT tool is mentioned in candidate profile above
3. Multiple choice -> reply with EXACT option text only
4. Yes/no -> reply "Yes" or "No" only
5. Be concise - max 5 words

ANSWER:"""

    def __init__(self):
        self._available = None
    
    @property
    def available(self) -> bool:
        if self._available is None:
            try:
                resp = requests.get("http://localhost:11434/api/tags", timeout=2)
                self._available = resp.status_code == 200
            except:
                self._available = False
        return self._available
    
    def generate(self, question: str, profile_context: str, options: List[str] = None) -> Optional[str]:
        if not self.available:
            return None
        
        options_text = ""
        if options:
            options_text = f"OPTIONS (choose one): {options}"
        
        prompt = self.PROMPT_TEMPLATE.format(
            profile_context=profile_context,
            question=question,
            options_text=options_text
        )
        
        try:
            resp = requests.post(
                self.OLLAMA_URL,
                json={
                    "model": self.MODEL,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.1, "num_predict": 50}
                },
                timeout=60
            )
            
            if resp.status_code == 200:
                answer = resp.json().get("response", "").strip()
                answer = answer.split("\n")[0].strip()
                answer = re.sub(r'^[*\-]\s*', '', answer)
                answer = re.sub(r'^(Answer:|ANSWER:|A:)\s*', '', answer)
                return answer
        except Exception as e:
            print(f"   ⚠️ Ollama error: {e}")
        return None
    
    def match_option(self, answer: str, options: List[str]) -> Optional[str]:
        if not options:
            return answer
        answer_lower = answer.lower().strip()
        for opt in options:
            if opt.lower() == answer_lower:
                return opt
        for opt in options:
            if answer_lower in opt.lower() or opt.lower() in answer_lower:
                return opt
        return options[0] if options else answer


# ═══════════════════════════════════════════════════════════════════════════
# AI HELPER - Claude API for verification
# ═══════════════════════════════════════════════════════════════════════════

class AIHelper:
    """AI Helper using Claude API for verification."""
    
    def __init__(self):
        self._vision_ai = None
    
    @property
    def vision_ai(self):
        """Lazy-load Vision AI."""
        if self._vision_ai is None:
            from .vision_ai import VisionAI
            self._vision_ai = VisionAI()
        return self._vision_ai
    
    @property
    def available(self) -> bool:
        """Check if Claude API is configured."""
        return self.vision_ai.available
    
    def generate(self, question: str, context: str, max_length: int = 150) -> str:
        """Generate answer for custom question using Claude."""
        if not self.available:
            print("   ⚠️ Claude API not configured. Set ANTHROPIC_API_KEY.")
            return ""
        
        result = self.vision_ai.generate_custom_answer(
            question=question,
            profile_context=context,
            max_words=max_length // 5
        )
        
        if result.get("success") and result.get("answer"):
            return result["answer"]
        
        if result.get("error"):
            print(f"   ⚠️ Claude error: {result['error']}")
        
        return ""
    
    def choose_option(self, question: str, options: List[str], context: str) -> Optional[str]:
        """Choose best option from list using Claude."""
        if not options:
            return None
        
        if not self.available:
            print("   ⚠️ Claude API not configured. Set ANTHROPIC_API_KEY.")
            return None
        
        try:
            prompt = f"""Job application dropdown:
Question: {question}
Options: {options}
Profile: {context}

Which option should be selected? Return ONLY the exact option text, nothing else."""
            
            response = self.vision_ai.client.messages.create(
                model=self.vision_ai.config.model,
                max_tokens=100,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            
            answer = response.content[0].text.strip()
            
            # Find matching option (exact match first)
            for opt in options:
                if opt.lower() == answer.lower():
                    return opt
            
            # Partial match
            for opt in options:
                if opt.lower() in answer.lower() or answer.lower() in opt.lower():
                    return opt
            
            # Word overlap match
            answer_words = set(answer.lower().split())
            for opt in options:
                opt_words = set(opt.lower().split())
                if answer_words & opt_words:
                    return opt
                    
        except Exception as e:
            print(f"   ⚠️ Claude error: {e}")
        
        return None
    
    def analyze_field_screenshot(self, screenshot_path: str, field_description: str = "") -> Dict[str, Any]:
        """Analyze form field from screenshot using Claude Vision."""
        if not self.available:
            return {"success": False, "error": "Claude API not configured"}
        return self.vision_ai.analyze_field(screenshot_path, field_description)
    
    def verify_field_filled(self, screenshot_path: str, expected_value: str, field_label: str = "") -> Dict[str, Any]:
        """Verify field was filled correctly using Claude Vision."""
        if not self.available:
            return {"success": False, "verified": False, "error": "Claude API not configured"}
        return self.vision_ai.verify_field_filled(screenshot_path, expected_value, field_label)
    
    def analyze_full_form(self, screenshot_path: str) -> Dict[str, Any]:
        """Analyze entire form from screenshot."""
        if not self.available:
            return {"success": False, "error": "Claude API not configured"}
        return self.vision_ai.analyze_form(screenshot_path)

class FormFillerV5:
    """
    Universal Form Filler V5
    
    Usage:
        # Pre-flight check
        filler = FormFillerV5()
        report = filler.analyze("https://greenhouse.io/...")
        print(report.summary())
        
        # Fill form
        filler.fill("https://greenhouse.io/...", mode=FillMode.INTERACTIVE)
    """
    
    # Repeatable sections configuration (Work Experience, Education)
    # Greenhouse uses patterns like: company-name-0, title-0, school--0, degree--0
    REPEATABLE_SECTIONS = {
        'work_experience': {
            'profile_key': 'work_experience',
            'add_button_selectors': [
                'button:has-text("Add another")',
                'a:has-text("Add another")',
                '.add-section a', '.add-section button',
                '#add_work_experience',
            ],
            'section_text': 'work experience',
            'field_patterns': {
                'company-name-{N}': 'company',
                'title-{N}': 'title',
                'start-date-month-{N}': 'start_month',
                'start-date-year-{N}': 'start_year',
                'end-date-month-{N}': 'end_month',
                'end-date-year-{N}': 'end_year',
                'current-role-{N}_1': 'current',
            },
            'skip_end_date_if_current': True,
        },
        'education': {
            'profile_key': 'education',
            'add_button_selectors': [
                'button:has-text("Add another")',
                'a:has-text("Add another")',
                '.add-section a', '.add-section button',
                '#add_education',
            ],
            'section_text': 'education',
            # Greenhouse education uses two ID formats:
            # - Select2: s2id_education_school_name_{N}, s2id_education_degree_{N}
            # - Legacy: school--{N}, degree--{N}
            'field_patterns': {
                'education_school_name_{N}': 'school',
                'education_degree_{N}': 'degree',
                'education_discipline_{N}': 'discipline',
                # React Select date fields (modern Greenhouse forms)
                'start-month--{N}': 'start_month',
                'start-year--{N}': 'start_year',
                'end-month--{N}': 'end_month',
                'end-year--{N}': 'end_year',
            },
            # Date fields fallback — legacy name-based selectors
            'date_selectors': {
                'start_month': "input[name='job_application[educations][][start_date][month]']",
                'start_year': "input[name='job_application[educations][][start_date][year]']",
                'end_month': "input[name='job_application[educations][][end_date][month]']",
                'end_year': "input[name='job_application[educations][][end_date][year]']",
            },
            'legacy_field_patterns': {
                'school--{N}': 'school',
                'degree--{N}': 'degree',
                'discipline--{N}': 'discipline',
                'start-date-month--{N}': 'start_month',
                'start-date-year--{N}': 'start_year',
                'end-date-month--{N}': 'end_month',
                'end-date-year--{N}': 'end_year',
            },
            'skip_end_date_if_current': False,
        }
    }
    
    def __init__(self, browser_mode: BrowserMode = BrowserMode.PERSISTENT):
        self.browser_mode = browser_mode
        self.profile = Profile()
        self.learned_db = LearnedDB()
        self.schema_db = FormSchemaDB()
        self.kb = KnowledgeBase()
        self.ai = AIHelper()
        self.ollama = OllamaHelper()
        self.logger = FormLogger()

        self.browser: Optional[BrowserManager] = None
        self.page: Optional[Page] = None
        self.fields: List[FormField] = []
        self._seen_selectors: set = set()
        self._ats_type: str = ""  # Detected ATS type for schema lookups
        self._section_filled_ids: set = set()  # Element IDs filled by repeatable section handler
    
    # ─────────────────────────────────────────────────────────────────────
    # PUBLIC API
    # ─────────────────────────────────────────────────────────────────────
    
    def analyze(self, url: str) -> FillReport:
        """
        Pre-flight analysis: scan form, find answers, generate readiness report.
        Does NOT fill any fields.
        """
        with BrowserManager(mode=self.browser_mode) as browser:
            self.browser = browser
            self.page = browser.page
            
            browser.goto(url)
            browser.wait_for_stable()

            # Detect ATS for schema lookups
            self._ats_type = self._detect_ats(url)

            # Scan and analyze
            self._scan_fields()
            self._resolve_all_answers()

            # Generate report
            return self._generate_report(url)
    
    def fill(self, url: str, mode: FillMode = FillMode.INTERACTIVE, keep_open: bool = False) -> FillReport:
        """
        Fill form with specified mode.
        Includes re-scan logic for dynamic forms (fields that appear after selection).

        Args:
            keep_open: If True, keeps browser open for manual review (CDP: just disconnect,
                      PERSISTENT/FRESH: wait for ENTER)
        """
        with BrowserManager(mode=self.browser_mode) as browser:
            self.browser = browser
            self.page = browser.page

            browser.goto(url)
            browser.wait_for_stable()

            # Start logging session
            company = self._extract_company_from_url(url)
            self.logger.start_session(url=url, company=company)

            # Detect ATS type early for schema lookups
            self._ats_type = self._detect_ats(url)
            schema = self.schema_db.get_schema(self._ats_type)
            if schema:
                print(f"   📐 Schema loaded: {self._ats_type} ({len(schema.get('fields', {}))} fields, fill #{schema.get('fill_count', 0)})")
            else:
                print(f"   📐 No schema yet for {self._ats_type} — will learn from this fill")

            # Wait for iframes to load (Greenhouse, Lever forms are in iframes)
            self._wait_for_iframes()
            
            # Extract job info BEFORE Apply click (full JD available on description page)
            self._extract_job_info()
            jd_before_apply = getattr(self, 'job_description', '') or ''

            # Try to find and click Apply button if on job description page
            self._find_and_click_apply_button()

            # Handle login page if needed
            if self._handle_login_page():
                # After login, wait and rescan
                browser.wait_for_stable()

            # Re-extract job info after Apply (may get additional info from form page)
            self._extract_job_info()
            # Keep the longer JD (description page usually has more than form page)
            jd_after_apply = getattr(self, 'job_description', '') or ''
            if len(jd_before_apply) > len(jd_after_apply):
                self.job_description = jd_before_apply
                print(f"   📄 Using pre-Apply JD: {len(self.job_description)} chars (form page had {len(jd_after_apply)})")

            # Initial scan, prescan dropdowns, resolve, fill
            self._scan_fields()
            self._prescan_all_options()
            self._resolve_all_answers()
            
            if mode == FillMode.PRE_FLIGHT:
                return self._generate_report(url)
            
            # Fill repeatable sections first (work experience, education)
            try:
                self.fill_all_repeatable_sections()
            except Exception as e:
                print(f"⚠️ Repeatable sections error (continuing): {e}")

            # Mark fields already filled by section handler so main loop skips them
            self._mark_section_filled_fields()

            # Fill fields with re-scan for dynamic forms
            # Loop until no new fields appear
            max_iterations = 5
            for iteration in range(max_iterations):
                fields_before = len(self.fields)
                
                # Fill current fields
                self._fill_all_fields(mode)
                
                # Wait for potential new fields to appear
                time.sleep(0.5)
                browser.wait_for_stable()
                
                # Re-scan for new fields
                new_fields = self._scan_for_new_fields()
                
                if new_fields:
                    # Filter out fields already filled by repeatable section handler
                    # or belonging to empty education/work_experience slots
                    unfilled_new = []
                    edu_count = len(self.profile.data.get('education', []))
                    for f in new_fields:
                        f_eid = f.element_id or ""
                        f_label = f.label.lower()

                        # Fast path: field was explicitly filled by section handler (tracked by ID)
                        if f_eid and f_eid in self._section_filled_ids:
                            print(f"   ⏭️ {f.label[:40]} (already filled by section handler)")
                            f.status = FillStatus.VERIFIED
                            f.answer_source = AnswerSource.PROFILE
                            try:
                                context2 = getattr(self, '_active_frame', self.page)
                                dom_val = self._read_section_field_value(context2, f)
                                if dom_val:
                                    f.answer = dom_val[:50]
                            except:
                                pass
                            continue

                        # Check if field already has a value (filled by repeatable sections)
                        try:
                            context = getattr(self, '_active_frame', self.page)
                            el = context.query_selector(f.selector)
                            if el:
                                # For Select2/autocomplete, check the display text
                                tag = el.evaluate("e => e.tagName.toLowerCase()")
                                if tag == 'input' and el.get_attribute('type') == 'hidden':
                                    val = el.input_value()
                                else:
                                    val = el.input_value() if tag in ('input', 'textarea') else el.evaluate("e => e.value || ''")
                                if val and val.strip():
                                    print(f"   ⏭️ {f.label[:40]} (already filled by section handler)")
                                    f.status = FillStatus.VERIFIED
                                    f.answer = val[:50]
                                    f.answer_source = AnswerSource.PROFILE
                                    continue

                                # Skip empty education/repeatable section fields
                                # (Greenhouse pre-renders extra slots, but we only fill as many as profile has)
                                is_section_field = f_label in ('school', 'degree', 'discipline',
                                                                'start date month', 'start date year',
                                                                'end date month', 'end date year') or \
                                    any(kw in f_eid for kw in
                                        ('education_school', 'education_degree', 'education_discipline',
                                         'school--', 'degree--', 'discipline--',
                                         'start-month--', 'start-year--', 'end-month--', 'end-year--'))
                                if is_section_field:
                                    print(f"   ⏭️ {f.label[:40]} (extra education slot, no profile data)")
                                    f.status = FillStatus.SKIPPED
                                    continue
                        except:
                            pass
                        unfilled_new.append(f)

                    if unfilled_new:
                        print(f"\n🔄 Found {len(unfilled_new)} new unfilled fields after filling")
                        for f in unfilled_new:
                            print(f"   + {f.label[:40]}")

                        # Resolve answers for new fields
                        for f in unfilled_new:
                            self._resolve_field_answer(f)
                    else:
                        print(f"\n✅ {len(new_fields)} dynamic fields already filled by section handler")
                        break

                    # Continue loop to fill new fields
                else:
                    # No new fields, we're done
                    break
            
            # Blur all fields to trigger validation
            self._blur_all_fields()
            self._validate_all_fields()
            
            # Feedback loop: save verified AI answers to learned DB
            self._save_verified_ai_answers()

            # Save form schema for this ATS type
            if self._ats_type:
                self.schema_db.save_schema_from_fill(self._ats_type, self.fields, url)

            # Save duration before end_session clears it
            self._fill_duration = 0.0
            if self.logger.start_time:
                from datetime import datetime as _dt
                self._fill_duration = (_dt.now() - self.logger.start_time).total_seconds()

            # End logging session
            log_path = self.logger.end_session(status="completed")
            if log_path:
                print(f"   📄 Log saved: {log_path}")

            # Generate report + save JSON/PDF (before browser might close)
            report = self._generate_report(url)
            self._save_application_report(report)

            # Keep browser open for review
            if keep_open or mode == FillMode.INTERACTIVE:
                if self.browser_mode == BrowserMode.CDP:
                    # CDP: Chrome stays open after disconnect — just inform user
                    print("\n👀 Form filled! Browser stays open for review.")
                    print("   Close the tab manually when done.")
                else:
                    # PERSISTENT/FRESH: Need to hold the session open
                    # Save form screenshots before waiting
                    try:
                        self.page.evaluate('window.scrollTo(0, 0)')
                        time.sleep(0.3)
                        self.page.screenshot(path='/tmp/form_screenshot_top.png', full_page=False)
                        self.page.evaluate('window.scrollTo(0, document.body.scrollHeight / 2)')
                        time.sleep(0.3)
                        self.page.screenshot(path='/tmp/form_screenshot_mid.png', full_page=False)
                        self.page.evaluate('window.scrollTo(0, document.body.scrollHeight)')
                        time.sleep(0.3)
                        self.page.screenshot(path='/tmp/form_screenshot_bot.png', full_page=False)
                        print("   📸 Screenshots: /tmp/form_screenshot_*.png")
                    except:
                        pass
                    print("\n👀 Review the form in browser.")
                    print("   Press ENTER when done...")
                    try:
                        input()
                    except:
                        # Fallback: keep open for 5 minutes when input() not available
                        print("   (Browser will stay open for 5 minutes)")
                        time.sleep(300)

            print(report.detailed_report())
            return report

    def _scan_for_new_fields(self) -> List[FormField]:
        """
        Scan for fields that appeared after initial scan.
        Returns list of new fields not seen before.
        """
        new_fields = []
        elements = self.page.query_selector_all("input, select, textarea")
        
        for el in elements:
            field = self._detect_field(el)
            if field and field.selector not in self._seen_selectors:
                self._seen_selectors.add(field.selector)
                self.fields.append(field)
                new_fields.append(field)
        
        return new_fields
    
    # ─────────────────────────────────────────────────────────────────────
    # LAYER 0: FIND AND CLICK APPLY BUTTON
    # ─────────────────────────────────────────────────────────────────────
    
    def _wait_for_iframes(self, timeout=8):
        """
        Wait for iframes to load (Greenhouse, Lever, etc.).
        Many ATS embed their forms in iframes that load async.
        """
        import time
        
        start = time.time()
        initial_frames = len(self.page.frames)
        
        # Wait up to timeout seconds for more frames to appear
        while time.time() - start < timeout:
            current_frames = len(self.page.frames)
            
            # Check if we have Greenhouse or Lever iframe
            for frame in self.page.frames:
                if frame.url and any(ats in frame.url for ats in ['greenhouse', 'lever', 'workday', 'icims']):
                    print(f"   \u2705 Found ATS iframe: {frame.url[:50]}")
                    time.sleep(1)  # Extra wait for iframe content
                    return
            
            # If frames increased, wait a bit more
            if current_frames > initial_frames:
                time.sleep(1)
                continue
            
            time.sleep(0.5)
        
        print(f"   \ud83d\udcc4 {len(self.page.frames)} frames loaded")
    
    def _find_and_click_apply_button(self):
        """
        Find and click Apply button on job description pages.
        Many ATS (iCIMS, Workday, etc.) show job description first,
        then require clicking Apply to open the form.
        Supports both main page and iframes.
        """
        import time
        
        # Check if we're on a job description page (no form fields visible)
        initial_fields = self.page.query_selector_all("input[type='text'], input[type='email'], select, textarea")
        visible_fields = [f for f in initial_fields if f.is_visible()]
        
        if len(visible_fields) > 3:
            print("\n✅ Already on application form")
            return
        
        print("\n🔎 Looking for Apply button...")
        
        apply_selectors = [
            "a.iCIMS_ApplyButton",
            "a[title='Apply for this job online']",
            "a.iCIMS_Action_Button.iCIMS_ApplyButton",
            "a.iCIMS_PrimaryButton",
            "a#apply_button",
            "a.postings-btn",
            "button[data-automation-id='jobPostingApplyButton']",
            "a:has-text('Apply')",
            "button:has-text('Apply')",
            "a:has-text('Apply Now')",
            "a:has-text('Apply for this job')",
            "a[class*='apply']",
            "button[class*='apply']",
        ]
        
        # First try main page
        if self._try_click_apply_in_context(self.page, apply_selectors, "main page"):
            return
        
        # Then try each iframe
        frames = self.page.frames
        print(f"   📝 Checking {len(frames)} frames for Apply button...")
        
        for i, frame in enumerate(frames):
            if frame == self.page.main_frame:
                continue
            
            frame_url = frame.url[:50] if frame.url else "(empty)"
            if self._try_click_apply_in_context(frame, apply_selectors, f"frame {i}: {frame_url}"):
                return
        
        print("   ⚠️ No Apply button found, proceeding with current page")
    
    def _try_click_apply_in_context(self, context, selectors, context_name):
        """Try to find and click Apply button in a given context (page or frame)."""
        import time
        
        for selector in selectors:
            try:
                btn = context.query_selector(selector)
                if btn and btn.is_visible():
                    btn_text = btn.inner_text().strip()[:50]
                    print(f"   ✅ Found Apply button in {context_name}: '{btn_text}'")
                    btn.click()
                    print(f"   🖱️ Clicked Apply button")
                    time.sleep(2)
                    self.browser.wait_for_stable()
                    print(f"   📄 Current URL: {self.page.url[:80]}...")
                    return True
            except:
                continue
        
        # Try text-based search
        try:
            all_clickables = context.query_selector_all("a, button")
            for el in all_clickables:
                try:
                    if not el.is_visible():
                        continue
                    text = el.inner_text().strip().lower()
                    if any(kw in text for kw in ['apply', 'submit application', 'start application']):
                        if 'applied' not in text and "don't apply" not in text:
                            print(f"   ✅ Found Apply button in {context_name} (text): '{text[:50]}'")
                            el.click()
                            print(f"   🖱️ Clicked Apply button")
                            time.sleep(2)
                            self.browser.wait_for_stable()
                            return True
                except:
                    continue
        except:
            pass
        
        return False

    def _handle_login_page(self):
        """
        Handle login/authentication pages.
        Supports: Google OAuth, LinkedIn OAuth, Email login.
        """
        import time
        
        # Check if we're on a login page
        url = self.page.url.lower()
        if 'login' not in url and 'signin' not in url and 'auth' not in url:
            return False
        
        print("\n🔐 Login page detected...")
        
        # First try to find Google/LinkedIn login buttons (preferred)
        for frame in self.page.frames:
            try:
                # Look for Google login button
                google_btn = frame.query_selector(
                    "a[href*='google'], button[data-provider='google'], "
                    "a:has-text('Google'), button:has-text('Google'), "
                    "a:has-text('Sign in with Google'), button:has-text('Sign in with Google'), "
                    "a[class*='google'], button[class*='google'], "
                    "div[data-provider='google']"
                )
                if google_btn and google_btn.is_visible():
                    print("   🔵 Found Google login button")
                    print("   👉 Clicking Google login...")
                    google_btn.click()
                    time.sleep(3)
                    self.browser.wait_for_stable()
                    
                    # Wait for Google OAuth popup or redirect
                    print("   👉 Please complete Google sign-in in the browser...")
                    print("   Press ENTER when logged in...")
                    try:
                        input()
                    except:
                        time.sleep(30)
                    
                    self.browser.wait_for_stable()
                    return True
                
                # Look for LinkedIn login button
                linkedin_btn = frame.query_selector(
                    "a[href*='linkedin'], button[data-provider='linkedin'], "
                    "a:has-text('LinkedIn'), button:has-text('LinkedIn'), "
                    "a[class*='linkedin'], button[class*='linkedin']"
                )
                if linkedin_btn and linkedin_btn.is_visible():
                    print("   🔵 Found LinkedIn login button")
                    print("   👉 Clicking LinkedIn login...")
                    linkedin_btn.click()
                    time.sleep(3)
                    self.browser.wait_for_stable()
                    
                    print("   👉 Please complete LinkedIn sign-in...")
                    print("   Press ENTER when logged in...")
                    try:
                        input()
                    except:
                        time.sleep(30)
                    
                    self.browser.wait_for_stable()
                    return True
            except:
                continue
        
        # Fallback: email login
        email_value = self.profile.get('personal.email') or ''
        if not email_value or email_value == 'YOUR_EMAIL':
            print("   ⚠️ No email in profile. Please set 'email' in your profile.")
            return False
        
        for frame in self.page.frames:
            try:
                email_input = frame.query_selector("input[type='email'], input[name*='email'], input[name*='loginName']")
                if email_input and email_input.is_visible():
                    print(f"   📧 Found email field, entering: {email_value}")
                    email_input.fill(email_value)
                    time.sleep(0.5)
                    
                    # Look for submit button
                    submit_btn = frame.query_selector("input[type='submit'], button[type='submit'], button:has-text('Continue'), button:has-text('Next')")
                    if submit_btn and submit_btn.is_visible():
                        print("   👉 Clicking submit...")
                        submit_btn.click()
                        time.sleep(2)
                        self.browser.wait_for_stable()
                    
                    # Check for captcha
                    captcha = frame.query_selector("[class*='captcha'], [class*='hcaptcha'], [class*='recaptcha'], iframe[src*='captcha']")
                    if captcha:
                        print("\n   🧩 CAPTCHA detected!")
                        print("   👉 Please solve the captcha in the browser...")
                        print("   Press ENTER when done...")
                        try:
                            input()
                        except:
                            time.sleep(30)
                        
                        self.browser.wait_for_stable()
                    
                    return True
            except:
                continue
        
        return False

    # ─────────────────────────────────────────────────────────────────────
    # LAYER 1: DETECTION
    # ─────────────────────────────────────────────────────────────────────
    
    def _scan_fields(self):
        """Scan all form fields on page and all iframes."""
        print("\n🔍 Scanning form fields...")
        self.fields = []
        self._seen_selectors = set()
        self._section_filled_ids = set()  # Reset for new form scan
        self._active_frame = self.page  # Track which frame has the form

        # Scan main page first
        print("   Scanning main page...", flush=True)
        elements = self.page.query_selector_all("input, select, textarea")
        main_count = self._scan_elements(elements, "main")
        print(f"   Main page: {main_count} fields", flush=True)

        # Scan all iframes (important for Greenhouse, Lever embedded forms)
        frames = self.page.frames
        if len(frames) > 1:
            print(f"   📄 Checking {len(frames)} frames...", flush=True)

            for i, frame in enumerate(frames):
                if frame == self.page.main_frame:
                    continue  # Already scanned

                try:
                    frame_url = frame.url[:40] if frame.url else "(empty)"
                    print(f"      Frame {i}: {frame_url}...", flush=True)
                    elements = frame.query_selector_all("input, select, textarea")
                    count = self._scan_elements(elements, f"frame[{i}]")

                    if count > 0:
                        print(f"      ✅ Frame {i}: {count} fields")
                        # Remember which frame has the form
                        if count > main_count:
                            self._active_frame = frame
                except Exception as e:
                    print(f"      ⚠️ Frame {i} error: {e}", flush=True)
                    continue

        # Greenhouse file upload fieldsets (Resume + Cover Letter)
        # Greenhouse hides <input type="file"> — detect by fieldset structure instead
        self._scan_greenhouse_file_fieldsets()

        # Print scan summary
        print(f"   Total fields: {len(self.fields)}", flush=True)
        for f in self.fields:
            print(f"      {f.field_type.value:12} | {f.label[:40]:<40} | {f.selector[:30]}", flush=True)
    
    def _scan_greenhouse_file_fieldsets(self):
        """Detect Greenhouse file upload fieldsets (Resume/CV, Cover Letter).

        Greenhouse hides <input type='file'> inside S3 upload forms.
        The visible UI is a fieldset with an 'Attach' button.
        We detect by fieldset ID and create FILE fields pointing to the fieldset.
        """
        context = getattr(self, '_active_frame', self.page)

        fieldsets = [
            ("resume_fieldset", "Resume/CV"),
            ("cover_letter_fieldset", "Cover Letter"),
        ]

        for fieldset_id, label in fieldsets:
            try:
                fieldset = context.query_selector(f"#{fieldset_id}")
                if not fieldset:
                    continue

                # Check if we already have a file field for this fieldset
                # (might have been detected from the hidden <input type="file">)
                already_found = False
                fieldset_key = fieldset_id.replace("_fieldset", "")  # "resume" or "cover_letter"
                for f in self.fields:
                    if f.field_type == FieldType.FILE:
                        sel_lower = (f.selector or "").lower()
                        fl_lower = (f.label or "").lower()
                        fid_lower = (f.element_id or "").lower()
                        if fieldset_key in sel_lower or \
                           fieldset_key in fid_lower or \
                           label.lower() in fl_lower:
                            already_found = True
                            # Update label and selector if they were generic
                            if f.label.lower() in ("file", ""):
                                f.label = label
                                print(f"      📎 Updated file field label: '{label}'")
                            # Also update element_id to point to fieldset for Attach button lookup
                            if not f.element_id or f.element_id.lower() == "file":
                                f.element_id = fieldset_id
                            break

                if already_found:
                    continue

                # Also check: is the first generic file field likely this one?
                # On Greenhouse, first [name='file'] is Resume, but we need to
                # upgrade it rather than create a duplicate
                if fieldset_key == "resume":
                    for f in self.fields:
                        if f.field_type == FieldType.FILE and \
                           f.label.lower() in ("file", "") and \
                           "[name='file']" in (f.selector or "").lower():
                            # This generic file field is likely the Resume — upgrade it
                            f.label = label
                            f.element_id = fieldset_id
                            f.selector = f"#{fieldset_id} input[type='file']"
                            self._seen_selectors.add(f.selector)
                            print(f"      📎 Upgraded generic file → {label} ({f.selector})")
                            already_found = True
                            break

                if already_found:
                    continue

                # Check there's an Attach button (confirms this is a file upload fieldset)
                attach_btn = context.query_selector(
                    f"#{fieldset_id} button[data-source='attach']"
                )
                if not attach_btn:
                    continue

                # Find the hidden <input type="file"> inside the S3 form
                file_input = context.query_selector(
                    f"#{fieldset_id} input[type='file']"
                )

                # Create selector — prefer the file input if found, else use fieldset
                if file_input:
                    selector = f"#{fieldset_id} input[type='file']"
                else:
                    selector = f"#{fieldset_id}"

                if selector in self._seen_selectors:
                    continue

                # Create FILE field
                field = FormField(
                    selector=selector,
                    element_id=fieldset_id,
                    name="file",
                    label=label,
                    field_type=FieldType.FILE,
                    detection_method=DetectionMethod.HTML,
                    html_tag="fieldset",
                    input_type="file",
                    required=(fieldset_id == "resume_fieldset"),  # Resume is usually required
                )

                self._seen_selectors.add(selector)
                self.fields.append(field)
                print(f"      📎 Greenhouse file fieldset: {label} → {selector}")

            except Exception as e:
                print(f"      ⚠️ Fieldset scan error ({fieldset_id}): {e}")

    def _scan_elements(self, elements, source):
        """Scan elements from a specific source (main or frame)."""
        count = 0
        for el in elements:
            field = self._detect_field(el)
            if field and field.selector not in self._seen_selectors:
                self._seen_selectors.add(field.selector)
                self.fields.append(field)
                count += 1
        return count
        
        # Print summary by type
        type_counts = {}
        for f in self.fields:
            type_counts[f.field_type.value] = type_counts.get(f.field_type.value, 0) + 1
        
        print(f"   Found {len(self.fields)} fields: {type_counts}")
    
    def _detect_field(self, el: ElementHandle) -> Optional[FormField]:
        """Detect field type and metadata."""
        try:
            # Skip hidden/invisible
            if not el.is_visible():
                return None
            
            # Basic attributes
            tag = el.evaluate("e => e.tagName.toLowerCase()")
            input_type = el.get_attribute("type") or "text"
            el_id = el.get_attribute("id") or ""
            el_name = el.get_attribute("name") or ""
            
            # Skip system fields
            if input_type in ("hidden", "submit", "button"):
                return None
            
            # Build selector (escape special CSS chars in IDs like question_123[])
            if el_id:
                escaped_id = re.sub(r'([\[\](){}!@#$%^&*+=|~`<>?,/\\])', r'\\\1', el_id)
                selector = f"#{escaped_id}"
            elif el_name:
                # For file inputs, use parent form ID to disambiguate (Resume vs Cover Letter)
                if input_type == "file":
                    form_id = el.evaluate("e => e.form ? e.form.id : ''")
                    if form_id:
                        selector = f"#{form_id} input[type='file']"
                    else:
                        selector = f"[name='{el_name}']"
                else:
                    selector = f"[name='{el_name}']"
            else:
                return None
            
            # Get label
            label = self._find_label(el, el_id)
            
            # Get current value
            current_value = self._get_value(el, tag, input_type)
            
            # Required?
            required = el.get_attribute("required") is not None or \
                       el.get_attribute("aria-required") == "true"
            
            # Detect type
            field_type, detection_method = self._detect_type(el, tag, input_type)
            
            # Get options for select fields only during scan
            # Autocomplete options will be discovered in prescan phase (faster)
            options = []
            if field_type == FieldType.SELECT:
                try:
                    options = el.evaluate(
                        "e => Array.from(e.options).map(o => o.text).filter(t => t && t !== 'Select...')"
                    )
                except:
                    pass
            # NOTE: autocomplete options are NOT probed here - they're done in _prescan_all_options()
            
            # ── DOM constraints (placeholder, maxlength, pattern) ──
            # Engine reads these to understand expected format BEFORE filling
            placeholder = el.get_attribute("placeholder") or ""
            maxlength_str = el.get_attribute("maxlength") or ""
            maxlength = int(maxlength_str) if maxlength_str.isdigit() else 0
            html_pattern = el.get_attribute("pattern") or ""

            return FormField(
                selector=selector,
                element_id=el_id,
                name=el_name,
                label=label,
                field_type=field_type,
                detection_method=detection_method,
                html_tag=tag,
                input_type=input_type,
                options=options,
                required=required,
                current_value=current_value,
                placeholder=placeholder,
                maxlength=maxlength,
                pattern=html_pattern,
            )
            
        except Exception as e:
            return None
    
    def _find_label(self, el: ElementHandle, el_id: str) -> str:
        """
        Find label text for field using multiple strategies.

        Priority:
        1. Standard label[for] attribute (in active frame + page)
        1b. Parent <label> direct text (Greenhouse wraps fields in <label> tags)
        1c. div.field > label sibling (Greenhouse custom questions in div.field)
        2. aria-label / placeholder
        3. Context Discovery (traverse DOM for nearby text)
        4. Name/ID as fallback
        """
        label = ""
        el_name = el.get_attribute("name") or ""

        # Strategy 1: By for attribute - search in element's own frame first
        if el_id:
            # Try element's owner frame first (handles iframes)
            frame = el.owner_frame() or self.page
            label_el = frame.query_selector(f"label[for='{el_id}']")
            if not label_el and frame != self.page:
                label_el = self.page.query_selector(f"label[for='{el_id}']")
            if label_el:
                label = label_el.inner_text().strip()

        # Strategy 1b: Parent <label> direct text nodes
        # In Greenhouse, fields (inputs, textareas, Select2) are inside a <label> tag
        # whose direct text nodes contain the actual question text.
        # e.g. <label>Question text? *<span class="asterisk">*</span><br><input...></label>
        # This works for ALL field types, not just Select2.
        if not label:
            try:
                parent_label_text = el.evaluate('''el => {
                    // For Select2 inputs, start from the container
                    let start = el.closest('.select2-container') || el;

                    // Walk up to find a <label> parent (max 4 levels)
                    let parent = start.parentElement;
                    for (let i = 0; i < 4 && parent; i++) {
                        if (parent.tagName === 'LABEL') {
                            // Extract only direct text nodes (not children text)
                            let text = '';
                            for (const node of parent.childNodes) {
                                if (node.nodeType === 3) { // TEXT_NODE
                                    text += node.textContent.trim() + ' ';
                                }
                            }
                            text = text.trim();
                            // Filter out noise: just asterisks, very short text
                            if (text && text !== '*' && text.length > 3) {
                                return text;
                            }
                        }
                        parent = parent.parentElement;
                    }
                    return '';
                }''')
                if parent_label_text and len(parent_label_text) > 3:
                    label = parent_label_text
            except Exception:
                pass

        # Strategy 1c: div.field > label sibling
        # Some Greenhouse forms use: <div class="field"><label>Question</label><input></div>
        if not label:
            try:
                sibling_label = el.evaluate('''el => {
                    // Walk up to find div.field
                    let parent = el.parentElement;
                    for (let i = 0; i < 5 && parent; i++) {
                        if (parent.classList && parent.classList.contains('field')) {
                            // Get first direct label child that doesn't contain our input
                            const labels = parent.querySelectorAll(':scope > label');
                            for (const lbl of labels) {
                                if (!lbl.contains(el)) {
                                    let text = '';
                                    for (const node of lbl.childNodes) {
                                        if (node.nodeType === 3) text += node.textContent.trim() + ' ';
                                    }
                                    text = text.trim();
                                    if (text && text !== '*' && text.length > 3) return text;
                                }
                            }
                            break;
                        }
                        parent = parent.parentElement;
                    }
                    return '';
                }''')
                if sibling_label and len(sibling_label) > 3:
                    label = sibling_label
            except Exception:
                pass

        # Strategy 1d: fieldset > legend > label (Greenhouse file uploads)
        if not label:
            try:
                fieldset_label = el.evaluate('''el => {
                    let parent = el.parentElement;
                    for (let i = 0; i < 6 && parent; i++) {
                        if (parent.tagName === 'FIELDSET') {
                            const legend = parent.querySelector('legend label, legend');
                            if (legend) return legend.textContent.trim();
                        }
                        parent = parent.parentElement;
                    }
                    return '';
                }''')
                if fieldset_label and len(fieldset_label) > 2:
                    label = fieldset_label
            except:
                pass

        # Strategy 2: aria/placeholder
        if not label:
            label = el.get_attribute("aria-label") or \
                    el.get_attribute("placeholder") or ""

        # Strategy 3: Context Discovery (for Shadow DOM)
        if not label or len(label) < 5:
            try:
                context = self._discover_field_context(el_id)
                if context:
                    label = context
            except Exception as e:
                pass  # Fallback to other methods

        # Strategy 4: Name/ID fallback
        if not label:
            label = el_name or el_id or ""

        # Append name attribute if different (helps with mapping)
        if el_name and el_name.lower() not in label.lower():
            label = f"{label} [{el_name}]"

        return label
    
    def _discover_field_context(self, field_id: str) -> str:
        """Use Context Discovery to find label from surrounding DOM."""
        if not field_id:
            return ""
        
        result = self.page.evaluate('''(fieldId) => {
            function findInShadow(root, selector) {
                let el = root.querySelector(selector);
                if (el) return el;
                const shadows = root.querySelectorAll('*');
                for (const s of shadows) {
                    if (s.shadowRoot) {
                        el = findInShadow(s.shadowRoot, selector);
                        if (el) return el;
                    }
                }
                return null;
            }
            
            const input = findInShadow(document, '#' + fieldId);
            if (!input) return '';
            
            const context = [];
            let container = input;
            
            for (let i = 0; i < 10 && container; i++) {
                container = container.parentElement || 
                           (container.getRootNode && container.getRootNode().host);
                
                if (!container || !container.querySelectorAll) continue;
                
                const textEls = container.querySelectorAll(
                    'label, legend, p, span, h3, h4, div'
                );
                
                for (const el of textEls) {
                    if (el.contains(input)) continue;
                    const txt = el.textContent.trim();
                    if (txt && txt.length > 3 && txt.length < 150) {
                        if (!context.some(c => c.includes(txt) || txt.includes(c))) {
                            context.push(txt);
                        }
                    }
                }
                
                if (context.length >= 1) break;
            }
            
            return context[0] || '';
        }''', field_id)
        
        return result or ""
    
    def _get_value(self, el: ElementHandle, tag: str, input_type: str) -> str:
        """Get current field value."""
        try:
            if tag == "select":
                return el.evaluate("e => e.options[e.selectedIndex]?.text || ''")
            elif input_type == "checkbox":
                return "checked" if el.is_checked() else ""
            elif input_type != "file":
                return el.input_value() or ""
        except:
            pass
        return ""
    
    def _detect_type(self, el: ElementHandle, tag: str, input_type: str) -> Tuple[FieldType, DetectionMethod]:
        """Detect field type using cascade."""
        
        # Layer 1: HTML standard
        if tag == "select":
            return FieldType.SELECT, DetectionMethod.HTML
        if tag == "textarea":
            return FieldType.TEXTAREA, DetectionMethod.HTML
        if input_type == "file":
            return FieldType.FILE, DetectionMethod.HTML
        if input_type == "checkbox":
            return FieldType.CHECKBOX, DetectionMethod.HTML
        if input_type == "email":
            return FieldType.EMAIL, DetectionMethod.HTML
        if input_type == "tel":
            return FieldType.PHONE, DetectionMethod.HTML
        if input_type == "date":
            return FieldType.DATE, DetectionMethod.HTML
        
        # Layer 2: ARIA
        role = el.get_attribute("role") or ""
        aria_haspopup = el.get_attribute("aria-haspopup") or ""
        
        if role == "combobox" or aria_haspopup in ("true", "listbox"):
            return FieldType.AUTOCOMPLETE, DetectionMethod.ARIA
        if role == "listbox":
            return FieldType.SELECT, DetectionMethod.ARIA
        
        # Default: text
        return FieldType.TEXT, DetectionMethod.DEFAULT
    
    def _probe_autocomplete_options(self, el: ElementHandle, selector: str) -> List[str]:
        """Click autocomplete to read options."""
        options = []
        try:
            el.click()
            time.sleep(0.4)
            
            # Wait for menu
            try:
                self.page.wait_for_selector('.select__menu', timeout=1000)
            except:
                self.page.keyboard.press("Escape")
                return []
            
            # Read options
            opt_elements = self.page.query_selector_all('.select__option')
            for opt in opt_elements:
                text = opt.inner_text().strip()
                if text and text != 'No options':
                    options.append(text)
            
            # Close menu
            self.page.keyboard.press("Escape")
            time.sleep(0.1)
            
        except:
            pass
        
        return options[:50]  # Limit
    
    # ─────────────────────────────────────────────────────────────────────
    # PRESCAN (ported from V3.5)
    # ─────────────────────────────────────────────────────────────────────

    def _prescan_all_options(self):
        """
        Pre-scan all autocomplete/select fields to discover options BEFORE filling.
        Ported from V3.5's prescan_options() — key to higher fill rates.

        Opens each dropdown, reads options, finds exact matches, then closes.
        This prevents the issue of typing wrong text and getting no results.
        """
        print("\n🔍 Pre-scanning dropdown options...")
        # Search in active frame AND all frames
        context = getattr(self, '_active_frame', self.page)
        prescan_count = 0

        # Build lookup of field selectors to frames for fill later
        self._field_frames = {}

        for field in self.fields:
            if field.field_type not in (FieldType.AUTOCOMPLETE, FieldType.SELECT):
                continue

            # Skip if already has options (from initial scan)
            if field.options:
                prescan_count += 1
                continue

            # Skip location/school - these are SEARCH type, don't prescan
            label_lower = field.label.lower()
            if any(kw in label_lower for kw in ['location', 'city', 'school', 'university']):
                print(f"   [SEARCH] {field.label[:35]}: skipping prescan (API-driven)")
                continue

            # Skip Select2 fields - they need special _fill_select2 handling
            if 's2id' in (field.element_id or ''):
                print(f"   [SELECT2] {field.label[:35]}: will use Select2 handler")
                continue

            print(f"   Prescanning: {field.label[:35]}...", flush=True)

            try:
                el = context.query_selector(field.selector)
                if not el:
                    # Try all frames
                    for frame in self.page.frames:
                        try:
                            el = frame.query_selector(field.selector)
                            if el:
                                break
                        except:
                            continue
                if not el or not el.is_visible():
                    print(f"      ⚠️ Not visible, skipping")
                    continue

                # Close any open dropdowns
                self.page.keyboard.press('Escape')
                time.sleep(0.1)

                try:
                    el.scroll_into_view_if_needed(timeout=3000)
                except:
                    pass
                el.click(timeout=3000)  # Short timeout for prescan clicks
                time.sleep(0.5)

                # Read options via aria-controls (V5 method)
                controls_id = el.get_attribute('aria-controls')
                options = []

                if controls_id:
                    listbox = context.query_selector(f'#{controls_id}')
                    if not listbox:
                        # Try in all frames
                        for frame in self.page.frames:
                            try:
                                listbox = frame.query_selector(f'#{controls_id}')
                                if listbox:
                                    break
                            except:
                                continue
                    if listbox:
                        opt_els = listbox.query_selector_all('[role="option"]')
                        for opt in opt_els[:100]:
                            text = opt.inner_text().strip()
                            if text and text not in ('No options', 'No results'):
                                options.append(text)

                # Fallback: global selectors with SHORT timeout
                if not options:
                    try:
                        self.page.wait_for_selector('.select__menu, [role="listbox"]', timeout=1000)
                        opt_els = self.page.query_selector_all('.select__option, [role="option"]')
                        for opt in opt_els[:100]:
                            text = opt.inner_text().strip()
                            if text and text not in ('No options', 'No results'):
                                options.append(text)
                    except:
                        pass

                # Close dropdown
                self.page.keyboard.press('Escape')
                time.sleep(0.1)

                if options:
                    field.options = options
                    is_fixed = len(options) <= 25
                    status = "FIXED" if is_fixed else "SEARCH"
                    print(f"      [{status}] {len(options)} options found")
                    prescan_count += 1
                else:
                    print(f"      No options found")

            except Exception as e:
                print(f"      ⚠️ Error: {str(e)[:50]}")
                try:
                    self.page.keyboard.press('Escape')
                except:
                    pass

        print(f"   📊 Pre-scanned {prescan_count} dropdowns")

    # ─────────────────────────────────────────────────────────────────────
    # LAYER 2: RESOLUTION
    # ─────────────────────────────────────────────────────────────────────

    def _resolve_all_answers(self):
        """Find answers for all fields."""
        print("\n📋 Resolving answers...")
        
        for field in self.fields:
            if field.field_type == FieldType.FILE:
                self._resolve_file_field(field)
            else:
                self._resolve_field_answer(field)
        
        # Summary
        ready = sum(1 for f in self.fields if f.status == FillStatus.READY)
        needs = sum(1 for f in self.fields if f.status == FillStatus.NEEDS_INPUT)
        print(f"   ✅ Ready: {ready}, ⚠️ Needs input: {needs}")
    
    def _resolve_field_answer(self, field: FormField):
        """Find answer for a field using cascade."""
        
        # Skip if already filled
        if field.current_value and field.current_value not in ("", "Select...", "Select"):
            field.status = FillStatus.FILLED
            return
        
        label_lower = field.label.lower()
        
        # Skip end date fields if current role (work_experience.0.current == true)
        # But NOT for education fields — those always have end dates
        is_education_field = "education" in label_lower or "educations" in (field.selector or "").lower()
        if not is_education_field and any(kw in label_lower for kw in ["end date", "end month", "end year"]):
            work_exp = self.profile.data.get("work_experience", [{}])
            if work_exp and work_exp[0].get("current", False):
                field.status = FillStatus.SKIPPED
                field.answer = ""
                field.error_message = "Skipped - current role"
                return
        
        # Cascade resolution
        answer, source, confidence = None, AnswerSource.NONE, 0.0
        is_dropdown = field.field_type in (FieldType.SELECT, FieldType.AUTOCOMPLETE)

        # 0. Schema database — fast lookup from previous successful fills
        if self._ats_type:
            schema_result = self.schema_db.resolve_from_schema(
                self._ats_type, field, self.profile
            )
            if schema_result:
                answer, source, confidence = schema_result
                # Preserve profile_key from schema for future saves
                mapping = self.schema_db.get_field_mapping(self._ats_type, field.element_id or "")
                if mapping and mapping.get("profile_key"):
                    field.profile_key = mapping["profile_key"]
                print(f"   📐 Schema: '{field.label[:30]}' → '{answer[:30]}' ({source.value})")

        # 1. Learned database
        if not answer:
            saved = self.learned_db.find(field.label, is_dropdown)
            if saved:
                answer, source, confidence = saved, AnswerSource.LEARNED, 0.95
        
        # 2. Yes/No patterns (checked BEFORE profile mapping for dropdowns
        #    to avoid "city" in "...same city...willing to relocate?" matching personal.city)
        if not answer and is_dropdown:
            yn = self.profile.find_yes_no(field.label)
            if yn:
                answer, source, confidence = yn, AnswerSource.DEFAULT, 0.85

        # 3. Profile mapping
        if not answer:
            val, key = self.profile.find_by_label(field.label)
            if val:
                answer, source, confidence = val, AnswerSource.PROFILE, 0.9
                field.profile_key = key

        # 3.5 Common answers from KnowledgeBase (salary, why interested, etc.)
        if not answer and not is_dropdown:
            common = self.kb.find_common_answer(field.label)
            if common:
                answer, source, confidence = common, AnswerSource.DEFAULT, 0.88
                print(f"   📚 KB common: '{field.label[:30]}' → '{common[:40]}...'")

        # 4. Yes/No patterns (for text fields — fallback)
        if not answer and not is_dropdown:
            yn = self.profile.find_yes_no(field.label)
            if yn:
                answer, source, confidence = yn, AnswerSource.DEFAULT, 0.85
        
        # 5. Demographic defaults
        if not answer:
            demo = self.profile.find_demographic(field.label)
            if demo:
                answer, source, confidence = demo, AnswerSource.DEFAULT, 0.8

        # 6. Option matching for dropdowns
        if not answer and field.options:
            matched = self._match_option(field)
            if matched:
                answer, source, confidence = matched, AnswerSource.DEFAULT, 0.7

        # 7. Text defaults (years of experience, salary, etc.)
        if not answer:
            text_default = self.profile.find_text_default(field.label)
            if text_default:
                answer, source, confidence = text_default, AnswerSource.DEFAULT, 0.75

        # 8. Ollama for custom questions (with KB context)
        if not answer and self.ollama.available:
            profile_context = self._get_profile_context_for_ai()
            kb_context = self.kb.get_context_for_question(field.label)
            if kb_context:
                profile_context += f"\n\n{kb_context}"
            ollama_answer = self.ollama.generate(field.label, profile_context, field.options)
            if ollama_answer:
                if field.options:
                    ollama_answer = self.ollama.match_option(ollama_answer, field.options)
                answer, source, confidence = ollama_answer, AnswerSource.AI, 0.6
                print(f"   🤖 Ollama: '{field.label[:30]}' → '{ollama_answer[:30]}'")

        # 8. Claude AI fallback for remaining unknown fields (with KB context)
        if not answer and self.ai.available:
            profile_context = self._get_profile_context_for_ai()
            kb_context = self.kb.get_context_for_question(field.label)
            if kb_context:
                profile_context += f"\n\n{kb_context}"
            try:
                if field.options:
                    claude_answer = self.ai.choose_option(field.label, field.options, profile_context)
                else:
                    claude_answer = self.ai.generate(field.label, profile_context)
                if claude_answer:
                    answer, source, confidence = claude_answer, AnswerSource.AI, 0.55
                    print(f"   🧠 Claude: '{field.label[:30]}' → '{claude_answer[:30]}'")
            except Exception as e:
                print(f"   ⚠️ Claude fallback error: {e}")

        # Set result — apply DOM-aware format adaptation
        if answer:
            # DOM-aware: adapt value to field's placeholder/maxlength/type
            # e.g., "September" → "09" if placeholder="MM" or maxlength=2
            if field.field_type in (FieldType.TEXT, FieldType.DATE):
                original = answer
                answer = self._adapt_value_to_dom(answer, field)
                if answer != original:
                    print(f"   🔧 DOM adapt: '{original}' → '{answer}' (placeholder='{field.placeholder}', maxlen={field.maxlength})")
            field.answer = answer
            field.answer_source = source
            field.confidence = confidence
            field.status = FillStatus.READY
        else:
            field.status = FillStatus.NEEDS_INPUT
    
    def _get_profile_context_for_ai(self) -> str:
        """Get rich profile context for AI questions."""
        p = self.profile.data.get("personal", {})
        w = self.profile.data.get("work_experience", [{}])[0]
        certs = self.profile.data.get("certifications", [])
        common = self.profile.data.get("common_answers", {})
        demo = self.profile.data.get("demographics", {})
        wa = self.profile.data.get("work_authorization", {})

        # Job context (JD + title + company)
        jd = getattr(self, 'job_description', '') or ''
        jt = getattr(self, 'job_title', '') or ''
        cn = getattr(self, 'company_name', '') or ''
        job_context = ""
        if jt or cn:
            job_context = f"\n\nAPPLYING FOR: {jt} at {cn}"
        if jd:
            job_context += f"\n\nJOB DESCRIPTION (use this to tailor answers):\n{jd[:5000]}"

        return f"""Name: {p.get('first_name', '')} {p.get('last_name', '')}
Location: {p.get('location', '')}, {p.get('state', '')}, US
Current Role: {w.get('title', '')} at {w.get('company', '')}
Experience: {w.get('description', '')}
Years: 15+ years in PM/TPM
Tools: GCP, AWS, Jira, Confluence, SharePoint, Python, SQL
Certifications: {', '.join(certs) if certs else 'SAFe, PSM, GCP'}
Work Authorization: Authorized to work in US: {wa.get('authorized_us', 'Yes')}, Needs sponsorship: {wa.get('requires_sponsorship', 'No')}
Willing to relocate: Yes
18 or older: Yes
Previously employed at company: No
Government official: No
Gender: {demo.get('gender', 'Decline')}
Veteran: {demo.get('veteran_status', 'Not a protected veteran')}
Disability: {demo.get('disability_status', 'Prefer not to answer')}{job_context}"""
    
    def _match_option(self, field: FormField) -> Optional[str]:
        """Try to match profile data to dropdown options."""
        label_lower = field.label.lower()
        options_lower = [opt.lower() for opt in field.options]

        # ── Demographics: use profile values first, then fallback to defaults ──
        demo = self.profile.data.get("demographics", {}) if self.profile else {}

        # Gender
        if "gender" in label_lower:
            profile_val = demo.get("gender", "")
            if profile_val:
                for i, opt_lower in enumerate(options_lower):
                    if profile_val.lower() in opt_lower or opt_lower in profile_val.lower():
                        return field.options[i]
            # Fallback: decline
            for pattern in ["decline", "prefer not", "do not wish"]:
                for i, opt_lower in enumerate(options_lower):
                    if pattern in opt_lower:
                        return field.options[i]

        # Hispanic/Latino
        if "hispanic" in label_lower or "latino" in label_lower:
            profile_val = demo.get("hispanic_latino", "")
            if profile_val:
                for i, opt_lower in enumerate(options_lower):
                    if opt_lower.startswith(profile_val.lower()) or profile_val.lower() == opt_lower:
                        return field.options[i]
            for pattern in ["decline", "prefer not", "no"]:
                for i, opt_lower in enumerate(options_lower):
                    if pattern in opt_lower:
                        return field.options[i]

        # Race/Ethnicity
        if "race" in label_lower or "ethnicity" in label_lower:
            profile_val = demo.get("race_ethnicity", "")
            if profile_val:
                for i, opt_lower in enumerate(options_lower):
                    if profile_val.lower() in opt_lower or opt_lower in profile_val.lower():
                        return field.options[i]
            for pattern in ["decline", "prefer not", "two or more", "do not wish"]:
                for i, opt_lower in enumerate(options_lower):
                    if pattern in opt_lower:
                        return field.options[i]

        # Veteran status
        if "veteran" in label_lower:
            profile_val = demo.get("veteran_status", "")
            if profile_val:
                for i, opt_lower in enumerate(options_lower):
                    if profile_val.lower() in opt_lower or opt_lower in profile_val.lower():
                        return field.options[i]
            for pattern in ["not a protected veteran", "i am not a", "decline", "no"]:
                for i, opt_lower in enumerate(options_lower):
                    if pattern in opt_lower:
                        return field.options[i]

        # Disability
        if "disability" in label_lower:
            profile_val = demo.get("disability_status", "")
            if profile_val:
                for i, opt_lower in enumerate(options_lower):
                    if profile_val.lower() in opt_lower or opt_lower in profile_val.lower():
                        return field.options[i]
            for pattern in ["do not want to answer", "prefer not", "decline", "no, i do not"]:
                for i, opt_lower in enumerate(options_lower):
                    if pattern in opt_lower:
                        return field.options[i]

        # Country - prefer United States
        if "country" in label_lower:
            for i, opt_lower in enumerate(options_lower):
                if "united states" in opt_lower or opt_lower == "usa" or opt_lower == "us":
                    return field.options[i]

        # State - match from profile
        if "state" in label_lower and "united" not in label_lower:
            profile_state = self.profile.get("personal.state")
            if profile_state:
                state_lower = profile_state.lower()
                for i, opt_lower in enumerate(options_lower):
                    if state_lower in opt_lower or opt_lower in state_lower:
                        return field.options[i]

        # Yes/No questions - match from YES_NO_PATTERNS
        yes_no = self.profile.find_yes_no(field.label)
        if yes_no:
            for i, opt_lower in enumerate(options_lower):
                if yes_no.lower() == opt_lower or yes_no.lower() in opt_lower:
                    return field.options[i]

        return None
    
    def _ensure_compatible_extension(self, file_path: Path) -> Path:
        """Convert .dotx/.dotm to .docx for ATS compatibility.

        Greenhouse, Lever, and most ATS only accept: pdf, doc, docx, txt, rtf.
        .dotx (Word template) is functionally identical to .docx but rejected by extension filter.
        """
        ext = file_path.suffix.lower()
        incompatible = {'.dotx': '.docx', '.dotm': '.docm'}

        if ext in incompatible:
            import shutil
            new_ext = incompatible[ext]
            new_path = file_path.with_suffix(new_ext)

            # Copy to compatible extension (in same directory)
            if not new_path.exists():
                shutil.copy2(file_path, new_path)
                print(f"   📎 Converted {file_path.name} → {new_path.name} (ATS compatibility)")
            else:
                print(f"   📎 Using existing {new_path.name}")

            return new_path

        return file_path

    def _resolve_file_field(self, field: FormField):
        """Resolve file upload field based on job title.
        Tries personalized documents first (AI-generated CL, tailored CV),
        falls back to static files from profile.
        """
        label_lower = field.label.lower()
        field_id = (field.element_id or "").lower()
        field_name = (field.name or "").lower()

        # Get job title from page
        job_title_page = self.page.title() if self.page else ""

        # Get static CV and Cover Letter paths based on role (fallback)
        cv_path, cover_letter_path = self.profile.get_files_for_role(job_title_page)

        # Ensure compatible file extensions (.dotx → .docx)
        if cv_path:
            cv_path = self._ensure_compatible_extension(cv_path)
        if cover_letter_path:
            cover_letter_path = self._ensure_compatible_extension(cover_letter_path)

        # Try personalized documents if we have job info
        jt = getattr(self, 'job_title', '') or ''
        cn = getattr(self, 'company_name', '') or ''
        jd = getattr(self, 'job_description', '') or ''

        # Check what type of file is requested (check ID, selector, then label)
        selector_lower = (field.selector or "").lower()
        is_cover_letter = (
            "cover_letter" in field_id or
            "coverletter" in field_id or
            "cover_letter" in selector_lower or
            any(kw in label_lower for kw in ["cover letter", "cover_letter", "coverletter"])
        )

        is_resume = (
            "resume" in field_id or
            "cv" in field_id or
            any(kw in label_lower for kw in ["resume", "cv"])
        )

        if is_cover_letter:
            # Try personalized CL first
            if jt and cn:
                personalized_cl = self._generate_personalized_cover_letter(jt, cn, jd)
                if personalized_cl and personalized_cl.exists():
                    field.answer = str(personalized_cl)
                    field.answer_source = AnswerSource.AI
                    field.status = FillStatus.READY
                    print(f"   📄 Cover Letter (personalized): {personalized_cl.name}")
                    return

            # Fallback to static CL
            if cover_letter_path and cover_letter_path.exists():
                field.answer = str(cover_letter_path)
                field.answer_source = AnswerSource.PROFILE
                field.status = FillStatus.READY
                print(f"   📄 Cover Letter (static): {cover_letter_path.name}")
            else:
                field.status = FillStatus.NEEDS_INPUT
                field.error_message = "Cover letter not found for this role"

        elif is_resume or "attach" in label_lower or "upload" in label_lower or "browse" in label_lower or field_name == "file":
            # Try tailored CV first
            if jt and cn and jd:
                tailored_cv = self._create_tailored_cv(jt, cn, jd)
                if tailored_cv and tailored_cv.exists():
                    field.answer = str(tailored_cv)
                    field.answer_source = AnswerSource.AI
                    field.status = FillStatus.READY
                    print(f"   📄 Resume/CV (tailored): {tailored_cv.name}")
                    return

            # Fallback to static CV
            if cv_path and cv_path.exists():
                field.answer = str(cv_path)
                field.answer_source = AnswerSource.PROFILE
                field.status = FillStatus.READY
                print(f"   📄 Resume/CV (static): {cv_path.name}")
            else:
                field.status = FillStatus.ERROR
                field.error_message = "CV not found for this role"
        else:
            # Unknown file field — assume resume as default
            if cv_path and cv_path.exists():
                field.answer = str(cv_path)
                field.answer_source = AnswerSource.PROFILE
                field.status = FillStatus.READY
                print(f"   📄 Default file → Resume: {cv_path.name}")
            else:
                field.status = FillStatus.NEEDS_INPUT
    
    # ─────────────────────────────────────────────────────────────────────
    # LAYER 3: INPUT
    # ─────────────────────────────────────────────────────────────────────
    
    def _fill_all_fields(self, mode: FillMode):
        """Fill all fields with answers."""
        print("\n📝 Filling fields...")

        for field in self.fields:
            if field.status == FillStatus.READY:
                self._fill_field(field)
            elif field.status == FillStatus.NEEDS_INPUT and mode == FillMode.INTERACTIVE:
                self._interactive_fill(field)
            elif field.status == FillStatus.NEEDS_INPUT:
                field.status = FillStatus.SKIPPED
            # Don't overwrite FILLED/VERIFIED/ERROR status from previous iteration
    
    def _fill_field(self, field: FormField) -> bool:
        """Fill single field. Uses _active_frame to support iframes. Timeout protected."""
        try:
            # Use active frame (main page or iframe with form)
            context = getattr(self, '_active_frame', self.page)
            el = context.query_selector(field.selector)

            # Fallback: try all frames if not found
            if not el:
                el = self.page.query_selector(field.selector)
            if not el:
                for frame in self.page.frames:
                    try:
                        el = frame.query_selector(field.selector)
                        if el:
                            break
                    except:
                        continue
            if not el:
                field.status = FillStatus.ERROR
                field.error_message = "Element not found"
                return False

            try:
                el.scroll_into_view_if_needed(timeout=3000)
            except:
                pass
            time.sleep(0.1)

            success = False

            if field.field_type == FieldType.FILE:
                success = self._fill_file(el, field)
            elif field.field_type == FieldType.SELECT:
                success = self._fill_select(el, field)
            elif field.field_type == FieldType.AUTOCOMPLETE:
                success = self._fill_autocomplete(el, field)
            elif field.field_type == FieldType.CHECKBOX:
                success = self._fill_checkbox(el, field)
            elif field.field_type == FieldType.PHONE:
                success = self._fill_phone(el, field)
            else:
                success = self._fill_text(el, field)

            if success:
                field.status = FillStatus.FILLED
                print(f"   ✅ {field.label[:35]:<35} = {field.answer[:20]} [{field.answer_source.value}]")
            else:
                field.status = FillStatus.ERROR
                print(f"   ❌ {field.label[:35]} - fill failed")

            # Log field fill
            self.logger.log_field(
                field_id=field.element_id or field.selector,
                field_type=field.field_type.value,
                question=field.label,
                value=field.answer or "",
                source=field.answer_source.value if field.answer_source else "none",
                success=success,
                error=field.error_message
            )

            return success

        except Exception as e:
            field.status = FillStatus.ERROR
            field.error_message = str(e)[:100]
            print(f"   ❌ {field.label[:35]} - error: {str(e)[:50]}")
            self.logger.log_field(
                field_id=field.element_id or field.selector,
                field_type=field.field_type.value,
                question=field.label,
                value="",
                source="error",
                success=False,
                error=str(e)[:100]
            )
            return False
    
    def _fill_text(self, el: ElementHandle, field: FormField) -> bool:
        """Fill text/email/phone field. Uses keyboard for React compatibility."""
        try:
            # Try standard fill first
            el.fill(field.answer)
            el.evaluate("e => e.blur()")
            return True
        except:
            pass
        
        # Fallback: click and type (for React forms)
        try:
            el.click(click_count=3)  # Select all
            time.sleep(0.1)
            self.page.keyboard.type(field.answer, delay=5)
            el.evaluate("e => e.blur()")
            return True
        except:
            return False
    
    def _fill_phone(self, el: ElementHandle, field: FormField) -> bool:
        """
        Fill phone field with delay for intl-tel-input compatibility.
        Greenhouse uses intl-tel-input which needs slower typing.
        """
        try:
            el.click()
            time.sleep(0.2)
            # Use slower delay (30ms) for phone input formatting
            self.page.keyboard.type(field.answer, delay=30)
            time.sleep(0.2)
            el.evaluate("e => e.blur()")
            return True
        except:
            return False
    
    def _fill_select(self, el: ElementHandle, field: FormField) -> bool:
        """Fill native select with fuzzy matching."""
        try:
            # Try exact label match
            el.select_option(label=field.answer)
            el.evaluate("e => e.blur()")
            return True
        except:
            pass

        # Fuzzy match: get all options and find best match
        try:
            options = el.evaluate(
                "e => Array.from(e.options).map(o => ({value: o.value, text: o.text}))"
            )
            answer_lower = field.answer.lower().strip()

            for opt in options:
                opt_text = opt.get('text', '').lower()
                if answer_lower in opt_text or opt_text in answer_lower:
                    el.select_option(value=opt['value'])
                    el.evaluate("e => e.blur()")
                    return True

            # Word overlap matching
            answer_words = set(answer_lower.split())
            for opt in options:
                opt_text = opt.get('text', '').lower()
                opt_words = set(opt_text.split())
                if len(answer_words & opt_words) >= 1:
                    el.select_option(value=opt['value'])
                    el.evaluate("e => e.blur()")
                    return True

            # Last resort: first non-empty option
            for opt in options:
                if opt.get('value') and opt.get('text', '').strip():
                    el.select_option(value=opt['value'])
                    el.evaluate("e => e.blur()")
                    return True
        except Exception as e:
            print(f"      ⚠️ Select error: {e}")

        return False
    
    def _fill_autocomplete(self, el: ElementHandle, field: FormField) -> bool:
        """
        Fill autocomplete/combobox. Supports:
        - React Select (aria-controls)
        - Select2 (#s2id_autogen* selectors)
        - Location fields (API-driven)
        - School fields (search with fallback)
        """
        frame = el.owner_frame() or self.page
        label_lower = field.label.lower()

        # Check field type for special handling
        is_location = 'location' in label_lower or 'city' in label_lower
        is_school = 'school' in label_lower or 'university' in label_lower or 'college' in label_lower

        # ── SELECT2 DETECTION ──
        # Select2 uses hidden inputs like #s2id_autogen1
        # BUT on newer forms, the input may have a regular ID (e.g. question_11097823007)
        # while still being wrapped in a .select2-container
        is_select2 = 's2id' in (field.element_id or '') or 's2id' in field.selector
        if not is_select2:
            # Check DOM: is this element inside a .select2-container?
            try:
                is_select2 = frame.evaluate(f'''() => {{
                    const el = document.querySelector('{field.selector}');
                    if (!el) return false;
                    let parent = el.parentElement;
                    for (let i = 0; i < 5; i++) {{
                        if (!parent) break;
                        if (parent.classList && parent.classList.contains('select2-container')) return true;
                        parent = parent.parentElement;
                    }}
                    return false;
                }}''')
            except:
                is_select2 = False
        if is_select2:
            return self._fill_select2(el, field, frame)

        # Close any open dropdowns first
        self.page.keyboard.press('Escape')
        time.sleep(0.1)

        # Scroll into view
        try:
            el.scroll_into_view_if_needed()
        except:
            pass
        time.sleep(0.1)

        # For Location: type first, then wait for API
        if is_location:
            print(f"      📍 Location: typing '{field.answer[:30]}'...")
            el.click()
            time.sleep(0.3)
            self.page.keyboard.type(field.answer[:30], delay=30)
            time.sleep(2.0)  # Wait for API response
            
            # Get options via aria-controls
            controls_id = el.get_attribute('aria-controls')
            if controls_id:
                listbox = frame.query_selector(f'#{controls_id}')
                if listbox:
                    options = listbox.query_selector_all('[role="option"]')
                    if options:
                        options[0].click()
                        time.sleep(0.2)
                        return True
            
            # Fallback: just click away to confirm text
            self.page.keyboard.press('Tab')
            return True

        # For School fields: search with fallback to "0 - Other"
        if is_school:
            return self._fill_school_autocomplete(el, field, frame)

        # ── Pre-resolve: adapt answer to match actual form options (state codes, etc.) ──
        if field.options and len(field.options) >= 3:
            resolved = self._resolve_select2_from_options(field, field.options)
            if resolved:
                field.answer = resolved
                print(f"      🔄 Answer adapted to option: '{resolved}'")

        # STRATEGY: Use prescan data if we have options from prescan
        # Threshold 100: includes US states (~70 options) in the fixed dropdown path
        if field.options and len(field.options) <= 100:
            # Fixed dropdown — find exact match from prescan data and click it
            try:
                el.click(force=True, timeout=5000)
            except:
                try:
                    frame.evaluate(f'''() => {{
                        const el = document.querySelector('{field.selector}');
                        if (el) {{ el.focus(); el.dispatchEvent(new MouseEvent('mousedown', {{bubbles:true}})); }}
                    }}''')
                except:
                    pass
            time.sleep(0.4)

            # Read live options
            controls_id = el.get_attribute('aria-controls')
            live_options = []
            if controls_id:
                listbox = frame.query_selector(f'#{controls_id}')
                if listbox:
                    live_options = listbox.query_selector_all('[role="option"]')
            if not live_options:
                try:
                    frame.wait_for_selector('[role="listbox"], .select__menu', timeout=1500)
                    live_options = frame.query_selector_all('[role="option"], .select__option')
                except:
                    pass

            answer_lower = field.answer.lower().strip()

            # Score-based matching (like V3.5)
            best_score = 0
            best_opt = None
            for opt in live_options:
                opt_text = opt.inner_text().strip()
                opt_lower = opt_text.lower()
                score = 0
                if opt_lower == answer_lower:
                    score = 100  # Exact match
                elif answer_lower in opt_lower:
                    # Substring match — but for short answers (Yes/No) require word boundary
                    # "No" in "now" is NOT a match; "No" in "No, I do not..." IS a match
                    if len(answer_lower) <= 4:
                        # Short answer: must start the option or be a separate word
                        import re
                        if re.search(r'(?:^|[\s,])' + re.escape(answer_lower) + r'(?:[\s,.]|$)', opt_lower):
                            score = 85 if opt_lower.startswith(answer_lower) else 80
                    else:
                        score = 85 if opt_lower.startswith(answer_lower) else 80
                elif opt_lower in answer_lower:
                    score = 70  # Option is substring of answer
                else:
                    # Word overlap
                    answer_words = set(answer_lower.replace('-', ' ').split())
                    opt_words = set(opt_lower.replace('-', ' ').split())
                    overlap = len(answer_words & opt_words)
                    if overlap >= 2:
                        score = 60
                    elif overlap >= 1:
                        score = 40
                if score > best_score:
                    best_score = score
                    best_opt = opt

            if best_opt and best_score >= 70:
                opt_text = best_opt.inner_text().strip()
                print(f"      ✓ Clicking option: '{opt_text[:40]}' (score={best_score})")
                try:
                    best_opt.click(force=True, timeout=5000)
                except:
                    try:
                        best_opt.evaluate("e => e.click()")
                    except:
                        pass
                time.sleep(0.2)
                return True

            # No good match in live options — type answer as filter to narrow down
            # (React Select virtualizes options, so not all may be visible)
            print(f"      🔎 Typing filter '{field.answer[:30]}' (best_score={best_score}, {len(live_options)} live opts)")
            el.type(field.answer[:30], delay=20)
            time.sleep(0.5)

            # Re-read filtered options
            filtered = []
            if controls_id:
                listbox = frame.query_selector(f'#{controls_id}')
                if listbox:
                    filtered = listbox.query_selector_all('[role="option"]')
            if not filtered:
                filtered = frame.query_selector_all('[role="option"], .select__option')

            if filtered:
                first_text = filtered[0].inner_text().strip().lower()
                if 'no result' not in first_text and 'no option' not in first_text:
                    try:
                        filtered[0].click(force=True, timeout=5000)
                    except:
                        filtered[0].evaluate("e => e.click()")
                    time.sleep(0.2)
                    return True

            self.page.keyboard.press('Escape')
            return True

        # SEARCH dropdown or no prescan data — open and type to filter
        print(f"      🔎 Search path: typing '{field.answer[:30]}' into {field.selector}")
        # Close any previously open dropdowns
        self.page.keyboard.press('Escape')
        time.sleep(0.2)
        # Use JavaScript click to avoid Playwright actionability timeouts on hidden/overlay elements
        try:
            frame.evaluate(f'''() => {{
                const el = document.querySelector('{field.selector}');
                if (el) {{
                    el.scrollIntoView({{block: 'center'}});
                    el.focus();
                    el.click();
                    el.dispatchEvent(new MouseEvent('mousedown', {{bubbles: true}}));
                }}
            }}''')
        except:
            # Fallback: try Playwright click
            try:
                el.click(force=True, timeout=5000)
            except:
                pass
        time.sleep(0.4)
        controls_id = el.get_attribute('aria-controls')
        print(f"      🔎 controls_id={controls_id}")

        # Type to filter — use el.type() to ensure typing goes to the right input in iframe
        el.type(field.answer[:30], delay=20)
        time.sleep(0.8)

        # Re-read controls_id (React Select sets it after interaction)
        if not controls_id:
            controls_id = el.get_attribute('aria-controls')

        # Read filtered options
        live_options = []
        if controls_id:
            listbox = frame.query_selector(f'#{controls_id}')
            if listbox:
                live_options = listbox.query_selector_all('[role="option"]')
        if not live_options:
            live_options = frame.query_selector_all('[role="option"], .select__option')

        if live_options:
            first_text = live_options[0].inner_text().strip()
            print(f"      🔎 Found {len(live_options)} options, first: '{first_text[:40]}'")
            # Check first option isn't "no results"
            if 'no result' not in first_text.lower() and 'no option' not in first_text.lower():
                # Click the option — try multiple methods for React Select compatibility
                clicked = False
                try:
                    live_options[0].click(force=True, timeout=5000)
                    clicked = True
                except:
                    pass
                if not clicked:
                    try:
                        live_options[0].evaluate("e => e.click()")
                        clicked = True
                    except:
                        pass
                # Fallback: keyboard — ArrowDown selects first, Enter confirms
                if not clicked or not controls_id:
                    self.page.keyboard.press('ArrowDown')
                    time.sleep(0.1)
                    self.page.keyboard.press('Enter')
                    print(f"      🔎 Used keyboard fallback (ArrowDown+Enter)")
                time.sleep(0.2)
                return True

        # Fallback: keyboard navigation
        self.page.keyboard.press('ArrowDown')
        self.page.keyboard.press('Tab')
        time.sleep(0.2)
        return True

    def _fill_select2(self, el: ElementHandle, field: FormField, frame) -> bool:
        """
        Fill Select2 dropdown (Greenhouse EEO/demographics/custom questions).

        Select2 DOM structure:
        - .select2-container (id=s2id_job_application_gender)
          - a.select2-choice  ← CLICK THIS
          - input#s2id_autogen2 ← the input we detected
        - .select2-drop (appears after click)
          - .select2-results
            - li.select2-result ← the options to click
        """
        print(f"      🔽 Select2: {field.label[:30]}...", flush=True)

        is_school_field = 'school' in field.label.lower()

        try:
            # Open dropdown by clicking .select2-choice via JavaScript
            opened = frame.evaluate(f'''() => {{
                const input = document.querySelector('{field.selector}');
                if (!input) return false;
                // Find parent select2-container
                let container = input.closest('.select2-container');
                if (!container) {{
                    // Walk up manually
                    let parent = input.parentElement;
                    for (let i = 0; i < 5; i++) {{
                        if (!parent) break;
                        if (parent.classList && parent.classList.contains('select2-container')) {{
                            container = parent;
                            break;
                        }}
                        parent = parent.parentElement;
                    }}
                }}
                if (!container) return false;
                const choice = container.querySelector('a.select2-choice, .select2-choices');
                if (choice) {{
                    choice.dispatchEvent(new MouseEvent('mousedown', {{bubbles: true}}));
                    return true;
                }}
                return false;
            }}''')

            if not opened:
                print(f"      ⚠️ Could not find Select2 container")
                return False

            time.sleep(0.5)

            # Wait for and read options from .select2-drop
            try:
                frame.wait_for_selector('.select2-drop:not(.select2-display-none) .select2-results', timeout=2000)
            except:
                # Try alternative: just check if results exist
                pass

            options = frame.query_selector_all('.select2-drop:not(.select2-display-none) .select2-results li.select2-result')
            if not options:
                options = frame.query_selector_all('.select2-results li.select2-result')
            if not options:
                options = frame.query_selector_all('.select2-results li')

            # Filter out disabled/header items
            valid_options = []
            for opt in options:
                try:
                    cls = opt.get_attribute('class') or ''
                    if 'select2-disabled' in cls or 'select2-result-unselectable' in cls:
                        continue
                    text = opt.inner_text().strip()
                    if text:
                        valid_options.append((opt, text))
                except:
                    continue

            if not valid_options:
                print(f"      ⚠️ No Select2 options found")
                self.page.keyboard.press('Escape')
                return False

            option_texts = [t for _, t in valid_options]
            answer_lower = field.answer.lower().strip()
            print(f"      Looking for: '{answer_lower[:30]}' in {len(valid_options)} options")

            # Score-based matching
            best_score, best_opt, best_text = self._select2_match(answer_lower, valid_options)

            # School fields need higher confidence — weak matches mean school isn't in DB
            min_score = 80 if is_school_field else 40

            if best_score >= min_score:
                try:
                    best_opt.click(force=True, timeout=5000)
                except:
                    best_opt.evaluate("e => e.click()")
                time.sleep(0.3)
                print(f"      ✅ Select2 matched (score={best_score})")
                return True

            # No match from pre-resolved answer — try re-resolving with actual options
            # This handles cases where label was wrong during resolution or AI gave freetext
            # Skip re-resolve for school fields — they should fallback to "Other" if not matched
            if not is_school_field:
                print(f"      🔄 Re-resolving with actual {len(option_texts)} options...")
                new_answer = self._resolve_select2_from_options(field, option_texts)
                if new_answer:
                    new_lower = new_answer.lower().strip()
                    best_score2, best_opt2, best_text2 = self._select2_match(new_lower, valid_options)
                    if best_score2 >= min_score:
                        try:
                            best_opt2.click(force=True, timeout=5000)
                        except:
                            best_opt2.evaluate("e => e.click()")
                        time.sleep(0.3)
                        field.answer = new_answer
                        print(f"      ✅ Select2 re-resolved: '{best_text2[:30]}' (score={best_score2})")
                        return True

            # For school fields: search for "Other" instead of picking first match
            if is_school_field:
                print(f"      🎓 School not matched — searching for 'Other'...")
                self.page.keyboard.press('Escape')
                time.sleep(0.2)

                # Re-open and search for "Other"
                opened2 = frame.evaluate(f'''() => {{
                    const input = document.querySelector('{field.selector}');
                    if (!input) return false;
                    let container = input.closest('.select2-container');
                    if (!container) {{
                        let parent = input.parentElement;
                        for (let i = 0; i < 5; i++) {{
                            if (!parent) break;
                            if (parent.classList && parent.classList.contains('select2-container')) {{
                                container = parent; break;
                            }}
                            parent = parent.parentElement;
                        }}
                    }}
                    if (!container) return false;
                    const choice = container.querySelector('a.select2-choice, .select2-choices');
                    if (choice) {{
                        choice.dispatchEvent(new MouseEvent('mousedown', {{bubbles: true}}));
                        return true;
                    }}
                    return false;
                }}''')

                if opened2:
                    time.sleep(0.5)
                    search_input = frame.query_selector('.select2-drop:not(.select2-display-none) .select2-input')
                    if not search_input:
                        search_input = frame.query_selector('.select2-search input')
                    if search_input:
                        try:
                            search_input.fill('', timeout=3000)
                        except:
                            pass
                        self.page.keyboard.type('Other', delay=20)
                        time.sleep(0.8)

                    fb_opts = frame.query_selector_all('.select2-drop:not(.select2-display-none) .select2-results li.select2-result')
                    if not fb_opts:
                        fb_opts = frame.query_selector_all('.select2-results li')

                    for opt in fb_opts:
                        try:
                            text = opt.inner_text().strip()
                            if text.lower() == 'other':
                                try:
                                    opt.click(force=True, timeout=5000)
                                except:
                                    opt.evaluate("e => e.click()")
                                time.sleep(0.3)
                                print(f"      ✅ School fallback: 'Other'")
                                return True
                        except:
                            continue

                    # No exact "Other" — pick first valid result
                    for opt in fb_opts:
                        try:
                            text = opt.inner_text().strip()
                            if text and 'no result' not in text.lower():
                                try:
                                    opt.click(force=True, timeout=5000)
                                except:
                                    opt.evaluate("e => e.click()")
                                time.sleep(0.3)
                                print(f"      ✅ School fallback: '{text[:30]}'")
                                return True
                        except:
                            continue

            # Final fallback — try rule-based resolve (no AI needed)
            rule_answer = self._resolve_select2_from_options(field, option_texts)
            if rule_answer:
                ra_lower = rule_answer.lower().strip()
                for opt_el, opt_text in valid_options:
                    if opt_text.lower().strip() == ra_lower or ra_lower in opt_text.lower():
                        try:
                            opt_el.click(force=True, timeout=5000)
                        except:
                            opt_el.evaluate("e => e.click()")
                        time.sleep(0.3)
                        field.answer = rule_answer
                        print(f"      ✅ Select2 rule-based: '{opt_text[:30]}'")
                        return True

            # Last resort — click first non-placeholder option
            fallback_idx = 0
            for idx, (_, text) in enumerate(valid_options):
                tl = text.lower()
                if tl not in ('--', 'please select', 'select', 'select...', ''):
                    fallback_idx = idx
                    break
            try:
                valid_options[fallback_idx][0].click(force=True, timeout=5000)
            except:
                valid_options[fallback_idx][0].evaluate("e => e.click()")
            time.sleep(0.3)
            print(f"      ⚠️ Select2 fallback: '{valid_options[fallback_idx][1][:30]}'")
            return True

        except Exception as e:
            print(f"      ⚠️ Select2 error: {str(e)[:60]}")
            try:
                self.page.keyboard.press('Escape')
            except:
                pass
            return False

    def _select2_match(self, answer_lower: str, valid_options: list) -> tuple:
        """Score-based matching for Select2 options. Returns (score, element, text)."""
        best_score = 0
        best_opt = None
        best_text = ""
        for opt_el, opt_text in valid_options:
            opt_lower = opt_text.lower()
            score = 0
            if opt_lower == answer_lower:
                score = 100
            elif answer_lower in opt_lower:
                score = 80
            elif opt_lower in answer_lower:
                score = 70
            else:
                answer_words = set(answer_lower.split())
                opt_words = set(opt_lower.split())
                overlap = len(answer_words & opt_words)
                if overlap >= 2:
                    score = 60
                elif overlap >= 1:
                    score = 40
            if score > best_score:
                best_score = score
                best_opt = opt_el
                best_text = opt_text
        return best_score, best_opt, best_text

    def _resolve_select2_from_options(self, field: FormField, option_texts: list) -> str:
        """
        Re-resolve a Select2 answer using the actual dropdown options.
        Uses profile matching, defaults, and AI to pick the best option.
        """
        label = field.label.lower()
        opt_lower = [t.lower() for t in option_texts]

        # 1. Yes/No questions — check known patterns
        yes_patterns = ['authorize', 'authorized', 'legal right', 'legally',
                       '18 years', 'age', 'willing to relocat', 'relocated',
                       'background check', 'drug test', 'submit verification']
        no_patterns = ['require sponsor', 'sponsorship', 'non-compete',
                      'previously been employed', 'worked for our company',
                      'ever worked for', 'currently or have you ever worked',
                      'worked for', 'different name', 'convicted']

        has_yes = any('yes' in o for o in opt_lower)
        has_no = any('no' in o for o in opt_lower)

        if has_yes and has_no:
            if any(p in label for p in yes_patterns):
                return 'Yes'
            if any(p in label for p in no_patterns):
                return 'No'

        # 2. State/location questions
        if any(k in label for k in ['state', 'province', 'reside', 'location', 'where']):
            profile_state = ""
            if self.profile:
                profile_state = self.profile.get("personal.state")
            if profile_state:
                # Bidirectional state mapping (full name ↔ abbreviation)
                state_to_abbr = {
                    "alabama": "AL", "alaska": "AK", "arizona": "AZ", "arkansas": "AR",
                    "california": "CA", "colorado": "CO", "connecticut": "CT", "delaware": "DE",
                    "florida": "FL", "georgia": "GA", "hawaii": "HI", "idaho": "ID",
                    "illinois": "IL", "indiana": "IN", "iowa": "IA", "kansas": "KS",
                    "kentucky": "KY", "louisiana": "LA", "maine": "ME", "maryland": "MD",
                    "massachusetts": "MA", "michigan": "MI", "minnesota": "MN", "mississippi": "MS",
                    "missouri": "MO", "montana": "MT", "nebraska": "NE", "nevada": "NV",
                    "new hampshire": "NH", "new jersey": "NJ", "new mexico": "NM", "new york": "NY",
                    "north carolina": "NC", "north dakota": "ND", "ohio": "OH", "oklahoma": "OK",
                    "oregon": "OR", "pennsylvania": "PA", "rhode island": "RI", "south carolina": "SC",
                    "south dakota": "SD", "tennessee": "TN", "texas": "TX", "utah": "UT",
                    "vermont": "VT", "virginia": "VA", "washington": "WA", "west virginia": "WV",
                    "wisconsin": "WI", "wyoming": "WY"
                }
                abbr_to_state = {v: k for k, v in state_to_abbr.items()}

                ps_lower = profile_state.lower().strip()
                # Determine both full name and abbreviation
                if len(ps_lower) == 2:
                    state_abbr = ps_lower.upper()
                    state_full = abbr_to_state.get(state_abbr, "").lower()
                else:
                    state_full = ps_lower
                    state_abbr = state_to_abbr.get(ps_lower, "").upper()

                # Try matching: full name, abbreviation, "USA-XX" pattern
                for t in option_texts:
                    tl = t.lower()
                    if state_full and state_full in tl:
                        return t
                    if state_abbr and (state_abbr.lower() in tl.split('-') or
                                       tl.endswith(state_abbr.lower()) or
                                       f"usa-{state_abbr.lower()}" == tl):
                        return t
                    if ps_lower in tl or tl in ps_lower:
                        return t

                # Options may be virtualized (React Select) — not all visible in prescan.
                # If we see USA-XX or CAN-XX pattern in existing options, generate the code directly.
                if state_abbr:
                    has_usa_pattern = any(t.startswith('USA-') for t in option_texts[:20])
                    if has_usa_pattern:
                        return f"USA-{state_abbr}"

                # Also try city + state combo
                profile_city = self.profile.get("personal.city") if self.profile else ""
                if profile_city:
                    for t in option_texts:
                        if profile_city.lower() in t.lower():
                            return t

        # 3. Pronouns
        if 'pronoun' in label:
            for t in option_texts:
                if 'he/him' in t.lower():
                    return t

        # 4. How did you hear / source
        if any(k in label for k in ['how did you', 'learn about', 'hear about', 'source']):
            for t in option_texts:
                tl = t.lower()
                if any(k in tl for k in ['linkedin', 'job board', 'online']):
                    return t
            # Return "Other" if available
            for t in option_texts:
                if t.lower() in ('other', 'other (please specify)'):
                    return t

        # 5. Previously employed
        if any(k in label for k in ['previously', 'former', 'employed']):
            for t in option_texts:
                if t.lower().startswith('no'):
                    return t

        # 5b. Years of experience questions — JD-aware selection
        # For "How many years..." with options like "No experience", "1-3", "4-7", "8+"
        if any(k in label for k in ['how many years', 'years of experience', 'years experience']):
            # Build ranked options list
            ranked = []
            for t in option_texts:
                tl = t.lower()
                if 'no experience' in tl or 'none' in tl or '0' == tl.strip():
                    ranked.append((0, t))
                elif '8+' in tl or '10+' in tl or '15+' in tl:
                    ranked.append((100, t))
                elif '4' in tl and '7' in tl:
                    ranked.append((70, t))
                elif '3' in tl and '7' in tl:
                    ranked.append((70, t))
                elif '1' in tl and ('3' in tl or '2' in tl):
                    ranked.append((30, t))
                elif 'advanced' in tl:
                    ranked.append((100, t))
                elif 'intermediate' in tl:
                    ranked.append((60, t))
                elif 'beginner' in tl:
                    ranked.append((20, t))
                else:
                    ranked.append((50, t))
            if ranked:
                ranked.sort(key=lambda x: x[0], reverse=True)
                # Detect if this is a niche/unrelated domain where we have NO experience
                # Core competencies: PM, TPM, product management, agile, scrum, SAFe, technical,
                # developer, API, platform, consumer, marketplace, engineering, software, data
                core_keywords = ['product manag', 'program manag', 'technical', 'agile', 'scrum',
                                 'safe', 'developer', 'api', 'platform', 'consumer', 'marketplace',
                                 'engineering', 'software', 'data', 'project manag', 'delivery',
                                 'stakeholder', 'cross-functional', 'roadmap', 'strategy']
                # Domains where user explicitly has NO experience
                # NOTE: avoid short substrings that match verbs (e.g. "design" in "designing")
                no_experience_keywords = ['product marketing', 'marketing manager', 'sales ',
                                         'ux design', 'ui design', 'visual design',
                                         'ux research', 'human resources', 'legal ',
                                         'finance ', 'financial analyst',
                                         'accounting', 'recruiting', 'content writing',
                                         'graphic design', 'public relations']

                is_no_experience = any(nk in label for nk in no_experience_keywords)
                is_core = any(ck in label for ck in core_keywords)

                if is_no_experience and not is_core:
                    # Pick "No experience" for domains we don't have
                    for score, text in ranked:
                        if score == 0:
                            print(f"   📊 Years: '{text}' (no experience in this domain)")
                            return text
                    # If no "No experience" option, pick lowest non-zero
                    ranked.sort(key=lambda x: x[0])
                    for score, text in ranked:
                        if score > 0:
                            return text

                # For core competencies or ambiguous — pick highest
                return ranked[0][1]

        # 5c. Proficiency questions — pick highest
        if any(k in label for k in ['level of proficiency', 'proficiency level', 'how proficient']):
            for t in option_texts:
                if t.lower() in ('advanced', 'expert'):
                    return t
            # Fallback to second-highest
            for t in option_texts:
                if t.lower() == 'intermediate':
                    return t

        # 6. Demographics — use profile values first, then fallback to decline
        demo = self.profile.data.get("demographics", {}) if self.profile else {}
        demo_map = {
            'gender': demo.get('gender', ''),
            'hispanic': demo.get('hispanic_latino', ''),
            'latino': demo.get('hispanic_latino', ''),
            'race': demo.get('race_ethnicity', ''),
            'ethnic': demo.get('race_ethnicity', ''),
            'veteran': demo.get('veteran_status', ''),
            'disability': demo.get('disability_status', ''),
        }
        if any(k in label for k in demo_map):
            # Find matching profile value
            profile_val = ''
            for key, val in demo_map.items():
                if key in label and val:
                    profile_val = val
                    break
            # Try profile value first
            if profile_val:
                for t in option_texts:
                    tl = t.lower()
                    pv = profile_val.lower()
                    if pv in tl or tl in pv or tl.startswith(pv) or pv.startswith(tl):
                        return t
            # Fallback to decline/prefer not
            for t in option_texts:
                tl = t.lower()
                if 'decline' in tl or 'prefer not' in tl or 'do not want' in tl or 'not a protected' in tl:
                    return t

        # 7. Claude AI as last resort — ask to pick from specific options
        if self.ai and self.ai.available:
            try:
                clean_options = [t for t in option_texts if t.lower() not in ('--', 'please select', 'select', '')]
                if clean_options:
                    context = self._get_profile_context_for_ai()
                    ai_pick = self.ai.choose_option(field.label, clean_options, context)
                    if ai_pick:
                        return ai_pick
            except Exception as e:
                print(f"      ⚠️ AI re-resolve error: {str(e)[:40]}")

        return ""

    def _fill_school_autocomplete(self, el: ElementHandle, field: FormField, frame) -> bool:
        """
        Fill school/university searchable dropdown with fallback.

        PROVEN METHOD (from V6):
        1. Click dropdown and type search text
        2. Wait for search API results
        3. If good match found - select it
        4. Otherwise use fallback "0 - Other"
        """
        FALLBACK_OPTIONS = ["0 - Other", "Other", "Not Listed"]

        # Close any open dropdowns
        self.page.keyboard.press('Escape')
        time.sleep(0.1)

        # Open school dropdown — JS-first to avoid Playwright actionability timeouts
        try:
            frame.evaluate(f'''() => {{
                const el = document.querySelector('{field.selector}');
                if (el) {{
                    el.scrollIntoView({{block: 'center'}});
                    el.focus();
                    el.click();
                    el.dispatchEvent(new MouseEvent('mousedown', {{bubbles: true}}));
                }}
            }}''')
        except:
            try:
                el.click(force=True, timeout=5000)
            except:
                return False
        time.sleep(0.3)

        # Type search - use shorter text for better matches
        search_text = field.answer[:30] if len(field.answer) > 30 else field.answer
        self.page.keyboard.type(search_text, delay=20)
        time.sleep(1.0)  # Wait for search API

        # Get options via aria-controls
        controls_id = el.get_attribute('aria-controls')
        options = []

        if controls_id:
            listbox = frame.query_selector(f'#{controls_id}')
            if listbox:
                options = listbox.query_selector_all('[role="option"]')

        # Fallback to global selectors
        if not options:
            options = frame.query_selector_all('[role="option"], .select__option')

        # Check if we found good results
        if options:
            first_text = options[0].inner_text().strip().lower()
            # Check if first option is valid (not "no results" message)
            if 'no result' not in first_text and 'no option' not in first_text:
                # Check for match with our search
                search_lower = field.answer.lower()
                for opt in options:
                    opt_text = opt.inner_text().strip()
                    opt_lower = opt_text.lower()
                    # Match if search text appears in option
                    if search_lower[:15] in opt_lower or opt_lower in search_lower:
                        try:
                            opt.click(force=True, timeout=5000)
                        except:
                            opt.evaluate("e => e.click()")
                        time.sleep(0.2)
                        print(f"      🎓 School matched: {opt_text[:40]}")
                        return True
                # No exact match but results exist - take first
                try:
                    options[0].click(force=True, timeout=5000)
                except:
                    options[0].evaluate("e => e.click()")
                time.sleep(0.2)
                return True

        # No results from search - try fallback options
        print(f"      ⚠️ School not found, trying fallback...")
        self.page.keyboard.press('Escape')
        time.sleep(0.1)
        try:
            frame.evaluate(f'''() => {{
                const el = document.querySelector('{field.selector}');
                if (el) {{ el.focus(); el.click(); }}
            }}''')
        except:
            try:
                el.click(force=True, timeout=5000)
            except:
                pass
        time.sleep(0.2)

        # Clear and try each fallback
        for fallback in FALLBACK_OPTIONS:
            self.page.keyboard.press('Control+a')
            self.page.keyboard.press('Backspace')
            self.page.keyboard.type(fallback, delay=20)
            time.sleep(0.8)

            if controls_id:
                listbox = frame.query_selector(f'#{controls_id}')
                if listbox:
                    options = listbox.query_selector_all('[role="option"]')
                    if options:
                        first_text = options[0].inner_text().strip()
                        if 'no result' not in first_text.lower():
                            try:
                                options[0].click(force=True, timeout=5000)
                            except:
                                options[0].evaluate("e => e.click()")
                            time.sleep(0.2)
                            print(f"      🎓 School fallback: {first_text[:40]}")
                            return True

        # Last resort: just tab out
        self.page.keyboard.press('Escape')
        self.page.keyboard.press('Tab')
        print(f"      ⚠️ School: typed but no selection")
        return True

    def _fill_checkbox(self, el: ElementHandle, field: FormField) -> bool:
        """Fill checkbox."""
        should_check = field.answer.lower() in ("yes", "true", "1", "checked")
        if should_check != el.is_checked():
            el.click()
        return True
    
    def _fill_file(self, el: ElementHandle, field: FormField) -> bool:
        """Upload file via Attach button (Greenhouse S3) or direct set_input_files."""
        context = getattr(self, '_active_frame', self.page)

        # Strategy 1: Greenhouse S3 upload — click "Attach" button inside fieldset
        # Greenhouse hides <input type="file"> inside S3 forms and uses JS upload handlers
        # The correct approach is: click "Attach" → intercept file_chooser → set file
        try:
            # Find parent fieldset (resume_fieldset, cover_letter_fieldset)
            fieldset_id = el.evaluate("""el => {
                let parent = el;
                for (let i = 0; i < 8 && parent; i++) {
                    if (parent.tagName === 'FIELDSET' && parent.id) return parent.id;
                    parent = parent.parentElement;
                }
                // Also check if selector points to fieldset itself
                if (el.tagName === 'FIELDSET') return el.id || '';
                return '';
            }""")

            if fieldset_id:
                # Found fieldset — look for Attach button inside it
                attach_btn = context.query_selector(
                    f"#{fieldset_id} button[data-source='attach'], "
                    f"#{fieldset_id} button:has-text('Attach')"
                )
                if attach_btn and attach_btn.is_visible():
                    try:
                        with self.page.expect_file_chooser(timeout=5000) as fc_info:
                            attach_btn.click(timeout=3000)
                        file_chooser = fc_info.value
                        file_chooser.set_files(field.answer)
                        time.sleep(1.0)
                        print(f"      📎 Uploaded via Attach button (fieldset: {fieldset_id})")
                        return True
                    except Exception as e:
                        print(f"      ⚠️ Attach button upload failed: {e}")
                        # Fall through to Strategy 2
        except Exception as e:
            print(f"      ⚠️ Fieldset detection error: {e}")

        # Strategy 2: Find Attach button near the element (broader search)
        try:
            # Check selector for fieldset hints (e.g. #s3_upload_for_cover_letter input[type='file'])
            selector_lower = (field.selector or "").lower()
            fieldset_hint = ""
            if "cover_letter" in selector_lower:
                fieldset_hint = "cover_letter_fieldset"
            elif "resume" in selector_lower or "s3_upload_for_resume" in selector_lower:
                fieldset_hint = "resume_fieldset"

            if fieldset_hint:
                attach_btn = context.query_selector(
                    f"#{fieldset_hint} button[data-source='attach']"
                )
                if attach_btn and attach_btn.is_visible():
                    try:
                        with self.page.expect_file_chooser(timeout=5000) as fc_info:
                            attach_btn.click(timeout=3000)
                        file_chooser = fc_info.value
                        file_chooser.set_files(field.answer)
                        time.sleep(1.0)
                        print(f"      📎 Uploaded via Attach button (hint: {fieldset_hint})")
                        return True
                    except Exception as e:
                        print(f"      ⚠️ Attach button (hint) failed: {e}")
        except Exception as e:
            pass

        # Strategy 3: Direct set_input_files (works for standard file inputs)
        try:
            el.set_input_files(field.answer)
            time.sleep(0.5)
            print(f"      📎 Uploaded via direct set_input_files")
            return True
        except Exception as e:
            print(f"      ❌ Direct file upload failed: {e}")
            return False
    
    def _interactive_fill(self, field: FormField):
        """Interactive fill for unknown fields."""
        self.browser.highlight_element(field.selector, "orange")
        
        print(f"\n📌 {field.label}")
        print(f"   Type: {field.field_type.value}")
        
        if field.options:
            print(f"   Options: {field.options[:5]}...")
        
        # Try AI
        if self.ai.available:
            if field.options:
                suggested = self.ai.choose_option(field.label, field.options, self.profile.get_context())
            else:
                suggested = self.ai.generate(field.label, self.profile.get_context())
            
            if suggested:
                print(f"   💡 AI suggests: {suggested[:50]}")
                field.answer = suggested
                field.answer_source = AnswerSource.AI
        
        print(f"\n   Type answer (or press Enter to skip):")
        user_input = input("   > ").strip()
        
        if user_input:
            field.answer = user_input
            field.answer_source = AnswerSource.HUMAN
            self._fill_field(field)
            
            # Learn for next time
            is_dropdown = field.field_type in (FieldType.SELECT, FieldType.AUTOCOMPLETE)
            self.learned_db.save_answer(field.label, user_input, is_dropdown)
        else:
            field.status = FillStatus.SKIPPED
        
        self.browser.unhighlight_element(field.selector)
    
    # ─────────────────────────────────────────────────────────────────────
    # LAYER 4: VALIDATION
    # ─────────────────────────────────────────────────────────────────────
    
    def _blur_all_fields(self):
        """Click outside all fields to trigger blur/validation."""
        try:
            context = getattr(self, '_active_frame', self.page)
            # Click on body to blur any focused field
            try:
                context.click('body', position={'x': 10, 'y': 10})
            except:
                self.page.click('body', position={'x': 10, 'y': 10})
            time.sleep(0.3)

            # Also blur each field explicitly
            for field in self.fields:
                if field.selector:
                    el = context.query_selector(field.selector)
                    if not el:
                        el = self.page.query_selector(field.selector)
                    if el:
                        try:
                            el.evaluate("e => e.blur()")
                        except:
                            pass
            time.sleep(0.3)
        except:
            pass
    
    def _validate_all_fields(self):
        """Validate all filled fields."""
        print("\n✓ Validating...")
        
        for field in self.fields:
            if field.status == FillStatus.FILLED:
                self._validate_field(field)
        
        verified = sum(1 for f in self.fields if f.status == FillStatus.VERIFIED)
        print(f"   Verified: {verified}/{sum(1 for f in self.fields if f.status in (FillStatus.FILLED, FillStatus.VERIFIED))}")
    
    def _validate_field(self, field: FormField) -> bool:
        """Validate single field by reading back value."""
        try:
            context = getattr(self, '_active_frame', self.page)
            el = context.query_selector(field.selector)
            if not el:
                el = self.page.query_selector(field.selector)
            if not el:
                return False
            
            actual = self._get_value(el, field.html_tag, field.input_type)
            
            # For file fields, check if Greenhouse shows the filename
            if field.field_type == FieldType.FILE:
                # Try to verify by checking if filename is visible in the fieldset
                try:
                    context = getattr(self, '_active_frame', self.page)
                    fid = (field.element_id or "").lower()
                    # Greenhouse shows filename in #resume_filename or #cover_letter_filename
                    for name_id in ["resume_filename", "cover_letter_filename"]:
                        if fid.replace("_fieldset", "") in name_id or \
                           name_id.replace("_filename", "") in fid:
                            fn_el = context.query_selector(f"#{name_id}")
                            if fn_el:
                                fn_text = fn_el.inner_text().strip()
                                if fn_text:
                                    print(f"   📎 File verified: {fn_text}")
                                    field.status = FillStatus.VERIFIED
                                    return True
                except:
                    pass
                # Fallback: assume verified (file inputs can't be read back)
                field.status = FillStatus.VERIFIED
                return True
            
            # Compare
            if actual and (
                actual.lower() == field.answer.lower() or
                field.answer.lower() in actual.lower() or
                actual.lower() in field.answer.lower()
            ):
                field.status = FillStatus.VERIFIED
                return True
            
            # Check for errors
            aria_invalid = el.get_attribute("aria-invalid")
            if aria_invalid == "true":
                field.status = FillStatus.ERROR
                field.error_message = "Field marked as invalid"
                return False
            
            # Partial match is ok for autocomplete
            if field.field_type == FieldType.AUTOCOMPLETE:
                field.status = FillStatus.VERIFIED
                return True
            
            return False
            
        except:
            return False
    
    # ─────────────────────────────────────────────────────────────────────
    # REPORT
    # ─────────────────────────────────────────────────────────────────────
    
    # ─────────────────────────────────────────────────────────────────────
    # REPEATABLE SECTIONS (Work Experience, Education)
    # ─────────────────────────────────────────────────────────────────────
    
    def fill_section_entry(self, section_name: str, entry_index: int, form_index: int) -> bool:
        """Fill one entry of a repeatable section (work experience, education)."""
        config = self.REPEATABLE_SECTIONS.get(section_name)
        if not config:
            return False

        profile_key = config['profile_key']
        field_patterns = config['field_patterns']
        skip_end_if_current = config.get('skip_end_date_if_current', False)

        entries = self.profile.data.get(profile_key, [])
        if entry_index >= len(entries):
            return False

        entry = entries[entry_index]
        is_current = entry.get('current', False)

        company_or_school = entry.get('company', entry.get('school', ''))
        print(f"   📝 {section_name}[{entry_index}]: {company_or_school[:40]}")

        filled_any = False
        context = getattr(self, '_active_frame', self.page)

        # ── Phase 1: Fill ID-based fields (school, degree, discipline, company, title) ──
        # Try primary patterns first, then legacy patterns
        all_patterns = dict(field_patterns)
        legacy = config.get('legacy_field_patterns', {})

        for pattern, field_name in all_patterns.items():
            selector = '#' + pattern.replace('{N}', str(form_index))

            if skip_end_if_current and is_current and 'end' in field_name:
                print(f"      ⏭️ {field_name} (skipped - current role)")
                continue

            value = entry.get(field_name, '')
            if not value and field_name != 'current':
                continue

            search_value = entry.get(f'{field_name}_search', '')
            fallback_value = entry.get(f'{field_name}_fallback', '')

            if field_name == 'current':
                if value == True:
                    value = 'checked'
                else:
                    continue

            # Find element — try primary selector, then with s2id_ prefix (Select2), then legacy
            el = context.query_selector(selector)
            if not el and context != self.page:
                el = self.page.query_selector(selector)

            # If element is hidden (Greenhouse wraps selects in Select2), prefer the s2id_ container
            if el:
                try:
                    el_visible = el.is_visible()
                    el_type_check = el.get_attribute('type') or ''
                    if not el_visible or el_type_check == 'hidden':
                        # Found hidden element — try Select2 container by known prefix
                        s2id_selector = '#s2id_' + pattern.replace('{N}', str(form_index))
                        s2id_el = context.query_selector(s2id_selector)
                        if not s2id_el and context != self.page:
                            s2id_el = self.page.query_selector(s2id_selector)
                        if s2id_el:
                            el = s2id_el  # Use visible Select2 container
                        else:
                            # Select2 may use auto-generated IDs (s2id_autogen14)
                            # Find Select2 container via DOM: sibling .select2-container
                            try:
                                s2_el = context.evaluate_handle(f'''() => {{
                                    const hidden = document.querySelector('{selector}');
                                    if (!hidden) return null;
                                    // Select2 adds .select2-container as next sibling of the hidden <select>
                                    let sib = hidden.nextElementSibling;
                                    if (sib && sib.classList.contains('select2-container')) return sib;
                                    // Or look in parent
                                    let parent = hidden.parentElement;
                                    if (parent) {{
                                        const s2 = parent.querySelector('.select2-container');
                                        if (s2) return s2;
                                    }}
                                    return null;
                                }}''')
                                if s2_el:
                                    el = s2_el.as_element()
                                    if el:
                                        print(f"      🔍 Found Select2 container via DOM for {field_name}")
                            except:
                                pass
                except:
                    pass

            # For Greenhouse Select2: the actual element is s2id_{id}
            if not el:
                s2id_selector = '#s2id_' + pattern.replace('{N}', str(form_index))
                el = context.query_selector(s2id_selector)
                if not el and context != self.page:
                    el = self.page.query_selector(s2id_selector)

            # Try legacy patterns
            if not el and field_name in {fp_fn for fp_fn in legacy.values()}:
                for leg_pat, leg_fn in legacy.items():
                    if leg_fn == field_name:
                        leg_sel = '#' + leg_pat.replace('{N}', str(form_index))
                        el = context.query_selector(leg_sel)
                        if not el and context != self.page:
                            el = self.page.query_selector(leg_sel)
                        if el:
                            break

            if not el:
                print(f"      ⚠️ {field_name}: element not found")
                continue

            try:
                ok = self._fill_section_element(el, field_name, value, search_value, fallback_value)
                if ok:
                    filled_any = True
                    print(f"      ✅ {field_name}: {str(value)[:25]}")
                    # Track filled element IDs so main loop skips them
                    # Add primary pattern ID
                    actual_id = pattern.replace('{N}', str(form_index))
                    self._section_filled_ids.add(actual_id)
                    self._section_filled_ids.add(f's2id_{actual_id}')
                    # Also track the actual element's DOM ID (may differ from pattern)
                    try:
                        el_dom_id = el.get_attribute('id')
                        if el_dom_id:
                            self._section_filled_ids.add(el_dom_id)
                            self._section_filled_ids.add(f's2id_{el_dom_id}')
                    except:
                        pass
                    # Add all legacy patterns for this field_name
                    for leg_pat, leg_fn in legacy.items():
                        if leg_fn == field_name:
                            leg_id = leg_pat.replace('{N}', str(form_index))
                            self._section_filled_ids.add(leg_id)
                            self._section_filled_ids.add(f's2id_{leg_id}')
                time.sleep(0.1)
            except Exception as e:
                print(f"      ❌ {field_name}: {e}")

        # ── Phase 2: Fill date fields (name-based selectors, no ID) ──
        date_selectors = config.get('date_selectors', {})
        if date_selectors:
            for field_name, name_selector in date_selectors.items():
                if skip_end_if_current and is_current and 'end' in field_name:
                    print(f"      ⏭️ {field_name} (skipped - current role)")
                    continue

                value = entry.get(field_name, '')
                if not value:
                    continue

                # Find all matching elements (one per education entry) and pick by index
                els = context.query_selector_all(name_selector)
                if not els and context != self.page:
                    els = self.page.query_selector_all(name_selector)

                if form_index < len(els):
                    el = els[form_index]
                    try:
                        fill_value = str(value)
                        # DOM-aware format adaptation
                        fill_value = self._adapt_value_to_dom_element(fill_value, el)
                        el.fill(fill_value)
                        filled_any = True
                        print(f"      ✅ {field_name}: {fill_value}")
                        # Track by element ID so _mark_section_filled_fields can skip it
                        el_id = el.get_attribute('id')
                        if el_id:
                            self._section_filled_ids.add(el_id)
                        time.sleep(0.1)
                    except Exception as e:
                        print(f"      ❌ {field_name}: {e}")
                else:
                    print(f"      ⚠️ {field_name}: element[{form_index}] not found (only {len(els)} elements)")

        return filled_any

    def _fill_section_element(self, el, field_name: str, value: str,
                               search_value: str = '', fallback_value: str = '') -> bool:
        """Fill a single element within a repeatable section. Handles Select2, select, combobox, text."""
        tag = el.evaluate("e => e.tagName.toLowerCase()")
        el_type = el.get_attribute('type') or 'text'
        role = el.get_attribute('role') or ''
        aria = el.get_attribute('aria-haspopup') or ''
        el_class = el.get_attribute('class') or ''

        # ── Select2 container (Greenhouse education school/degree/discipline) ──
        if 'select2-container' in el_class or (tag == 'div' and 's2id' in (el.get_attribute('id') or '')):
            return self._fill_section_select2(el, field_name, value, search_value, fallback_value)

        # ── Native <select> ──
        if tag == 'select':
            el.select_option(label=str(value))
            return True

        # ── Combobox / aria autocomplete ──
        if role == 'combobox' or aria in ('true', 'listbox'):
            return self._fill_section_autocomplete(el, field_name, value, search_value, fallback_value)

        # ── Checkbox ──
        if el_type == 'checkbox':
            if value == 'checked' and not el.is_checked():
                el.click()
            return True

        # ── Text input ──
        fill_value = str(value)
        fill_value = self._adapt_value_to_dom_element(fill_value, el)
        el.fill(fill_value)
        return True

    def _fill_section_select2(self, container, field_name: str, value: str,
                               search_value: str = '', fallback_value: str = '') -> bool:
        """Fill Select2 dropdown within a repeatable section (e.g., Greenhouse education)."""
        frame = container.owner_frame() or self.page

        # Open Select2 by clicking the choice element (always via JS to avoid actionability timeouts)
        try:
            container.evaluate('''el => {
                const choice = el.querySelector('a.select2-choice, .select2-choices');
                if (choice) {
                    choice.dispatchEvent(new MouseEvent('mousedown', {bubbles: true}));
                } else {
                    // Fallback: try clicking the container itself
                    el.dispatchEvent(new MouseEvent('mousedown', {bubbles: true}));
                    el.click();
                }
            }''')
        except:
            try:
                container.click(force=True, timeout=5000)
            except:
                return False

        time.sleep(0.5)

        # Type search text
        search_text = str(search_value or value)[:40]

        # Find the search input inside the open dropdown
        search_input = frame.query_selector('.select2-drop:not(.select2-display-none) .select2-input')
        if not search_input:
            search_input = frame.query_selector('.select2-search input')

        if search_input:
            try:
                search_input.fill('', timeout=3000)
            except:
                pass  # fill may fail on hidden inputs — keyboard will handle it
            self.page.keyboard.type(search_text, delay=20)
        else:
            self.page.keyboard.type(search_text, delay=20)

        time.sleep(0.8)

        # Read options
        opts = frame.query_selector_all('.select2-drop:not(.select2-display-none) .select2-results li.select2-result')
        if not opts:
            opts = frame.query_selector_all('.select2-results li.select2-result')
        if not opts:
            opts = frame.query_selector_all('.select2-results li')

        # Filter valid options
        valid_opts = []
        for opt in opts:
            try:
                cls = opt.get_attribute('class') or ''
                if 'select2-disabled' in cls or 'select2-result-unselectable' in cls:
                    continue
                text = opt.inner_text().strip()
                if text and 'no result' not in text.lower():
                    valid_opts.append((opt, text))
            except:
                continue

        # Score-based matching
        value_lower = str(value).lower()
        best_match = None
        best_score = 0

        for opt, opt_text in valid_opts:
            opt_lower = opt_text.lower()
            score = 0
            if opt_lower == value_lower:
                score = 100
            elif value_lower in opt_lower:
                score = 90
            elif opt_lower in value_lower:
                score = 85
            elif 'school' in field_name.lower():
                # School: strict word overlap
                value_words = set(value_lower.split())
                opt_words = set(opt_lower.split())
                overlap = len(value_words & opt_words)
                total = max(len(value_words), len(opt_words))
                if total > 0 and overlap / total >= 0.6:
                    score = 70
            else:
                value_words = set(value_lower.split())
                opt_words = set(opt_lower.split())
                overlap = len(value_words & opt_words)
                if overlap >= 2:
                    score = 60
            if score > best_score:
                best_score = score
                best_match = opt

        if best_match and best_score >= 70:
            best_match.click()
            print(f"      ✅ Select2 matched (score={best_score})")
            time.sleep(0.2)
            return True

        # No confident match → fallback (e.g., "Other" for schools)
        if fallback_value:
            self.page.keyboard.press('Escape')
            time.sleep(0.2)

            # Re-open dropdown
            try:
                container.evaluate('''el => {
                    const choice = el.querySelector('a.select2-choice, .select2-choices');
                    if (choice) choice.dispatchEvent(new MouseEvent('mousedown', {bubbles: true}));
                }''')
            except:
                container.click()
            time.sleep(0.5)

            # Type fallback
            search_input = frame.query_selector('.select2-drop:not(.select2-display-none) .select2-input')
            if not search_input:
                search_input = frame.query_selector('.select2-search input')
            if search_input:
                try:
                    search_input.fill('', timeout=3000)
                except:
                    pass
                self.page.keyboard.type(fallback_value, delay=20)
            else:
                self.page.keyboard.type(fallback_value, delay=20)
            time.sleep(0.8)

            # Find exact "Other" match
            fb_opts = frame.query_selector_all('.select2-drop:not(.select2-display-none) .select2-results li.select2-result')
            if not fb_opts:
                fb_opts = frame.query_selector_all('.select2-results li')

            for opt in fb_opts:
                try:
                    text = opt.inner_text().strip()
                    if text.lower() == fallback_value.lower():
                        opt.click()
                        print(f"      ✅ Used fallback: '{fallback_value}'")
                        time.sleep(0.2)
                        return True
                except:
                    continue

            # Click first valid option
            for opt in fb_opts:
                try:
                    text = opt.inner_text().strip()
                    if text and 'no result' not in text.lower():
                        opt.click()
                        print(f"      ✅ Used fallback (first match): '{text[:30]}'")
                        time.sleep(0.2)
                        return True
                except:
                    continue

        # Last resort — click first valid option or escape
        if valid_opts:
            valid_opts[0][0].click()
            print(f"      ⚠️ No match, using first: '{valid_opts[0][1][:30]}'")
            time.sleep(0.2)
            return True

        self.page.keyboard.press('Escape')
        return False

    def _fill_section_autocomplete(self, el, field_name: str, value: str,
                                    search_value: str = '', fallback_value: str = '') -> bool:
        """Fill combobox/autocomplete within a repeatable section."""
        frame = el.owner_frame() or self.page

        self.page.keyboard.press('Escape')
        time.sleep(0.1)

        try:
            el.scroll_into_view_if_needed(timeout=3000)
        except:
            pass
        try:
            el.click(force=True, timeout=5000)
        except:
            # Fallback: JS dispatch
            try:
                el.evaluate("e => { e.focus(); e.dispatchEvent(new MouseEvent('mousedown', {bubbles:true})); }")
            except:
                return False
        time.sleep(0.3)

        search_text = str(search_value or value)[:40]
        try:
            el.fill('', timeout=3000)
        except:
            pass
        self.page.keyboard.type(search_text, delay=15)
        time.sleep(0.5)

        controls_id = el.get_attribute('aria-controls')
        opts = []
        if controls_id:
            listbox = frame.query_selector(f'#{controls_id}')
            if listbox:
                opts = listbox.query_selector_all('[role="option"]')
        if not opts:
            opts = frame.query_selector_all('[role="option"], .select__option')

        value_lower = str(value).lower()
        best_match = None
        best_score = 0

        for opt in opts:
            opt_text = opt.inner_text().strip().lower()
            score = 0
            if opt_text == value_lower:
                score = 100
            elif value_lower in opt_text:
                score = 90
            elif opt_text in value_lower:
                score = 85
            elif 'school' in field_name.lower():
                value_words = set(value_lower.split())
                opt_words = set(opt_text.split())
                overlap = len(value_words & opt_words)
                total = max(len(value_words), len(opt_words))
                if total > 0 and overlap / total >= 0.5:
                    score = 70
            else:
                value_words = set(value_lower.split())
                opt_words = set(opt_text.split())
                overlap = len(value_words & opt_words)
                if overlap >= 3:
                    score = 70
                elif overlap >= 2:
                    score = 60
            if score > best_score:
                best_score = score
                best_match = opt

        if best_match and best_score >= 70:
            best_match.click()
            print(f"      ✅ Autocomplete matched (score={best_score})")
            time.sleep(0.2)
            return True

        # Fallback
        if fallback_value:
            self.page.keyboard.press('Escape')
            time.sleep(0.1)
            try:
                el.fill('', timeout=3000)
            except:
                pass
            self.page.keyboard.type(fallback_value, delay=15)
            time.sleep(0.5)
            new_opts = []
            if controls_id:
                listbox = frame.query_selector(f'#{controls_id}')
                if listbox:
                    new_opts = listbox.query_selector_all('[role="option"]')
            if not new_opts:
                new_opts = frame.query_selector_all('[role="option"], .select__option')
            for opt in new_opts:
                if opt.inner_text().strip().lower() == fallback_value.lower():
                    opt.click()
                    print(f"      ✅ Used fallback: '{fallback_value}'")
                    time.sleep(0.2)
                    return True
            if new_opts:
                new_opts[-1].click()
                print(f"      ✅ Used fallback (last option)")
                time.sleep(0.2)
                return True

        if opts:
            opts[0].click()
            print(f"      ⚠️ No match — using first option")
            time.sleep(0.2)
            return True

        self.page.keyboard.press('Tab')
        return False
    
    def click_add_another(self, section_name: str) -> bool:
        """Click 'Add another' button for a section. Searches in iframe too."""
        config = self.REPEATABLE_SECTIONS.get(section_name)
        if not config:
            return False

        contexts = [getattr(self, '_active_frame', self.page), self.page]
        section_text = config.get('section_text', section_name.replace('_', ' '))

        for context in contexts:
            # Try configured selectors
            for selector in config.get('add_button_selectors', []):
                try:
                    btn = context.query_selector(selector)
                    if btn and btn.is_visible():
                        btn.click()
                        time.sleep(0.5)
                        self.browser.wait_for_stable()
                        return True
                except:
                    continue

            # Text-based search — must match section context
            try:
                all_links = context.query_selector_all('a, button')
                for link in all_links:
                    try:
                        if not link.is_visible():
                            continue
                        text = link.inner_text().strip().lower()
                        if ('add another' in text or 'add a' in text) and \
                           section_text.lower() in text:
                            link.click()
                            time.sleep(0.5)
                            self.browser.wait_for_stable()
                            return True
                    except:
                        continue
            except:
                continue

        return False
    
    def fill_repeatable_section(self, section_name: str):
        """Fill all entries for a repeatable section."""
        config = self.REPEATABLE_SECTIONS.get(section_name)
        if not config:
            return
        
        entries = self.profile.data.get(config['profile_key'], [])
        if not entries:
            print(f"   ⚠️ No {section_name} entries in profile")
            return
        
        print(f"\n🔄 Filling {section_name}: {len(entries)} entries")
        
        # Fill first entry — if no fields found, skip the whole section
        first_ok = False
        try:
            first_ok = self.fill_section_entry(section_name, entry_index=0, form_index=0)
        except Exception as e:
            print(f"   ❌ Error filling {section_name}[0]: {e}")

        if not first_ok:
            print(f"   ⚠️ No {section_name} fields found on this form, skipping remaining entries")
            return

        # Add and fill additional entries
        for i in range(1, len(entries)):
            try:
                print(f"\n   ➕ Adding {section_name} entry {i+1}...")

                # Check if slot already exists (Greenhouse pre-renders 4 education slots)
                first_pattern = list(config['field_patterns'].keys())[0]
                new_id = first_pattern.replace('{N}', str(i))
                slot_exists = self.page.query_selector(f'#{new_id}') or \
                              self.page.query_selector(f'#s2id_{new_id}')

                if not slot_exists:
                    # Try clicking "Add another" button
                    if self.click_add_another(section_name):
                        time.sleep(1.0)  # Wait for React to render new fields
                    else:
                        print(f"   ⚠️ Could not add entry {i+1} and slot not pre-rendered")
                        break
                else:
                    print(f"      ℹ️ Slot {i} already exists (pre-rendered)")

                # Scroll to new section to ensure visibility
                new_el = self.page.query_selector(f'#{new_id}') or \
                         self.page.query_selector(f'#s2id_{new_id}')
                if new_el:
                    try:
                        new_el.scroll_into_view_if_needed(timeout=3000)
                    except:
                        pass
                    time.sleep(0.3)

                self.fill_section_entry(section_name, entry_index=i, form_index=i)
            except Exception as e:
                print(f"   ❌ Error filling {section_name}[{i}]: {e}")
                continue  # Try next entry instead of crashing
    
    def fill_all_repeatable_sections(self):
        """Fill all repeatable sections (work experience, education)."""
        print("\n" + "="*60)
        print("📋 FILLING REPEATABLE SECTIONS")
        print("="*60)
        
        for section_name in self.REPEATABLE_SECTIONS.keys():
            try:
                self.fill_repeatable_section(section_name)
            except Exception as e:
                print(f"   ❌ Error in {section_name} section: {e}")
    
    def _mark_section_filled_fields(self):
        """After repeatable sections are filled, mark overlapping fields in self.fields as VERIFIED.
        This prevents the main fill loop from overwriting section-filled values."""
        context = getattr(self, '_active_frame', self.page)
        marked = 0

        # Build set of element ID patterns from repeatable sections
        # e.g., "school--", "degree--", "discipline--", "company-name-", "title-"
        section_id_prefixes = set()
        for sec_config in self.REPEATABLE_SECTIONS.values():
            for pattern in sec_config.get('field_patterns', {}):
                prefix = pattern.replace('{N}', '')
                section_id_prefixes.add(prefix)
            for pattern in sec_config.get('legacy_field_patterns', {}):
                prefix = pattern.replace('{N}', '')
                section_id_prefixes.add(prefix)

        for f in self.fields:
            if f.status in (FillStatus.VERIFIED, FillStatus.FILLED, FillStatus.ERROR):
                continue
            f_label = f.label.lower().strip()
            f_eid = f.element_id or ""

            # ── Fast path: field was explicitly filled by section handler (tracked by ID) ──
            # Check for READY, NEEDS_INPUT, SKIPPED — section handler may have filled
            # fields that answer resolver couldn't resolve (e.g., education dates)
            if f_eid and f_eid in self._section_filled_ids:
                f.status = FillStatus.VERIFIED
                f.answer_source = AnswerSource.PROFILE
                # Read actual DOM value so report shows what section handler wrote (not answer resolver guess)
                try:
                    dom_val = self._read_section_field_value(context, f)
                    if dom_val:
                        f.answer = dom_val[:50]
                except:
                    pass
                marked += 1
                continue

            # Check if element ID matches any repeatable section pattern
            is_section_by_id = any(f_eid.startswith(prefix) or f_eid.startswith(f's2id_{prefix}')
                                   for prefix in section_id_prefixes if prefix)

            # Check if this is an education/section field by label
            is_section_field = is_section_by_id or \
                f_label in ('school', 'degree', 'discipline') or \
                'education' in f_label or \
                any(kw in f_label for kw in ('start month', 'start year', 'end month', 'end year',
                                                'start date', 'end date'))
            if not is_section_field:
                continue
            # Check if the underlying element already has a value (set by section handler)
            try:
                el = context.query_selector(f.selector)
                if not el:
                    continue
                tag = el.evaluate("e => e.tagName.toLowerCase()")
                if tag == 'div':
                    # Select2 container — check the display text
                    chosen_text = el.evaluate('''e => {
                        const chosen = e.querySelector('.select2-chosen');
                        return chosen ? chosen.textContent.trim() : '';
                    }''')
                    if chosen_text and chosen_text != '' and chosen_text.lower() not in ('', 'select...', 'select', '—'):
                        f.status = FillStatus.VERIFIED
                        f.answer = chosen_text[:50]
                        f.answer_source = AnswerSource.PROFILE
                        marked += 1
                        continue
                # Check hidden input value
                if tag == 'input':
                    val = el.input_value()
                    if val and val.strip():
                        f.status = FillStatus.VERIFIED
                        f.answer = val[:50]
                        f.answer_source = AnswerSource.PROFILE
                        marked += 1
                # Check React Select combobox — value stored in aria-label or display text
                if role := el.get_attribute('role'):
                    if role == 'combobox':
                        # React Select stores display value in a sibling/parent span
                        react_val = el.evaluate('''e => {
                            const parent = e.closest('.select2-container') || e.parentElement;
                            if (!parent) return '';
                            const span = parent.querySelector('[class*="singleValue"], .select2-chosen');
                            return span ? span.textContent.trim() : '';
                        }''')
                        if react_val and react_val.lower() not in ('', 'select...', 'select', '—'):
                            f.status = FillStatus.VERIFIED
                            f.answer = react_val[:50]
                            f.answer_source = AnswerSource.PROFILE
                            marked += 1
            except:
                pass
        if marked:
            print(f"   ✅ Marked {marked} fields as already filled by section handler")

    def _read_section_field_value(self, context, field) -> str:
        """Read the visible DOM value of a section field (Select2, React Select, or plain input)."""
        el = context.query_selector(field.selector)
        if not el:
            return ""

        tag = el.evaluate("e => e.tagName.toLowerCase()")
        role = el.get_attribute('role') or ''
        el_type = el.get_attribute('type') or ''
        el_id = field.element_id or ''

        # Hidden input (Select2 pattern) — look for s2id_ container's chosen text
        if el_type == 'hidden' or (tag == 'input' and not el.is_visible()):
            s2id = f's2id_{el_id}' if el_id else ''
            if s2id:
                s2_el = context.query_selector(f'#{s2id}')
                if s2_el:
                    chosen = s2_el.evaluate('''e => {
                        const c = e.querySelector('.select2-chosen');
                        return c ? c.textContent.trim() : '';
                    }''')
                    if chosen and chosen.lower() not in ('select...', 'select', '—', ''):
                        return chosen
            # Also try reading hidden input value directly
            val = el.input_value()
            return val if val else ""

        # React Select combobox input — read singleValue from parent
        if role == 'combobox' or 'select__input' in (el.get_attribute('class') or ''):
            val = el.evaluate('''e => {
                let c = e.parentElement;
                for (let i = 0; i < 6; i++) {
                    if (!c) break;
                    const sv = c.querySelector('[class*="singleValue"]');
                    if (sv) return sv.textContent.trim();
                    const sc = c.querySelector('.select2-chosen');
                    if (sc) return sc.textContent.trim();
                    c = c.parentElement;
                }
                return '';
            }''')
            return val if val and val.lower() not in ('select...', 'select', '') else ""

        # Regular text input
        if tag in ('input', 'textarea'):
            return el.input_value()

        # Div container (Select2 or React Select wrapper)
        if tag == 'div':
            val = el.evaluate('''e => {
                const v = e.querySelector('[class*="singleValue"], .select2-chosen');
                return v ? v.textContent.trim() : '';
            }''')
            return val if val and val.lower() not in ('select...', 'select', '') else ""

        return ""

    def _safe_page_title(self) -> str:
        """Get page title safely (page may be closed)."""
        try:
            return self.page.title() if self.page else ""
        except Exception:
            return ""

    def _generate_report(self, url: str) -> FillReport:
        """Generate fill report."""
        report = FillReport(
            url=url,
            title=getattr(self, 'job_title', '') or self._safe_page_title(),
            ats_type=self._detect_ats(url),
            total_fields=len(self.fields),
            fields=self.fields,
        )
        
        for f in self.fields:
            if f.status == FillStatus.READY:
                report.ready_fields += 1
            elif f.status == FillStatus.FILLED:
                report.filled_fields += 1
            elif f.status == FillStatus.VERIFIED:
                report.verified_fields += 1
            elif f.status == FillStatus.NEEDS_INPUT:
                report.needs_input += 1
            elif f.status == FillStatus.SKIPPED:
                report.skipped += 1
            elif f.status == FillStatus.ERROR:
                report.errors += 1
        
        return report

    def _save_application_report(self, report: FillReport):
        """Save JSON + PDF report to Applications folder.

        JSON: full field data (no truncation), JD, file paths, source stats
        PDF: formatted table with header, summary, field details
        """
        from datetime import datetime as _dt
        from pathlib import Path as _Path

        # Determine app_folder (same logic as _generate_personalized_cover_letter)
        files_config = self.profile.data.get("files", {})
        base_path = _Path(files_config.get("base_path", ""))

        jt = getattr(self, 'job_title', '') or ''
        cn = getattr(self, 'company_name', '') or ''
        jd = getattr(self, 'job_description', '') or ''

        if jt and cn and base_path.exists():
            safe_company = re.sub(r'[^\w\s-]', '', cn).strip().replace(' ', '_')[:30]
            safe_title = re.sub(r'[^\w\s-]', '', jt).strip().replace(' ', '_')[:40]
            folder_name = f"{safe_company}_{safe_title}_{_dt.now().strftime('%Y%m%d')}"
            app_folder = base_path / "Applications" / folder_name
        else:
            # Fallback: save alongside form logs
            app_folder = _Path(__file__).parent.parent.parent / "logs" / "form_fills"

        app_folder.mkdir(parents=True, exist_ok=True)
        timestamp = _dt.now().isoformat()

        # ── Build source summary ──
        source_counts = {}
        for f in report.fields:
            src = f.answer_source.value if f.answer_source else "none"
            source_counts[src] = source_counts.get(src, 0) + 1

        # ── Find file paths from filled fields ──
        cv_path = ""
        cl_path = ""
        cv_personalized = False
        cl_personalized = False
        for f in report.fields:
            if f.field_type == FieldType.FILE:
                answer = f.answer or ""
                label_lower = (f.label or "").lower()
                if "cover" in label_lower or "letter" in label_lower:
                    cl_path = answer
                    cl_personalized = f.answer_source == AnswerSource.AI
                else:
                    cv_path = answer
                    cv_personalized = f.answer_source == AnswerSource.AI

        # ── JSON Report (full, no truncation) ──
        json_data = {
            "meta": {
                "timestamp": timestamp,
                "url": report.url,
                "ats_type": report.ats_type,
                "job_title": jt,
                "company": cn,
                "status": "completed" if report.errors == 0 else "with_errors",
                "duration_seconds": round(getattr(self, '_fill_duration', 0), 1),
                "fields_total": report.total_fields,
                "fields_filled": report.verified_fields + report.filled_fields,
                "fields_skipped": report.skipped,
                "fields_error": report.errors,
            },
            "files": {
                "cv_path": cv_path,
                "cover_letter_path": cl_path,
                "cv_personalized": cv_personalized,
                "cover_letter_personalized": cl_personalized,
            },
            "job_description": jd,
            "fields": [],
            "source_summary": source_counts,
            "errors": [f.error_message for f in report.fields if f.status == FillStatus.ERROR],
        }

        for f in report.fields:
            question = (f.label or "").split(" [")[0].strip()  # Remove [name=...] suffix
            json_data["fields"].append({
                "id": f.element_id or f.selector or "",
                "type": f.field_type.value if f.field_type else "unknown",
                "question": question,
                "answer": f.answer or "",
                "source": f.answer_source.value if f.answer_source else "none",
                "status": f.status.value if f.status else "unknown",
            })

        json_path = app_folder / "application_report.json"
        with open(json_path, "w", encoding="utf-8") as fp:
            json.dump(json_data, fp, indent=2, ensure_ascii=False)

        # ── PDF Report ──
        try:
            self._generate_pdf_report(json_data, app_folder)
        except Exception as e:
            print(f"   ⚠️ PDF report error: {e}")

        print(f"   📋 Report saved: {app_folder.name}/")

    def _generate_pdf_report(self, data: dict, app_folder):
        """Generate a formatted PDF report from JSON data."""
        from fpdf import FPDF
        from pathlib import Path as _Path

        def _latin(text: str) -> str:
            """Sanitize text for Helvetica (latin-1) encoding."""
            if not text:
                return ""
            return text.encode('latin-1', 'replace').decode('latin-1')

        def _row_height(pdf_obj, text: str, col_w: float, font_size: int = 8) -> float:
            """Calculate row height needed for multi_cell text."""
            pdf_obj.set_font("Helvetica", "", font_size)
            # Approximate: chars per line based on column width
            char_w = font_size * 0.22  # approximate char width for Helvetica 8pt
            effective_w = col_w - 2  # padding
            if effective_w <= 0:
                effective_w = col_w
            chars_per_line = max(1, int(effective_w / char_w))
            lines = 1
            for paragraph in text.split('\n'):
                if len(paragraph) == 0:
                    lines += 1
                else:
                    lines += max(1, -(-len(paragraph) // chars_per_line))  # ceil division
            return max(5, lines * 4.2)

        # ── Landscape A4 for wider table ──
        pdf = FPDF(orientation='L', format='A4')
        pdf.set_auto_page_break(auto=True, margin=15)
        page_w = pdf.w - pdf.l_margin - pdf.r_margin  # usable width (~257mm)

        pdf.add_page()

        # ── Title ──
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 12, "Application Report", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        # ── Header info ──
        meta = data.get("meta", {})

        info_lines = [
            ("Position", meta.get("job_title", "N/A")),
            ("Company", meta.get("company", "N/A")),
            ("Date", meta.get("timestamp", "")[:19].replace("T", " ")),
            ("ATS", meta.get("ats_type", "N/A")),
            ("URL", meta.get("url", "N/A")),
        ]

        for label, value in info_lines:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(25, 6, f"{label}:", new_x="END")
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 6, f"  {_latin(value)}", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(3)

        # ── Summary box ──
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, "  Summary", fill=True, new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "", 10)
        total = meta.get("fields_total", 0)
        filled = meta.get("fields_filled", 0)
        skipped = meta.get("fields_skipped", 0)
        errors = meta.get("fields_error", 0)
        duration = meta.get("duration_seconds", 0)

        pdf.cell(0, 6, f"  Fields: {filled}/{total} filled   |   Skipped: {skipped}   |   Errors: {errors}   |   Duration: {duration:.0f}s", new_x="LMARGIN", new_y="NEXT")

        # Source breakdown
        src = data.get("source_summary", {})
        if src:
            src_text = "  Sources: " + ", ".join(f"{k}: {v}" for k, v in sorted(src.items(), key=lambda x: -x[1]))
            pdf.cell(0, 6, _latin(src_text), new_x="LMARGIN", new_y="NEXT")

        # Files
        files = data.get("files", {})
        if files.get("cv_path"):
            cv_name = _Path(files["cv_path"]).name if files["cv_path"] else "N/A"
            cv_tag = " (personalized)" if files.get("cv_personalized") else ""
            pdf.cell(0, 6, f"  CV: {_latin(cv_name)}{cv_tag}", new_x="LMARGIN", new_y="NEXT")
        if files.get("cover_letter_path"):
            cl_name = _Path(files["cover_letter_path"]).name if files["cover_letter_path"] else "N/A"
            cl_tag = " (personalized)" if files.get("cover_letter_personalized") else ""
            pdf.cell(0, 6, f"  Cover Letter: {_latin(cl_name)}{cl_tag}", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(5)

        # ── Fields table (landscape — full width) ──
        # Column widths: #(10) | Question(80) | Answer(120) | Source(27) | Status(20) = ~257
        col_w = [10, 80, page_w - 10 - 80 - 27 - 20, 27, 20]
        headers = ["#", "Question", "Answer", "Source", "Status"]

        status_marks = {
            "verified": "OK",
            "filled": "OK",
            "skipped": "SKIP",
            "error": "ERR",
            "needs_input": "?",
            "ready": "-",
        }

        # Table header
        pdf.set_fill_color(50, 50, 80)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 9)
        for header, w in zip(headers, col_w):
            pdf.cell(w, 8, f"  {header}", border=1, fill=True, new_x="END")
        pdf.ln()

        # Table rows with multi_cell wrapping
        pdf.set_text_color(0, 0, 0)

        for i, field in enumerate(data.get("fields", []), 1):
            question = _latin(field.get("question", ""))
            answer = _latin(field.get("answer", ""))
            source = field.get("source", "")
            status = status_marks.get(field.get("status", ""), "?")

            # Calculate row height based on longest content
            h_q = _row_height(pdf, question, col_w[1])
            h_a = _row_height(pdf, answer, col_w[2])
            row_h = max(6, h_q, h_a)

            # Check if row fits on page, add new page + header if not
            if pdf.get_y() + row_h > pdf.h - 15:
                pdf.add_page()
                pdf.set_fill_color(50, 50, 80)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("Helvetica", "B", 9)
                for header, w in zip(headers, col_w):
                    pdf.cell(w, 8, f"  {header}", border=1, fill=True, new_x="END")
                pdf.ln()
                pdf.set_text_color(0, 0, 0)

            # Alternate row colors
            if i % 2 == 0:
                pdf.set_fill_color(243, 243, 248)
            else:
                pdf.set_fill_color(255, 255, 255)

            y_start = pdf.get_y()
            x_start = pdf.l_margin

            # Draw row background + borders first
            for w in col_w:
                pdf.rect(x_start, y_start, w, row_h, style='DF')
                x_start += w

            # Fill cells with text
            x = pdf.l_margin
            pdf.set_font("Helvetica", "", 8)

            # #
            pdf.set_xy(x + 1, y_start + 1)
            pdf.cell(col_w[0] - 2, 4, str(i), align="C")
            x += col_w[0]

            # Question (multi_cell for wrapping)
            pdf.set_xy(x + 1, y_start + 1)
            pdf.multi_cell(col_w[1] - 2, 4, question, new_x="LEFT", new_y="TOP")
            x += col_w[1]

            # Answer (multi_cell for wrapping)
            pdf.set_xy(x + 1, y_start + 1)
            pdf.multi_cell(col_w[2] - 2, 4, answer, new_x="LEFT", new_y="TOP")
            x += col_w[2]

            # Source
            pdf.set_xy(x + 1, y_start + 1)
            pdf.cell(col_w[3] - 2, 4, source)
            x += col_w[3]

            # Status
            pdf.set_xy(x + 1, y_start + 1)
            # Color-code status
            if status == "OK":
                pdf.set_text_color(0, 128, 0)
            elif status == "ERR":
                pdf.set_text_color(200, 0, 0)
            elif status == "SKIP":
                pdf.set_text_color(180, 140, 0)
            pdf.cell(col_w[4] - 2, 4, status, align="C")
            pdf.set_text_color(0, 0, 0)

            # Move to next row
            pdf.set_y(y_start + row_h)

        # ── Job Description (if available) ──
        jd_text = data.get("job_description", "")
        if jd_text and len(jd_text) > 100:
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "Job Description", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

            pdf.set_font("Helvetica", "", 9)
            clean_jd = _latin(jd_text[:4000])
            pdf.multi_cell(0, 5, clean_jd)

        # Save PDF
        pdf_path = app_folder / "Application_Report.pdf"
        pdf.output(str(pdf_path))
        print(f"   📄 PDF report: {pdf_path.name}")

    MONTH_MAP = {
        'january': '01', 'february': '02', 'march': '03', 'april': '04',
        'may': '05', 'june': '06', 'july': '07', 'august': '08',
        'september': '09', 'october': '10', 'november': '11', 'december': '12',
        'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
        'jun': '06', 'jul': '07', 'aug': '08', 'sep': '09',
        'oct': '10', 'nov': '11', 'dec': '12',
    }

    def _month_to_number(self, value: str) -> str:
        """Convert month name to number (e.g., 'September' → '09', '06' → '06')."""
        v = value.strip().lower()
        # Already numeric
        if v.isdigit():
            return v.zfill(2)
        return self.MONTH_MAP.get(v, value)

    def _adapt_value_to_dom(self, value: str, field: 'FormField') -> str:
        """
        DOM-aware value adaptation.

        Reads field's placeholder, maxlength, pattern, and input_type
        to understand what format the field expects, then converts the value.

        Examples:
        - placeholder="MM", value="September" → "09"
        - placeholder="YYYY", value="September" → "September" (no change, not a year)
        - maxlength=2, value="September" → "09" (month name in short field → number)
        - maxlength=4, value="2025" → "2025" (year fits fine)
        - placeholder="MM/YYYY", value="September" → "09" (extract month part)
        - input_type="number", value="September" → "09"
        """
        if not value or not value.strip():
            return value

        v = value.strip()
        placeholder = field.placeholder.strip().upper() if field.placeholder else ""
        maxlen = field.maxlength
        input_type = field.input_type.lower() if field.input_type else ""

        # ── Rule 1: Placeholder indicates numeric month (MM) ──
        # Covers: placeholder="MM", "mm", "MM/YYYY" (month part)
        if placeholder in ("MM", "M"):
            converted = self._month_to_number(v)
            if converted != v:
                return converted
            # Value is already numeric or unknown — truncate to maxlen
            if maxlen and len(v) > maxlen:
                return v[:maxlen]
            return v

        # ── Rule 2: Placeholder indicates year (YYYY) ──
        if placeholder in ("YYYY", "YY"):
            # Year should be numeric — if it's already digits, fine
            if v.isdigit():
                if placeholder == "YY" and len(v) == 4:
                    return v[2:]  # 2025 → 25
                return v
            return v

        # ── Rule 3: maxlength constraint ──
        # If field only accepts 2 chars and value is a month name → convert
        if maxlen and maxlen <= 2 and not v.isdigit():
            # Likely a month field that expects numeric format
            converted = self._month_to_number(v)
            if converted != v:
                return converted

        # If field only accepts 4 chars and value is longer text → might be year
        if maxlen and maxlen == 4 and v.isdigit():
            return v[:4]

        # ── Rule 4: input type="number" or type="tel" ──
        # Field expects numeric input — try to convert month names
        if input_type in ("number", "tel"):
            converted = self._month_to_number(v)
            if converted != v:
                return converted

        # ── Rule 5: HTML5 pattern attribute ──
        # e.g., pattern="[0-9]{2}" means 2-digit number expected
        if field.pattern:
            pat = field.pattern
            if re.match(r'^\[0-9\]', pat):
                # Numeric pattern — convert month names if applicable
                converted = self._month_to_number(v)
                if converted != v:
                    return converted

        # ── Rule 6: General maxlength truncation ──
        if maxlen and len(v) > maxlen:
            # Don't silently truncate meaningful text — log a warning
            print(f"      ⚠️ DOM: value '{v[:20]}...' exceeds maxlength={maxlen}, truncating")
            return v[:maxlen]

        return v

    def _adapt_value_to_dom_element(self, value: str, el: 'ElementHandle') -> str:
        """
        DOM-aware adaptation for elements without FormField (e.g., repeatable sections).
        Reads placeholder/maxlength directly from the DOM element.
        """
        if not value or not value.strip():
            return value

        try:
            placeholder = el.get_attribute("placeholder") or ""
            maxlength_str = el.get_attribute("maxlength") or ""
            maxlength = int(maxlength_str) if maxlength_str.isdigit() else 0
            input_type = el.get_attribute("type") or "text"
            html_pattern = el.get_attribute("pattern") or ""
        except:
            return value

        # Create a minimal FormField-like object for the adapter
        class _DOMHint:
            pass
        hint = _DOMHint()
        hint.placeholder = placeholder
        hint.maxlength = maxlength
        hint.input_type = input_type
        hint.pattern = html_pattern

        return self._adapt_value_to_dom(value, hint)

    # ─────────────────────────────────────────────────────────────────────
    # PERSONALIZED DOCUMENTS (Cover Letter + CV)
    # ─────────────────────────────────────────────────────────────────────

    def _extract_job_info(self):
        """Extract job title, company, and description from the page.
        Saves to self.job_title, self.company_name, self.job_description.
        Ported from V3.5 extract_job_info() + improved.
        """
        self.job_title = ""
        self.company_name = ""
        self.job_description = ""

        try:
            title_text = self.page.title() if self.page else ""

            # Greenhouse format: "Job Application for {TITLE} at {COMPANY}"
            if " at " in title_text:
                parts = title_text.split(" at ", 1)
                self.job_title = parts[0].replace("Job Application for ", "").strip()
                self.company_name = parts[1].strip()
                self.company_name = self.company_name.replace(" Careers Page", "").replace(" Careers", "").strip()
            elif " - " in title_text:
                self.job_title = title_text.split(" - ")[0].replace("Job Application for ", "").strip()

            # Fallback: company from URL
            if not self.company_name:
                url = self.page.url if self.page else ""
                m = re.search(r'[?&]for=([^&]+)', url)
                if m:
                    self.company_name = m.group(1).replace('-', ' ').replace('_', ' ').title()

            # Fallback: company from DOM
            if not self.company_name:
                try:
                    context = getattr(self, '_active_frame', self.page)
                    company_el = context.query_selector('h1, .company-name, [class*="company"]')
                    if company_el:
                        self.company_name = company_el.inner_text().strip()[:50]
                except:
                    pass

            # Job description: search in main page first, then iframes
            jd_selectors = [
                '.job-description', '[class*="description"]',
                '[class*="job-details"]', '[class*="job-content"]',
                '[class*="posting-page"]', '.content', '#content',
                'article', '[role="main"]',
            ]
            # Search main page first (description pages have JD in main content)
            contexts_to_try = [self.page]
            # Then iframes
            for frame in self.page.frames:
                if frame != self.page.main_frame:
                    contexts_to_try.append(frame)

            for ctx in contexts_to_try:
                if self.job_description:
                    break
                for selector in jd_selectors:
                    try:
                        desc_el = ctx.query_selector(selector)
                        if desc_el:
                            text = desc_el.inner_text().strip()
                            if len(text) > 200:
                                self.job_description = text[:4000]
                                break
                    except:
                        continue

            # Fallback: If DOM JD is short (<500 chars), try ATS API (Greenhouse, Workday, etc.)
            # Use cached API JD if already fetched (avoid redundant API calls on re-extract)
            if len(self.job_description or '') < 500:
                cached_api_jd = getattr(self, '_api_jd_cache', None)
                if cached_api_jd:
                    self.job_description = cached_api_jd
                    print(f"   📄 JD from API (cached): {len(self.job_description)} chars")
                else:
                    try:
                        from parsers.jd_parser import fetch_jd_from_url
                        page_url = self.page.url if self.page else ""
                        if page_url:
                            api_jd = fetch_jd_from_url(page_url)
                            if api_jd and len(api_jd) > len(self.job_description or ''):
                                self.job_description = api_jd[:6000]
                                self._api_jd_cache = self.job_description
                                print(f"   📄 JD from API: {len(self.job_description)} chars")
                    except Exception as e:
                        print(f"   ⚠️ JD API fallback error: {e}")

            if self.job_title:
                print(f"   📋 Job: {self.job_title[:60]}")
            if self.company_name:
                print(f"   🏢 Company: {self.company_name[:40]}")
            if self.job_description:
                print(f"   📄 JD: {len(self.job_description)} chars")
            else:
                print(f"   ⚠️ No job description found on page")

        except Exception as e:
            print(f"   ⚠️ Job info extraction error: {e}")

    def _generate_personalized_cover_letter(self, job_title: str, company: str,
                                             jd: str) -> Optional[Path]:
        """Generate personalized cover letter: AI text → DOCX template.

        Flow:
        1. Check cache (Applications/{company}_{title}/)
        2. Generate CL body via Claude (personalized to JD)
        3. Insert into role-specific DOCX template (preserving header/signature formatting)
        4. Save to Applications folder
        """
        from pathlib import Path as _Path

        files_config = self.profile.data.get("files", {})
        base_path = _Path(files_config.get("base_path", ""))
        applications_dir = base_path / "Applications"

        # Build folder name
        safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')[:30]
        safe_title = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')[:40]
        from datetime import datetime as _dt
        folder_name = f"{safe_company}_{safe_title}_{_dt.now().strftime('%Y%m%d')}"
        app_folder = applications_dir / folder_name

        # Check cache
        cached_cl = app_folder / f"Cover_Letter_{safe_company}.docx"
        if cached_cl.exists():
            print(f"   📄 Using cached CL: {cached_cl.name}")
            return cached_cl

        # Get role-specific template
        page_title = self.page.title() if self.page else ""
        _, cl_template_path = self.profile.get_files_for_role(page_title)
        if not cl_template_path or not cl_template_path.exists():
            print(f"   ⚠️ No CL template found")
            return None

        # Ensure .docx (not .dotx)
        cl_template_path = self._ensure_compatible_extension(cl_template_path)

        # Generate personalized body via Claude
        cl_body = ""
        if jd and self.ai.available:
            try:
                profile_context = self._get_profile_context_for_ai()
                kb_context = ""
                if self.kb:
                    kb_context = self.kb.get_context_for_question("cover letter experience achievements")

                prompt = f"""Write the BODY of a cover letter (without header/salutation/signature — just the paragraphs).

Position: {job_title}
Company: {company}

Job Description (key parts):
{jd[:2500]}

Candidate Profile:
{profile_context}

Relevant Experience:
{kb_context}

Requirements:
- 3-4 paragraphs, 300-400 words total
- Professional but not generic — reference SPECIFIC requirements from JD
- Highlight matching experience with concrete results (numbers, metrics)
- Show genuine enthusiasm for this specific role and company
- Do NOT include salutation (Dear...) or signature (Sincerely...) — just body paragraphs
- Do NOT start with "I am excited to apply" — be more creative

Write ONLY the paragraphs, nothing else."""

                response = self.ai.vision_ai.client.messages.create(
                    model=self.ai.vision_ai.config.model,
                    max_tokens=800,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.4,
                )
                cl_body = response.content[0].text.strip()
                print(f"   ✍️ AI generated CL body: {len(cl_body)} chars")
            except Exception as e:
                print(f"   ⚠️ AI CL generation failed: {e}")

        # Load DOCX template and modify
        try:
            from docx import Document
            doc = Document(str(cl_template_path))

            if cl_body:
                # Strategy: replace body paragraphs (between "Dear..." and "Sincerely")
                # Find boundaries
                dear_idx = None
                sincerely_idx = None
                for i, para in enumerate(doc.paragraphs):
                    text_lower = para.text.strip().lower()
                    if text_lower.startswith("dear"):
                        dear_idx = i
                    if text_lower.startswith("sincerely") or text_lower.startswith("best regards"):
                        sincerely_idx = i

                if dear_idx is not None and sincerely_idx is not None:
                    # Clear body paragraphs (between Dear and Sincerely)
                    body_start = dear_idx + 1
                    body_end = sincerely_idx

                    # Remove old body paragraphs (in reverse to preserve indices)
                    for i in range(body_end - 1, body_start - 1, -1):
                        p = doc.paragraphs[i]._element
                        p.getparent().remove(p)

                    # Insert new body paragraphs after "Dear..."
                    from docx.oxml.ns import qn
                    import copy
                    # Get reference to Dear paragraph element
                    dear_element = doc.paragraphs[dear_idx]._element

                    # Split AI body into paragraphs
                    new_paragraphs = [p.strip() for p in cl_body.split('\n\n') if p.strip()]

                    for para_text in reversed(new_paragraphs):
                        # Create new paragraph element
                        new_p = copy.deepcopy(dear_element)
                        # Clear and set text
                        for child in list(new_p):
                            new_p.remove(child)
                        run = doc.paragraphs[dear_idx].runs[0] if doc.paragraphs[dear_idx].runs else None
                        new_run = copy.deepcopy(run._element) if run else None
                        if new_run is not None:
                            new_run.text = para_text
                            new_p.append(new_run)
                        else:
                            from docx.oxml import OxmlElement
                            r = OxmlElement('w:r')
                            t = OxmlElement('w:t')
                            t.text = para_text
                            r.append(t)
                            new_p.append(r)

                        dear_element.addnext(new_p)

                    print(f"   ✅ Inserted {len(new_paragraphs)} AI paragraphs into template")
                else:
                    print(f"   ⚠️ Could not find Dear/Sincerely boundaries, using marker replacement")
                    cl_body = ""  # Fall through to marker replacement

            # Always do marker replacement (company name, position title)
            for para in doc.paragraphs:
                for run in para.runs:
                    if "[COMPANY NAME]" in run.text:
                        run.text = run.text.replace("[COMPANY NAME]", company)
                    if "[POSITION TITLE]" in run.text:
                        run.text = run.text.replace("[POSITION TITLE]", job_title)
                    if "[COMPANY MISSION]" in run.text:
                        run.text = run.text.replace("[COMPANY MISSION]",
                                                     f"contribute to {company}'s continued success")

            # Save
            app_folder.mkdir(parents=True, exist_ok=True)
            output_path = app_folder / f"Cover_Letter_{safe_company}.docx"
            doc.save(str(output_path))
            print(f"   📄 Saved personalized CL: {output_path.name}")
            return output_path

        except Exception as e:
            print(f"   ❌ DOCX CL generation error: {e}")
            return None

    def _create_tailored_cv(self, job_title: str, company: str, jd: str) -> Optional[Path]:
        """Create tailored CV with keywords from JD injected into TECHNICAL SKILLS section.

        Flow:
        1. Check cache
        2. Extract keywords from JD (via Claude or regex)
        3. Load role-specific CV DOCX
        4. Add keywords to TECHNICAL SKILLS / COMPETENCIES section
        5. Save to Applications folder
        """
        from pathlib import Path as _Path

        files_config = self.profile.data.get("files", {})
        base_path = _Path(files_config.get("base_path", ""))
        applications_dir = base_path / "Applications"

        safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')[:30]
        safe_title = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')[:40]
        from datetime import datetime as _dt
        folder_name = f"{safe_company}_{safe_title}_{_dt.now().strftime('%Y%m%d')}"
        app_folder = applications_dir / folder_name

        # Check cache
        cached_cv = app_folder / f"CV_Anton_Kondakov_{safe_company}.docx"
        if cached_cv.exists():
            print(f"   📄 Using cached CV: {cached_cv.name}")
            return cached_cv

        # Find role-specific CV DOCX (not PDF)
        page_title = self.page.title() if self.page else ""
        by_role = files_config.get("by_role", {})
        default_role = files_config.get("default_role", "TPM")

        # Detect role
        job_lower = page_title.lower()
        detected_role = default_role
        role_patterns = [
            ("TPM", ["technical program manager", "tpm"]),
            ("Product Manager", ["product manager"]),
            ("Product Owner", ["product owner"]),
            ("Project Manager", ["project manager"]),
            ("Scrum Master", ["scrum master", "agile coach"]),
            ("Delivery Lead", ["delivery lead", "delivery manager"]),
        ]
        for role_name, patterns in role_patterns:
            if any(p in job_lower for p in patterns):
                detected_role = role_name
                break

        role_files = by_role.get(detected_role, by_role.get(default_role, {}))
        cv_filename = role_files.get("cv", "")

        if not cv_filename:
            return None

        # Need DOCX version (not PDF)
        cv_docx_name = cv_filename.replace('.pdf', '.docx')
        cv_docx_path = base_path / cv_docx_name

        if not cv_docx_path.exists():
            print(f"   ⚠️ CV DOCX not found: {cv_docx_name}")
            return None

        # Extract keywords from JD
        keywords = []
        if jd and self.ai.available:
            try:
                prompt = f"""Extract 5-10 key technical skills/tools/technologies from this job description
that are NOT already in the candidate's CV.

The candidate already has: AWS, GCP, Azure, Jira, Confluence, ServiceNow, Python, Java, Angular,
Terraform, Jenkins, GitHub Actions, SQL, Tableau, Power BI, SAFe, Scrum, Kanban.

Job Description:
{jd[:2500]}

Return ONLY a comma-separated list of missing keywords (max 10). No explanations.
Example: Kubernetes, Docker, Snowflake, dbt, Airflow"""

                response = self.ai.vision_ai.client.messages.create(
                    model=self.ai.vision_ai.config.model,
                    max_tokens=200,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1,
                )
                kw_text = response.content[0].text.strip()
                keywords = [k.strip() for k in kw_text.split(',') if k.strip() and len(k.strip()) < 30]
                keywords = keywords[:8]  # Cap at 8
                if keywords:
                    print(f"   🔑 Keywords from JD: {', '.join(keywords)}")
            except Exception as e:
                print(f"   ⚠️ Keyword extraction failed: {e}")

        if not keywords:
            print(f"   ⚠️ No keywords to add — using base CV")
            return None

        # Load CV and inject keywords
        try:
            from docx import Document
            doc = Document(str(cv_docx_path))

            # Find TECHNICAL SKILLS section (preferred) or CORE COMPETENCIES (fallback)
            injected = False
            skills_section_idx = None
            # First pass: find TECHNICAL SKILLS (preferred)
            for i, para in enumerate(doc.paragraphs):
                if "TECHNICAL SKILLS" in para.text.upper():
                    skills_section_idx = i
                    break
            # Second pass: fallback to CORE COMPETENCIES
            if skills_section_idx is None:
                for i, para in enumerate(doc.paragraphs):
                    if "CORE COMPETENCIES" in para.text.upper():
                        skills_section_idx = i
                        break

            if skills_section_idx is not None:
                # Find last paragraph in this section (before next section header or end)
                last_content_idx = skills_section_idx
                for j in range(skills_section_idx + 1, min(skills_section_idx + 20, len(doc.paragraphs))):
                    next_text = doc.paragraphs[j].text.strip().upper()
                    if not next_text:
                        continue
                    # Stop at next section header (all-caps line that's a known section)
                    if any(kw in next_text for kw in [
                        "EXPERIENCE", "EDUCATION", "CERTIFICATION",
                        "ACHIEVEMENT", "CORE COMPETENCIES", "TECHNICAL SKILLS"
                    ]) and j != skills_section_idx:
                        break
                    last_content_idx = j

                # Append keywords to the last content paragraph in the section
                target = doc.paragraphs[last_content_idx]
                current = target.text.rstrip('.')
                kw_str = ', '.join(keywords)
                if target.runs:
                    target.runs[-1].text = f"{current}, {kw_str}"
                else:
                    target.clear()
                    target.add_run(f"{current}, {kw_str}")
                injected = True
                print(f"   📝 Keywords injected at paragraph [{last_content_idx}]: ...{target.text[-60:]}")

            if not injected:
                print(f"   ⚠️ Could not find TECHNICAL SKILLS section for keyword injection")
                return None

            # Save
            app_folder.mkdir(parents=True, exist_ok=True)
            output_path = app_folder / f"CV_Anton_Kondakov_{safe_company}.docx"
            doc.save(str(output_path))
            print(f"   📄 Saved tailored CV: {output_path.name} (+{len(keywords)} keywords)")
            return output_path

        except Exception as e:
            print(f"   ❌ CV tailoring error: {e}")
            return None

    def _extract_company_from_url(self, url: str) -> str:
        """Extract company name from Greenhouse/Lever URL."""
        url_lower = url.lower()
        # Greenhouse: ?for=companyname or boards.greenhouse.io/companyname
        import re as _re
        m = _re.search(r'for=([a-z0-9_-]+)', url_lower)
        if m:
            return m.group(1)
        m = _re.search(r'boards\.greenhouse\.io/([a-z0-9_-]+)', url_lower)
        if m:
            return m.group(1)
        m = _re.search(r'jobs\.lever\.co/([a-z0-9_-]+)', url_lower)
        if m:
            return m.group(1)
        return ""

    def _save_verified_ai_answers(self):
        """Feedback loop: save verified AI answers to learned DB for future use."""
        saved_count = 0
        for field in self.fields:
            if (field.status == FillStatus.VERIFIED and
                    field.answer_source == AnswerSource.AI and
                    field.answer and
                    len(field.label) > 5):
                # Save AI-generated answer that was verified as correct
                is_dropdown = field.field_type in (FieldType.SELECT, FieldType.AUTOCOMPLETE)
                self.learned_db.save_answer(field.label, field.answer, is_dropdown)
                saved_count += 1
        if saved_count:
            print(f"   🔄 Feedback: saved {saved_count} verified AI answers to learned DB")

    def _detect_ats(self, url: str) -> str:
        """Detect ATS type from URL."""
        url_lower = url.lower()
        if "greenhouse" in url_lower or "grnhse" in url_lower:
            return "Greenhouse"
        if "lever" in url_lower or "jobs.lever.co" in url_lower:
            return "Lever"
        if "workday" in url_lower or "myworkdayjobs" in url_lower:
            return "Workday"
        if "ashby" in url_lower:
            return "Ashby"
        if "smartrecruiters" in url_lower:
            return "SmartRecruiters"
        if "icims" in url_lower:
            return "iCIMS"
        if "jobvite" in url_lower:
            return "Jobvite"
        if "taleo" in url_lower:
            return "Taleo"
        if "breezy" in url_lower:
            return "BreezyHR"
        return "Unknown"

    # ATS-specific selectors for common fields
    ATS_SELECTORS = {
        "Greenhouse": {
            "first_name": "#first_name",
            "last_name": "#last_name",
            "email": "#email",
            "phone": "#phone",
            "location": "#candidate-location, #location",
            "resume": "input[type='file'][name*='resume'], input[type='file']:first-of-type",
            "cover_letter": "input[type='file'][name*='cover'], input[type='file']:nth-of-type(2)",
            "linkedin": "#job_application_answers_attributes_0_text_value, input[name*='linkedin']",
            "country": "#country",
            "work_company": "#company-name-{N}",
            "work_title": "#title-{N}",
            "school": "#school--{N}",
            "degree": "#degree--{N}",
            "apply_button": "button:has-text('Apply'), a.postings-btn",
        },
        "Lever": {
            "first_name": "input[name='name']",
            "email": "input[name='email']",
            "phone": "input[name='phone']",
            "location": "input[name='location']",
            "resume": "input[type='file']",
            "linkedin": "input[name='urls[LinkedIn]']",
            "apply_button": "button[type='submit'], button:has-text('Submit')",
        },
        "Workday": {
            "first_name": "input[data-automation-id='legalNameSection_firstName']",
            "last_name": "input[data-automation-id='legalNameSection_lastName']",
            "email": "input[data-automation-id='email']",
            "phone": "input[data-automation-id='phone']",
            "resume": "input[data-automation-id='file-upload-input-ref']",
            "country": "button[data-automation-id='countryDropdown']",
            "apply_button": "button[data-automation-id='jobPostingApplyButton']",
        },
        "Ashby": {
            "first_name": "input[name='_systemfield_first_name']",
            "last_name": "input[name='_systemfield_last_name']",
            "email": "input[name='_systemfield_email']",
            "phone": "input[name='_systemfield_phone']",
            "resume": "input[type='file']",
            "linkedin": "input[name='_systemfield_linkedin']",
            "apply_button": "button[type='submit']",
        },
        "iCIMS": {
            "first_name": "input[name='firstName']",
            "last_name": "input[name='lastName']",
            "email": "input[name='email']",
            "phone": "input[name='phoneNumber']",
            "resume": "input[type='file']",
            "apply_button": "a.iCIMS_ApplyButton, button:has-text('Apply')",
        },
    }

    def get_ats_selector(self, ats_type: str, field_name: str, index: int = 0) -> str:
        """Get ATS-specific selector for a field."""
        selectors = self.ATS_SELECTORS.get(ats_type, {})
        selector = selectors.get(field_name, "")
        if "{N}" in selector:
            selector = selector.replace("{N}", str(index))
        return selector


# ═══════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════

def main():
    import sys
    
    url = sys.argv[1] if len(sys.argv) > 1 else \
        "https://job-boards.greenhouse.io/embed/job_app?token=7404427&for=coinbase&gh_jid=7404427"
    
    mode_arg = sys.argv[2] if len(sys.argv) > 2 else "interactive"
    keep_open = "--keep-open" in sys.argv or "-k" in sys.argv

    mode_map = {
        "preflight": FillMode.PRE_FLIGHT,
        "interactive": FillMode.INTERACTIVE,
        "auto": FillMode.AUTONOMOUS,
    }
    mode = mode_map.get(mode_arg, FillMode.INTERACTIVE)

    print("\n" + "="*70)
    print("🚀 SMART FORM FILLER V5.0")
    print("="*70)
    print(f"URL: {url[:60]}...")
    print(f"Mode: {mode.value}")
    if keep_open:
        print("Keep open: Yes (browser stays open after fill)")
    print("="*70)

    filler = FormFillerV5(browser_mode=BrowserMode.CDP)

    if mode == FillMode.PRE_FLIGHT:
        report = filler.analyze(url)
    else:
        report = filler.fill(url, mode=mode, keep_open=keep_open)
    
    print("\n" + report.summary())


if __name__ == "__main__":
    main()
