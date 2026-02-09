# I have access to - main.py

from __future__ import annotations

import os

# Load .env file
from dotenv import load_dotenv
load_dotenv()

import asyncio
from datetime import datetime, timezone, timedelta
import json
from collections import Counter
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Optional

# ========== UNIVERSAL PATHS (work on any machine via iCloud) ==========
def get_icloud_path() -> Path:
    """Get iCloud Drive path for current user."""
    return Path.home() / "Library" / "Mobile Documents" / "com~apple~CloudDocs"

def get_ai_projects_path() -> Path:
    """Get AI_projects folder path."""
    return get_icloud_path() / "Dev" / "AI_projects"

def get_gold_cv_path() -> Path:
    """Get Gold CV folder path."""
    return get_ai_projects_path() / "Gold CV"

def get_applications_path() -> Path:
    """Get Applications folder for cover letters."""
    return get_gold_cv_path() / "Applications"

# Pre-compute paths
ICLOUD_PATH = get_icloud_path()
AI_PROJECTS_PATH = get_ai_projects_path()
GOLD_CV_PATH = get_gold_cv_path()
APPLICATIONS_PATH = get_applications_path()

# Environment: PROD or DEV (auto-detect from folder name)
_current_dir = Path(__file__).parent.name
ENV = os.getenv("JOB_TRACKER_ENV", "DEV" if "dev" in _current_dir.lower() else "PROD")

from fastapi import FastAPI, Query
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from parsers.greenhouse import fetch_greenhouse
from parsers.lever import fetch_lever
from parsers.smartrecruiters import fetch_smartrecruiters
from parsers.ashby import fetch_ashby_jobs
from parsers.workday_v2 import fetch_workday_v2
from parsers.atlassian import fetch_atlassian
from parsers.phenom import fetch_phenom_jobs
from ats_detector import try_repair_company, verify_ats_url
from company_storage import load_profile
from utils.normalize import normalize_location, STATE_MAP
from utils.cache_manager import load_cache, save_cache, clear_cache, get_cache_info, load_stats
from utils.job_utils import generate_job_id, classify_role, find_similar_jobs

# ATS parser mapping - these ATS support automatic job fetching
from parsers.icims import fetch_icims
from parsers.jibe import fetch_jibe

ATS_PARSERS = {
    "greenhouse": lambda url: fetch_greenhouse("", url),
    "lever": lambda url: fetch_lever("", url),
    "smartrecruiters": lambda url: fetch_smartrecruiters("", url),
    "ashby": fetch_ashby_jobs,
    "workday": lambda url: fetch_workday_v2("", url),
    "atlassian": lambda url: fetch_atlassian("", url),
    "phenom": fetch_phenom_jobs,
    "icims": lambda url: fetch_icims("", url),
    "jibe": lambda url: fetch_jibe("", url),
}

# List of supported ATS for UI display
SUPPORTED_ATS = list(ATS_PARSERS.keys())

from storage.job_storage import (
    get_all_jobs, get_jobs_by_status, get_jobs_by_statuses,
    get_active_jobs, get_archive_jobs, get_all_job_ids,
    add_job, add_jobs_bulk, update_status as job_update_status,
    update_last_seen, update_last_seen_bulk, mark_missing_jobs,
    get_stats as get_job_stats, get_job_by_id, job_exists,
    STATUS_NEW, STATUS_APPLIED, STATUS_INTERVIEW, STATUS_OFFER,
    STATUS_REJECTED, STATUS_WITHDRAWN, STATUS_CLOSED, STATUS_EXCLUDED,
    ACTIVE_STATUSES, ARCHIVE_STATUSES,
)


# -----------------------------
# Runtime/local state files
# -----------------------------
JOB_STATUS_FILE = Path("job_status.json")

VALID_APPLICATION_STATUSES = ["New", "Applied", "Interview", "Offer", "Rejected", "Withdrawn", "Closed"]

# My Roles families for filtering
MY_ROLE_FAMILIES = {"product", "tpm_program", "project"}


def sync_cache_to_pipeline(jobs: list) -> dict:
    """
    Sync parsed jobs from cache to jobs.json (pipeline storage).
    Only adds jobs that match My Roles and are new.
    Returns stats: {added: N, updated: N, total: N}
    """
    # Filter to My Roles only
    my_roles_jobs = [
        j for j in jobs 
        if j.get("role_family") in MY_ROLE_FAMILIES
    ]
    
    # Get existing job IDs
    existing_ids = get_all_job_ids()
    
    # Separate new vs existing
    new_jobs = [j for j in my_roles_jobs if j.get("id") not in existing_ids]
    existing_jobs = [j for j in my_roles_jobs if j.get("id") in existing_ids]
    
    # Add new jobs in bulk
    added = add_jobs_bulk(new_jobs, STATUS_NEW) if new_jobs else 0
    
    # Update last_seen for existing jobs
    existing_job_ids = {j.get("id") for j in existing_jobs}
    updated = update_last_seen_bulk(existing_job_ids) if existing_job_ids else 0
    
    return {
        "added": added,
        "updated": updated,
        "total_my_roles": len(my_roles_jobs),
        "total_parsed": len(jobs)
    }


def _safe_read_json(path: Path) -> dict:
    try:
        if not path.exists():
            return {}
        txt = path.read_text(encoding="utf-8").strip()
        if not txt:
            return {}
        return json.loads(txt)
    except Exception:
        return {}


def _safe_write_json(path: Path, data: dict) -> None:
    try:
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:  # noqa: BLE001
        print(f"Failed to write {path}: {e}")


def _load_job_status_map(profile: str) -> dict[str, str]:
    """
    Returns dict: job_key -> status
    Stored format:
      {
        "profiles": {
          "all": {"<job_key>":"Applied", ...},
          "fintech": {...}
        }
      }
    """
    root = _safe_read_json(JOB_STATUS_FILE)
    profiles = root.get("profiles") if isinstance(root, dict) else None
    if not isinstance(profiles, dict):
        return {}
    mp = profiles.get(profile)
    if isinstance(mp, dict):
        # ensure values are strings
        out: dict[str, str] = {}
        for k, v in mp.items():
            if isinstance(k, str) and isinstance(v, str):
                out[k] = v
        return out
    return {}


def _set_job_status(profile: str, job_key: str, status: str) -> None:
    status_norm = status if status in VALID_APPLICATION_STATUSES else "New"
    root = _safe_read_json(JOB_STATUS_FILE)
    if not isinstance(root, dict):
        root = {}
    profiles = root.get("profiles")
    if not isinstance(profiles, dict):
        profiles = {}
        root["profiles"] = profiles
    mp = profiles.get(profile)
    if not isinstance(mp, dict):
        mp = {}
        profiles[profile] = mp
    mp[job_key] = status_norm
    _safe_write_json(JOB_STATUS_FILE, root)


def compute_job_key(job: dict) -> str:
    """
    Stable key for status tracking.
    Preference:
      1) job_url
      2) url
      3) (company|title|location)
    """
    url = (job.get("job_url") or job.get("url") or "").strip()
    if url:
        return url
    company = (job.get("company") or "").strip()
    title = (job.get("title") or "").strip()
    location = (job.get("location") or "").strip()
    return f"{company}|{title}|{location}".strip("|")


# Кэш статуса по компаниям: "profile:company" -> {ok, error, checked_at, ats, url}
COMPANY_STATUS_FILE = Path("data/company_status.json")

def load_company_status() -> dict:
    """Load company fetch status from file"""
    if COMPANY_STATUS_FILE.exists():
        try:
            with open(COMPANY_STATUS_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_company_status(status: dict):
    """Save company fetch status to file"""
    try:
        with open(COMPANY_STATUS_FILE, "w") as f:
            json.dump(status, f, indent=2)
    except Exception as e:
        print(f"Failed to save company status: {e}")

# Load on startup
company_fetch_status: dict[str, dict] = load_company_status()

# Geo scoring/bucketing configuration
TARGET_STATE = "NC"
NEIGHBOR_STATES = {"VA", "SC", "GA", "TN"}
LOCAL_CITIES = {"raleigh", "durham", "cary", "chapel hill", "morrisville"}


def compute_geo_bucket_and_score(loc_norm: dict | None) -> tuple[str, int]:
    if not loc_norm:
        return "unknown", 0

    city = (loc_norm.get("city") or "").lower()
    state = (loc_norm.get("state") or "").upper() if loc_norm.get("state") else None
    remote = bool(loc_norm.get("remote"))
    remote_scope = (loc_norm.get("remote_scope") or "").lower()

    if city in LOCAL_CITIES and state == TARGET_STATE:
        return "local", 100
    if state == TARGET_STATE:
        return "nc", 80
    if state in NEIGHBOR_STATES:
        return "neighbor", 60
    if remote and remote_scope in ["usa", "us"]:
        return "remote_usa", 50
    if state or city or remote:
        return "other", 0
    return "unknown", 0


app = FastAPI(
    title="Job Tracker",
    description="Simple job aggregator for Product / PM roles from ATS",
    version="0.3.0",
)

# Gzip compression for large responses (like 8MB jobs cache)
app.add_middleware(GZipMiddleware, minimum_size=1000)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")

# ========== BACKGROUND REFRESH DAEMON ==========

# Lock file path (in data/ folder, synced via iCloud)
DAEMON_LOCK_FILE = Path("data/daemon.lock")

def get_machine_id() -> str:
    """Get unique machine identifier (hostname + username)"""
    import socket
    import getpass
    return f"{getpass.getuser()}@{socket.gethostname()}"

def check_daemon_lock() -> dict | None:
    """Check if daemon is locked by another machine. Returns lock info or None."""
    if not DAEMON_LOCK_FILE.exists():
        return None
    try:
        lock_data = json.loads(DAEMON_LOCK_FILE.read_text())
        # Check if lock is stale (older than 10 minutes)
        lock_time = datetime.fromisoformat(lock_data.get("timestamp", "1970-01-01"))
        if datetime.now(timezone.utc) - lock_time > timedelta(minutes=10):
            # Stale lock - can be overwritten
            return None
        return lock_data
    except:
        return None

def acquire_daemon_lock() -> tuple[bool, str]:
    """Try to acquire daemon lock. Returns (success, message)."""
    machine_id = get_machine_id()
    existing_lock = check_daemon_lock()
    
    if existing_lock:
        lock_machine = existing_lock.get("machine", "unknown")
        if lock_machine != machine_id:
            return False, f"Daemon already running on {lock_machine}"
    
    # Create/update lock
    lock_data = {
        "machine": machine_id,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "pid": os.getpid()
    }
    DAEMON_LOCK_FILE.write_text(json.dumps(lock_data, indent=2))
    return True, "Lock acquired"

def release_daemon_lock():
    """Release daemon lock if we own it."""
    machine_id = get_machine_id()
    existing_lock = check_daemon_lock()
    
    if existing_lock and existing_lock.get("machine") == machine_id:
        try:
            DAEMON_LOCK_FILE.unlink()
        except:
            pass

def update_daemon_lock():
    """Update lock timestamp (heartbeat)."""
    machine_id = get_machine_id()
    existing_lock = check_daemon_lock()
    
    if existing_lock and existing_lock.get("machine") == machine_id:
        lock_data = {
            "machine": machine_id,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "pid": os.getpid()
        }
        DAEMON_LOCK_FILE.write_text(json.dumps(lock_data, indent=2))

# Daemon status (global)
# NOTE: Daemon disabled by default to avoid conflicts when running on multiple machines via iCloud
DAEMON_STATUS = {
    "enabled": True,  # Auto-start daemon on server startup
    "running": False,
    "locked_by": None,  # Which machine has the lock
    "current_company": None,
    "last_company": None,
    "last_updated": None,
    "next_company": None,
    "companies_in_cycle": 0,
    "current_index": 0,
    "batch_size": 3,
    "pause_seconds": 45,
    "cycle_count": 0,
    "jobs_added_this_cycle": 0,
    "jobs_added_total": 0,
    "last_cycle_jobs_added": 0,
    "companies_refreshed_this_cycle": 0,
    "refresh_log": []  # List of {company, status, jobs_count, time, error}
}

async def refresh_company_async(company: dict) -> dict:
    """Parse a single company (runs in thread pool to not block async)"""
    import concurrent.futures
    
    loop = asyncio.get_event_loop()
    with concurrent.futures.ThreadPoolExecutor() as pool:
        result = await loop.run_in_executor(pool, refresh_company_sync, company)
    return result

def refresh_company_sync(company: dict) -> dict:
    """Synchronous company refresh (called from thread pool)"""
    company_id = company.get("id", company.get("name", "unknown"))
    ats = company.get("ats", "unknown")
    board_url = company.get("board_url", "")
    
    result = {
        "company": company_id,
        "ok": False,
        "jobs": 0,
        "jobs_added": 0,
        "error": None
    }
    
    try:
        fetcher = ATS_PARSERS.get(ats)
        if not fetcher:
            result["error"] = f"Unknown ATS: {ats}"
            return result
        
        jobs = fetcher(board_url)
        result["jobs"] = len(jobs) if jobs else 0
        result["ok"] = True
        print(f"[refresh_company_sync] {company_id}: fetched {len(jobs) if jobs else 0} jobs")

        # Update company status in storage
        from company_storage import update_company_status
        update_company_status(company_id, ok=True, jobs_count=len(jobs) if jobs else 0)

        # Update cache with new jobs and track added count
        print(f"[refresh_company_sync] {company_id}: calling update_cache_for_company...")
        added = update_cache_for_company(company_id, jobs or [])
        print(f"[refresh_company_sync] {company_id}: update_cache returned {added}")
        result["jobs_added"] = added or 0
        
    except Exception as e:
        result["error"] = str(e)
        from company_storage import update_company_status
        update_company_status(company_id, ok=False, error=str(e))
    
    return result

def update_cache_for_company(company_id: str, new_jobs: list) -> int:
    """Update jobs_all.json cache for a specific company. Returns count of new jobs added to pipeline."""
    from utils.cache_manager import get_cache_path
    from utils.job_utils import classify_role, generate_job_id
    from utils.normalize import normalize_location

    cache_path = get_cache_path("all")
    if not cache_path.exists():
        print(f"[Daemon] Cache file not found: {cache_path}")
        return 0

    import time
    max_retries = 3
    for attempt in range(max_retries):
        try:
            print(f"[update_cache] {company_id}: reading cache (attempt {attempt + 1})...")
            with open(cache_path, "r") as f:
                cache_data = json.load(f)
            break  # Success, exit retry loop
        except json.JSONDecodeError as e:
            if attempt < max_retries - 1:
                print(f"[update_cache] {company_id}: JSON error, retrying in 1s... ({e})")
                time.sleep(1)
                continue
            else:
                print(f"[update_cache] {company_id}: JSON error after {max_retries} attempts: {e}")
                return 0

    try:

        all_jobs = cache_data.get("jobs", [])
        print(f"[update_cache] {company_id}: cache has {len(all_jobs)} jobs")

        # Remove old jobs from this company
        all_jobs = [j for j in all_jobs if j.get("company_id") != company_id]
        print(f"[update_cache] {company_id}: after removing old jobs: {len(all_jobs)}")

        # Resolve company name from companies.json
        company_name = ""
        try:
            from company_storage import load_companies
            for c in load_companies():
                if c.get("id") == company_id:
                    company_name = c.get("name", "")
                    break
        except Exception:
            pass

        # Normalize and classify new jobs
        print(f"[update_cache] {company_id}: normalizing {len(new_jobs)} new jobs...")
        for job in new_jobs:
            job["company_id"] = company_id
            if not job.get("company") and company_name:
                job["company"] = company_name
            if "location_norm" not in job:
                job["location_norm"] = normalize_location(job.get("location", ""))
            if "role_category" not in job:
                role_info = classify_role(job.get("title", ""), "")
                job["role_category"] = role_info.get("role_category", "unknown")
                job["role_family"] = role_info.get("role_family", "other")
                job["role_id"] = role_info.get("role_id")
            # Generate job ID if missing (required for pipeline)
            if not job.get("id"):
                job["id"] = generate_job_id(job)

        # Add new jobs
        all_jobs.extend(new_jobs)
        print(f"[Daemon] Cache update: {company_id} - adding {len(new_jobs)} jobs, total will be {len(all_jobs)}")

        # Save back using atomic write (write to temp file, then rename)
        cache_data["jobs"] = all_jobs
        cache_data["jobs_count"] = len(all_jobs)
        cache_data["last_updated"] = datetime.now(timezone.utc).isoformat()

        temp_path = cache_path.with_suffix('.tmp')
        with open(temp_path, "w") as f:
            json.dump(cache_data, f)
        temp_path.rename(cache_path)  # Atomic on most filesystems
        print(f"[Daemon] Cache saved successfully for {company_id}")

        # Also update pipeline (jobs.json) with relevant jobs
        added = update_pipeline_for_company(company_id, new_jobs)
        return added

    except Exception as e:
        import traceback
        print(f"[Daemon] Cache update error for {company_id}: {e}")
        traceback.print_exc()
        return 0

def update_pipeline_for_company(company_id: str, new_jobs: list) -> int:
    """Update pipeline jobs.json with new relevant jobs from company. Returns count added."""
    from storage.job_storage import get_all_jobs, add_job
    
    # Get existing pipeline jobs
    existing = get_all_jobs()
    existing_urls = {j.get("job_url") or j.get("url") for j in existing}
    
    # Add new relevant jobs (primary/adjacent only)
    added = 0
    for job in new_jobs:
        if job.get("role_category") in ["primary", "adjacent"]:
            job_url = job.get("job_url") or job.get("url")
            if job_url and job_url not in existing_urls:
                if add_job(job):
                    added += 1
    
    return added
    
    if added > 0:
        print(f"[Daemon] Added {added} new jobs to pipeline from {company_id}")

# Self-healing: consecutive error tracking per company
_COMPANY_ERRORS: dict[str, int] = {}  # company_id → consecutive error count
_MAX_CONSECUTIVE_ERRORS = 3  # disable company after this many consecutive failures


def _track_company_error(company: dict, error: str):
    """
    Track consecutive errors for a company.
    After _MAX_CONSECUTIVE_ERRORS consecutive failures → try repair first, then auto-disable.
    """
    company_id = company.get("id", "")
    if not company_id:
        return

    _COMPANY_ERRORS[company_id] = _COMPANY_ERRORS.get(company_id, 0) + 1
    count = _COMPANY_ERRORS[company_id]

    if count >= _MAX_CONSECUTIVE_ERRORS:
        print(f"[Daemon] ⚠️ {company_id}: {count} consecutive errors → attempting repair...")

        # Try repair before disabling
        repair_result = try_repair_company({
            "name": company.get("name", company_id),
            "ats": company.get("ats", ""),
            "board_url": company.get("board_url", ""),
        })

        if repair_result and repair_result.get("verified"):
            new_ats = repair_result["ats"]
            new_url = repair_result["board_url"]
            print(f"[Daemon] 🔧 Repair found new URL for {company_id}: {new_ats} → {new_url}")

            # Update companies.json with repaired data
            try:
                companies_path = Path("data/companies.json")
                with open(companies_path, "r", encoding="utf-8") as f:
                    companies = json.load(f)

                for c in companies:
                    if c.get("id") == company_id:
                        c["ats"] = new_ats
                        c["board_url"] = new_url
                        c["repaired_at"] = datetime.now(timezone.utc).isoformat()
                        break

                with open(companies_path, "w", encoding="utf-8") as f:
                    json.dump(companies, f, indent=2, ensure_ascii=False)

                print(f"[Daemon] ✅ Repaired {company_id} → {new_ats}")
                _COMPANY_ERRORS.pop(company_id, None)  # reset errors after repair
                return  # Don't disable — repaired successfully
            except Exception as e:
                print(f"[Daemon] ❌ Failed to save repair for {company_id}: {e}")

        # Repair failed → auto-disable
        print(f"[Daemon] ❌ Repair failed for {company_id} → auto-disabling")
        _auto_disable_company(company_id, error, count)
        _COMPANY_ERRORS.pop(company_id, None)  # cleanup


def _reset_company_errors(company_id: str):
    """Reset consecutive error counter on successful parse."""
    if company_id in _COMPANY_ERRORS:
        del _COMPANY_ERRORS[company_id]


def _auto_disable_company(company_id: str, last_error: str, error_count: int):
    """
    Auto-disable a company after consecutive failures.
    Sets enabled=false, status="auto_disabled" in companies.json.
    """
    companies_path = Path("data/companies.json")
    if not companies_path.exists():
        return

    try:
        with open(companies_path, "r", encoding="utf-8") as f:
            companies = json.load(f)

        for c in companies:
            if c.get("id") == company_id:
                c["enabled"] = False
                c["status"] = "auto_disabled"
                c["auto_disabled_at"] = datetime.now(timezone.utc).isoformat()
                c["auto_disabled_reason"] = f"{error_count} consecutive errors: {last_error[:200]}"
                break
        else:
            return  # company not found

        with open(companies_path, "w", encoding="utf-8") as f:
            json.dump(companies, f, indent=2, ensure_ascii=False)

        print(f"[Daemon] 🔒 Auto-disabled company: {company_id}")
    except Exception as e:
        print(f"[Daemon] Error auto-disabling {company_id}: {e}")


async def background_refresh_daemon():
    """Background task that continuously refreshes companies"""
    global DAEMON_STATUS
    
    # Wait for app to fully start
    await asyncio.sleep(5)
    
    print("[Daemon] Background refresh daemon started")
    DAEMON_STATUS["running"] = True
    
    while DAEMON_STATUS["enabled"]:
        # Update lock heartbeat every iteration
        update_daemon_lock()
        
        try:
            # Load all companies from JSON
            companies_path = Path("data/companies.json")
            if companies_path.exists():
                companies = json.loads(companies_path.read_text())
            else:
                companies = []
            
            # Filter to enabled companies only
            companies = [c for c in companies if c.get("enabled", True)]
            
            # Sort by last_checked (oldest first)
            companies.sort(key=lambda c: c.get("last_checked") or "1970-01-01")
            
            DAEMON_STATUS["companies_in_cycle"] = len(companies)
            DAEMON_STATUS["cycle_count"] += 1
            DAEMON_STATUS["last_cycle_jobs_added"] = DAEMON_STATUS["jobs_added_this_cycle"]
            DAEMON_STATUS["jobs_added_this_cycle"] = 0
            DAEMON_STATUS["companies_refreshed_this_cycle"] = 0
            DAEMON_STATUS["refresh_log"] = []  # Clear log for new cycle
            
            print(f"[Daemon] Starting cycle #{DAEMON_STATUS['cycle_count']} with {len(companies)} companies")
            
            # Process in batches
            batch_size = DAEMON_STATUS["batch_size"]
            for i in range(0, len(companies), batch_size):
                if not DAEMON_STATUS["enabled"]:
                    break
                
                batch = companies[i:i+batch_size]
                DAEMON_STATUS["current_index"] = i
                
                for company in batch:
                    if not DAEMON_STATUS["enabled"]:
                        break
                    
                    company_name = company.get("name", company.get("id", "unknown"))
                    DAEMON_STATUS["current_company"] = company_name
                    
                    # Set next company
                    next_idx = companies.index(company) + 1
                    if next_idx < len(companies):
                        DAEMON_STATUS["next_company"] = companies[next_idx].get("name")
                    else:
                        DAEMON_STATUS["next_company"] = companies[0].get("name") if companies else None
                    
                    print(f"[Daemon] Refreshing: {company_name}")
                    
                    result = await refresh_company_async(company)
                    
                    DAEMON_STATUS["last_company"] = company_name
                    DAEMON_STATUS["last_updated"] = datetime.now(timezone.utc).isoformat()
                    DAEMON_STATUS["companies_refreshed_this_cycle"] += 1
                    
                    # Track jobs added
                    jobs_added = result.get("jobs_added", 0)
                    if jobs_added > 0:
                        DAEMON_STATUS["jobs_added_this_cycle"] += jobs_added
                        DAEMON_STATUS["jobs_added_total"] += jobs_added
                    
                    # Add to refresh log
                    log_entry = {
                        "company": company_name,
                        "ats": company.get("ats", ""),
                        "ok": result["ok"],
                        "jobs": result.get("jobs", 0),
                        "jobs_added": jobs_added,
                        "error": result.get("error"),
                        "time": datetime.now(timezone.utc).isoformat(),
                        "index": DAEMON_STATUS["companies_refreshed_this_cycle"],
                        "total": len(companies)
                    }
                    DAEMON_STATUS["refresh_log"].append(log_entry)
                    # Keep only last 100 entries
                    if len(DAEMON_STATUS["refresh_log"]) > 100:
                        DAEMON_STATUS["refresh_log"] = DAEMON_STATUS["refresh_log"][-100:]
                    
                    if result["ok"]:
                        added_str = f" (+{jobs_added} new)" if jobs_added > 0 else ""
                        print(f"[Daemon] ✓ {company_name}: {result['jobs']} jobs{added_str}")
                        # Reset consecutive error counter on success
                        _reset_company_errors(company.get("id", ""))
                    else:
                        print(f"[Daemon] ✗ {company_name}: {result['error']}")
                        # Track consecutive errors, auto-disable after threshold
                        _track_company_error(company, result.get("error", ""))
                
                DAEMON_STATUS["current_company"] = None
                
                # Pause between batches
                if DAEMON_STATUS["enabled"] and i + batch_size < len(companies):
                    await asyncio.sleep(DAEMON_STATUS["pause_seconds"])
            
            # Pause before next cycle (5 minutes)
            print(f"[Daemon] Cycle #{DAEMON_STATUS['cycle_count']} complete. Waiting 5 minutes...")
            await asyncio.sleep(300)
            
        except Exception as e:
            print(f"[Daemon] Error: {e}")
            await asyncio.sleep(60)  # Wait on error
    
    DAEMON_STATUS["running"] = False
    release_daemon_lock()  # Release lock when stopping
    print("[Daemon] Background refresh daemon stopped")

# Start daemon on app startup
@app.on_event("startup")
async def startup_event():
    asyncio.create_task(background_refresh_daemon())

@app.get("/daemon/status")
def get_daemon_status():
    """Get background refresh daemon status"""
    # Check lock status
    lock = check_daemon_lock()
    DAEMON_STATUS["locked_by"] = lock.get("machine") if lock else None
    return DAEMON_STATUS

@app.post("/daemon/toggle")
def toggle_daemon(enabled: bool = Query(...)):
    """Enable or disable background refresh daemon"""
    global DAEMON_STATUS
    
    if enabled:
        # Try to acquire lock before enabling
        success, message = acquire_daemon_lock()
        if not success:
            return {"ok": False, "error": message}
        DAEMON_STATUS["locked_by"] = get_machine_id()
    else:
        # Release lock when disabling
        release_daemon_lock()
        DAEMON_STATUS["locked_by"] = None
    
    DAEMON_STATUS["enabled"] = enabled
    return {"ok": True, "enabled": enabled}


@app.get("/")
async def root():
    """Redirect to UI"""
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url="/static/index.html")


def _is_us_location(location: str | None) -> bool:
    if not location:
        return False
    loc = location.lower()
    if "united states" in loc or "usa" in loc or "us" in loc:
        return True
    # очень грубый хак по штатам
    us_markers = [
        ", ca", ", ny", ", wa", ", ma", ", tx", ", co", ", il", ", ga", ", nc",
        "washington, dc", "new york, ny", "san francisco", "remote - us",
    ]
    return any(m in loc for m in us_markers)


def _mark_company_status(profile: str, cfg: dict, ok: bool, error: str | None = None):
    key = f"{profile}:{cfg.get('company', '')}"
    company_fetch_status[key] = {
        "ok": ok,
        "error": error or "",
        "checked_at": datetime.utcnow().isoformat() + "Z",
        "ats": cfg.get("ats", ""),
        "url": cfg.get("url", ""),
    }
    # Save to file for persistence
    save_company_status(company_fetch_status)


def _fetch_for_company(profile: str, cfg: dict, _retry: bool = False) -> list[dict]:
    """
    Унифицированный вызов парсеров + запись статуса компании.
    Также добавляет нормализованную локацию, классификацию роли, geo bucket/score, company_data.
    _retry: internal flag to prevent infinite recursion
    """
    company = cfg.get("company", "")
    ats = cfg.get("ats", "")
    url = cfg.get("url", "")

    try:
        if ats == "greenhouse":
            jobs = fetch_greenhouse(company, url)
        elif ats == "lever":
            jobs = fetch_lever(company, url)
        elif ats == "smartrecruiters":
            # SmartRecruiters needs api_url, not board_url
            api_url = cfg.get("api_url") or url
            jobs = fetch_smartrecruiters(company, api_url)
        elif ats == "ashby":
            jobs = fetch_ashby_jobs(url)
        elif ats == "workday":
            jobs = fetch_workday_v2(company, url)
        elif ats == "atlassian":
            jobs = fetch_atlassian(company, url)
        elif ats == "phenom":
            jobs = fetch_phenom_jobs(url)
        elif ats == "icims":
            jobs = fetch_icims(company, url)
        elif ats == "jibe":
            jobs = fetch_jibe(company, url)
        else:
            jobs = []

        # записываем успех
        _mark_company_status(profile, cfg, ok=True)

        # добавляем мета-инфу к каждой вакансии
        for j in jobs:
            j["company"] = company
            j["industry"] = cfg.get("industry", "")
            if not j.get("ats"):
                j["ats"] = ats

            # Генерируем уникальный ID (используя ats_job_id из парсера)
            j["id"] = generate_job_id(j)

            # нормализация локации
            loc_norm = normalize_location(j.get("location"))
            j["location_norm"] = loc_norm

            # классификация роли (улучшенная, с roles.json)
            role = classify_role(j.get("title"), j.get("description") or j.get("jd") or "")
            j["role_family"] = role.get("role_family")
            j["role_category"] = role.get("role_category")  # primary/adjacent/unknown/excluded
            j["role_id"] = role.get("role_id")
            j["role_confidence"] = role.get("confidence")
            j["role_reason"] = role.get("reason")
            j["role_excluded"] = role.get("excluded", False)
            j["role_exclude_reason"] = role.get("exclude_reason")

            # company data (for scoring/prioritization)
            j["company_data"] = {
                "priority": cfg.get("priority", 0),
                "hq_state": cfg.get("hq_state", None),
                "region": cfg.get("region", None),
                "tags": cfg.get("tags", []),
            }

            # geo bucket + score
            bucket, score = compute_geo_bucket_and_score(loc_norm)
            j["geo_bucket"] = bucket
            j["geo_score"] = score

        return jobs

    except Exception as e:  # noqa: BLE001
        error_str = str(e)
        print(f"Error for {company}: {error_str}")
        
        # Try auto-repair for 404 errors (only on first attempt, not retry)
        if "404" in error_str and not _retry:
            print(f"  🔧 Attempting auto-repair for {company}...")
            repair_result = try_repair_company({
                "name": company,
                "ats": ats,
                "board_url": url,
            })
            
            if repair_result and repair_result.get("verified"):
                new_ats = repair_result["ats"]
                new_url = repair_result["board_url"]
                print(f"  ✅ Found new URL: {new_ats} → {new_url}")
                
                # Update companies.json
                try:
                    companies_path = Path("data/companies.json")
                    with open(companies_path, "r") as f:
                        companies = json.load(f)
                    
                    for c in companies:
                        if c.get("name", "").lower() == company.lower():
                            c["ats"] = new_ats
                            c["board_url"] = new_url
                            print(f"  ✅ Updated companies.json for {company}")
                            break
                    
                    with open(companies_path, "w") as f:
                        json.dump(companies, f, indent=2, ensure_ascii=False)
                    
                    # Retry fetch with new URL (with _retry=True to prevent loop)
                    new_cfg = cfg.copy()
                    new_cfg["ats"] = new_ats
                    new_cfg["url"] = new_url
                    return _fetch_for_company(profile, new_cfg, _retry=True)
                    
                except Exception as update_err:
                    print(f"  ❌ Failed to update companies.json: {update_err}")
            else:
                print(f"  ❌ Auto-repair failed for {company}")
        
        # Mark as failed
        _mark_company_status(profile, cfg, ok=False, error=error_str)
        return []


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/env")
def get_env():
    """Return current environment (PROD/DEV)"""
    return {"env": ENV}


class StatusUpdate(BaseModel):
    profile: str = "all"
    job_key: str
    status: str


@app.get("/job_status")
def get_job_status(profile: str = Query("all")):
    """
    Return status map for a profile:
      { "count": N, "statuses": { job_key: status } }
    """
    mp = _load_job_status_map(profile)
    return {"count": len(mp), "statuses": mp}


@app.post("/job_status")
def update_job_status(payload: StatusUpdate):
    """
    Update status for a job_key under a profile.
    """
    status = payload.status if payload.status in VALID_APPLICATION_STATUSES else "New"
    _set_job_status(payload.profile, payload.job_key, status)
    return {"ok": True, "profile": payload.profile, "job_key": payload.job_key, "status": status}


@app.get("/jobs")
async def get_jobs(
    profile: str = Query("all", description="Имя профиля из папки profiles/*.json"),
    ats_filter: str = Query("all", description="all / greenhouse / lever / smartrecruiters"),
    role_filter: str = Query("all", description="all / product / tpm_program / project / other"),
    location_filter: str = Query("all", description="all / us / nonus"),
    company_filter: str = Query("", description="подстрока в названии компании"),
    search: str = Query("", description="поиск по title+location"),
    states: str = Query("", description="Comma-separated US state codes or full names, e.g. NC,VA,South Carolina"),
    include_remote_usa: bool = Query(False, description="Include Remote-USA roles in addition to state selection"),
    state: str = Query("", description="(deprecated) Filter by state substring"),
    city: str = Query("", description="Filter by city substring"),
    geo_mode: str = Query("all", description="all / nc_priority / local_only / neighbor_only / remote_usa"),
    refresh: bool = Query(False, description="Force refresh from ATS, ignore cache"),
):
    """
    Основной эндпоинт: собирает вакансии по профилю и фильтрам.
    """
    # NEW: Check cache first (unless refresh=True)
    cache_key = profile
    cached = None if refresh else load_cache(cache_key, ignore_ttl=True)
    
    if cached:
        print(f"✅ Using cached data ({cached['jobs_count']} jobs)")
        all_jobs = cached["jobs"]
    else:
        # Parse from companies
        companies_cfg = load_profile(profile)
        all_jobs: list[dict] = []
        
        # Filter companies first
        companies_to_fetch = []
        for cfg in companies_cfg:
            if cfg.get("enabled") == False:
                continue
            ats = cfg.get("ats", "")
            if ats_filter != "all" and ats_filter != ats:
                continue
            companies_to_fetch.append(cfg)
        
        # Parallel fetch with ThreadPoolExecutor
        def fetch_company(cfg):
            return _fetch_for_company(profile, cfg)
        
        with ThreadPoolExecutor(max_workers=8) as executor:
            futures = {executor.submit(fetch_company, cfg): cfg for cfg in companies_to_fetch}
            for future in as_completed(futures):
                try:
                    jobs = future.result()
                    all_jobs.extend(jobs)
                except Exception as e:
                    cfg = futures[future]
                    print(f"Error fetching {cfg.get('company', 'unknown')}: {e}")
        
        # NEW: Save to cache after parsing all companies
        save_cache(cache_key, all_jobs)
    
    # Load status map once
    status_map = _load_job_status_map(profile)

    # --- фильтры на уровне вакансий ---

    # parse states CSV to normalized list of 2-letter codes
    raw_states = [s.strip() for s in states.split(",") if s.strip()]
    normalized_states: list[str] = []
    for s in raw_states:
        s_low = s.lower()
        if s.upper() in STATE_MAP.values():
            normalized_states.append(s.upper())
        elif s_low in STATE_MAP:
            normalized_states.append(STATE_MAP[s_low])
        else:
            normalized_states.append(s.upper())

    states_set_upper = set(ns.upper() for ns in normalized_states)
    cities_set = set([city.lower()]) if city else set()

    def match_role(job: dict) -> bool:
        if role_filter == "all":
            return True
        return job.get("role_family") == role_filter

    def match_location(loc: str | None) -> bool:
        if location_filter == "all":
            return True
        is_us = _is_us_location(loc)
        if location_filter == "us":
            return is_us
        if location_filter == "nonus":
            return not is_us
        return True

    def match_company(name: str | None) -> bool:
        if not company_filter:
            return True
        if not name:
            return False
        return company_filter.lower() in name.lower()

    def match_search(job: dict) -> bool:
        if not search:
            return True
        s = search.lower()
        haystack = f"{job.get('title', '')} {job.get('location', '')} {job.get('company', '')}".lower()
        return s in haystack

    def match_states(job: dict) -> bool:
        loc_norm = job.get("location_norm", {}) or {}
        # Collect job states as 2-letter codes where possible
        job_states = []
        if isinstance(loc_norm.get("states"), list):
            job_states.extend([str(st).upper() for st in loc_norm.get("states") if st])
        if loc_norm.get("state"):
            job_states.append(str(loc_norm.get("state")).upper())
        if loc_norm.get("state_full"):
            sf = str(loc_norm.get("state_full")).lower()
            if sf in STATE_MAP:
                job_states.append(STATE_MAP[sf])

        # Remote-USA flag
        remote_usa = bool(loc_norm.get("remote")) and (str(loc_norm.get("remote_scope") or "").lower() in ["usa", "us"])
        state_matches = any(ns in job_states for ns in states_set_upper)

        # New behavior: states selection and remote toggle are independent.
        if normalized_states and include_remote_usa:
            return state_matches or remote_usa
        if normalized_states and not include_remote_usa:
            return state_matches
        if not normalized_states and include_remote_usa:
            return remote_usa
        return True

    def match_old_state(job: dict) -> bool:
        # fallback old state substring filter
        if not state:
            return True
        loc = job.get("location", "") or ""
        return state.lower() in loc.lower()

    def match_city(job: dict) -> bool:
        if not city:
            return True
        loc_norm = job.get("location_norm", {}) or {}
        if loc_norm:
            return city.lower() == str(loc_norm.get("city") or "").lower()
        loc = job.get("location", "") or ""
        return city.lower() in loc.lower()

    def match_geo(job: dict) -> bool:
        bucket = job.get("geo_bucket", "unknown")
        if geo_mode == "all":
            return True
        elif geo_mode == "nc_priority":
            return bucket in ["local", "nc", "neighbor", "remote_usa"]
        elif geo_mode == "local_only":
            return bucket == "local"
        elif geo_mode == "neighbor_only":
            return bucket == "neighbor"
        elif geo_mode == "remote_usa":
            return bucket == "remote_usa"
        return True

    filtered: list[dict] = []
    for j in all_jobs:
        if not match_role(j):
            continue
        if not match_states(j):
            continue
        if not match_old_state(j):
            continue
        if not match_city(j):
            continue
        if not match_geo(j):
            continue
        if not match_location(j.get("location")):
            continue
        if not match_company(j.get("company")):
            continue
        if not match_search(j):
            continue
        filtered.append(j)

    # compute score and sort
    def parse_date(datestr: str | None):
        if not datestr:
            return None
        try:
            ds = datestr
            if ds.endswith("Z"):
                ds = ds[:-1] + "+00:00"
            dt = datetime.fromisoformat(ds)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt.astimezone(timezone.utc)
        except Exception:
            return None

    now = datetime.now(timezone.utc)
    for job in filtered:
        score = 0
        loc_norm = job.get("location_norm", {}) or {}

        job_state_upper = (loc_norm.get("state") or "").upper()
        job_city_lower = (loc_norm.get("city") or "").lower()
        job_remote_scope = (loc_norm.get("remote_scope") or "").lower()
        job_remote = bool(loc_norm.get("remote"))

        # Prefer explicit states/cities selections
        if states_set_upper and job_state_upper in states_set_upper:
            score += 30
        if cities_set and job_city_lower in cities_set:
            score += 15

        # If include_remote_usa requested, give a boost
        if include_remote_usa and job_remote_scope == "usa":
            score += 20
        if not states_set_upper and not city and job_remote:
            score += 5

        # Company priority
        company_data = job.get("company_data") or {}
        score += int(company_data.get("priority") or 0)

        # Freshness penalty
        updated = parse_date(job.get("updated_at"))
        if updated:
            age_days = (now - updated).days
            if age_days > 60:
                score -= 20
            elif age_days > 30:
                score -= 10

        # Add geo_score as primary weight
        score += int(job.get("geo_score", 0))

        job["score"] = score

        # Attach job_key + status
        job_key = compute_job_key(job)
        job["job_key"] = job_key
        job["application_status"] = status_map.get(job_key, "New")

    filtered.sort(key=lambda j: (j.get("score", 0), str(j.get("updated_at") or "")), reverse=True)

    # ========== PIPELINE SYNC ==========
    # Sync relevant jobs with pipeline storage
    try:
        known_ids = get_all_job_ids()
        active_ids = set()
        
        for job in filtered:
            job_id = job.get("id")
            if not job_id:
                continue
            
            active_ids.add(job_id)
            
            # Only sync jobs with relevant roles
            role_family = job.get("role_family", "other")
            if role_family not in ["product", "tpm_program", "project"]:
                continue
            
            # Skip if role was excluded
            if job.get("role_excluded"):
                continue
            
            if job_id in known_ids:
                # Already known - update last_seen
                update_last_seen(job_id, is_active=True)
            else:
                # New job - add to inbox
                add_job(job)
        
        # Mark missing jobs as potentially closed
        mark_missing_jobs(active_ids, days_threshold=3)
        
    except Exception as e:
        print(f"Pipeline sync error: {e}")
    # ========== END PIPELINE SYNC ==========

    return {"count": len(filtered), "jobs": filtered}


@app.get("/companies")
def get_companies(
    profile: str = Query("all", description="Имя профиля из profiles/*.json"),
    my_roles: bool = Query(False, description="Filter by My Roles (Product, TPM, Program)"),
    my_location: bool = Query(False, description="Filter by My Location (US + Remote USA + NC,VA,SC,GA,TN)"),
):
    """
    Возвращает список ВСЕХ компаний профиля + статус последней попытки fetch'а + статистика jobs.
    """
    companies_cfg = load_profile(profile)
    
    # My Roles filter
    MY_ROLES = ["product", "tpm_program", "project"]
    
    # My Location filter states
    MY_LOCATION_STATES = ["NC", "VA", "SC", "GA", "TN"]
    
    # Load all jobs from pipeline for counting
    all_pipeline_jobs = get_active_jobs()
    
    # Load cache to get total jobs per company (all jobs from ATS)
    cache_data = load_cache(profile, ignore_ttl=True)  # Cache key is just 'all', not 'jobs_all'
    cache_jobs = cache_data.get("jobs", []) if cache_data else []
    
    # Build total_jobs per company from cache
    total_jobs_by_company = {}
    for job in cache_jobs:
        company_name = job.get("company", "")
        total_jobs_by_company[company_name] = total_jobs_by_company.get(company_name, 0) + 1
    
    # Filter jobs by role if my_roles is enabled
    def job_matches_my_roles(job):
        if not my_roles:
            return True
        role_family = job.get("role_family", "other")
        return role_family in MY_ROLES
    
    # Filter jobs by geo if my_location is enabled
    def job_matches_my_location(job):
        if not my_location:
            return True
        
        loc_norm = job.get("location_norm", {})
        state = loc_norm.get("state", "")
        remote = loc_norm.get("remote", False)
        remote_scope = loc_norm.get("remote_scope", "")
        
        # Check if Remote USA
        if remote and remote_scope == "usa":
            return True
        
        # Check if in My Location states
        if state in MY_LOCATION_STATES:
            return True
        
        return False
    
    filtered_jobs = [j for j in all_pipeline_jobs if job_matches_my_roles(j) and job_matches_my_location(j)]
    
    # Build company stats: company_name -> {jobs_count, applied_count, new_count, status_counts}
    company_stats = {}
    for job in filtered_jobs:
        company_name = job.get("company", "")
        if company_name not in company_stats:
            company_stats[company_name] = {
                "jobs_count": 0,
                "new_count": 0,
                "applied_count": 0,
                "interview_count": 0,
            }
        
        stats = company_stats[company_name]
        stats["jobs_count"] += 1
        
        status = job.get("status", "New")
        if status == STATUS_NEW:
            stats["new_count"] += 1
        elif status == STATUS_APPLIED:
            stats["applied_count"] += 1
        elif status == STATUS_INTERVIEW:
            stats["interview_count"] += 1
    
    items: list[dict] = []

    for cfg in companies_cfg:
        company_name = cfg.get("company", "") or cfg.get("name", "")
        key = f"{profile}:{company_name}"
        st = company_fetch_status.get(key, {})
        
        # For disabled companies, override status
        is_disabled = cfg.get("enabled") == False
        company_status = cfg.get("status", "active")
        
        # Get stats for this company
        stats = company_stats.get(company_name, {
            "jobs_count": 0,
            "new_count": 0, 
            "applied_count": 0,
            "interview_count": 0,
        })

        ats_type = cfg.get("ats", "")
        items.append(
            {
                "company": company_name,
                "id": cfg.get("id", ""),
                "industry": cfg.get("industry", ""),
                "tags": cfg.get("tags", []),
                "ats": ats_type,
                "ats_supported": ats_type in SUPPORTED_ATS,  # Can auto-fetch jobs
                "url": cfg.get("url", "") or cfg.get("board_url", ""),
                "enabled": cfg.get("enabled", True),
                "status": company_status,
                "last_ok": "disabled" if is_disabled else st.get("ok", None),
                "last_error": cfg.get("status", "") if is_disabled else st.get("error", ""),
                "last_checked": st.get("checked_at", ""),
                # Total jobs from cache (all jobs from ATS)
                "total_jobs": total_jobs_by_company.get(company_name, 0),
                # Stats from pipeline (filtered PM/TPM jobs)
                "jobs_count": stats["jobs_count"],
                "new_count": stats["new_count"],
                "applied_count": stats["applied_count"],
                "interview_count": stats["interview_count"],
            }
        )

    # Sort by total_jobs desc, then by name
    items.sort(key=lambda x: (-x["total_jobs"], x["company"].lower()))
    
    # Summary stats
    total_jobs = sum(c["jobs_count"] for c in items)
    total_new = sum(c["new_count"] for c in items)
    total_applied = sum(c["applied_count"] for c in items)
    total_interview = sum(c["interview_count"] for c in items)

    return {
        "count": len(items),
        "companies": items,
        "supported_ats": SUPPORTED_ATS,  # List of ATS that support auto-fetching
        "summary": {
            "total_jobs": total_jobs,
            "total_new": total_new,
            "total_applied": total_applied,
            "total_interview": total_interview,
        }
    }


@app.get("/ats-info")
def get_ats_info():
    """
    Get information about supported ATS systems.
    Returns which ATS support automatic job fetching and collected unsupported ATS.
    """
    # Load unsupported ATS data
    unsupported_ats = {}
    unsupported_file = Path("data/unsupported_ats.json")
    if unsupported_file.exists():
        try:
            with open(unsupported_file, "r") as f:
                data = json.load(f)
                unsupported_ats = data.get("ats_systems", {})
        except:
            pass

    # Format unsupported ATS for response (skip those already in SUPPORTED_ATS)
    unsupported_list = []
    for ats_key, ats_data in unsupported_ats.items():
        if ats_key in ATS_PARSERS:
            continue  # Already supported, skip
        unsupported_list.append({
            "name": ats_data.get("name", ats_key),
            "companies_count": len(ats_data.get("companies_using", [])),
            "endpoints_found": len(ats_data.get("api_endpoints", [])),
            "parser_status": ats_data.get("parser_status", "not_started"),
            "first_seen": ats_data.get("first_seen"),
        })

    return {
        "supported_ats": SUPPORTED_ATS,
        "descriptions": {
            "greenhouse": "Greenhouse - Full API support, auto-fill available",
            "lever": "Lever - Full API support, auto-fill available",
            "ashby": "Ashby - Full API support, auto-fill available",
            "workday": "Workday - API support (limited)",
            "smartrecruiters": "SmartRecruiters - API support",
            "atlassian": "Atlassian - API support",
            "phenom": "Phenom - API support",
            "icims": "iCIMS - HTML parsing",
            "jibe": "Jibe (Google Hire) - API support",
        },
        "auto_fill_supported": ["greenhouse", "lever", "ashby"],
        "unsupported_ats": unsupported_list,
        "unsupported_count": len(unsupported_list),
    }


@app.post("/ats-discovery/{ats_name}")
def trigger_ats_discovery_endpoint(ats_name: str, url: str = Query(..., description="Careers page URL to analyze")):
    """
    Manually trigger ATS discovery for a specific ATS system.
    Analyzes the careers page and collects API endpoints for parser generation.
    """
    from tools.ats_discovery import load_unsupported_ats

    # Check if already has enough data
    data = load_unsupported_ats()
    ats_key = ats_name.lower().replace(" ", "_")
    existing = data.get("ats_systems", {}).get(ats_key, {})

    trigger_ats_discovery_background(ats_name, url, "manual_trigger")

    return {
        "status": "discovery_started",
        "ats": ats_name,
        "url": url,
        "existing_endpoints": len(existing.get("api_endpoints", [])),
        "message": "Discovery running in background. Check /ats-info for results."
    }


@app.post("/ats-generate/{ats_name}")
def trigger_parser_generation(ats_name: str):
    """
    Trigger parser generation for an ATS system using Claude API.
    Requires collected discovery data and ANTHROPIC_API_KEY.
    """
    import threading
    from tools.ats_discovery import load_unsupported_ats

    # Check if we have data
    data = load_unsupported_ats()
    ats_key = ats_name.lower().replace(" ", "_")
    ats_data = data.get("ats_systems", {}).get(ats_key)

    if not ats_data:
        return {"error": f"No discovery data for ATS '{ats_name}'. Run discovery first."}

    if len(ats_data.get("api_endpoints", [])) == 0:
        return {"error": f"No API endpoints found for '{ats_name}'. Need more discovery data."}

    # Check for API key
    if not os.getenv("ANTHROPIC_API_KEY"):
        return {"error": "ANTHROPIC_API_KEY not set. Cannot generate parser."}

    def generate_task():
        try:
            from tools.ats_parser_generator import generate_parser, test_parser, update_parser_status
            print(f"[Parser Generator] Starting generation for {ats_name}")

            result = generate_parser(ats_name, save=True)
            if result.get("error"):
                update_parser_status(ats_name, "generation_failed", result["error"])
                print(f"[Parser Generator] Generation failed: {result['error']}")
                return

            # Test the parser
            test_result = test_parser(ats_name)
            if test_result.get("error"):
                update_parser_status(ats_name, "test_failed", test_result["error"])
            elif test_result.get("jobs_count", 0) > 0:
                update_parser_status(ats_name, "completed", f"Found {test_result['jobs_count']} jobs")
            else:
                update_parser_status(ats_name, "needs_review", "Parser runs but found 0 jobs")

            print(f"[Parser Generator] Completed for {ats_name}")

        except Exception as e:
            print(f"[Parser Generator] Error: {e}")

    thread = threading.Thread(target=generate_task, daemon=True)
    thread.start()

    return {
        "status": "generation_started",
        "ats": ats_name,
        "endpoints_available": len(ats_data.get("api_endpoints", [])),
        "message": "Parser generation running in background. Check /ats-info for status."
    }


class CompanyCreate(BaseModel):
    name: str
    ats: str  # greenhouse, lever, smartrecruiters
    board_url: str
    industry: str = ""
    tags: list[str] = []


def trigger_ats_discovery_background(ats_name: str, careers_url: str, company_name: str):
    """
    Trigger ATS discovery in background thread when new unsupported ATS is detected.
    Collects API endpoints and data for future parser generation.
    """
    import threading

    def discover_task():
        try:
            from tools.ats_discovery import discover_api_endpoints, load_unsupported_ats, save_unsupported_ats
            import asyncio

            print(f"[ATS Discovery] Starting discovery for {ats_name} from {careers_url}")

            # Run async discovery
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            result = loop.run_until_complete(discover_api_endpoints(careers_url))
            loop.close()

            if result.get("error"):
                print(f"[ATS Discovery] Error: {result['error']}")
                return

            # Update unsupported_ats.json
            data = load_unsupported_ats()
            ats_key = ats_name.lower().replace(" ", "_")

            if ats_key not in data.get("ats_systems", {}):
                data.setdefault("ats_systems", {})[ats_key] = {
                    "name": ats_name,
                    "first_seen": datetime.now().isoformat(),
                    "companies_using": [],
                    "api_endpoints": [],
                    "sample_responses": [],
                    "parser_status": "not_started"
                }

            ats_data = data["ats_systems"][ats_key]

            # Add company if not already there
            if company_name not in ats_data.get("companies_using", []):
                ats_data.setdefault("companies_using", []).append(company_name)

            # Merge discovered endpoints
            existing_urls = {ep.get("url") if isinstance(ep, dict) else ep for ep in ats_data.get("api_endpoints", [])}
            for ep in result.get("api_endpoints", []):
                ep_url = ep.get("url") if isinstance(ep, dict) else ep
                if ep_url not in existing_urls:
                    ats_data.setdefault("api_endpoints", []).append(ep)

            # Add sample responses
            for sample in result.get("json_responses", [])[:3]:
                ats_data.setdefault("sample_responses", []).append(sample)

            # Update status
            if ats_data.get("parser_status") == "not_started" and len(ats_data.get("api_endpoints", [])) > 0:
                ats_data["parser_status"] = "data_collected"

            save_unsupported_ats(data)
            print(f"[ATS Discovery] Completed for {ats_name}: {len(result.get('api_endpoints', []))} endpoints found")

        except Exception as e:
            print(f"[ATS Discovery] Error in background task: {e}")

    # Run in background thread
    thread = threading.Thread(target=discover_task, daemon=True)
    thread.start()


@app.post("/companies")
def add_company(company: CompanyCreate):
    """
    Add a new company to companies.json.
    Validates: duplicate check by name/id AND board_url.
    Auto-detects ATS from board_url if ats is "universal" or empty.
    Normalizes board_url (strip trailing slash).
    If ATS is unsupported, triggers background discovery for future parser generation.
    """
    companies_path = Path("data/companies.json")

    # Load existing
    if companies_path.exists():
        with open(companies_path, "r") as f:
            companies = json.load(f)
    else:
        companies = []

    # Generate id from name
    company_id = company.name.lower().replace(" ", "_").replace("-", "_")

    # Normalize board_url (strip trailing slash, whitespace)
    board_url = company.board_url.strip().rstrip("/") if company.board_url else ""

    # Check if already exists by name/id
    for c in companies:
        if c.get("id") == company_id or c.get("name", "").lower() == company.name.lower():
            return {"error": f"Company '{company.name}' already exists (by name/id)", "status": "exists"}

    # Check duplicate by board_url (normalized comparison)
    if board_url:
        norm_url = board_url.lower().rstrip("/")
        for c in companies:
            existing_url = (c.get("board_url") or "").lower().rstrip("/")
            if existing_url and existing_url == norm_url:
                return {
                    "error": f"Company with board_url '{board_url}' already exists as '{c.get('name')}' (id={c.get('id')})",
                    "status": "exists"
                }

    # Auto-detect ATS if universal or empty
    ats_type = (company.ats or "").lower().strip()
    auto_detected = False
    if ats_type in ("", "universal", "other", "unknown") and board_url:
        detected = detect_ats_from_url(board_url)
        detected_ats = detected.get("ats", "universal")
        if detected_ats in SUPPORTED_ATS:
            ats_type = detected_ats
            # Use the normalized board_url from detection
            if detected.get("board_url"):
                board_url = detected["board_url"]
            auto_detected = True
            print(f"[AddCompany] Auto-detected ATS: {ats_type} for {board_url}")
        else:
            ats_type = detected_ats  # Keep detected ATS even if unsupported

    # Create new company entry
    new_company = {
        "id": company_id,
        "name": company.name,
        "ats": ats_type,
        "board_url": board_url,
        "api_url": None,
        "tags": company.tags,
        "industry": company.industry,
        "priority": 0,
        "hq_state": None,
        "region": "us",
        "enabled": ats_type in SUPPORTED_ATS,
    }

    companies.append(new_company)

    # Save
    with open(companies_path, "w", encoding="utf-8") as f:
        json.dump(companies, f, indent=2, ensure_ascii=False)

    # Trigger ATS discovery if unsupported ATS
    if ats_type not in SUPPORTED_ATS and board_url:
        print(f"[AddCompany] New unsupported ATS '{ats_type}' detected, triggering discovery...")
        trigger_ats_discovery_background(ats_type, board_url, company.name)

    return {
        "status": "ok",
        "company": new_company,
        "ats_auto_detected": auto_detected,
        "ats_discovery_triggered": ats_type not in SUPPORTED_ATS,
    }


@app.delete("/companies/{company_id}")
def remove_company(company_id: str):
    """
    Remove a company from companies.json.
    """
    companies_path = Path("data/companies.json")
    
    if not companies_path.exists():
        return {"error": "No companies file", "status": "error"}
    
    with open(companies_path, "r") as f:
        companies = json.load(f)
    
    # Find and remove
    original_len = len(companies)
    companies = [c for c in companies if c.get("id") != company_id]
    
    if len(companies) == original_len:
        return {"error": f"Company '{company_id}' not found", "status": "not_found"}
    
    # Save
    with open(companies_path, "w") as f:
        json.dump(companies, f, indent=2)
    
    return {"status": "ok", "removed": company_id}

@app.post("/companies/{company_id}/refresh")
def refresh_single_company(company_id: str, profile: str = Query("all")):
    """
    Refresh jobs for a single company.
    Parses ATS, updates cache and pipeline.
    """
    # Load company config
    companies_cfg = load_profile(profile)
    cfg = None
    for c in companies_cfg:
        # Match by id, company name, or name field
        cid = c.get("id", "") or ""
        cname = c.get("company", "") or c.get("name", "") or ""
        if cid == company_id or cname.lower() == company_id.lower():
            cfg = c
            break
    
    if not cfg:
        return {"ok": False, "error": f"Company '{company_id}' not found"}
    
    company_name = cfg.get("company", "")
    
    # Fetch jobs for this company
    try:
        jobs = _fetch_for_company(profile, cfg)
        
        # Update cache - load existing, replace this company's jobs, save
        cached = load_cache(profile, ignore_ttl=True) or {"jobs": []}
        other_jobs = [j for j in cached.get("jobs", []) if j.get("company") != company_name]
        all_jobs = other_jobs + jobs
        save_cache(profile, all_jobs)
        
        # Sync to pipeline (add new My Roles jobs)
        sync_result = sync_cache_to_pipeline(jobs)
        
        return {
            "ok": True,
            "company": company_name,
            "jobs_count": len(jobs),
            "total_cache": len(all_jobs),
            "pipeline_added": sync_result["added"],
            "pipeline_updated": sync_result["updated"]
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}

@app.get("/companies/{company_id}/refresh/stream")
async def refresh_single_company_stream(company_id: str, profile: str = Query("all")):
    """
    Streaming refresh for a single company.
    Sends progress events as jobs are parsed.
    """
    import asyncio
    
    async def generate():
        # Find company config
        companies_cfg = load_profile(profile)
        cfg = None
        for c in companies_cfg:
            cid = c.get("id", "") or ""
            cname = c.get("company", "") or c.get("name", "") or ""
            if cid == company_id or cname.lower() == company_id.lower():
                cfg = c
                break
        
        if not cfg:
            yield f"data: {json.dumps({'type': 'error', 'error': f'Company {company_id} not found'})}\n\n"
            return
        
        company_name = cfg.get("company", "") or cfg.get("name", "")
        ats = cfg.get("ats", "")
        url = cfg.get("url", "") or cfg.get("board_url", "")
        
        yield f"data: {json.dumps({'type': 'start', 'company': company_name, 'ats': ats})}\n\n"
        
        try:
            if ats == "workday":
                # Workday: stream progress page by page
                from parsers.workday_v2 import fetch_workday_v2_streaming
                
                raw_jobs = []
                for event in fetch_workday_v2_streaming(company_name, url):
                    if event.get("type") == "progress":
                        yield f"data: {json.dumps({'type': 'progress', 'jobs': event['jobs']})}\n\n"
                    elif event.get("type") == "done":
                        raw_jobs = event.get("jobs", [])
                    elif event.get("type") == "error":
                        yield f"data: {json.dumps({'type': 'error', 'error': event['error']})}\n\n"
                        return
                
                # Enrich jobs (same as _fetch_for_company)
                jobs = []
                for j in raw_jobs:
                    j["company"] = company_name
                    j["industry"] = cfg.get("industry", "")
                    if not j.get("ats"):
                        j["ats"] = ats
                    j["id"] = generate_job_id(j)
                    loc_norm = normalize_location(j.get("location"))
                    j["location_norm"] = loc_norm
                    role = classify_role(j.get("title"), j.get("description") or "")
                    j["role_family"] = role.get("role_family")
                    j["role_category"] = role.get("role_category")
                    j["role_id"] = role.get("role_id")
                    j["role_confidence"] = role.get("confidence")
                    j["role_reason"] = role.get("reason")
                    j["role_excluded"] = role.get("excluded", False)
                    j["role_exclude_reason"] = role.get("exclude_reason")
                    j["company_data"] = {
                        "priority": cfg.get("priority", 0),
                        "hq_state": cfg.get("hq_state"),
                        "region": cfg.get("region"),
                        "tags": cfg.get("tags", []),
                    }
                    bucket, score = compute_geo_bucket_and_score(loc_norm)
                    j["geo_bucket"] = bucket
                    j["geo_score"] = score
                    jobs.append(j)
                
                # Mark company status
                _mark_company_status(profile, cfg, ok=True)
            else:
                # Other ATS: single fetch
                yield f"data: {json.dumps({'type': 'progress', 'jobs': 0, 'message': 'Fetching...'})}\n\n"
                loop = asyncio.get_event_loop()
                jobs = await loop.run_in_executor(None, lambda: _fetch_for_company(profile, cfg))
            
            jobs_count = len(jobs)
            
            # Update cache
            cached = load_cache(profile, ignore_ttl=True) or {"jobs": []}
            other_jobs = [j for j in cached.get("jobs", []) if j.get("company") != company_name]
            all_jobs = other_jobs + jobs
            save_cache(profile, all_jobs)
            
            yield f"data: {json.dumps({'type': 'done', 'jobs': jobs_count, 'total_cache': len(all_jobs)})}\n\n"
            
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)[:200]})}\n\n"
    
    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
        }
    )


# ===== Discovery endpoints =====
# Auto-discovery of companies with relevant PM/TPM roles

@app.post("/discovery/search")
def discovery_search(ai: bool = True, seed: bool = True):
    """
    Run company discovery pipeline: AI search + seed list.
    Found companies go to data/discovered_companies.json staging area.
    """
    from tools.company_discovery import (
        load_companies, load_staging, save_staging,
        get_existing_ids, get_staging_ids,
        discover_via_ai, discover_from_seed_list,
    )

    existing_ids = get_existing_ids()
    existing_names = {c.get("name", "") for c in load_companies()}
    candidates = load_staging()
    initial_count = len(candidates)

    new_candidates = []
    ai_count = 0
    seed_count = 0

    if ai:
        ai_candidates = discover_via_ai(existing_ids, existing_names)
        new_candidates.extend(ai_candidates)
        ai_count = len(ai_candidates)

    if seed:
        seed_candidates = discover_from_seed_list(existing_ids)
        new_candidates.extend(seed_candidates)
        seed_count = len(seed_candidates)

    # Deduplicate with staging
    staging_ids = {c.get("id") for c in candidates}
    added = 0
    for nc in new_candidates:
        if nc["id"] not in staging_ids:
            candidates.append(nc)
            staging_ids.add(nc["id"])
            added += 1

    save_staging(candidates)

    return {
        "ok": True,
        "ai_suggested": ai_count,
        "seed_suggested": seed_count,
        "added_to_staging": added,
        "total_staging": len(candidates),
    }


@app.get("/discovery/candidates")
def discovery_candidates(status: str = None):
    """
    List discovered companies from staging area.
    Optional filter by status: pending_validation, validated, ready_to_approve, etc.
    """
    staging_path = Path("data/discovered_companies.json")
    if not staging_path.exists():
        return {"candidates": [], "total": 0}

    with open(staging_path, "r", encoding="utf-8") as f:
        candidates = json.load(f)

    if status:
        candidates = [c for c in candidates if c.get("status") == status]

    # Group counts by status
    all_candidates = json.load(open(staging_path, "r", encoding="utf-8"))
    status_counts = {}
    for c in all_candidates:
        s = c.get("status", "unknown")
        status_counts[s] = status_counts.get(s, 0) + 1

    return {
        "candidates": candidates,
        "total": len(candidates),
        "status_counts": status_counts,
    }


@app.post("/discovery/validate")
def discovery_validate():
    """
    Validate ATS for pending candidates in staging.
    Detects ATS via URL patterns and HTTP checks.
    """
    from tools.company_discovery import load_staging, save_staging, validate_candidates

    candidates = load_staging()
    if not candidates:
        return {"ok": True, "validated": 0, "message": "Staging is empty"}

    changes = validate_candidates(candidates)
    save_staging(candidates)

    supported = sum(1 for c in candidates if c.get("supported"))
    unsupported = sum(1 for c in candidates if c.get("status") == "unsupported_ats")
    no_ats = sum(1 for c in candidates if c.get("status") == "no_ats_detected")

    return {
        "ok": True,
        "validated": changes,
        "supported_ats": supported,
        "unsupported_ats": unsupported,
        "no_ats": no_ats,
    }


@app.post("/discovery/preview")
def discovery_preview():
    """
    Preview relevant PM/TPM roles for validated candidates.
    Parses jobs from ATS and counts matching role titles.
    """
    from tools.company_discovery import load_staging, save_staging, preview_relevant_roles

    candidates = load_staging()
    if not candidates:
        return {"ok": True, "previewed": 0, "message": "Staging is empty"}

    changes = preview_relevant_roles(candidates)
    save_staging(candidates)

    ready = sum(1 for c in candidates if c.get("status") == "ready_to_approve")
    no_roles = sum(1 for c in candidates if c.get("status") == "no_relevant_roles")

    return {
        "ok": True,
        "previewed": changes,
        "ready_to_approve": ready,
        "no_relevant_roles": no_roles,
    }


@app.post("/discovery/approve/{candidate_id}")
def discovery_approve(candidate_id: str):
    """
    Approve a discovered company — move from staging to companies.json.
    Triggers initial parsing via refresh_company_sync() (same as /onboard auto-add).
    """
    staging_path = Path("data/discovered_companies.json")
    companies_path = Path("data/companies.json")

    # Load staging
    if not staging_path.exists():
        return {"ok": False, "error": "No staging file"}

    with open(staging_path, "r", encoding="utf-8") as f:
        candidates = json.load(f)

    # Find candidate
    candidate = None
    for c in candidates:
        if c.get("id") == candidate_id:
            candidate = c
            break

    if not candidate:
        return {"ok": False, "error": f"Candidate '{candidate_id}' not found in staging"}

    if not candidate.get("supported"):
        return {"ok": False, "error": f"Candidate has unsupported ATS: {candidate.get('ats')}"}

    # Check not already in companies.json
    companies = json.load(open(companies_path, "r", encoding="utf-8")) if companies_path.exists() else []
    for c in companies:
        if c.get("id") == candidate_id or c.get("name", "").lower() == candidate.get("name", "").lower():
            return {"ok": False, "error": f"Company '{candidate_id}' already exists"}

    # Create company entry (same format as /onboard auto-add: main.py:3004-3015)
    new_company = {
        "id": candidate_id,
        "name": candidate.get("name", candidate_id),
        "ats": candidate.get("ats"),
        "board_url": candidate.get("board_url"),
        "industry": candidate.get("industry", ""),
        "tags": candidate.get("tags", []),
        "priority": 0,
        "hq_state": candidate.get("hq_state"),
        "region": "us",
        "enabled": True,
    }

    companies.append(new_company)
    with open(companies_path, "w", encoding="utf-8") as f:
        json.dump(companies, f, indent=2, ensure_ascii=False)

    # Trigger initial parsing (same as /onboard: main.py:3021-3033)
    parsing_result = {"ok": False, "jobs": 0}
    try:
        fetch_result = refresh_company_sync(new_company)
        parsing_result = {
            "ok": fetch_result.get("ok", False),
            "jobs": fetch_result.get("jobs", 0),
            "jobs_added": fetch_result.get("jobs_added", 0),
        }
    except Exception as e:
        parsing_result["error"] = str(e)

    # Update staging status
    candidate["status"] = "approved"
    candidate["approved_at"] = datetime.now().isoformat()
    with open(staging_path, "w", encoding="utf-8") as f:
        json.dump(candidates, f, indent=2, ensure_ascii=False)

    return {
        "ok": True,
        "company": new_company,
        "parsing": parsing_result,
    }


@app.post("/discovery/reject/{candidate_id}")
def discovery_reject(candidate_id: str):
    """Reject a discovered company — mark as rejected in staging."""
    staging_path = Path("data/discovered_companies.json")

    if not staging_path.exists():
        return {"ok": False, "error": "No staging file"}

    with open(staging_path, "r", encoding="utf-8") as f:
        candidates = json.load(f)

    found = False
    for c in candidates:
        if c.get("id") == candidate_id:
            c["status"] = "rejected"
            c["rejected_at"] = datetime.now().isoformat()
            found = True
            break

    if not found:
        return {"ok": False, "error": f"Candidate '{candidate_id}' not found"}

    with open(staging_path, "w", encoding="utf-8") as f:
        json.dump(candidates, f, indent=2, ensure_ascii=False)

    return {"ok": True, "rejected": candidate_id}


@app.post("/discovery/auto")
def discovery_auto():
    """
    Full discovery pipeline: search → validate → preview.
    Combines all steps into a single endpoint.
    """
    from tools.company_discovery import (
        load_companies, load_staging, save_staging,
        get_existing_ids, discover_via_ai, discover_from_seed_list,
        validate_candidates, preview_relevant_roles,
    )

    results = {"steps": []}

    # Step 1: Search
    existing_ids = get_existing_ids()
    existing_names = {c.get("name", "") for c in load_companies()}
    candidates = load_staging()

    new_candidates = []
    ai_candidates = discover_via_ai(existing_ids, existing_names)
    new_candidates.extend(ai_candidates)
    seed_candidates = discover_from_seed_list(existing_ids)
    new_candidates.extend(seed_candidates)

    staging_ids = {c.get("id") for c in candidates}
    added = 0
    for nc in new_candidates:
        if nc["id"] not in staging_ids:
            candidates.append(nc)
            staging_ids.add(nc["id"])
            added += 1

    save_staging(candidates)
    results["steps"].append({"search": {"ai": len(ai_candidates), "seed": len(seed_candidates), "added": added}})

    # Step 2: Validate
    validated = validate_candidates(candidates)
    save_staging(candidates)
    supported = sum(1 for c in candidates if c.get("supported"))
    results["steps"].append({"validate": {"validated": validated, "supported": supported}})

    # Step 3: Preview
    previewed = preview_relevant_roles(candidates)
    save_staging(candidates)
    ready = sum(1 for c in candidates if c.get("status") == "ready_to_approve")
    results["steps"].append({"preview": {"previewed": previewed, "ready_to_approve": ready}})

    results["ok"] = True
    results["ready_to_approve"] = ready
    results["total_staging"] = len(candidates)

    return results


@app.get("/profiles/{name}")
async def get_profile_companies(name: str):
    companies = load_profile(name)
    result_companies = []
    for c in companies:
        result_companies.append({
            "id": c.get("id"),
            "name": c.get("name"),
            "ats": c.get("ats", ""),
            "board_url": c.get("board_url", ""),
            "tags": c.get("tags", []),
            "priority": c.get("priority", 0),
            "hq_state": c.get("hq_state", None),
            "region": c.get("region", None)
        })
    return {"count": len(result_companies), "companies": result_companies}


@app.get("/debug/location_stats")
async def location_stats(profile: str = Query("all")):
    """
    Возвращает статистику по нормализованным локациям вакансий для указанного профиля.
    """
    companies_cfg = load_profile(profile)

    all_jobs: list[dict] = []
    for cfg in companies_cfg:
        jobs = _fetch_for_company(profile, cfg)
        all_jobs.extend(jobs)

    total_jobs = len(all_jobs)
    remote_usa_count = 0
    remote_global_count = 0
    jobs_with_states_count = 0
    states_counter = Counter()

    for job in all_jobs:
        loc_norm = job.get("location_norm", {}) or {}
        if loc_norm.get("remote") and str(loc_norm.get("remote_scope")).lower() == "usa":
            remote_usa_count += 1
        if loc_norm.get("remote") and str(loc_norm.get("remote_scope")).lower() == "global":
            remote_global_count += 1

        states = loc_norm.get("states") or []
        if states:
            jobs_with_states_count += 1
            states_counter.update([str(s).upper() for s in states if s])

    top_20_states = states_counter.most_common(20)

    return {
        "total_jobs": total_jobs,
        "remote_usa_count": remote_usa_count,
        "remote_global_count": remote_global_count,
        "jobs_with_states_count": jobs_with_states_count,
        "top_20_states": top_20_states,
    }


# ============= NEW CACHE ENDPOINTS =============

@app.get("/cache/info")
def cache_info_endpoint(cache_key: str = Query("all")):
    """Get cache information"""
    info = get_cache_info(cache_key)
    return info


@app.post("/cache/refresh")
def cache_refresh_endpoint(cache_key: str = Query("all")):
    """Force refresh cache"""
    clear_cache(cache_key)
    return {"ok": True, "message": f"Cache cleared for '{cache_key}'. Next /jobs request will refresh."}


@app.delete("/cache/clear")
def cache_clear_all_endpoint():
    """Clear all caches"""
    clear_cache()
    return {"ok": True, "message": "All caches cleared"}


@app.get("/refresh/stream")
async def refresh_stream(profile: str = Query("all")):
    """
    Two-wave streaming refresh:
    Wave 1: Fast ATS (greenhouse, lever, ashby, smartrecruiters) - quick results
    Wave 2: Slow ATS (workday) - parallel in background
    """
    import asyncio
    from concurrent.futures import ThreadPoolExecutor, as_completed
    
    FAST_ATS = {"greenhouse", "lever", "ashby", "smartrecruiters"}
    SLOW_ATS = {"workday"}
    
    async def generate():
        companies_cfg = load_profile(profile)
        companies_cfg = [c for c in companies_cfg if c.get("enabled", True) != False]
        
        # Split into waves
        wave1 = [c for c in companies_cfg if c.get("ats", "") in FAST_ATS]
        wave2 = [c for c in companies_cfg if c.get("ats", "") in SLOW_ATS]
        
        total = len(companies_cfg)
        all_jobs = []
        idx = 0
        
        # Send start event
        yield f"data: {json.dumps({'type': 'start', 'total': total, 'wave1': len(wave1), 'wave2': len(wave2)})}\n\n"
        
        # === WAVE 1: Fast ATS (sequential, quick) ===
        yield f"data: {json.dumps({'type': 'wave', 'wave': 1, 'message': 'Fast ATS (Greenhouse, Lever, Ashby)'})}\n\n"
        
        for cfg in wave1:
            company_name = cfg.get("company", "") or cfg.get("name", "")
            yield f"data: {json.dumps({'type': 'loading', 'company': company_name, 'index': idx, 'total': total})}\n\n"
            
            try:
                jobs = _fetch_for_company(profile, cfg)
                all_jobs.extend(jobs)
                yield f"data: {json.dumps({'type': 'ok', 'company': company_name, 'jobs': len(jobs), 'index': idx, 'total': total})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'company': company_name, 'error': str(e)[:100], 'index': idx, 'total': total})}\n\n"
            
            idx += 1
            await asyncio.sleep(0.01)
        
        # Save intermediate cache (Wave 1 complete)
        save_cache(profile, all_jobs)
        
        # Sync wave 1 to pipeline
        sync_result = sync_cache_to_pipeline(all_jobs)
        yield f"data: {json.dumps({'type': 'wave_complete', 'wave': 1, 'jobs': len(all_jobs), 'pipeline_added': sync_result['added']})}"
        yield "\n\n"
        
        # === WAVE 2: Slow ATS (parallel) ===
        if wave2:
            yield f"data: {json.dumps({'type': 'wave', 'wave': 2, 'message': 'Slow ATS (Workday) - parallel'})}\n\n"
            
            def fetch_slow(cfg):
                return cfg, _fetch_for_company(profile, cfg)
            
            # Process in parallel with ThreadPoolExecutor
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor(max_workers=4) as executor:
                futures = {executor.submit(fetch_slow, cfg): cfg for cfg in wave2}
                
                for future in as_completed(futures):
                    cfg = futures[future]
                    company_name = cfg.get("company", "") or cfg.get("name", "")
                    
                    try:
                        _, jobs = future.result()
                        all_jobs.extend(jobs)
                        yield f"data: {json.dumps({'type': 'ok', 'company': company_name, 'jobs': len(jobs), 'index': idx, 'total': total})}\n\n"
                    except Exception as e:
                        yield f"data: {json.dumps({'type': 'error', 'company': company_name, 'error': str(e)[:100], 'index': idx, 'total': total})}\n\n"
                    
                    idx += 1
            
            # Save final cache
            save_cache(profile, all_jobs)
        
        # Save stats
        from utils.cache_manager import save_stats
        role_jobs = [j for j in all_jobs if j.get("role_family") in ["product", "tpm_program", "project"]]
        us_jobs = [j for j in role_jobs if j.get("location_norm", {}).get("state") or j.get("location_norm", {}).get("remote")]
        my_area_jobs = [j for j in us_jobs if j.get("geo_bucket") in ["local", "nc_other", "neighbor", "remote_usa"]]
        save_stats(len(all_jobs), len(role_jobs), len(us_jobs), len(my_area_jobs))
        
        # Sync to pipeline (jobs.json)
        sync_result = sync_cache_to_pipeline(all_jobs)
        yield f"data: {json.dumps({'type': 'sync', 'added': sync_result['added'], 'updated': sync_result['updated']})}"
        yield "\n\n"
        
        # Send complete event
        yield f"data: {json.dumps({'type': 'complete', 'total_jobs': len(all_jobs), 'companies': total, 'pipeline_added': sync_result['added']})}"
        yield "\n\n"
    
    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
        }
    )


@app.get("/stats")
def get_funnel_stats():
    """Get funnel stats from pipeline data (always fresh, daemon updates it)."""
    from datetime import datetime, timezone
    from storage.pipeline_storage import load_new_jobs
    
    # Load pipeline jobs (актуальные данные от daemon)
    jobs = load_new_jobs()
    
    # Pipeline already contains only relevant roles
    total = len(jobs)
    
    # Filter by US location
    us_jobs = [j for j in jobs if _is_us_location(j.get("location", ""))]
    
    # Filter by my area (local states + remote USA)
    my_area_jobs = [j for j in us_jobs if j.get("geo_bucket") in ["local", "nc_other", "neighbor", "remote_usa"]]
    
    return {
        "total": total,
        "role": total,  # Pipeline уже отфильтрован по ролям
        "us": len(us_jobs),
        "my_area": len(my_area_jobs),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "is_stale": False,  # Pipeline всегда актуален (daemon обновляет)
        "age_hours": 0
    }



@app.get("/stats/by-date")
def get_stats_by_date(days: int = Query(14, ge=1, le=60)):
    """Get job statistics grouped by first_seen date from pipeline (unique new jobs only)."""
    from collections import defaultdict
    from datetime import datetime, timedelta
    
    # Load from pipeline (unique jobs only, not cache)
    from storage.job_storage import get_all_jobs
    pipeline_jobs = get_all_jobs()
    if not pipeline_jobs:
        return {"error": "Pipeline not loaded", "dates": []}
    
    # Group by first_seen date (when job was first added to pipeline)
    by_date = defaultdict(lambda: {
        "total": 0,
        "primary": 0,
        "adjacent": 0,
        "unknown": 0,
        "excluded": 0,
        "us": 0,
        "remote": 0,
        "nc": 0,
        "neighbor": 0,
        "nonus": 0
    })
    
    neighbor_states = {"VA", "SC", "GA", "TN"}
    
    for job in pipeline_jobs:
        # Use first_seen date (when job was first discovered/added)
        first_seen = job.get("first_seen", "")
        if first_seen:
            date_str = str(first_seen)[:10]
        else:
            date_str = "unknown"
        
        stats = by_date[date_str]
        stats["total"] += 1
        
        # Category
        cat = job.get("role_category", "unknown")
        if cat in stats:
            stats[cat] += 1
        
        # Location
        ln = job.get("location_norm", {}) or {}
        state = (ln.get("state") or "").upper()
        is_remote = ln.get("remote", False)
        remote_scope = (ln.get("remote_scope") or "").upper()
        
        if is_remote and remote_scope == "USA":
            stats["remote"] += 1
            stats["us"] += 1
        elif state == "NC":
            stats["nc"] += 1
            stats["us"] += 1
        elif state in neighbor_states:
            stats["neighbor"] += 1
            stats["us"] += 1
        elif state:
            stats["us"] += 1
        elif job.get("location"):
            stats["nonus"] += 1
    
    # Sort by date descending, limit to days
    sorted_dates = sorted(by_date.items(), key=lambda x: x[0], reverse=True)
    
    # Filter to recent days only
    result = []
    for date_str, stats in sorted_dates[:days]:
        if date_str == "unknown":
            continue
        result.append({
            "date": date_str,
            **stats
        })
    
    # Get last refresh from cache for display
    cached = load_cache("all", ignore_ttl=True)
    last_refresh = cached.get("last_updated") if cached else None
    
    return {
        "dates": result,
        "last_refresh": last_refresh,
        "source": "pipeline"
    }

# ============= PIPELINE ENDPOINTS =============


@app.get("/cache/browse")
def browse_cache_jobs(
    date: str = Query(None, description="Filter by date (YYYY-MM-DD)"),
    category: str = Query(None, description="Filter by role_category"),
    location: str = Query(None, description="Filter by location (us/nc/neighbor/remote)"),
    page: int = Query(1, ge=1),
    limit: int = Query(100, ge=10, le=500)
):
    """Browse ALL cached jobs with filters (not just pipeline)"""
    from datetime import datetime
    
    cached = load_cache("all", ignore_ttl=True)
    if not cached:
        return {"error": "Cache not loaded", "jobs": [], "total": 0}
    
    all_jobs = cached.get("jobs", [])
    neighbor_states = {"VA", "SC", "GA", "TN"}
    
    # Apply date filter
    if date:
        def match_date(j):
            updated = j.get("updated_at", "")
            if isinstance(updated, int):
                if updated > 10000000000:
                    updated = updated / 1000
                return datetime.fromtimestamp(updated).strftime("%Y-%m-%d") == date
            return str(updated)[:10] == date
        all_jobs = [j for j in all_jobs if match_date(j)]
    
    # Apply category filter
    if category:
        all_jobs = [j for j in all_jobs if j.get("role_category") == category]
    
    # Apply location filter
    if location:
        filtered = []
        for j in all_jobs:
            ln = j.get("location_norm", {}) or {}
            state = (ln.get("state") or "").upper()
            is_remote = ln.get("remote", False)
            
            if location == "us" and (state or is_remote):
                filtered.append(j)
            elif location == "nc" and state == "NC":
                filtered.append(j)
            elif location == "neighbor" and state in neighbor_states:
                filtered.append(j)
            elif location == "remote" and is_remote:
                filtered.append(j)
        all_jobs = filtered
    
    # Check which are in pipeline
    pipeline_ids = get_all_job_ids()
    for j in all_jobs:
        j["in_pipeline"] = j.get("id") in pipeline_ids
    
    # Pagination
    total = len(all_jobs)
    start = (page - 1) * limit
    end = start + limit
    
    return {
        "jobs": all_jobs[start:end],
        "total": total,
        "page": page,
        "has_prev": page > 1,
        "has_next": end < total
    }

@app.get("/pipeline/stats")
def pipeline_stats_endpoint():
    """Get pipeline statistics"""
    return get_job_stats()


@app.get("/jobs/review")
def get_review_jobs(
    date: str = Query(None, description="Filter by date (YYYY-MM-DD)"),
    category: str = Query("unknown", description="unknown / excluded / all"),
    search: str = Query("", description="Search in title, company, location"),
    page: int = Query(1, ge=1, description="Page number (1-indexed)"),
    limit: int = Query(100, ge=10, le=500, description="Jobs per page"),
):
    """
    Get Unknown + Excluded jobs with server-side pagination.
    Much faster than loading all 7500 jobs.
    """
    # Load from cache
    cache_key = "all"
    cached = load_cache(cache_key, ignore_ttl=True)
    
    if not cached:
        return {"error": "Cache not loaded. Run /jobs?refresh=true first.", "jobs": [], "total": 0}
    
    all_jobs = cached.get("jobs", [])

    # Apply date filter
    if date:
        all_jobs = [j for j in all_jobs if str(j.get("updated_at", ""))[:10] == date]
    
    # Filter by role_category (unknown or excluded)
    def get_category(job):
        if job.get("role_category"):
            return job["role_category"]
        if job.get("role_excluded"):
            return "excluded"
        if job.get("role_id"):
            return "primary"
        return "unknown"
    
    if category == "all":
        filtered = [j for j in all_jobs if get_category(j) in ["unknown", "excluded"]]
    else:
        filtered = [j for j in all_jobs if get_category(j) == category]
    
    # Filter by search
    if search:
        search_lower = search.lower()
        filtered = [
            j for j in filtered
            if search_lower in (j.get("title", "") + " " + j.get("company", "") + " " + j.get("location", "")).lower()
        ]
    
    # Pagination
    total = len(filtered)
    total_pages = (total + limit - 1) // limit  # ceiling division
    start = (page - 1) * limit
    end = start + limit
    page_jobs = filtered[start:end]
    
    # Check which jobs are already in pipeline
    pipeline_ids = get_all_job_ids()
    for job in page_jobs:
        job["in_pipeline"] = job.get("id") in pipeline_ids
    
    return {
        "jobs": page_jobs,
        "total": total,
        "page": page,
        "limit": limit,
        "total_pages": total_pages,
        "has_next": page < total_pages,
        "has_prev": page > 1,
    }


@app.get("/pipeline/all")
def pipeline_all_endpoint(
    date: str = Query(None, description="Filter by first_seen date (YYYY-MM-DD)"),
    category: str = Query(None, description="Filter by role_category (primary/adjacent)"),
    location: str = Query(None, description="Filter by location (us/nc/neighbor/remote)")
):
    """Get ALL jobs from storage with optional filters"""
    all_jobs = get_all_jobs()
    
    # Apply date filter (by first_seen - when job was added to pipeline)
    if date:
        all_jobs = [j for j in all_jobs if str(j.get("first_seen", ""))[:10] == date]
    
    # Apply category filter
    if category:
        all_jobs = [j for j in all_jobs if j.get("role_category") == category]
    
    # Apply location filter
    if location:
        neighbor_states = {"VA", "SC", "GA", "TN"}
        filtered = []
        for j in all_jobs:
            ln = j.get("location_norm", {}) or {}
            state = (ln.get("state") or "").upper()
            is_remote = ln.get("remote", False)
            
            if location == "us" and (state or is_remote):
                filtered.append(j)
            elif location == "nc" and state == "NC":
                filtered.append(j)
            elif location == "neighbor" and state in neighbor_states:
                filtered.append(j)
            elif location == "remote" and is_remote:
                filtered.append(j)
        all_jobs = filtered
    
    # Normalize folder_path to use ~/ for portability
    import os
    home = os.path.expanduser("~")
    for j in all_jobs:
        fp = j.get("folder_path", "")
        if fp:
            # Replace any /Users/xxx/ with ~/
            if fp.startswith("/Users/"):
                parts = fp.split("/")
                if len(parts) > 2:
                    j["folder_path"] = "~/" + "/".join(parts[3:])
    
    stats = get_job_stats()
    return {
        "count": len(all_jobs), 
        "jobs": all_jobs,
        "breakdown": stats["status_breakdown"]
    }



@app.get("/pipeline/new")
def pipeline_new_endpoint():
    """Get new (inbox) jobs"""
    jobs = get_jobs_by_status(STATUS_NEW)
    return {"count": len(jobs), "jobs": jobs}


@app.get("/pipeline/active")
def pipeline_active_endpoint():
    """Get active pipeline jobs (Applied, Interview)"""
    jobs = get_jobs_by_statuses({STATUS_APPLIED, STATUS_INTERVIEW, STATUS_CLOSED})
    return {"count": len(jobs), "jobs": jobs}


@app.get("/pipeline/archive")
def pipeline_archive_endpoint():
    """Get archived jobs (Rejected, Offer, Withdrawn)"""
    jobs = get_archive_jobs()
    return {"count": len(jobs), "jobs": jobs}


class PipelineAddJob(BaseModel):
    job: dict


@app.post("/pipeline/add")
def pipeline_add_job_endpoint(payload: PipelineAddJob):
    """
    Manually add a job to pipeline (for Unknown/Excluded jobs).
    """
    job = payload.job.copy()  # Don't modify original
    job_id = job.get("id")
    
    if not job_id:
        return {"ok": False, "error": "Job must have an id"}
    
    # Check if already in pipeline
    existing = get_job_by_id(job_id)
    if existing:
        return {"ok": False, "error": "Job already in pipeline"}
    
    # Mark as manually added
    job["source"] = "manual"
    
    # Add to pipeline
    added = add_job(job)
    
    if added:
        return {"ok": True, "job": job}
    else:
        return {"ok": False, "error": "Job already exists"}


@app.delete("/pipeline/remove/{job_id}")
def pipeline_remove_job_endpoint(job_id: str):
    """
    Remove a job from pipeline (for manual jobs).
    Only removes jobs with source='manual'.
    """
    job = get_job_by_id(job_id)
    
    if not job:
        return {"ok": False, "error": "Job not found"}
    
    # Only allow removing manual jobs
    if job.get("source") != "manual":
        return {"ok": False, "error": "Can only remove manually added jobs"}
    
    # Remove from jobs_new.json
    try:
        jobs_new_path = Path("data/jobs_new.json")
        with open(jobs_new_path, "r") as f:
            jobs = json.load(f)
        
        original_len = len(jobs)
        jobs = [j for j in jobs if j.get("id") != job_id]
        
        if len(jobs) == original_len:
            return {"ok": False, "error": "Job not found in storage"}
        
        with open(jobs_new_path, "w") as f:
            json.dump(jobs, f, indent=2)
        
        return {"ok": True, "removed": job_id}
    except Exception as e:
        return {"ok": False, "error": str(e)}


class PipelineStatusUpdate(BaseModel):
    job_id: str
    status: str
    notes: str = ""
    folder_path: str = ""


@app.post("/pipeline/status")
def pipeline_status_update_endpoint(payload: PipelineStatusUpdate):
    """
    Update job status in pipeline.
    Valid statuses: New, Selected, Ready, Applied, Interview, Offer, Rejected, Withdrawn, Closed
    """
    # Accept both capitalized and lowercase status values
    status_map = {
        "new": "new", "New": "new",
        "selected": "Selected", "Selected": "Selected",
        "ready": "Ready", "Ready": "Ready",
        "applied": "applied", "Applied": "applied",
        "interview": "interview", "Interview": "interview",
        "offer": "offer", "Offer": "offer",
        "rejected": "rejected", "Rejected": "rejected",
        "withdrawn": "withdrawn", "Withdrawn": "withdrawn",
        "closed": "closed", "Closed": "closed",
    }
    
    normalized_status = status_map.get(payload.status)
    if not normalized_status:
        return {"ok": False, "error": f"Invalid status: {payload.status}"}
    
    job = job_update_status(payload.job_id, normalized_status, payload.notes, payload.folder_path)
    
    if job:
        # Auto-parse JD when status changes to Selected and JD not yet parsed
        if normalized_status == "Selected" and not job.get("jd_summary"):
            try:
                from parsers.jd_parser import parse_and_store_jd, has_jd
                job_url = job.get("job_url") or job.get("url") or ""
                if job_url and not has_jd(payload.job_id):
                    print(f"[Pipeline] Auto-parsing JD for {job.get('company')} - {job.get('title')}")
                    result = parse_and_store_jd(
                        job_id=payload.job_id,
                        url=job_url,
                        title=job.get("title", ""),
                        company=job.get("company", ""),
                        ats=job.get("ats", "greenhouse")
                    )
                    if result.get("ok") and result.get("summary"):
                        job["jd_summary"] = result["summary"]
                        # Save updated job
                        job_update_status(payload.job_id, normalized_status, payload.notes, payload.folder_path, jd_summary=result["summary"])
            except Exception as e:
                print(f"[Pipeline] JD parsing failed: {e}")
        
        return {"ok": True, "job": job}
    else:
        return {"ok": False, "error": "Job not found"}


@app.get("/pipeline/job/{job_id}")
def pipeline_get_job_endpoint(job_id: str):
    """Get job by ID from any storage"""
    job = get_job_by_id(job_id)
    if job:
        return {"ok": True, "job": job}
    else:
        return {"ok": False, "error": "Job not found"}


class ParseJDRequest(BaseModel):
    job_id: str
    url: str
    title: str
    company: str
    ats: str = "greenhouse"


@app.post("/jd/parse")
def parse_jd_endpoint(payload: ParseJDRequest):
    """
    Parse job description from URL and extract structured summary.
    Saves full text to data/jd/{job_id}.txt and returns summary.
    """
    try:
        from parsers.jd_parser import parse_and_store_jd
        
        result = parse_and_store_jd(
            job_id=payload.job_id,
            url=payload.url,
            title=payload.title,
            company=payload.company,
            ats=payload.ats
        )
        
        if result.get("ok"):
            # Initialize variables
            match_score = None
            analysis = None

            # Update job in storage with jd_summary
            job = get_job_by_id(payload.job_id)
            if job:
                # Use storage function to save jd_summary
                from storage.job_storage import update_jd_summary, _load_jobs, _save_jobs
                update_jd_summary(payload.job_id, result["summary"])

                jd_text = result.get("jd_text", "")
                if jd_text and len(jd_text) > 100:
                    try:
                        from api.prepare_application import analyze_job_with_ai
                        role_family = job.get("role_family", "tpm_program")
                        analysis = analyze_job_with_ai(payload.title, payload.company, jd_text, role_family)
                        if analysis and "match_score" in analysis:
                            match_score = analysis["match_score"]
                            # Save match_score to job
                            jobs = _load_jobs()
                            for j in jobs:
                                if j.get("id") == payload.job_id:
                                    j["match_score"] = match_score
                                    j["analysis"] = analysis
                                    break
                            _save_jobs(jobs)
                    except Exception as e:
                        print(f"AI analysis error: {e}")

            return {
                "ok": True,
                "summary": result["summary"],
                "jd_preview": result.get("jd_text", "")[:500] + "...",
                "match_score": match_score,
                "analysis": analysis
            }
        else:
            return {"ok": False, "error": result.get("error", "Unknown error")}
            
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/jd/{job_id}")
def get_jd_endpoint(job_id: str):
    """Get stored JD for a job"""
    try:
        from parsers.jd_parser import get_stored_jd, has_jd
        
        if not has_jd(job_id):
            return {"ok": False, "error": "JD not found"}
        
        jd_text = get_stored_jd(job_id)
        job = get_job_by_id(job_id)
        
        return {
            "ok": True,
            "jd_text": jd_text,
            "summary": job.get("jd_summary") if job else None
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.get("/job/find-by-url")
def find_job_by_url_endpoint(url: str = Query(..., description="Job URL to search")):
    """
    Search for a job by URL in all storages (pipeline, cache, etc.)
    Returns job if found, or not_found status.
    """
    url = url.strip()
    if not url:
        return {"ok": False, "error": "URL is required"}
    
    # Normalize URL for comparison
    from urllib.parse import urlparse, parse_qs
    parsed = urlparse(url)
    
    # Extract job ID from URL based on ATS patterns
    job_id = None
    qs = parse_qs(parsed.query)
    
    # Greenhouse: gh_jid parameter, token parameter, or /jobs/NUMBER
    if "gh_jid" in url:
        job_id = qs.get("gh_jid", [None])[0]
    elif "token" in qs:
        # Greenhouse embed format: ?token=7404427&for=coinbase
        job_id = qs.get("token", [None])[0]
    elif "/jobs/" in parsed.path:
        # Extract number from path like /jobs/12345
        import re
        match = re.search(r'/jobs/(\d+)', parsed.path)
        if match:
            job_id = match.group(1)
    
    # Lever: last segment of path
    if "lever.co" in parsed.netloc:
        job_id = parsed.path.rstrip('/').split('/')[-1]
    
    # SmartRecruiters: last segment
    if "smartrecruiters.com" in parsed.netloc:
        job_id = parsed.path.rstrip('/').split('/')[-1]
    
    # Search in all jobs
    all_jobs = get_all_jobs()
    
    for job in all_jobs:
        job_url = job.get("job_url", "") or job.get("url", "")
        ats_job_id = str(job.get("ats_job_id", ""))
        
        # Skip jobs without URL
        if not job_url:
            continue
        
        # Match by exact URL
        if url in job_url or job_url in url:
            return {"ok": True, "found": True, "job": job}
        
        # Match by job ID
        if job_id and ats_job_id == job_id:
            return {"ok": True, "found": True, "job": job}
    
    return {"ok": True, "found": False, "message": "Job not found in database"}


@app.get("/pipeline/attention")
def pipeline_attention_endpoint():
    """Get jobs that need attention (Closed, etc.)"""
    all_jobs = get_all_jobs()
    attention = [j for j in all_jobs if j.get("needs_attention")]
    return {"count": len(attention), "jobs": attention}


# ============= SYNC DEV->PROD ENDPOINT =============

@app.post("/sync-to-prod")
def sync_to_prod_endpoint():
    """
    Sync data from DEV to PROD.
    Only available in DEV environment.
    """
    if ENV != "DEV":
        return {"ok": False, "error": "Only available in DEV environment"}
    
    try:
        from sync_to_prod import sync_companies, sync_jobs
        
        companies_result = sync_companies()
        jobs_result = sync_jobs()
        
        return {
            "ok": True,
            "companies": companies_result,
            "jobs": jobs_result,
            "synced_at": datetime.now(timezone.utc).isoformat()
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ============= ONBOARDING ENDPOINTS =============

import re
from urllib.parse import urlparse


def detect_ats_from_url(url: str) -> dict:
    """
    Detect ATS type and extract company info from job URL.
    Returns: {ats, company, board_url, job_id} or {error}
    """
    url = url.strip()
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    path = parsed.path
    
    # Greenhouse: boards.greenhouse.io/company/jobs/123 or company.com/jobs?gh_jid=123
    if "greenhouse.io" in host:
        # https://boards.greenhouse.io/stripe/jobs/7294977
        match = re.match(r"/([^/]+)/jobs/(\d+)", path)
        if match:
            company = match.group(1)
            job_id = match.group(2)
            return {
                "ats": "greenhouse",
                "company": company.replace("-", " ").title(),
                "company_slug": company,
                "board_url": f"https://boards.greenhouse.io/{company}",
                "job_id": job_id,
                "job_api_url": f"https://boards-api.greenhouse.io/v1/boards/{company}/jobs/{job_id}"
            }
    
    if "gh_jid" in url:
        # https://stripe.com/jobs/search?gh_jid=7294977
        match = re.search(r"gh_jid=(\d+)", url)
        if match:
            job_id = match.group(1)
            # Extract company from domain
            company = host.replace("www.", "").split(".")[0]
            return {
                "ats": "greenhouse",
                "company": company.title(),
                "company_slug": company,
                "board_url": f"https://boards.greenhouse.io/{company}",
                "job_id": job_id,
                "job_api_url": f"https://boards-api.greenhouse.io/v1/boards/{company}/jobs/{job_id}"
            }
    
    # Workday: company.wd1.myworkdayjobs.com/site/job/..._JOBID
    if "myworkdayjobs.com" in host:
        # https://hitachi.wd1.myworkdayjobs.com/en-US/hitachi/job/..._R0082977
        # https://proofpoint.wd5.myworkdayjobs.com/ProofpointCareers/job/...
        parts = host.split(".")
        company = parts[0] if parts else ""

        # Extract wd number (wd1, wd5, etc.) from host
        wd_match = re.search(r"\.(wd\d+)\.", host)
        wd_num = wd_match.group(1) if wd_match else "wd1"

        # Extract job ID from path (usually ends with _JOBID)
        job_id_match = re.search(r"_([A-Z0-9]+-?[0-9]*)$", path)
        job_id = job_id_match.group(1) if job_id_match else ""

        # Extract site name from path
        site_match = re.search(r"myworkdayjobs\.com/(?:[a-z][a-z]-[A-Z][A-Z]/)?([^/]+)", url)
        site = site_match.group(1) if site_match else company

        return {
            "ats": "workday",
            "company": company.title(),
            "company_slug": company,
            "wd_num": wd_num,
            "board_url": f"https://{company}.{wd_num}.myworkdayjobs.com/{site}",
            "job_id": job_id,
            "job_path": path,
            "job_url": url
        }
    
    # Lever: jobs.lever.co/company/job-uuid
    if "lever.co" in host:
        match = re.match(r"/([^/]+)/([a-f0-9-]+)", path)
        if match:
            company = match.group(1)
            job_id = match.group(2)
            return {
                "ats": "lever",
                "company": company.replace("-", " ").title(),
                "company_slug": company,
                "board_url": f"https://jobs.lever.co/{company}",
                "job_id": job_id,
                "job_api_url": f"https://api.lever.co/v0/postings/{company}/{job_id}"
            }
    
    # SmartRecruiters: jobs.smartrecruiters.com/Company/job-id
    if "smartrecruiters.com" in host:
        match = re.match(r"/([^/]+)/([^/]+)", path)
        if match:
            company = match.group(1)
            job_id = match.group(2)
            return {
                "ats": "smartrecruiters",
                "company": company.replace("-", " ").title(),
                "company_slug": company,
                "board_url": f"https://jobs.smartrecruiters.com/{company}",
                "job_id": job_id
            }
    
    # Ashby: jobs.ashbyhq.com/company/job-uuid
    if "ashbyhq.com" in host:
        match = re.match(r"/([^/]+)/([a-f0-9-]+)", path)
        if match:
            company = match.group(1)
            job_id = match.group(2)
            return {
                "ats": "ashby",
                "company": company.replace("-", " ").title(),
                "company_slug": company,
                "board_url": f"https://jobs.ashbyhq.com/{company}",
                "job_id": job_id
            }
    
    # Jibe (Google Hire): company.jibeapply.com/jobs/slug
    if "jibeapply.com" in host:
        company_slug = host.split(".")[0]
        job_id_match = re.match(r"/jobs/(\d+)", path)
        job_id = job_id_match.group(1) if job_id_match else ""
        return {
            "ats": "jibe",
            "company": company_slug.replace("-", " ").title(),
            "company_slug": company_slug,
            "board_url": f"https://{company_slug}.jibeapply.com/jobs",
            "job_id": job_id,
            "job_url": url
        }

    # iCIMS: external-company.icims.com/jobs/12345/title/job
    if "icims.com" in host:
        # Extract company from subdomain: external-firstcitizens.icims.com -> firstcitizens
        subdomain = host.split(".")[0]  # external-firstcitizens
        company_slug = subdomain.replace("external-", "").replace("careers-", "")
        
        # Extract job ID from path: /jobs/32378/title/job
        job_id_match = re.match(r"/jobs/(\d+)", path)
        job_id = job_id_match.group(1) if job_id_match else ""
        
        # Extract title from path
        title_match = re.match(r"/jobs/\d+/([^/]+)/job", path)
        title_slug = title_match.group(1) if title_match else ""
        
        return {
            "ats": "icims",
            "company": company_slug.replace("-", " ").title(),
            "company_slug": company_slug,
            "board_url": f"https://{subdomain}.icims.com/jobs",
            "job_id": job_id,
            "title_slug": title_slug,
            "job_url": url
        }
    
    # Unknown ATS - use universal parser
    # Extract company from domain (skip common prefixes like apply, careers, jobs)
    parts = host.replace("www.", "").split(".")
    skip_prefixes = {"apply", "careers", "jobs", "job", "hire", "recruiting", "talent"}
    company = parts[0]
    if company.lower() in skip_prefixes and len(parts) > 1:
        company = parts[1]
    return {
        "ats": "universal",
        "company": company.title(),
        "company_slug": company.lower(),
        "board_url": f"https://{host}",
        "job_url": url
    }


def fetch_single_job(ats_info: dict) -> dict:
    """
    Fetch single job details from ATS.
    Returns job dict or {error}
    """
    import requests
    
    ats = ats_info.get("ats")
    
    if ats == "greenhouse":
        api_url = ats_info.get("job_api_url")
        try:
            resp = requests.get(api_url, timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                return {
                    "title": data.get("title", ""),
                    "location": data.get("location", {}).get("name", ""),
                    "url": data.get("absolute_url", ""),
                    "updated_at": data.get("updated_at", ""),
                    "ats_job_id": str(data.get("id", "")),
                }
            else:
                return {"error": f"Greenhouse API returned {resp.status_code}"}
        except Exception as e:
            return {"error": str(e)}
    
    elif ats == "workday":
        # Use direct job API endpoint
        company_slug = ats_info.get("company_slug")
        board_url = ats_info.get("board_url")
        job_url = ats_info.get("job_url", "")
        wd_num = ats_info.get("wd_num", "wd1")  # wd1, wd5, etc.

        # Extract site and job path from URL
        # URL format: https://company.wd5.myworkdayjobs.com/site/job/location/title_JOBID
        site_match = re.search(r"myworkdayjobs\.com/(?:[a-z][a-z]-[A-Z][A-Z]/)?([^/]+)", board_url or job_url)
        site = site_match.group(1) if site_match else company_slug

        # Extract job path from original URL
        path_match = re.search(r"/job/(.+?)(?:\?|$)", job_url)
        job_path = path_match.group(1) if path_match else ""

        if job_path:
            # Direct API call for specific job
            api_url = f"https://{company_slug}.{wd_num}.myworkdayjobs.com/wday/cxs/{company_slug}/{site}/job/{job_path}"
            try:
                resp = requests.get(api_url, headers={"Accept": "application/json"}, timeout=15)
                if resp.status_code == 200:
                    data = resp.json()
                    job_info = data.get("jobPostingInfo", {})
                    return {
                        "title": job_info.get("title", ""),
                        "location": job_info.get("location", ""),
                        "url": job_url,
                        "ats_job_id": ats_info.get("job_id", ""),
                        "description": job_info.get("jobDescription", ""),
                    }
            except Exception as e:
                print(f"[Workday] Direct API error: {e}")

        # Fallback: search by job ID
        job_id = ats_info.get("job_id", "")
        search_url = f"https://{company_slug}.{wd_num}.myworkdayjobs.com/wday/cxs/{company_slug}/{site}/jobs"
        try:
            resp = requests.post(
                search_url,
                json={"appliedFacets": {}, "limit": 20, "offset": 0, "searchText": job_id},
                headers={"Content-Type": "application/json"},
                timeout=15
            )
            if resp.status_code == 200:
                data = resp.json()
                postings = data.get("jobPostings", [])

                # Find exact match by job ID in externalPath
                for posting in postings:
                    ext_path = posting.get("externalPath", "")
                    if job_id and job_id in ext_path:
                        return {
                            "title": posting.get("title", ""),
                            "location": posting.get("locationsText", ""),
                            "url": f"https://{company_slug}.{wd_num}.myworkdayjobs.com{ext_path}",
                            "ats_job_id": job_id,
                        }

                return {"error": f"Job {job_id} not found in search results"}
            else:
                return {"error": f"Workday API returned {resp.status_code}"}
        except Exception as e:
            return {"error": str(e)}
    
    elif ats == "icims":
        # iCIMS - extract from URL/title slug since no public API
        job_url = ats_info.get("job_url", "")
        job_id = ats_info.get("job_id", "")
        title_slug = ats_info.get("title_slug", "")
        
        # Decode title from URL slug
        title = title_slug.replace("-", " ").title() if title_slug else ""
        # Clean up common URL encoding
        title = title.replace("%26", "&").replace("%2f", "/")
        
        return {
            "title": title,
            "location": "",  # Would need to scrape page for location
            "url": job_url,
            "ats_job_id": job_id,
        }
    
    elif ats == "lever":
        api_url = ats_info.get("job_api_url")
        try:
            resp = requests.get(api_url, timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                return {
                    "title": data.get("text", ""),
                    "location": data.get("categories", {}).get("location", ""),
                    "url": data.get("hostedUrl", ""),
                    "ats_job_id": data.get("id", ""),
                }
            else:
                return {"error": f"Lever API returned {resp.status_code}"}
        except Exception as e:
            return {"error": str(e)}
    
    elif ats == "ashby":
        # Ashby API - fetch job details
        company_slug = ats_info.get("company_slug", "")
        job_id = ats_info.get("job_id", "")
        job_url = ats_info.get("job_url", "") or f"https://jobs.ashbyhq.com/{company_slug}/{job_id}"

        # Ashby has a GraphQL API, but we can also scrape the page
        # Try API first
        api_url = f"https://jobs.ashbyhq.com/api/non-user-graphql?op=ApiJobPosting"
        try:
            payload = {
                "operationName": "ApiJobPosting",
                "variables": {
                    "organizationHostedJobsPageName": company_slug,
                    "jobPostingId": job_id
                },
                "query": """
                    query ApiJobPosting($organizationHostedJobsPageName: String!, $jobPostingId: String!) {
                        jobPosting(organizationHostedJobsPageName: $organizationHostedJobsPageName, jobPostingId: $jobPostingId) {
                            id
                            title
                            locationName
                            descriptionHtml
                            publishedDate
                        }
                    }
                """
            }
            resp = requests.post(api_url, json=payload, timeout=15)
            if resp.status_code == 200:
                data = resp.json()
                job_posting = data.get("data", {}).get("jobPosting", {})
                if job_posting:
                    from bs4 import BeautifulSoup
                    desc_html = job_posting.get("descriptionHtml", "")
                    desc_text = BeautifulSoup(desc_html, "html.parser").get_text(separator="\n", strip=True) if desc_html else ""

                    return {
                        "title": job_posting.get("title", ""),
                        "location": job_posting.get("locationName", ""),
                        "url": job_url,
                        "ats_job_id": job_id,
                        "description": desc_text,
                        "updated_at": job_posting.get("publishedDate", ""),
                    }
        except Exception as e:
            print(f"[Ashby] API error: {e}")

        # Fallback: return basic info from URL
        return {
            "title": "",
            "location": "",
            "url": job_url,
            "ats_job_id": job_id,
        }

    elif ats == "universal":
        # Use Playwright-based universal parser
        try:
            from parsers.universal import extract_job_details
            result = extract_job_details(ats_info.get("job_url"))
            if result.get("error"):
                return {"error": result["error"]}
            return {
                "title": result.get("title", ""),
                "location": result.get("location", ""),
                "url": result.get("url", ats_info.get("job_url")),
                "description": result.get("description", ""),
                "salary": result.get("salary", ""),
            }
        except Exception as e:
            return {"error": f"Universal parser error: {str(e)}"}

    return {"error": f"Fetch not implemented for ATS: {ats}"}


class OnboardRequest(BaseModel):
    url: str
    match_score: Optional[int] = None
    analysis: Optional[dict] = None


@app.post("/onboard")
def onboard_job(payload: OnboardRequest):
    """
    Onboard a new job by URL.
    1. Detect ATS and extract company info
    2. Add company if new
    3. Fetch single job details
    4. Classify role
    5. Add to pipeline if relevant
    """
    url = payload.url.strip()
    
    # 1. Detect ATS
    ats_info = detect_ats_from_url(url)
    if "error" in ats_info:
        return {"ok": False, "error": ats_info["error"]}
    
    ats = ats_info["ats"]
    company_name = ats_info["company"]
    company_slug = ats_info["company_slug"]
    board_url = ats_info["board_url"]
    
    # Fix company name using AI if it looks like a URL slug
    if ats == "universal" and (company_name.lower() == company_slug or "careers" in company_name.lower()):
        try:
            from utils.ollama_ai import fix_company_name, is_ollama_available
            if is_ollama_available():
                fixed_name = fix_company_name(company_name, board_url)
                if fixed_name and fixed_name != company_name:
                    company_name = fixed_name
        except Exception as e:
            print(f"AI company name fix error: {e}")
    
    # 2. Check if company exists
    companies_path = Path("data/companies.json")
    companies = json.load(open(companies_path)) if companies_path.exists() else []
    
    company_exists = any(
        c.get("id") == company_slug or c.get("name", "").lower() == company_name.lower()
        for c in companies
    )
    
    new_company = None
    company_parsing_started = False
    total_jobs_found = 0

    if not company_exists:
        # Check if ATS is supported for automatic parsing
        ats_supported = ats in SUPPORTED_ATS

        new_company = {
            "id": company_slug,
            "name": company_name,
            "ats": ats,
            "board_url": board_url,
            "industry": "",
            "tags": [],
            "priority": 0,
            "hq_state": None,
            "region": "global",
            "enabled": ats_supported  # Only enable for supported ATS
        }
        companies.append(new_company)
        with open(companies_path, "w") as f:
            json.dump(companies, f, indent=2, ensure_ascii=False)

        # If ATS is supported, trigger initial parsing of all company jobs
        if ats_supported:
            try:
                print(f"[Onboard] New company with supported ATS ({ats}), starting initial parse...")
                # Use refresh_company_sync which handles all the parsing logic
                fetch_result = refresh_company_sync(new_company)
                if fetch_result.get("ok"):
                    company_parsing_started = True
                    total_jobs_found = fetch_result.get("jobs", 0)
                    print(f"[Onboard] Initial parse complete: {total_jobs_found} jobs found, {fetch_result.get('jobs_added', 0)} added to pipeline")
                else:
                    print(f"[Onboard] Initial parse failed: {fetch_result.get('error')}")
            except Exception as e:
                print(f"[Onboard] Initial parse error: {e}")
        else:
            # Register unsupported ATS for future implementation
            try:
                from tools.ats_discovery import register_unsupported_ats
                register_unsupported_ats(
                    ats_name=ats,
                    discovery_result={"api_endpoints": [], "json_responses": []},
                    company_url=url
                )
                print(f"[Onboard] Registered unsupported ATS '{ats}' for future implementation")
            except Exception as e:
                print(f"[Onboard] Could not register unsupported ATS: {e}")

    # 3. Fetch single job
    job_data = fetch_single_job(ats_info)
    if "error" in job_data:
        return {
            "ok": False, 
            "error": job_data["error"],
            "company": {"name": company_name, "new": new_company is not None}
        }
    
    # 4. Build full job object
    job = {
        "company": company_name,
        "ats": ats,
        "ats_job_id": job_data.get("ats_job_id", ats_info.get("job_id", "")),
        "title": job_data.get("title", ""),
        "location": job_data.get("location", ""),
        "job_url": job_data.get("url", url),
        "updated_at": job_data.get("updated_at", datetime.now(timezone.utc).isoformat()),
    }
    
    # Generate ID
    job["id"] = generate_job_id(job)
    
    # Normalize location
    loc_norm = normalize_location(job.get("location"))
    job["location_norm"] = loc_norm
    
    # Classify role (rule-based first, then AI fallback)
    role = classify_role(job.get("title"), job_data.get("description", ""))
    
    # If rule-based classification failed, try AI
    if role.get("role_family") == "other" and role.get("confidence", 0) < 60:
        try:
            from utils.ollama_ai import classify_role_ai, is_ollama_available
            if is_ollama_available():
                ai_role = classify_role_ai(job.get("title"), job_data.get("description", ""))
                if ai_role.get("confidence", 0) > role.get("confidence", 0):
                    role = ai_role
                    role["role_category"] = "ai_classified"
        except Exception as e:
            print(f"AI classification error: {e}")
    
    job["role_family"] = role.get("role_family")
    job["role_category"] = role.get("role_category")
    job["role_id"] = role.get("role_id")
    job["role_confidence"] = role.get("confidence")
    job["role_reason"] = role.get("reason")
    job["role_excluded"] = role.get("excluded", False)
    
    # Geo scoring
    bucket, score = compute_geo_bucket_and_score(loc_norm)
    job["geo_bucket"] = bucket
    job["geo_score"] = score
    
    # Company data
    job["company_data"] = {
        "priority": 0,
        "hq_state": None,
        "region": "global",
        "tags": []
    }
    
    job["source"] = "onboard"

    # Add match_score and analysis if provided (from pre-analysis)
    print(f"[Onboard] Received match_score: {payload.match_score}, has analysis: {payload.analysis is not None}")
    if payload.match_score is not None:
        job["match_score"] = payload.match_score
        print(f"[Onboard] Set match_score to {payload.match_score}")
    if payload.analysis is not None:
        job["analysis"] = payload.analysis

    # If no analysis provided, run AI analysis automatically
    if payload.analysis is None and job_data.get("description"):
        try:
            from api.prepare_application import analyze_job_with_ai
            jd_text = job_data.get("description", "")
            if jd_text and len(jd_text) > 100:
                role_family = job.get("role_family", "tpm_program")
                analysis = analyze_job_with_ai(job.get("title"), company_name, jd_text, role_family)
                if analysis and "match_score" in analysis:
                    job["match_score"] = analysis["match_score"]
                    job["analysis"] = analysis
                    print(f"[Onboard] AI analysis complete: {analysis.get('match_score')}% match")
        except Exception as e:
            print(f"[Onboard] AI analysis error: {e}")

    # 5. Add to pipeline - always add manual jobs (user explicitly added them)
    added_to_pipeline = False
    existing = get_job_by_id(job["id"])
    if not existing:
        add_job(job)
        added_to_pipeline = True
    
    return {
        "ok": True,
        "company": {
            "name": company_name,
            "new": new_company is not None,
            "ats": ats,
            "ats_supported": ats in SUPPORTED_ATS,
            "parsing_started": company_parsing_started,
            "total_jobs_found": total_jobs_found
        },
        "job": {
            "id": job["id"],
            "title": job["title"],
            "location": job["location"],
            "url": job["job_url"]
        },
        "classification": {
            "role_family": job.get("role_family"),
            "role_id": job.get("role_id"),
            "confidence": job.get("role_confidence"),
            "reason": job.get("role_reason")
        },
        "geo": {
            "bucket": bucket,
            "score": score
        },
        "added_to_pipeline": added_to_pipeline
    }


# ============= ANALYZE JOB BEFORE ADD =============

class AnalyzeJobUrlRequest(BaseModel):
    url: str

class ClearAnalysisCacheRequest(BaseModel):
    url: Optional[str] = None  # If None, clears all cache

# Cache for job analysis results (in-memory, clears on restart)
_analysis_cache = {}

@app.post("/analyze-job-url")
async def analyze_job_url_endpoint(payload: AnalyzeJobUrlRequest):
    """
    Analyze a job URL before adding to pipeline.
    Returns match score and recommendation.
    Results are cached to ensure consistent responses.
    """
    from api.prepare_application import analyze_job_with_ai
    import hashlib

    url = payload.url.strip()

    # Detect thank-you/confirmation/application pages - these are NOT job postings
    url_lower = url.lower()
    confirmation_indicators = [
        "/thank-you", "/thankyou", "/thanks", "/confirmation",
        "/success", "/applied", "/submitted", "/complete",
        "application_id=", "confirmation_id=", "applied=true",
        "mode=submit_apply", "mode=apply", "/apply"
    ]
    if any(indicator in url_lower for indicator in confirmation_indicators):
        # Try to extract job URL without apply mode
        clean_url = url
        if "mode=submit_apply" in url_lower or "mode=apply" in url_lower:
            clean_url = url.split("?")[0]  # Remove query params
            if "/apply" in clean_url.lower():
                clean_url = clean_url.replace("/apply", "")

        return {
            "ok": False,
            "error": "⚠️ This appears to be an application form page, not the job listing. Try using the job description URL instead.",
            "is_confirmation_page": True,
            "suggested_url": clean_url if clean_url != url else None,
            "url": url
        }

    # Normalize URL for cache key (remove tracking params)
    cache_url = url.split('?')[0].lower().rstrip('/')
    cache_key = hashlib.md5(cache_url.encode()).hexdigest()

    # Check cache first
    if cache_key in _analysis_cache:
        print(f"[AnalyzeJobUrl] Cache hit for {cache_url[:50]}...")
        return _analysis_cache[cache_key]

    # 1. Parse URL to get job data
    job_data = None
    ats = None
    
    # Try different parsers
    if "greenhouse" in url.lower() or "boards.greenhouse.io" in url:
        from parsers.greenhouse import parse_jobs_api
        ats = "greenhouse"
        # Extract company from URL
        if "boards.greenhouse.io" in url:
            parts = url.split("boards.greenhouse.io/")[1].split("/")
            company_slug = parts[0]
            job_data = {"board_url": f"https://boards.greenhouse.io/{company_slug}"}
    elif "lever.co" in url.lower():
        ats = "lever"
    elif "smartrecruiters" in url.lower():
        ats = "smartrecruiters"
    elif "myworkdayjobs" in url.lower() or "workday" in url.lower():
        ats = "workday"
    elif "ashby" in url.lower():
        ats = "ashby"
    else:
        ats = "universal"
    
    # 2. Fetch JD
    jd = ""
    title = ""
    company = ""
    location = ""
    
    try:
        import requests
        from bs4 import BeautifulSoup
        import re
        
        # Special handling for Greenhouse embedded (gh_jid in URL) - includes Epic Games, etc.
        gh_jid_match = re.search(r'gh_jid=(\d+)', url)
        if gh_jid_match:
            job_id = gh_jid_match.group(1)
            
            # Map known companies to their Greenhouse board slug
            company_slug = None
            if "epicgames" in url.lower():
                company_slug = "epicgames"
                company = "Epic Games"
            elif "stripe" in url.lower():
                company_slug = "stripe"
                company = "Stripe"
            elif "figma" in url.lower():
                company_slug = "figma"
                company = "Figma"
            elif "notion" in url.lower():
                company_slug = "notion"
                company = "Notion"
            elif "discord" in url.lower():
                company_slug = "discord"
                company = "Discord"
            elif "airbnb" in url.lower():
                company_slug = "airbnb"
                company = "Airbnb"
            # Try extracting from URL path for unknown companies
            else:
                # Try to get company from the domain
                import urllib.parse
                parsed = urllib.parse.urlparse(url)
                domain_parts = parsed.netloc.replace("www.", "").split(".")
                if domain_parts:
                    company_slug = domain_parts[0].lower()
                    company = domain_parts[0].replace("-", " ").title()
            
            # Fetch from Greenhouse API
            if job_id and company_slug:
                api_url = f"https://boards-api.greenhouse.io/v1/boards/{company_slug}/jobs/{job_id}"
                print(f"[AnalyzeJobUrl] Trying Greenhouse API: {api_url}")
                resp = requests.get(api_url, timeout=15)
                if resp.ok:
                    data = resp.json()
                    title = data.get("title", "")
                    if not company:
                        company = data.get("company", {}).get("name", company_slug.replace("-", " ").title())
                    loc_data = data.get("location", {})
                    location = loc_data.get("name", "") if isinstance(loc_data, dict) else str(loc_data)
                    
                    # Get full JD from content field
                    content = data.get("content", "")
                    if content:
                        soup = BeautifulSoup(content, "html.parser")
                        jd = soup.get_text(separator="\n", strip=True)
                    
                    ats = "greenhouse"
                    print(f"[AnalyzeJobUrl] Greenhouse API success: {title} at {company}, JD={len(jd)} chars")
                else:
                    print(f"[AnalyzeJobUrl] Greenhouse API failed: {resp.status_code}")
        
        # Standard Greenhouse URL (boards.greenhouse.io)
        if not jd and "boards.greenhouse.io" in url:
            parts = url.split("boards.greenhouse.io/")[1].split("/")
            company_slug = parts[0]
            job_id = None
            if len(parts) > 2:
                job_id = parts[2].split("?")[0]
            
            if job_id:
                api_url = f"https://boards-api.greenhouse.io/v1/boards/{company_slug}/jobs/{job_id}"
                resp = requests.get(api_url, timeout=15)
                if resp.ok:
                    data = resp.json()
                    title = data.get("title", "")
                    company = data.get("company", {}).get("name", company_slug.replace("-", " ").title())
                    loc_data = data.get("location", {})
                    location = loc_data.get("name", "") if isinstance(loc_data, dict) else str(loc_data)
                    content = data.get("content", "")
                    if content:
                        soup = BeautifulSoup(content, "html.parser")
                        jd = soup.get_text(separator="\n", strip=True)
                    ats = "greenhouse"
        
        # Special handling for Workday - use their API
        if not jd and "myworkdayjobs" in url:
            # Parse URL: https://company.wd1.myworkdayjobs.com/site/job/location/title_JOBID
            import re
            # Extract: company, site, and the job path (location/title_ID)
            match = re.match(r"https://([^.]+)\.wd\d\.myworkdayjobs\.com/(?:[a-z]{2}-[A-Z]{2}/)?([^/]+)/job/(.+)", url)
            if match:
                company_slug = match.group(1)
                site = match.group(2)
                job_path = match.group(3).split("?")[0]  # Remove query params
                
                # Fetch job details from Workday API - need full path
                api_url = f"https://{company_slug}.wd1.myworkdayjobs.com/wday/cxs/{company_slug}/{site}/job/{job_path}"
                headers = {"Content-Type": "application/json", "Accept": "application/json"}
                resp = requests.get(api_url, headers=headers, timeout=15)
                
                if resp.ok:
                    data = resp.json()
                    job_posting = data.get("jobPostingInfo", {})
                    title = job_posting.get("title", "")
                    company = job_posting.get("company", company_slug.replace("-", " ").title())
                    location = job_posting.get("location", "")
                    jd = job_posting.get("jobDescription", "")
                    
                    # Clean HTML from JD
                    if jd:
                        soup_jd = BeautifulSoup(jd, "html.parser")
                        jd = soup_jd.get_text(separator="\n", strip=True)
        
        # Standard HTML parsing for other ATS
        if not jd:
            headers = {"User-Agent": "Mozilla/5.0"}
            resp = requests.get(url, headers=headers, timeout=15)
            html_text = resp.text
            soup = BeautifulSoup(html_text, "html.parser")
            
            # Extract from og:tags first (works for Deloitte, etc)
            og_title = soup.find("meta", property="og:title")
            og_desc = soup.find("meta", property="og:description")
            og_site_name = soup.find("meta", property="og:site_name")

            # Try to get company from og:site_name (works for Phenom ATS like Cisco)
            if og_site_name and not company:
                company = og_site_name.get("content", "")

            # Try to extract clean title and company from JSON-LD
            json_ld_scripts = soup.find_all("script", type="application/ld+json")
            for script in json_ld_scripts:
                try:
                    ld_data = json.loads(script.string)
                    if ld_data.get("@type") == "JobPosting":
                        if not title and ld_data.get("title"):
                            title = ld_data.get("title")
                        if not company:
                            org = ld_data.get("hiringOrganization", {})
                            if isinstance(org, dict) and org.get("name"):
                                company = org.get("name")
                        break
                except:
                    continue

            if og_title and not title:
                title_content = og_title.get("content", "")
                # Clean up "Check out this job at Company, Title" format
                if "Check out this job at" in title_content:
                    parts = title_content.split(",", 1)
                    if len(parts) > 1:
                        title = parts[1].strip()
                        company = parts[0].replace("Check out this job at", "").strip()
                else:
                    title = title_content
            
            # Try to extract full JD from HTML - look for Position Summary section
            position_summary = soup.find("h3", string=lambda t: t and "Position Summary" in t)
            if position_summary:
                # Get the next sibling with content
                parent = position_summary.find_parent("article") or position_summary.find_parent("section")
                if parent:
                    jd = parent.get_text(separator="\n", strip=True)
            
            # Also try to find JSON-LD or embedded description
            if not jd or len(jd) < 200:
                import re
                # Look for "description": "..." pattern in HTML
                desc_match = re.search(r'"description"\s*:\s*"([^"]{200,})"', html_text)
                if desc_match:
                    jd_raw = desc_match.group(1)
                    # Unescape unicode and HTML
                    jd_raw = jd_raw.encode().decode('unicode_escape')
                    jd_soup = BeautifulSoup(jd_raw, "html.parser")
                    jd = jd_soup.get_text(separator="\n", strip=True)
            
            # Fallback to og:description
            if (not jd or len(jd) < 100) and og_desc:
                jd = og_desc.get("content", "")
            
            # Extract title from h1 or title tag
            if not title:
                title_el = soup.find("h1") or soup.find("title")
                if title_el:
                    title = title_el.get_text(strip=True)
            
            # Extract company from various sources
            if not company:
                company_el = soup.find(class_=lambda x: x and "company" in x.lower()) if soup else None
                if company_el:
                    company = company_el.get_text(strip=True)
                elif "greenhouse" in url:
                    parts = url.split("/")
                    for i, p in enumerate(parts):
                        if p == "boards.greenhouse.io" and i+1 < len(parts):
                            company = parts[i+1].replace("-", " ").title()
                            break
                elif "deloitte" in url.lower():
                    company = "Deloitte"
            
            # Extract JD text from page content as last resort
            if not jd or len(jd) < 100:
                content_selectors = [
                    "#content", ".job-description", ".description", 
                    "[class*='description']", "[class*='content']", "main", "article"
                ]
                for sel in content_selectors:
                    el = soup.select_one(sel)
                    if el:
                        jd_text = el.get_text(separator="\n", strip=True)
                        if len(jd_text) > 200:
                            jd = jd_text
                            break
            
            if not jd or len(jd) < 100:
                jd = soup.get_text(separator="\n", strip=True)[:5000]
            
    except Exception as e:
        print(f"[AnalyzeJobUrl] Standard fetch error: {e}, will try browser fallback")
        jd = ""  # Reset JD to trigger browser fallback

    # If standard parsing failed or JD is too short/low quality, try browser-based parsing
    # Check for indicators that we got navigation/boilerplate instead of real JD
    browser_result = None
    jd_seems_valid = jd and len(jd) > 500 and (
        "responsibilities" in jd.lower() or
        "requirements" in jd.lower() or
        "qualifications" in jd.lower() or
        "experience" in jd.lower() or
        "skills" in jd.lower()
    )

    if not jd_seems_valid:
        print(f"[AnalyzeJobUrl] Standard parsing insufficient (len={len(jd) if jd else 0}, valid={jd_seems_valid}), trying browser-based parsing...")
        try:
            from utils.browser_parser import parse_job_page_sync
            browser_result = parse_job_page_sync(url, take_screenshot=True)
            
            print(f"[AnalyzeJobUrl] Browser result: ok={browser_result.get('ok')}, jd_len={len(browser_result.get('jd', ''))}")
            if browser_result.get("ok"):
                jd = browser_result.get("jd", "")
                title = browser_result.get("title", title) or title
                company = browser_result.get("company", company) or company
                location = browser_result.get("location", location) or location

                # Special handling for dejobs.org - h1 contains company, first JD line is title
                if "dejobs.org" in url and jd:
                    jd_lines = jd.split("\n")
                    if jd_lines and not company:
                        # h1/title is actually company name on dejobs.org
                        company = title
                        # First non-empty line of JD is the actual job title
                        for line in jd_lines[:5]:
                            line = line.strip()
                            if line and len(line) > 5 and len(line) < 150:
                                title = line
                                break

                # Check if job is closed
                if browser_result.get("is_closed"):
                    return {
                        "ok": False, 
                        "error": f"⚠️ This job is no longer accepting applications ({browser_result.get('closed_reason', 'Position Closed')})",
                        "is_closed": True,
                        "title": title,
                        "company": company,
                        "screenshot_base64": browser_result.get("screenshot_base64")
                    }
                    
                print(f"[AnalyzeJobUrl] Browser parsing success: {title}, JD={len(jd)} chars")
        except Exception as e:
            print(f"[AnalyzeJobUrl] Browser parsing failed: {e}")
    
    if not jd or len(jd) < 100:
        return {"ok": False, "error": "Could not extract job description. The page may require login or the job may no longer be available."}
    
    # 3. Classify role
    from utils.job_utils import classify_role
    role = classify_role(title, jd)
    role_family = role.get("role_family", "other")
    
    # 4. Analyze with AI
    analysis = analyze_job_with_ai(title, company, jd, role_family)
    
    if "error" in analysis:
        # Fallback - still allow adding but with warning
        analysis = {
            "match_score": 50,
            "fit_level": "unknown",
            "analysis_summary": "Could not analyze. Add manually to review.",
            "key_requirements": [],
            "gaps": [],
            "red_flags": []
        }
    
    # 5. Generate recommendation
    score = analysis.get("match_score", 50)
    fit = analysis.get("fit_level", "unknown")
    
    if score >= 75:
        recommendation = "strong_match"
        rec_text = "✅ Strong match! Recommended to apply."
    elif score >= 60:
        recommendation = "good_match"
        rec_text = "👍 Good match. Worth applying."
    elif score >= 45:
        recommendation = "moderate_match"
        rec_text = "⚠️ Moderate match. Review before adding."
    else:
        recommendation = "low_match"
        rec_text = "❌ Low match. Consider skipping."
    
    # Use AI recommendation if available
    ai_rec = analysis.get("recommendation", "")
    if ai_rec:
        rec_text = f"{ai_rec}: {analysis.get('recommendation_reason', '')}"
    
    # Collect application flow info from browser parsing
    application_info = {}
    if browser_result:
        is_intermediate = browser_result.get("is_intermediate", False)
        apply_url = browser_result.get("apply_url")
        is_aggregator = browser_result.get("is_aggregator", False)

        if is_intermediate and apply_url:
            application_info = {
                "is_intermediate": True,
                "is_aggregator": is_aggregator,
                "apply_url": apply_url,
                "message": "⚠️ This is an aggregator page. Click 'Apply' to go to the actual company career page."
            }
        elif is_aggregator:
            application_info = {
                "is_intermediate": False,
                "is_aggregator": True,
                "message": "📋 This appears to be a job aggregator site."
            }

    result = {
        "ok": True,
        "url": url,
        "title": title,
        "company": company,
        "ats": ats,
        "role_family": role_family,
        "match_score": score,  # Also at top level for easy access
        "analysis": {
            "match_score": score,
            "fit_level": fit,
            "role_type": analysis.get("role_type", ""),
            "location_info": analysis.get("location_info", ""),
            "summary": analysis.get("analysis_summary", ""),
            "key_requirements": analysis.get("key_requirements", []),
            "matching_experience": analysis.get("matching_experience", []),
            "gaps": analysis.get("gaps", []),
            "red_flags": analysis.get("red_flags", []),
            "pros": analysis.get("pros", []),
            "cons": analysis.get("cons", [])
        },
        "recommendation": recommendation,
        "recommendation_text": rec_text,
        "application_info": application_info if application_info else None
    }
    
    # Cache successful result
    _analysis_cache[cache_key] = result
    print(f"[AnalyzeJobUrl] Cached result for {cache_url[:50]}... (score: {score}%)")

    return result


@app.post("/clear-analysis-cache")
def clear_analysis_cache_endpoint(payload: ClearAnalysisCacheRequest):
    """
    Clear the analysis cache for a specific URL or all cached results.
    Use this when re-analyzing a job or clearing stale cached results.
    """
    import hashlib
    global _analysis_cache

    if payload.url:
        # Clear specific URL
        cache_url = payload.url.split('?')[0].lower().rstrip('/')
        cache_key = hashlib.md5(cache_url.encode()).hexdigest()

        if cache_key in _analysis_cache:
            del _analysis_cache[cache_key]
            print(f"[ClearCache] Cleared cache for {cache_url[:50]}...")
            return {"ok": True, "cleared": 1, "message": f"Cleared cache for {cache_url}"}
        else:
            return {"ok": True, "cleared": 0, "message": "URL not in cache"}
    else:
        # Clear all cache
        count = len(_analysis_cache)
        _analysis_cache = {}
        print(f"[ClearCache] Cleared all {count} cached results")
        return {"ok": True, "cleared": count, "message": f"Cleared all {count} cached results"}


class CheckApplicationPageRequest(BaseModel):
    url: str


@app.post("/check-application-page")
async def check_application_page(payload: CheckApplicationPageRequest):
    """
    Check if a job URL leads to an application form or if it's an intermediate page.
    Use this before starting auto-fill to ensure we're on the right page.

    Returns:
        {
            "ok": bool,
            "url": str,              # Original URL
            "final_url": str,        # URL of actual application page
            "is_intermediate": bool, # True if original URL is aggregator/intermediate
            "has_form": bool,        # True if application form was found
            "form_fields_count": int,# Number of form fields detected
            "redirects": list,       # URLs navigated through
            "ready_for_autofill": bool,  # True if we can start auto-fill
            "message": str,          # Human-readable status
        }
    """
    from utils.browser_parser import navigate_to_application_form_sync, is_aggregator_url

    url = payload.url.strip()

    try:
        result = navigate_to_application_form_sync(url)

        is_intermediate = len(result.get("redirects", [])) > 0 or is_aggregator_url(url)
        has_form = result.get("has_form", False)
        ready = has_form and result.get("ok", False)

        # Generate status message
        if not result.get("ok"):
            message = f"❌ Error: {result.get('error', 'Unknown error')}"
        elif is_intermediate and not has_form:
            message = "⚠️ Intermediate page detected. Apply button may lead to external site."
        elif is_intermediate and has_form:
            message = f"✅ Navigated through {len(result.get('redirects', []))} redirect(s) to application form."
        elif has_form:
            message = "✅ Application form detected. Ready for auto-fill."
        else:
            message = "⚠️ No application form found. May need to click Apply button manually."

        return {
            "ok": result.get("ok", False),
            "url": url,
            "final_url": result.get("final_url", url),
            "is_intermediate": is_intermediate,
            "has_form": has_form,
            "form_fields_count": result.get("form_fields_count", 0),
            "redirects": result.get("redirects", []),
            "ready_for_autofill": ready,
            "message": message
        }

    except Exception as e:
        return {
            "ok": False,
            "url": url,
            "error": str(e),
            "message": f"❌ Error checking page: {str(e)}"
        }


@app.post("/analyze-missing-scores")
async def analyze_missing_scores():
    """
    Find jobs without match_score and analyze them.
    Returns list of job IDs that will be analyzed.
    """
    jobs = get_all_jobs()
    missing = [j for j in jobs if j.get("match_score") is None and j.get("status") in ["New", "Selected"]]

    return {
        "total_jobs": len(jobs),
        "missing_scores": len(missing),
        "jobs": [{"id": j["id"], "title": j.get("title"), "company": j.get("company"), "url": j.get("job_url")} for j in missing[:50]]
    }


class AnalyzeJobByIdRequest(BaseModel):
    job_id: str


@app.post("/analyze-job-by-id")
async def analyze_job_by_id(payload: AnalyzeJobByIdRequest):
    """
    Analyze a specific job by ID and update its match_score.
    """
    from api.prepare_application import analyze_job_with_ai

    job = get_job_by_id(payload.job_id)
    if not job:
        return {"ok": False, "error": "Job not found"}

    url = job.get("job_url")
    if not url:
        return {"ok": False, "error": "Job has no URL"}

    # Use existing analyze endpoint logic
    try:
        # Fetch JD
        ats_info = detect_ats_from_url(url)
        job_data = fetch_single_job(ats_info)

        if "error" in job_data:
            return {"ok": False, "error": job_data["error"]}

        jd = job_data.get("description", "")
        title = job_data.get("title") or job.get("title", "")
        company = job.get("company", "")

        if not jd:
            return {"ok": False, "error": "Could not fetch job description"}

        # Run AI analysis
        analysis = analyze_job_with_ai(jd, title, company)

        if not analysis or "error" in analysis:
            return {"ok": False, "error": analysis.get("error", "Analysis failed")}

        score = analysis.get("match_score", 0)

        # Update job in storage
        from storage.job_storage import _load_jobs, _save_jobs
        jobs = _load_jobs()
        for j in jobs:
            if j.get("id") == payload.job_id:
                j["match_score"] = score
                j["analysis"] = analysis
                break
        _save_jobs(jobs)

        return {
            "ok": True,
            "job_id": payload.job_id,
            "match_score": score,
            "analysis": analysis
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ============= APPLY AUTOMATION ENDPOINTS =============

class ApplyRequest(BaseModel):
    job_url: str
    profile: str = "anton_tpm"


@app.post("/apply/greenhouse")
def apply_greenhouse_endpoint(payload: ApplyRequest):
    """
    Open Greenhouse job application and auto-fill form using SmartFillerV35.
    """
    import subprocess
    import sys
    import os
    
    job_url = payload.job_url
    profile_name = payload.profile
    
    if "greenhouse" not in job_url.lower() and "gh_jid" not in job_url.lower():
        return {"ok": False, "error": "Only Greenhouse URLs supported"}
    
    profile_path = Path(f"browser/profiles/{profile_name}.json")
    if not profile_path.exists():
        return {"ok": False, "error": f"Profile '{profile_name}' not found"}
    
    # Use absolute path 
    cwd = os.path.dirname(os.path.abspath(__file__))
    
    # Write script to file to avoid shell escaping issues
    script_file = "/tmp/greenhouse_apply_script.py"
    with open(script_file, "w") as f:
        f.write(f'''
import sys
sys.path.insert(0, '{cwd}')
import os
os.chdir('{cwd}')

from browser.smart_filler_v35 import SmartFillerV35
import re

job_url = "{job_url}"

# Convert company career page URL to direct Greenhouse form URL
if 'gh_jid=' in job_url and 'job-boards.greenhouse.io' not in job_url:
    match = re.search(r'gh_jid=(\\d+)', job_url)
    if match:
        gh_jid = match.group(1)
        company_match = re.search(r'https?://(?:www\\.)?([^/]+)\\.com', job_url)
        company = company_match.group(1) if company_match else 'company'
        job_url = "https://job-boards.greenhouse.io/embed/job_app?token=" + gh_jid + "&for=" + company + "&gh_jid=" + gh_jid
        print("Converted to direct Greenhouse URL: " + job_url)

try:
    filler = SmartFillerV35(headless=False)
    filler.run(job_url, interactive=False)
    
    print("\\n" + "="*60)
    print("Browser will stay open for 60 seconds for review...")
    print("="*60)
    
    import time
    time.sleep(60)
except KeyboardInterrupt:
    pass
except Exception as e:
    print(f"Error: {{e}}")
    import traceback
    traceback.print_exc()
    import time
    time.sleep(10)
finally:
    if 'filler' in dir() and filler:
        filler.stop()
    print("Browser closed")
''')
    
    # Run script in background
    log_file = "/tmp/apply_greenhouse.log"
    with open(log_file, "w") as log:
        log.write(f"Starting apply for: {job_url}\n")
        log.write(f"Profile: {profile_name}\n")
        log.write("="*60 + "\n")
    
    # Start subprocess
    process = subprocess.Popen(
        [sys.executable, script_file],
        stdout=open(log_file, "a"),
        stderr=subprocess.STDOUT,
        cwd=cwd
    )
    
    return {
        "ok": True,
        "message": "Application form opened with SmartFiller V3.5",
        "pid": process.pid,
        "log_file": log_file
    }

@app.get("/answers")
def get_answer_library():
    """Get the full answer library."""
    path = Path("data/answer_library.json")
    if not path.exists():
        return {"personal": {}, "links": {}, "answers": {}, "cover_letter_template": {}}
    with open(path) as f:
        return json.load(f)


@app.put("/answers")
def update_answer_library(data: dict):
    """Update the answer library."""
    path = Path("data/answer_library.json")
    with open(path, "w") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    return {"ok": True}


@app.get("/answers/{category}/{key}")
def get_answer(category: str, key: str):
    """Get a specific answer."""
    path = Path("data/answer_library.json")
    if not path.exists():
        return {"error": "Answer library not found"}
    with open(path) as f:
        data = json.load(f)
    
    if category in data and key in data[category]:
        return {"value": data[category][key]}
    return {"error": f"Key {category}/{key} not found"}


@app.post("/generate-cover-letter")
def generate_cover_letter_endpoint(payload: dict):
    """
    Generate a personalized cover letter from DOCX template.
    Expects: {company, position, job_description?, role_family?}
    Returns: {ok, cover_letter, file_path}
    """
    from docx import Document
    import shutil
    
    company = payload.get("company", "Company")
    position = payload.get("position", "Position")
    job_description = payload.get("job_description", "")
    role_family = payload.get("role_family", "product")  # product, tpm_program, project
    
    # Map role_family to template
    cv_dir = GOLD_CV_PATH
    template_map = {
        "product": "Cover_Letter_Anton_Kondakov_ProductM.docx",
        "tpm_program": "Cover_Letter_Anton_Kondakov_Delivery Lead.docx",
        "project": "Cover_Letter_Anton_Kondakov_Project Manager.docx",
        "scrum": "Cover_Letter_Anton_Kondakov_Scrum Master.docx",
        "po": "Cover_Letter_Anton_Kondakov_PO.docx",
    }
    
    template_name = template_map.get(role_family, template_map["product"])
    template_path = cv_dir / template_name
    
    if not template_path.exists():
        return {"error": f"Template not found: {template_name}"}
    
    # Generate company mission using AI
    company_mission = ""
    try:
        from utils.ollama_ai import generate_company_mission, is_ollama_available
        if is_ollama_available():
            company_mission = generate_company_mission(company, job_description, position)
            print(f"Generated mission: {company_mission}")
    except Exception as e:
        print(f"AI mission generation error: {e}")
        company_mission = f"I'm excited about the opportunity to contribute to {company}'s continued success."
    
    # Load and modify template
    doc = Document(str(template_path))
    
    for para in doc.paragraphs:
        for run in para.runs:
            if "[COMPANY NAME]" in run.text:
                run.text = run.text.replace("[COMPANY NAME]", company)
            if "[POSITION TITLE]" in run.text:
                run.text = run.text.replace("[POSITION TITLE]", position)
            if "[COMPANY MISSION]" in run.text:
                run.text = run.text.replace("[COMPANY MISSION]", company_mission)
    
    # Create Applications folder
    applications_dir = cv_dir / "Applications"
    applications_dir.mkdir(exist_ok=True)
    
    # Create company folder
    safe_company = company.replace(" ", "_").replace("/", "_").replace("\\", "_")
    safe_position = position.replace(" ", "_").replace("/", "_").replace("\\", "_")[:50]
    job_folder = applications_dir / f"{safe_company}_{safe_position}"
    job_folder.mkdir(exist_ok=True)
    
    # Save cover letter
    output_filename = f"Cover_Letter_{safe_company}_{safe_position}.docx"
    output_path = job_folder / output_filename
    doc.save(str(output_path))
    
    # Also extract text for preview
    cover_letter_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    return {
        "ok": True,
        "cover_letter": cover_letter_text,
        "file_path": str(output_path),
        "folder_path": str(job_folder),
        "company_mission": company_mission,
        "template_used": template_name
    }

@app.post("/save-cover-letter")
def save_cover_letter(payload: dict):
    """
    Copy selected CV to the job application folder.
    Expects: {company, position, cv_filename}
    Returns: {ok, folder_path}
    """
    import shutil
    
    company = payload.get("company", "Unknown").replace(" ", "_").replace("/", "_").replace("\\", "_")
    position = payload.get("position", "Position").replace(" ", "_").replace("/", "_").replace("\\", "_")[:50]
    cv_filename = payload.get("cv_filename", "")
    
    cv_dir = GOLD_CV_PATH
    applications_dir = cv_dir / "Applications"
    job_folder = applications_dir / f"{company}_{position}"
    
    # Ensure folder exists
    job_folder.mkdir(parents=True, exist_ok=True)
    
    # Copy CV if specified
    if cv_filename:
        cv_source = cv_dir / cv_filename
        if cv_source.exists():
            cv_dest = job_folder / cv_filename
            shutil.copy2(cv_source, cv_dest)
            print(f"Copied CV: {cv_filename} -> {job_folder}")
    
    return {
        "ok": True,
        "folder_path": str(job_folder)
    }



@app.get("/available-cvs")
def get_available_cvs():
    """
    Get list of available CV files.
    """
    cv_dir = GOLD_CV_PATH
    cvs = []
    
    for f in cv_dir.glob("*.pdf"):
        cvs.append({"path": str(f), "name": f.name})
    
    # Also check for docx
    for f in cv_dir.glob("*CV*.docx"):
        if not f.name.startswith("~"):
            cvs.append({"path": str(f), "name": f.name})
    
    return {"cvs": cvs}


@app.post("/select-cv")
def select_cv_for_job(payload: dict):
    """
    Use AI to select best CV for a job.
    Expects: {job_title}
    Returns: {selected_cv, reason}
    """
    job_title = payload.get("job_title", "")
    
    # Get available CVs
    cv_dir = GOLD_CV_PATH
    available_cvs = [str(f) for f in cv_dir.glob("*.pdf")]
    
    if not available_cvs:
        return {"error": "No CV files found"}
    
    try:
        from utils.ollama_ai import select_cv_for_role, is_ollama_available
        
        if is_ollama_available():
            selected = select_cv_for_role(job_title, available_cvs)
            return {
                "ok": True,
                "selected_cv": selected,
                "filename": Path(selected).name if selected else None,
                "ai_selected": True
            }
    except Exception as e:
        print(f"AI CV selection error: {e}")
    
    # Fallback: simple keyword matching
    title_lower = job_title.lower()
    for cv in available_cvs:
        cv_lower = cv.lower()
        if 'tpm' in title_lower and 'tpm' in cv_lower:
            return {"ok": True, "selected_cv": cv, "filename": Path(cv).name, "ai_selected": False}
        if 'product' in title_lower and 'product' in cv_lower:
            return {"ok": True, "selected_cv": cv, "filename": Path(cv).name, "ai_selected": False}
        if 'program' in title_lower and 'tpm' in cv_lower:
            return {"ok": True, "selected_cv": cv, "filename": Path(cv).name, "ai_selected": False}
    
    return {"ok": True, "selected_cv": available_cvs[0], "filename": Path(available_cvs[0]).name, "ai_selected": False}


@app.get("/apply-log")
def get_apply_log():
    """
    Get the latest apply log content.
    """
    log_path = Path("/tmp/greenhouse_apply.log")
    if not log_path.exists():
        return {"ok": False, "log": "No log file found"}
    
    try:
        with open(log_path, "r") as f:
            content = f.read()
        return {"ok": True, "log": content}
    except Exception as e:
        return {"ok": False, "log": f"Error reading log: {e}"}


@app.post("/fetch-job-description")
def fetch_job_description(payload: dict):
    """
    Fetch job description from URL.
    Supports: Workday API, Greenhouse API, and scraping fallback.
    Expects: {url}
    Returns: {ok, description}
    """
    import requests
    from bs4 import BeautifulSoup
    import re
    import html
    
    url = payload.get("url", "")
    if not url:
        return {"ok": False, "error": "No URL provided"}
    
    try:
        # ============ WORKDAY API ============
        if "myworkdayjobs.com" in url:
            # URL formats:
            # 1. https://company.wd1.myworkdayjobs.com/en-US/site/job/Location/Title_JOBID
            # 2. https://company.wd1.myworkdayjobs.com/site/job/Location/Title_JOBID  
            # 3. https://company.wd1.myworkdayjobs.com/job/Location/Title_JOBID (no site)
            
            # Extract company slug
            company_match = re.match(r'https?://([^\.]+)\.wd\d+\.myworkdayjobs\.com', url)
            if not company_match:
                pass  # Will fall through to scraping
            else:
                company_slug = company_match.group(1)
                
                # Try to extract site and job path
                # Pattern with site: /site/job/path or /en-US/site/job/path
                site_job_match = re.search(r'myworkdayjobs\.com/(?:[a-z][a-z]-[A-Z][A-Z]/)?([^/]+)/job/(.+)', url)
                # Pattern without site: /job/path
                direct_job_match = re.search(r'myworkdayjobs\.com/job/(.+)', url)
                
                if site_job_match:
                    site = site_job_match.group(1)
                    job_path = site_job_match.group(2)
                elif direct_job_match:
                    # No site in URL - try common site names
                    job_path = direct_job_match.group(1)
                    site = f"{company_slug}_careers"  # Common pattern
                else:
                    site = None
                    job_path = None
                
                if site and job_path:
                    api_url = f"https://{company_slug}.wd1.myworkdayjobs.com/wday/cxs/{company_slug}/{site}/job/{job_path}"
                    
                    try:
                        resp = requests.get(api_url, timeout=15)
                        if resp.status_code == 200:
                            data = resp.json()
                            jp = data.get("jobPostingInfo", {})
                            desc_html = jp.get("jobDescription", "")
                            if desc_html:
                                soup = BeautifulSoup(desc_html, "html.parser")
                                text = soup.get_text(separator="\n", strip=True)
                                text = html.unescape(text)
                                lines = [l.strip() for l in text.split("\n") if l.strip()]
                                text = "\n".join(lines)[:5000]
                                return {"ok": True, "description": text, "source": "workday_api", "title": jp.get("title")}
                    except Exception as e:
                        print(f"Workday API error for {api_url}: {e}")
                    
                    # If first site didn't work, try without _careers suffix
                    if "_careers" in site:
                        alt_site = company_slug
                        api_url = f"https://{company_slug}.wd1.myworkdayjobs.com/wday/cxs/{company_slug}/{alt_site}/job/{job_path}"
                        try:
                            resp = requests.get(api_url, timeout=15)
                            if resp.status_code == 200:
                                data = resp.json()
                                jp = data.get("jobPostingInfo", {})
                                desc_html = jp.get("jobDescription", "")
                                if desc_html:
                                    soup = BeautifulSoup(desc_html, "html.parser")
                                    text = soup.get_text(separator="\n", strip=True)
                                    text = html.unescape(text)
                                    lines = [l.strip() for l in text.split("\n") if l.strip()]
                                    text = "\n".join(lines)[:5000]
                                    return {"ok": True, "description": text, "source": "workday_api", "title": jp.get("title")}
                        except Exception as e:
                            print(f"Workday API error (alt site) for {api_url}: {e}")
        
        # ============ GREENHOUSE API ============
        # Try to extract Greenhouse job ID from URL
        gh_job_id = None
        gh_board = None
        
        # Pattern 1: gh_jid parameter
        gh_match = re.search(r'gh_jid=(\d+)', url)
        if gh_match:
            gh_job_id = gh_match.group(1)
        
        # Pattern 2: /jobs/12345 in path
        if not gh_job_id:
            job_match = re.search(r'/jobs/(\d+)', url)
            if job_match:
                gh_job_id = job_match.group(1)
        
        # Try to get board name from URL or pipeline data
        # Common patterns: boards.greenhouse.io/BOARD, company.ai/careers
        if 'greenhouse.io' in url:
            board_match = re.search(r'greenhouse\.io/([^/]+)', url)
            if board_match:
                gh_board = board_match.group(1)
        
        # If we have job ID, try Greenhouse API
        if gh_job_id:
            # Try to find board from our companies data
            if not gh_board:
                # Look up in companies.json
                companies_path = Path("data/companies.json")
                if companies_path.exists():
                    companies = json.loads(companies_path.read_text())
                    for comp in companies:
                        if comp.get("ats") == "greenhouse" and comp.get("board_url"):
                            board_url = comp.get("board_url", "")
                            if board_url:
                                # Extract board name
                                match = re.search(r'greenhouse\.io/([^/]+)', board_url)
                                if match:
                                    test_board = match.group(1)
                                    # Try this board
                                    api_url = f"https://boards-api.greenhouse.io/v1/boards/{test_board}/jobs/{gh_job_id}"
                                    try:
                                        resp = requests.get(api_url, timeout=5)
                                        if resp.status_code == 200:
                                            gh_board = test_board
                                            break
                                    except:
                                        pass
            
            # Common board names to try
            boards_to_try = [gh_board] if gh_board else []
            
            # Extract potential board from URL domain
            domain_match = re.search(r'https?://([^/\.]+)', url)
            if domain_match:
                potential_board = domain_match.group(1).replace('-', '').lower()
                if potential_board not in boards_to_try:
                    boards_to_try.append(potential_board)
            
            # Try some common variations
            for board in boards_to_try:
                if not board:
                    continue
                api_url = f"https://boards-api.greenhouse.io/v1/boards/{board}/jobs/{gh_job_id}"
                try:
                    resp = requests.get(api_url, timeout=10)
                    if resp.status_code == 200:
                        data = resp.json()
                        content = data.get("content", "")
                        if content:
                            # Unescape HTML entities first (Greenhouse returns double-escaped)
                            content = html.unescape(content)
                            # Parse HTML content
                            soup = BeautifulSoup(content, "html.parser")
                            text = soup.get_text(separator="\n", strip=True)
                            # Clean up
                            lines = [l.strip() for l in text.split("\n") if l.strip()]
                            text = "\n".join(lines)
                            # Limit length
                            text = text[:5000]
                            return {"ok": True, "description": text, "source": "greenhouse_api"}
                except Exception as e:
                    pass
        
        # Fallback: try direct scraping
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
        }
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        
        soup = BeautifulSoup(resp.text, "html.parser")
        
        # Try different selectors for job description
        description = ""
        
        # Greenhouse
        content = soup.select_one("#content, .content, [data-automation='job-description']")
        if content:
            description = content.get_text(separator="\n", strip=True)
        
        # Lever
        if not description:
            content = soup.select_one(".posting-page, .section-wrapper")
            if content:
                description = content.get_text(separator="\n", strip=True)
        
        # Generic fallback - main content area
        if not description:
            for selector in ["main", "article", ".job-description", ".description", "#job-content"]:
                content = soup.select_one(selector)
                if content:
                    description = content.get_text(separator="\n", strip=True)
                    break
        
        # Clean up - remove excessive whitespace, limit length
        if description and len(description) > 50:
            lines = [line.strip() for line in description.split("\n") if line.strip()]
            description = "\n".join(lines[:100])  # Limit to ~100 lines
            description = description[:5000]  # Limit to 5000 chars
            return {"ok": True, "description": description, "source": "scraping"}
        
        return {"ok": False, "error": "Could not extract job description"}
            
    except Exception as e:
        return {"ok": False, "error": str(e)}

@app.post("/apply/vision")
def apply_with_vision(payload: ApplyRequest):
    """
    Apply to job using Vision AI Agent.
    AI looks at screenshots and fills form like a human.
    """
    import subprocess
    import sys
    import os
    
    job_url = payload.job_url
    profile_name = payload.profile
    
    profile_path = Path(f"browser/profiles/{profile_name}.json")
    if not profile_path.exists():
        return {"ok": False, "error": f"Profile '{profile_name}' not found"}
    
    # Load profile data
    with open(profile_path) as f:
        profile_data = json.load(f)
    
    cwd = os.path.dirname(os.path.abspath(__file__))
    if "Mobile Documents" in cwd:
        cwd = cwd.replace(
            str(AI_PROJECTS_PATH),
            str(ICLOUD_PATH)
        )
    
    script = f'''import sys
sys.path.insert(0, '{cwd}')
import os
os.chdir('{cwd}')

from browser.client import BrowserClient
from browser.vision_agent import VisionFormAgent
import json
import time

# Load profile from file
with open("browser/profiles/{profile_name}.json") as f:
    profile = json.load(f)

# Start browser
browser = BrowserClient()
browser.start()

try:
    # Open job page
    browser.open_job_page("{job_url}")
    time.sleep(3)
    
    # Start Vision Agent
    agent = VisionFormAgent(browser.page, profile)
    result = agent.fill_form()
    
    print("\\n" + "="*50)
    print(f"Result: {{result}}")
    print("="*50)
    
    # Keep open for review
    print("\\nBrowser stays open for 60 seconds...")
    time.sleep(60)
    
except Exception as e:
    print(f"Error: {{e}}")
    import traceback
    traceback.print_exc()
    time.sleep(10)
finally:
    browser.close()
'''
    
    # Write and execute script
    script_file = "/tmp/vision_apply_script.py"
    log_file = "/tmp/vision_apply.log"
    
    with open(script_file, "w") as f:
        f.write(script)
    
    # Run in background
    subprocess.Popen(
        [sys.executable, script_file],
        stdout=open(log_file, "w"),
        stderr=subprocess.STDOUT,
        start_new_session=True
    )
    
    return {
        "ok": True,
        "message": f"Vision AI Agent started for {job_url}",
        "log_file": log_file,
        "screenshots_dir": "/tmp/vision_agent"
    }


@app.post("/open-folder")
def open_folder(payload: dict):
    """Open folder in Finder (macOS)"""
    import subprocess
    import os
    folder_path = payload.get("path", "")
    
    # Expand ~ to home directory
    if folder_path.startswith("~/"):
        folder_path = os.path.expanduser(folder_path)
    
    if not folder_path or not Path(folder_path).exists():
        return {"error": f"Folder not found: {folder_path}"}
    
    try:
        subprocess.run(["open", folder_path], check=True)
        return {"ok": True}
    except Exception as e:
        return {"error": str(e)}


@app.get("/applications")
def list_applications():
    """List all prepared job applications"""
    applications_dir = APPLICATIONS_PATH
    
    if not applications_dir.exists():
        return {"applications": []}
    
    apps = []
    for folder in sorted(applications_dir.iterdir(), reverse=True):
        if folder.is_dir() and not folder.name.startswith("."):
            files = [f.name for f in folder.iterdir() if f.is_file()]
            apps.append({
                "name": folder.name,
                "path": str(folder),
                "files": files,
                "created": folder.stat().st_mtime
            })
    
    return {"applications": apps}



# ============= JOB ANALYSIS ENDPOINT =============

class AnalyzeJobRequest(BaseModel):
    job_description: str
    job_title: str
    company: str
    role_family: str = "product"

@app.post("/analyze-job")
async def analyze_job_endpoint(payload: AnalyzeJobRequest):
    """
    Analyze job description against candidate profile.
    Returns match score, missing keywords, ATS tips.
    """
    import re
    from pathlib import Path
    
    # Load candidate profile (Gold CV)
    gold_cv_path = GOLD_CV_PATH
    
    # Select CV based on role
    role_cv_map = {
        "product": "CV_Anton_Kondakov_Product Manager.pdf",
        "tpm_program": "CV_Anton_Kondakov_TPM_CV.pdf",
        "project": "CV_Anton_Kondakov_Project Manager.pdf",
    }
    cv_file = role_cv_map.get(payload.role_family, "CV_Anton_Kondakov_Product Manager.pdf")
    
    jd = payload.job_description.lower()
    title = payload.job_title.lower()
    
    # Extract keywords from JD
    # Common PM/TPM keywords
    skill_keywords = {
        "hard_skills": [
            "agile", "scrum", "kanban", "jira", "confluence", "sql", "python", "api", 
            "aws", "azure", "gcp", "kubernetes", "docker", "ci/cd", "devops",
            "data analysis", "analytics", "tableau", "looker", "amplitude",
            "a/b testing", "product analytics", "roadmap", "okr", "kpi",
            "user research", "ux", "figma", "prototyping", "wireframes",
            "technical specifications", "prd", "requirements", "stakeholder",
            "cross-functional", "program management", "project management",
            "release management", "sprint planning", "backlog", "prioritization",
            "saas", "b2b", "b2c", "enterprise", "platform", "infrastructure",
            "machine learning", "ml", "ai", "artificial intelligence",
            "payments", "fintech", "banking", "financial services",
            "security", "compliance", "gdpr", "sox", "pci",
            # Sales/Account Management (for non-PM roles)
            "sales", "account management", "client management", "customer success",
            "relationship building", "revenue", "quota", "pipeline",
            "crm", "salesforce", "hubspot", "negotiation", "closing",
            "territory", "prospecting", "lead generation", "business development"
        ],
        "soft_skills": [
            "leadership", "communication", "collaboration", "problem-solving",
            "strategic thinking", "decision-making", "influence", "negotiation",
            "mentoring", "coaching", "presentation", "executive"
        ],
        "experience": [
            "5+ years", "7+ years", "10+ years", "senior", "staff", "principal",
            "director", "lead", "manager", "head of"
        ]
    }
    
    # Find keywords in JD
    found_keywords = {"hard_skills": [], "soft_skills": [], "experience": []}
    missing_keywords = {"hard_skills": [], "soft_skills": [], "experience": []}
    
    # My profile keywords (from CV)
    my_keywords = {
        "agile", "scrum", "kanban", "jira", "confluence", "sql", "python",
        "aws", "data analysis", "analytics", "roadmap", "okr", "kpi",
        "cross-functional", "program management", "project management",
        "release management", "sprint planning", "backlog", "prioritization",
        "saas", "b2b", "enterprise", "platform", "stakeholder",
        "leadership", "communication", "collaboration", "strategic thinking",
        "prd", "requirements", "technical specifications", "api",
        "fintech", "payments", "banking", "financial services",
        # Extended skills from CV
        "gcp", "azure", "cloud", "microservices", "terraform",
        "ci/cd", "devops", "release management", "deployment",
        "regulatory", "compliance", "mifid", "sox", "gdpr",
        "uat", "testing", "quality", "integration",
        "machine learning", "ml", "ai", "data", "analytics",
        "tableau", "looker", "reporting", "metrics",
        "safe", "agile", "scrum", "kanban", "pi planning",
        "product strategy", "product vision", "product roadmap",
        "user stories", "acceptance criteria", "backlog management",
        "cross-functional", "influence", "negotiation", "executive",
        "client", "customer", "vendor", "partnership"
    }
    
    for category, keywords in skill_keywords.items():
        for kw in keywords:
            if kw in jd:
                found_keywords[category].append(kw)
                if kw.lower() not in my_keywords:
                    missing_keywords[category].append(kw)
    
    # Calculate match score
    total_found = len(found_keywords["hard_skills"]) + len(found_keywords["soft_skills"])
    total_missing = len(missing_keywords["hard_skills"]) + len(missing_keywords["soft_skills"])
    
    if total_found + total_missing > 0:
        match_score = int((total_found - total_missing * 0.5) / (total_found + total_missing) * 100)
        match_score = max(0, min(100, match_score + 50))  # Normalize to 0-100
    else:
        match_score = 70  # Default
    
    # ATS tips
    ats_tips = []
    
    # Check for exact title match
    if "product manager" in title and "product" in payload.role_family:
        ats_tips.append("✅ Title matches your target role")
    elif "program manager" in title and "tpm" in payload.role_family:
        ats_tips.append("✅ Title matches your target role")
    else:
        ats_tips.append("⚠️ Consider tailoring CV title to match job title")
    
    # Check years of experience
    exp_match = re.search(r'(\d+)\+?\s*years?', jd)
    if exp_match:
        years_required = int(exp_match.group(1))
        if years_required <= 12:  # Assuming 12+ years experience
            ats_tips.append(f"✅ You meet the {years_required}+ years requirement")
        else:
            ats_tips.append(f"⚠️ Position requires {years_required}+ years")
    
    # Check for missing critical keywords
    critical_missing = [kw for kw in missing_keywords["hard_skills"] if kw in ["machine learning", "ml", "ai", "kubernetes", "docker"]]
    if critical_missing:
        ats_tips.append(f"⚠️ Add if applicable: {', '.join(critical_missing[:3])}")
    
    if missing_keywords["hard_skills"]:
        top_missing = missing_keywords["hard_skills"][:5]
        ats_tips.append(f"💡 Consider adding: {', '.join(top_missing)}")
    
    # Red flags
    red_flags = []
    if "clearance" in jd or "security clearance" in jd:
        red_flags.append("🚨 Requires security clearance")
    if "relocation" in jd and "not" not in jd:
        red_flags.append("⚠️ May require relocation")
    if "visa" in jd and "sponsor" not in jd:
        red_flags.append("⚠️ Check visa sponsorship policy")
    
    # Response chance estimate
    if match_score >= 80:
        response_chance = "High (70-90%)"
        response_color = "#10b981"
    elif match_score >= 60:
        response_chance = "Medium (40-60%)"
        response_color = "#f59e0b"
    else:
        response_chance = "Low (10-30%)"
        response_color = "#ef4444"
    
    return {
        "ok": True,
        "match_score": match_score,
        "response_chance": response_chance,
        "response_color": response_color,
        "found_keywords": found_keywords,
        "missing_keywords": missing_keywords,
        "ats_tips": ats_tips,
        "red_flags": red_flags,
        "cv_file": cv_file,
        "keywords_to_add": missing_keywords["hard_skills"][:5]
    }


# ============= COMPREHENSIVE APPLICATION PREPARATION =============

class PrepareApplicationRequest(BaseModel):
    job_title: str
    company: str
    job_url: str
    job_description: str
    role_family: str = "product"
    force_regenerate: bool = False  # If True, regenerate even if files exist


class CheckExistingRequest(BaseModel):
    job_title: str
    company: str


@app.post("/check-existing-application")
async def check_existing_application(payload: CheckExistingRequest):
    """
    Check if application files already exist for this job.
    Returns paths to existing CV and Cover Letter if found.
    """
    print(f"[CheckExisting] Checking: {payload.company} - {payload.job_title}")
    from pathlib import Path
    import re
    
    gold_cv_path = GOLD_CV_PATH
    applications_path = gold_cv_path / "Applications"
    
    # Normalize company and position - remove special chars, replace spaces with underscores
    safe_company = re.sub(r'[^\w\s]', '', payload.company).strip().replace(' ', '_')
    safe_position = re.sub(r'[^\w\s]', '', payload.job_title).strip().replace(' ', '_')[:40]
    
    # For matching, also create a simplified version (just alphanumeric)
    match_company = re.sub(r'[^a-zA-Z0-9]', '', payload.company.lower())
    match_position = re.sub(r'[^a-zA-Z0-9]', '', payload.job_title.lower())[:30]
    
    # Look for existing application folders
    existing_folders = []
    if applications_path.exists():
        for folder in applications_path.iterdir():
            # Normalize folder name for matching
            folder_normalized = re.sub(r'[^a-zA-Z0-9]', '', folder.name.lower())
            
            if folder.is_dir() and match_company in folder_normalized:
                # Check if it matches position too
                if match_position[:20] in folder_normalized or len(match_position) < 10:
                    existing_folders.append(folder)
                    print(f"[CheckExisting] Found matching folder: {folder.name}")
    
    if not existing_folders:
        return {"exists": False}
    
    # Get most recent folder
    existing_folders.sort(key=lambda x: x.stat().st_mtime, reverse=True)
    latest_folder = existing_folders[0]
    
    # Find CV and Cover Letter in folder
    cv_file = None
    cl_file = None
    
    for f in latest_folder.iterdir():
        if f.is_file():
            if f.name.startswith("CV_") and f.suffix == ".docx":
                cv_file = f
            elif f.name.startswith("Cover_Letter_") and f.suffix == ".txt":
                cl_file = f
    
    if not cv_file and not cl_file:
        return {"exists": False}
    
    # Read cover letter content for preview
    cl_preview = None
    if cl_file:
        try:
            cl_content = cl_file.read_text()
            cl_preview = cl_content[:500] + "..." if len(cl_content) > 500 else cl_content
        except:
            cl_preview = "Could not read cover letter"
    
    # Read metadata.json for keywords and analysis info
    metadata = {}
    metadata_file = latest_folder / "metadata.json"
    if metadata_file.exists():
        try:
            import json
            metadata = json.load(open(metadata_file))
        except:
            pass
    
    return {
        "exists": True,
        "folder": str(latest_folder),
        "folder_name": latest_folder.name,
        "cv_path": str(cv_file) if cv_file else None,
        "cv_filename": cv_file.name if cv_file else None,
        "cover_letter_path": str(cl_file) if cl_file else None,
        "cover_letter_filename": cl_file.name if cl_file else None,
        "cover_letter_preview": cl_preview,
        "created_at": datetime.fromtimestamp(latest_folder.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
        "keywords_added": metadata.get("keywords_added", []),
        "match_score": metadata.get("match_score"),
        "fit_level": metadata.get("fit_level"),
        "cv_decision": metadata.get("cv_decision"),
        "cv_reason": metadata.get("cv_reason")
    }


@app.post("/prepare-application")
async def prepare_application_endpoint(payload: PrepareApplicationRequest):
    """
    Comprehensive application preparation using Claude API.
    
    1. Deep JD analysis
    2. CV decision (base vs optimize)
    3. Cover letter generation
    4. Returns paths to all documents
    """
    from api.prepare_application import prepare_application
    
    result = prepare_application(
        job_title=payload.job_title,
        company=payload.company,
        job_url=payload.job_url,
        jd=payload.job_description,
        role_family=payload.role_family
    )
    
    return result.to_dict()


@app.get("/open-file/{file_type}")
async def open_file_endpoint(file_type: str, path: str):
    """
    Open file in default application.
    file_type: cv, cover_letter, folder
    """
    import subprocess
    import os
    from pathlib import Path
    
    # Expand ~ to home directory
    if path.startswith("~/"):
        path = os.path.expanduser(path)
    
    file_path = Path(path)
    
    if not file_path.exists():
        return {"ok": False, "error": f"File not found: {path}"}
    
    try:
        if file_type == "folder":
            subprocess.run(["open", str(file_path)], check=True)
        else:
            subprocess.run(["open", str(file_path)], check=True)
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ============= CV PREVIEW & TAILORING =============

class CVPreviewRequest(BaseModel):
    job_title: str
    company: str
    role_family: str = "product"
    keywords_to_add: list = []
    matched_keywords: list = []
    cv_path: str = None  # Optional: path to specific CV (e.g., AI-optimized)

@app.post("/cv/preview")
async def cv_preview_endpoint(payload: CVPreviewRequest):
    """
    Generate CV preview with highlighted keywords.
    Returns HTML with:
    - Green highlights: matched keywords (already in CV and JD)
    - Yellow highlights: injected keywords (added to Skills section)
    """
    from docx import Document
    from pathlib import Path
    import re
    
    gold_cv_path = GOLD_CV_PATH
    
    # Use provided cv_path if specified, otherwise select by role
    print(f"DEBUG cv/preview: payload.cv_path = {payload.cv_path}")
    if payload.cv_path and Path(payload.cv_path).exists():
        cv_path = Path(payload.cv_path)
        cv_filename = cv_path.name
        print(f"DEBUG cv/preview: Using provided CV: {cv_path}")
    else:
        # Select CV based on role
        role_cv_map = {
            "product": "CV_Anton_Kondakov_Product Manager.docx",
            "tpm_program": "CV_Anton_Kondakov_TPM.docx",
            "project": "CV_Anton_Kondakov_Project Manager.docx",
        }
        cv_filename = role_cv_map.get(payload.role_family, "CV_Anton_Kondakov_Product Manager.docx")
        cv_path = gold_cv_path / cv_filename
        print(f"DEBUG cv/preview: Using role-based CV: {cv_path}")
    
    if not cv_path.exists():
        return {"ok": False, "error": f"CV not found: {cv_filename}"}
    
    doc = Document(cv_path)
    
    # Build HTML preview
    html_parts = []
    html_parts.append('<div class="cv-preview" style="font-family: Arial, sans-serif; font-size: 12px; line-height: 1.4; max-width: 800px;">')
    
    matched_kw = set(k.lower() for k in payload.matched_keywords)
    inject_kw = set(k.lower() for k in payload.keywords_to_add)
    
    # Add yellow banner at the TOP if there are keywords to add
    if inject_kw:
        html_parts.append('<div style="margin: 0 0 16px 0; padding: 12px; background-color: #fef08a; border-radius: 8px; border-left: 4px solid #eab308;">')
        html_parts.append('<strong style="color: #854d0e; font-size: 13px;">🔑 Keywords to be added to your CV:</strong><br>')
        html_parts.append('<div style="margin-top: 8px;">')
        html_parts.append(', '.join(f'<mark style="background-color: #facc15; padding: 2px 6px; border-radius: 3px; font-weight: 500;">{kw}</mark>' for kw in payload.keywords_to_add))
        html_parts.append('</div></div>')
    
    # Fallback: if no keywords provided, use common PM keywords for highlighting
    if not matched_kw:
        matched_kw = {
            "product strategy", "roadmap", "agile", "scrum", "stakeholder",
            "cross-functional", "backlog", "user stories", "sprint", "kpi",
            "okr", "prioritization", "requirements", "delivery", "release",
            "jira", "confluence", "aws", "sql", "data analysis"
        }
    
    def highlight_text(text: str, is_technical_section: bool = False) -> str:
        """Highlight matched and injected keywords in text."""
        result = text
        
        # First highlight matched keywords (green) - preserve original case
        for kw in matched_kw:
            pattern = re.compile(f'({re.escape(kw)})', re.IGNORECASE)
            result = pattern.sub(
                r'<mark style="background-color: #86efac !important; padding: 1px 3px; border-radius: 2px;">\1</mark>',
                result
            )
        
        # Add injected keywords to Technical section with yellow highlight
        if is_technical_section and inject_kw:
            injected_str = ', '.join(f'<mark style="background-color: #facc15 !important; padding: 1px 3px; border-radius: 2px; font-weight: 500;">{kw}</mark>' for kw in payload.keywords_to_add)
            result += f' <span style="color: #854d0e;">[+Added: {injected_str}]</span>'
        
        return result
    
    keywords_injected = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style = para.style.name if para.style else "Normal"
        
        # Check if this is Technical Delivery/Acumen line where we inject keywords
        # Must start with bullet point marker or "Technical Delivery" / "Technical Acumen"
        is_technical = (text.startswith("Technical Delivery") or text.startswith("Technical Acumen") or 
                       (text.startswith("•") and "Technical" in text)) and not keywords_injected
        highlighted = highlight_text(text, is_technical_section=is_technical)
        if is_technical and inject_kw:
            keywords_injected = True
        
        # Detect section headers
        if text.isupper() or style == "Heading 1" or text in ["CORE COMPETENCIES", "PROFESSIONAL EXPERIENCE", "EDUCATION", "CERTIFICATIONS"]:
            html_parts.append(f'<h3 style="margin: 16px 0 8px 0; color: #1e3a5f; border-bottom: 1px solid #ddd; padding-bottom: 4px;">{text}</h3>')
                
        elif style == "List Paragraph":
            html_parts.append(f'<div style="margin: 4px 0 4px 20px; padding-left: 10px; border-left: 2px solid #e5e7eb;">• {highlighted}</div>')
        else:
            # Check if it's a job title/company line
            if " | " in text or "–" in text:
                html_parts.append(f'<div style="margin: 12px 0 4px 0; font-weight: 600; color: #374151;">{highlighted}</div>')
            else:
                html_parts.append(f'<div style="margin: 4px 0;">{highlighted}</div>')
    
    html_parts.append('</div>')
    
    # Summary stats
    stats = {
        "matched_count": len(matched_kw),
        "injected_count": len(inject_kw),
        "cv_file": cv_filename
    }
    
    return {
        "ok": True,
        "html": "\n".join(html_parts),
        "stats": stats,
        "keywords_matched": list(matched_kw),
        "keywords_injected": list(inject_kw)
    }


class CVTailorRequest(BaseModel):
    company: str
    position: str  # Job title
    role_family: str = "product"
    keywords_to_add: list = []

@app.post("/cv/tailor")
async def cv_tailor_endpoint(payload: CVTailorRequest):
    """
    Create tailored CV with injected keywords.
    Saves to Applications folder.
    Returns path to new CV.
    """
    from docx import Document
    from docx.shared import RGBColor
    from pathlib import Path
    import re
    
    gold_cv_path = GOLD_CV_PATH
    apps_path = gold_cv_path / "Applications"
    
    # Select CV based on role
    role_cv_map = {
        "product": "CV_Anton_Kondakov_Product Manager.docx",
        "tpm_program": "CV_Anton_Kondakov_TPM.docx",
        "project": "CV_Anton_Kondakov_Project Manager.docx",
    }
    cv_filename = role_cv_map.get(payload.role_family, "CV_Anton_Kondakov_Product Manager.docx")
    cv_path = gold_cv_path / cv_filename
    
    if not cv_path.exists():
        return {"ok": False, "error": f"CV not found: {cv_filename}"}
    
    # Create application folder
    safe_company = re.sub(r'[^\w\s-]', '', payload.company).strip().replace(' ', '_')
    safe_position = re.sub(r'[^\w\s-]', '', payload.position).strip().replace(' ', '_')[:50]
    folder_name = f"{safe_company}_{safe_position}"
    app_folder = apps_path / folder_name
    app_folder.mkdir(parents=True, exist_ok=True)
    
    # Load and modify CV
    doc = Document(cv_path)
    
    keywords_to_add = payload.keywords_to_add
    
    if keywords_to_add:
        # Find CORE COMPETENCIES section and add keywords
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            if "COMPETENCIES" in text or "SKILLS" in text:
                # Find the next list paragraph and add keywords there
                for j in range(i+1, min(i+10, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    if next_para.style and "List" in next_para.style.name:
                        # Add keywords to Technical Acumen or create new line
                        if "technical" in next_para.text.lower() or "tools" in next_para.text.lower():
                            # Append to existing
                            current_text = next_para.text
                            if not current_text.endswith('.'):
                                current_text += '.'
                            new_keywords = ', '.join(keywords_to_add)
                            next_para.clear()
                            next_para.add_run(f"{current_text} Additional: {new_keywords}.")
                            break
                break
    
    # Save tailored CV
    output_filename = f"CV_Anton_Kondakov_{safe_company}_{safe_position}.docx"
    output_path = app_folder / output_filename
    doc.save(output_path)
    
    # Also try to create PDF (if possible)
    pdf_path = None
    try:
        import subprocess
        # Try using LibreOffice for conversion (if available)
        pdf_output = output_path.with_suffix('.pdf')
        result = subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            '--outdir', str(app_folder), str(output_path)
        ], capture_output=True, timeout=30)
        if pdf_output.exists():
            pdf_path = str(pdf_output)
    except Exception:
        pass  # PDF conversion optional
    
    return {
        "ok": True,
        "cv_path": str(output_path),
        "pdf_path": pdf_path,
        "folder": str(app_folder),
        "keywords_added": keywords_to_add
    }


class CVOptimizeRequest(BaseModel):
    job_title: str
    company: str
    job_description: str
    role_family: str = "product"


@app.post("/cv/optimize-ai")
async def cv_optimize_ai_endpoint(payload: CVOptimizeRequest):
    """
    Use Claude API to analyze JD and optimize CV.
    Extracts key requirements and tailors CV accordingly.
    """
    import os
    import re
    from pathlib import Path
    
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return {"ok": False, "error": "ANTHROPIC_API_KEY not set"}
    
    jd = payload.job_description
    if not jd or len(jd) < 50:
        return {"ok": False, "error": "Job description too short for analysis"}
    
    # Call Claude API to analyze JD
    try:
        import requests
        
        prompt = f"""Analyze this job description and extract:
1. Top 10 most important technical skills/tools required
2. Top 5 soft skills emphasized
3. Key experience requirements (years, domains)
4. Any specific keywords that should be in the CV

Job Title: {payload.job_title}
Company: {payload.company}

Job Description:
{jd[:4000]}

Respond in JSON format:
{{
  "technical_skills": ["skill1", "skill2", ...],
  "soft_skills": ["skill1", ...],
  "experience_requirements": ["req1", ...],
  "keywords_to_add": ["keyword1", ...],
  "cv_recommendations": ["recommendation1", ...]
}}"""
        
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 1000,
                "messages": [{"role": "user", "content": prompt}]
            },
            timeout=30
        )
        
        if response.status_code != 200:
            return {"ok": False, "error": f"Claude API error: {response.status_code}"}
        
        result = response.json()
        ai_text = result.get("content", [{}])[0].get("text", "{}")
        
        # Parse JSON from response
        import json
        # Extract JSON from potential markdown
        json_match = re.search(r'\{[\s\S]*\}', ai_text)
        if json_match:
            analysis = json.loads(json_match.group())
        else:
            analysis = {"error": "Could not parse AI response"}
        
        # Now tailor CV with extracted keywords
        keywords = analysis.get("keywords_to_add", []) + analysis.get("technical_skills", [])[:5]
        keywords = list(set(keywords))[:10]  # Dedupe and limit
        
        if keywords:
            # Call existing tailor endpoint logic
            from docx import Document
            
            gold_cv_path = GOLD_CV_PATH
            apps_path = gold_cv_path / "Applications"
            
            role_cv_map = {
                "product": "CV_Anton_Kondakov_Product Manager.docx",
                "tpm_program": "CV_Anton_Kondakov_TPM.docx",
                "project": "CV_Anton_Kondakov_Project Manager.docx",
            }
            cv_filename = role_cv_map.get(payload.role_family, "CV_Anton_Kondakov_Product Manager.docx")
            cv_path = gold_cv_path / cv_filename
            
            if not cv_path.exists():
                return {"ok": False, "error": f"Base CV not found: {cv_filename}"}
            
            # Create folder
            safe_company = re.sub(r'[^\w\s-]', '', payload.company).strip().replace(' ', '_')
            safe_position = re.sub(r'[^\w\s-]', '', payload.job_title).strip().replace(' ', '_')[:50]
            folder_name = f"{safe_company}_{safe_position}_AI"
            app_folder = apps_path / folder_name
            app_folder.mkdir(parents=True, exist_ok=True)
            
            # Load and modify CV
            doc = Document(cv_path)
            
            # Add keywords to Technical section
            for i, para in enumerate(doc.paragraphs):
                if "COMPETENCIES" in para.text.upper() or "SKILLS" in para.text.upper():
                    for j in range(i+1, min(i+15, len(doc.paragraphs))):
                        next_para = doc.paragraphs[j]
                        if "technical" in next_para.text.lower() or "acumen" in next_para.text.lower():
                            current = next_para.text.rstrip('.')
                            added_kw = ', '.join(keywords[:5])
                            next_para.clear()
                            next_para.add_run(f"{current} [+Added: {added_kw}]")
                            break
                    break
            
            # Save
            output_filename = f"CV_Anton_Kondakov_{safe_company}_AI_Optimized.docx"
            output_path = app_folder / output_filename
            doc.save(output_path)
            
            return {
                "ok": True,
                "cv_path": str(output_path),
                "cv_name": output_filename,
                "keywords_added": keywords[:5],
                "analysis": analysis,
                "folder": str(app_folder)
            }
        else:
            return {
                "ok": True,
                "cv_path": None,
                "cv_name": "No optimization needed",
                "keywords_added": [],
                "analysis": analysis
            }
            
    except Exception as e:
        import traceback
        return {"ok": False, "error": str(e), "traceback": traceback.format_exc()}


# ============= V5 FORM FILLER ENDPOINT =============


# ============= V5 FORM FILLER ENDPOINT =============

@app.post("/apply/v5")
def apply_v5_endpoint(payload: ApplyRequest):
    """
    Apply to job using V5 Form Filler with Claude AI.
    Auto-starts Chrome with debug port if not running.
    """
    import subprocess
    import sys
    
    job_url = payload.job_url
    profile_name = payload.profile
    
    # Check profile exists
    profile_path = Path(f"browser/profiles/{profile_name}.json")
    if not profile_path.exists():
        return {"ok": False, "error": f"Profile '{profile_name}' not found"}
    
    # Start Chrome with debug port using our helper
    cwd = Path(__file__).parent.resolve()
    
    try:
        result = subprocess.run(
            [sys.executable, str(cwd / "browser/start_chrome_debug.py")],
            capture_output=True,
            text=True,
            timeout=15
        )
        
        if "❌" in result.stdout:
            return {"ok": False, "error": "Failed to start Chrome with debug port"}
            
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "Chrome start timeout"}
    except Exception as e:
        return {"ok": False, "error": f"Chrome start error: {e}"}
    
    # Prepare V5 script
    script_content = f'''
import sys
sys.path.insert(0, '{cwd}')
import os
os.chdir('{cwd}')

# Tee output to log file
import io
class TeeWriter:
    def __init__(self, *writers):
        self.writers = writers
    def write(self, text):
        for w in self.writers:
            w.write(text)
            w.flush()
    def flush(self):
        for w in self.writers:
            w.flush()

log_file = open('/tmp/v5_apply.log', 'w')
sys.stdout = TeeWriter(sys.__stdout__, log_file)
sys.stderr = TeeWriter(sys.__stderr__, log_file)

from browser.v5.engine import FormFillerV5, FillMode
from browser.v5.browser_manager import BrowserMode
import time

job_url = "{job_url}"

print("="*60)
print("V5 Form Filler")
print(f"URL: {{job_url}}")
print("="*60)

try:
    # Initialize V5 with CDP (connects to existing Chrome)
    filler = FormFillerV5(browser_mode=BrowserMode.CDP)
    
    # Run in interactive mode
    result = filler.fill(job_url, mode=FillMode.INTERACTIVE)
    
    print("\\n" + "="*60)
    print(f"Filled: {{result.filled_fields}}/{{result.total_fields}}")
    print(f"Verified: {{result.verified_fields}}")
    if result.errors > 0:
        print(f"Errors: {{result.errors}}")
    print("="*60)
    
    # Keep browser open for review
    print("\\nBrowser stays open. Close manually when done.")
    input("Press Enter to close...")
    
except Exception as e:
    print(f"Error: {{e}}")
    import traceback
    traceback.print_exc()
    input("Press Enter to close...")
'''
    
    # Write and run script
    script_file = Path("/tmp/v5_apply_script.py")
    script_file.write_text(script_content)
    
    log_file = Path("/tmp/v5_apply.log")
    
    # Run in new Terminal window (so user can see output)
    apple_script = f'''
    tell application "Terminal"
        activate
        do script "cd {cwd} && {sys.executable} {script_file}"
    end tell
    '''
    
    try:
        subprocess.run(['osascript', '-e', apple_script], capture_output=True)
    except:
        # Fallback: run in background
        subprocess.Popen(
            [sys.executable, str(script_file)],
            stdout=open(log_file, "w"),
            stderr=subprocess.STDOUT,
            cwd=str(cwd)
        )
    
    return {
        "ok": True,
        "message": "V5 Form Filler started in Terminal",
        "log_file": str(log_file),
        "job_url": job_url
    }


@app.get("/apply/v5/log")
def get_v5_log():
    """Get V5 apply log content."""
    log_path = Path("/tmp/v5_apply.log")
    if not log_path.exists():
        return {"ok": False, "log": "No log file"}
    
    try:
        return {"ok": True, "log": log_path.read_text()}
    except Exception as e:
        return {"ok": False, "log": f"Error: {e}"}


# ============= V6 FORM FILLER ENDPOINT =============

@app.post("/apply/v6")
def apply_v6_endpoint(payload: ApplyRequest):
    """
    Apply to job using V6 Form Filler with Claude AI.
    Simpler engine, focused on Greenhouse forms.
    """
    import subprocess
    import sys
    
    job_url = payload.job_url
    profile_name = payload.profile
    
    # Start Chrome with debug port (uses separate profile, won't affect main browser)
    cwd = Path(__file__).parent.resolve()
    
    try:
        from browser.start_chrome_debug import start_chrome_debug
        
        result = start_chrome_debug()
        
        if not result["ok"]:
            return {"ok": False, "error": result.get("message", "Failed to start Chrome")}
            
    except Exception as e:
        return {"ok": False, "error": f"Chrome start error: {e}"}
    
    # Prepare V6 script
    script_content = f'''
import sys
sys.path.insert(0, '{cwd}')
import os
os.chdir('{cwd}')

# Tee output to log file
import io
class TeeWriter:
    def __init__(self, *writers):
        self.writers = writers
    def write(self, text):
        for w in self.writers:
            w.write(text)
            w.flush()
    def flush(self):
        for w in self.writers:
            w.flush()

log_file = open('/tmp/v6_apply.log', 'w')
sys.stdout = TeeWriter(sys.__stdout__, log_file)
sys.stderr = TeeWriter(sys.__stderr__, log_file)

from browser.v6.engine import FormFillerV6
import time

job_url = "{job_url}"

print("="*60)
print("V6 Form Filler (AI-Powered)")
print(f"URL: {{job_url}}")
print("="*60)

try:
    # Initialize V6 and connect to Chrome
    filler = FormFillerV6()
    filler.connect()
    
    # Navigate to job URL first!
    print(f"📍 Opening: {{job_url[:60]}}...")
    filler.page.goto(job_url, wait_until="domcontentloaded", timeout=20000)
    
    # Wait for Greenhouse iframe to appear
    print("⏳ Waiting for form to load...")
    try:
        filler.page.wait_for_selector('#grnhse_iframe', timeout=10000)
    except:
        pass
    
    if not filler.find_frame():
        print("❌ No form found on current page")
        print("Make sure you have the job application form open in Chrome")
        input("Press Enter to close...")
    else:
        # Fill the form
        filler.fill_greenhouse()
        
        print("\\n" + "="*60)
        print("✅ Form filling complete!")
        print("="*60)
        
        print("\\nBrowser stays open. Review and submit manually.")
        input("Press Enter to close...")
    
except Exception as e:
    print(f"Error: {{e}}")
    import traceback
    traceback.print_exc()
    input("Press Enter to close...")
'''
    
    # Write and run script
    script_file = Path("/tmp/v6_apply_script.py")
    script_file.write_text(script_content)
    
    log_file = Path("/tmp/v6_apply.log")
    
    # Open new terminal but don't steal focus and don't close old ones
    escaped_cwd = str(cwd).replace(' ', '\\\\ ')
    apple_script = f'''
    tell application "Terminal"
        do script "cd {escaped_cwd} && {sys.executable} {script_file}"
        delay 0.3
        set current settings of front window to settings set "Basic"
        set font size of front window to 13
    end tell
    '''
    
    try:
        subprocess.run(['osascript', '-e', apple_script], capture_output=True)
    except:
        # Fallback: run in background
        subprocess.Popen(
            [sys.executable, str(script_file)],
            stdout=open(log_file, "w"),
            stderr=subprocess.STDOUT,
            cwd=str(cwd)
        )
    
    return {
        "ok": True,
        "message": "V6 Form Filler started in Terminal",
        "log_file": str(log_file),
        "job_url": job_url
    }


@app.get("/apply/v6/log")
def get_v6_log():
    """Get V6 apply log content."""
    log_path = Path("/tmp/v6_apply.log")
    if not log_path.exists():
        return {"ok": False, "log": "No log file"}
    
    try:
        return {"ok": True, "log": log_path.read_text()}
    except Exception as e:
        return {"ok": False, "log": f"Error: {e}"}


# ============= V7 AGENT FORM FILLER ENDPOINT =============

@app.post("/apply/v7")
def apply_v7(payload: ApplyPayload):
    """
    V7 Agent Form Filler - Uses Claude Vision to fill forms like a human.
    """
    job_url = payload.job_url
    cwd = Path(__file__).parent.resolve()
    
    # Start Chrome with debug port
    try:
        from browser.start_chrome_debug import start_chrome_debug
        result = start_chrome_debug()
        if not result["ok"]:
            return {"ok": False, "error": result.get("message", "Failed to start Chrome")}
    except Exception as e:
        return {"ok": False, "error": f"Chrome start error: {e}"}
    
    # Create V7 script
    script_content = f'''
import sys
sys.path.insert(0, '{cwd}')

from browser.v7.agent import FormFillerAgent

job_url = "{job_url}"

print("=" * 60)
print("V7 AGENT Form Filler")
print(f"URL: {{job_url}}")
print("=" * 60)

agent = FormFillerAgent()

try:
    agent.connect()
    agent.fill_form(job_url)
except KeyboardInterrupt:
    print("\\n⚠️ Interrupted")
except Exception as e:
    print(f"\\n❌ Error: {{e}}")
    import traceback
    traceback.print_exc()
finally:
    input("\\nPress Enter to close...")
    agent.close()
'''
    
    script_file = Path("/tmp/v7_agent_script.py")
    script_file.write_text(script_content)
    
    # Run in Terminal
    escaped_cwd = str(cwd).replace(' ', '\\\\ ')
    apple_script = f'''
    tell application "Terminal"
        do script "cd {escaped_cwd} && {sys.executable} {script_file}"
        delay 0.3
        set current settings of front window to settings set "Basic"
        set font size of front window to 13
    end tell
    '''
    
    try:
        subprocess.run(['osascript', '-e', apple_script], capture_output=True)
    except:
        subprocess.Popen(
            [sys.executable, str(script_file)],
            stdout=open("/tmp/v7_agent.log", "w"),
            stderr=subprocess.STDOUT,
            cwd=str(cwd)
        )
    
    return {
        "ok": True,
        "message": "V7 Agent started in Terminal",
        "job_url": job_url
    }


@app.get("/chrome/status")
def chrome_debug_status():
    """Check if Chrome is running with debug port."""
    import socket
    
    def check_port(port):
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.settimeout(1)
                return s.connect_ex(('localhost', port)) == 0
        except:
            return False
    
    running = check_port(9222)
    return {
        "running": running,
        "port": 9222,
        "message": "Chrome debug ready" if running else "Chrome not running on debug port"
    }


@app.post("/chrome/start")
def start_chrome_debug_endpoint():
    """Start Chrome with debug port."""
    import subprocess
    import sys
    
    cwd = Path(__file__).parent.resolve()
    
    try:
        result = subprocess.run(
            [sys.executable, str(cwd / "browser/start_chrome_debug.py")],
            capture_output=True,
            text=True,
            timeout=15
        )
        
        if "✅" in result.stdout:
            return {"ok": True, "message": result.stdout.strip()}
        else:
            return {"ok": False, "error": result.stdout.strip()}
            
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─────────────────────────────────────────────────────────────────────
# Vision-based Form Filler
# ─────────────────────────────────────────────────────────────────────

@app.post("/apply/vision")
async def apply_vision(payload: dict):
    """
    Fill job application using Claude Vision API.
    Analyzes form screenshots and fills fields intelligently.
    """
    job_url = payload.get("job_url", "")
    if not job_url:
        return {"ok": False, "error": "job_url required"}
    
    # Run in background terminal
    script = f'''
import asyncio
import sys
sys.path.insert(0, str(Path(__file__).parent))

from browser.v5.vision_filler import VisionFormFiller
from playwright.async_api import async_playwright

async def main():
    filler = VisionFormFiller()
    
    async with async_playwright() as p:
        browser = await p.chromium.connect_over_cdp("http://localhost:9222")
        ctx = browser.contexts[0]
        
        print("Creating new page...")
        page = await ctx.new_page()
        await page.goto("{job_url}", wait_until="domcontentloaded", timeout=30000)
        await asyncio.sleep(6)
        
        await page.bring_to_front()
        print(f"Page loaded: {{await page.title()}}")
        
        print("\\nAnalyzing form with Claude Vision...")
        analysis = await filler.analyze_form(page, num_screenshots=3)
        print(analysis)
        
        input("\\nPress Enter to continue or Ctrl+C to cancel...")

asyncio.run(main())
'''
    
    # Save and run script
    script_path = "/tmp/vision_apply.py"
    with open(script_path, "w") as f:
        f.write(script)
    
    import subprocess
    subprocess.Popen([
        "osascript", "-e",
        f'tell application "Terminal" to do script "cd {str(Path(__file__).parent)} && source .venv/bin/activate && python {script_path}"'
    ])
    
    return {"ok": True, "message": "Vision Form Filler started in Terminal"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8001)

