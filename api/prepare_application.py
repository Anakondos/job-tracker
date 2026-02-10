"""
Comprehensive Application Preparation API
"""

import os
import re
import json
import requests
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field


# Ensure ANTHROPIC_API_KEY is loaded from .env
_env_file = Path(__file__).parent.parent / ".env"
if _env_file.exists():
    try:
        _content = _env_file.read_text()
        for _line in _content.strip().split("\n"):
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _v = _line.split("=", 1)
                os.environ.setdefault(_k.strip(), _v.strip())
    except Exception as _e:
        print(f"[PrepareApp] Error loading .env: {_e}")
_api_key_loaded = bool(os.environ.get("ANTHROPIC_API_KEY"))
print(f"[PrepareApp] Module init: API key loaded = {_api_key_loaded}")

GOLD_CV_PATH = Path("/Users/anton/Library/Mobile Documents/com~apple~CloudDocs/Dev/AI_projects/Gold CV")
# Fallback for other machine
if not GOLD_CV_PATH.exists():
    GOLD_CV_PATH = Path("/Users/antonkondakov/Library/Mobile Documents/com~apple~CloudDocs/Dev/AI_projects/Gold CV")
    
APPLICATIONS_PATH = GOLD_CV_PATH / "Applications"

ROLE_CV_MAP = {
    "product": "CV_Anton_Kondakov_Product Manager.docx",
    "tpm_program": "CV_Anton_Kondakov_TPM.docx",
    "project": "CV_Anton_Kondakov_Project Manager.docx",
    "scrum": "CV_Anton_Kondakov_Scrum Master.docx",
    "delivery": "CV_Anton_Kondakov_DeliveryLead.docx",
    "po": "CV_Anton_Kondakov_PO.docx",
    "other": "CV_Anton_Kondakov_Product Manager.docx",
}

# Additional CV source — latest tailored versions
TAILORED_CV_PATH = Path("/Users/antonkondakov/Library/Mobile Documents/com~apple~CloudDocs/CV/February_2026")


@dataclass
class PrepareResult:
    ok: bool = True
    error: str = None
    match_score: int = 0
    fit_level: str = "unknown"
    analysis_summary: str = ""
    key_requirements: List[str] = field(default_factory=list)
    matching_experience: List[str] = field(default_factory=list)
    gaps: List[str] = field(default_factory=list)
    red_flags: List[str] = field(default_factory=list)
    cv_decision: str = "base"
    cv_reason: str = ""
    cv_path: str = None
    cv_filename: str = None
    keywords_added: List[str] = field(default_factory=list)
    cover_letter_path: str = None
    cover_letter_filename: str = None
    cover_letter_preview: str = None
    application_url: str = None
    application_folder: str = None
    
    def to_dict(self):
        return {
            "ok": self.ok, "error": self.error,
            "match_score": self.match_score, "fit_level": self.fit_level,
            "analysis_summary": self.analysis_summary,
            "key_requirements": self.key_requirements or [],
            "matching_experience": self.matching_experience or [],
            "gaps": self.gaps or [], "red_flags": self.red_flags or [],
            "cv_decision": self.cv_decision, "cv_reason": self.cv_reason,
            "cv_path": self.cv_path, "cv_filename": self.cv_filename,
            "keywords_added": self.keywords_added or [],
            "cover_letter_path": self.cover_letter_path,
            "cover_letter_filename": self.cover_letter_filename,
            "cover_letter_preview": self.cover_letter_preview,
            "application_url": self.application_url,
            "application_folder": self.application_folder,
        }


def call_claude_api(prompt: str, max_tokens: int = 2000) -> Optional[str]:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        print("[PrepareApp] No ANTHROPIC_API_KEY")
        return None
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
            json={"model": "claude-3-5-haiku-20241022", "max_tokens": max_tokens, "messages": [{"role": "user", "content": prompt}]},
            timeout=60
        )
        if resp.status_code != 200:
            print(f"[PrepareApp] API error: {resp.status_code}")
            return None
        return resp.json().get("content", [{}])[0].get("text", "")
    except Exception as e:
        print(f"[PrepareApp] Exception: {e}")
        return None


def extract_cv_text(cv_path: Path) -> str:
    """Extract text from DOCX CV file for RAG"""
    try:
        from docx import Document
        doc = Document(str(cv_path))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # Also get tables (often used in CVs)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return '\n'.join(text_parts)
    except Exception as e:
        print(f"[PrepareApp] CV extraction error: {e}")
        return ""


def get_cv_for_role(role_family: str, job_title: str = "") -> tuple:
    """Get CV path and text for given role family, with title-based detection"""
    
    # Check if job title indicates specific role type
    title_lower = job_title.lower() if job_title else ""
    
    # Scrum Master detection - use Scrum CV even if role_family is tpm_program
    if any(x in title_lower for x in ["scrum master", "scrum coach", "agile coach", "agile delivery"]):
        cv_name = "CV_Anton_Kondakov_Scrum Master.docx"
        print(f"[PrepareApp] Detected Scrum Master role from title")
    # Product Owner detection
    elif "product owner" in title_lower:
        cv_name = "CV_Anton_Kondakov_PO.docx"
        print(f"[PrepareApp] Detected Product Owner role from title")
    # Delivery Manager/Lead detection  
    elif any(x in title_lower for x in ["delivery manager", "delivery lead"]):
        cv_name = "CV_Anton_Kondakov_DeliveryLead.docx"
        print(f"[PrepareApp] Detected Delivery role from title")
    else:
        cv_name = ROLE_CV_MAP.get(role_family, ROLE_CV_MAP["product"])
    
    cv_path = GOLD_CV_PATH / cv_name
    
    if cv_path.exists():
        cv_text = extract_cv_text(cv_path)
        return cv_path, cv_text
    
    # Fallback to product CV
    cv_path = GOLD_CV_PATH / ROLE_CV_MAP["product"]
    if cv_path.exists():
        return cv_path, extract_cv_text(cv_path)
    
    return None, ""


# Comprehensive candidate profile (synced from CV Optimisation project, Feb 2026)
CANDIDATE_PROFILE = """
CANDIDATE: Anton Kondakov
LinkedIn: https://www.linkedin.com/in/antonkondakov/
Location: Raleigh, NC | Remote / Relocation within US | US Green Card holder (ITAR/EAR eligible)

SUMMARY:
Senior Technical Program Manager and Product Manager with 15+ years leading complex programs
for Fortune 500 in financial services, insurance, and enterprise tech. $71M+ career project budgets.
Proven at cloud migrations, regulatory platforms, distributed team leadership.
Impact: $1.2M+ cost savings, 25% uptime improvement, 98% on-time delivery, zero production incidents.

WORK HISTORY (exact dates - no overlaps):
1. DXC Technology – Global Insurance Platform (Jun 2025 – Present)
   Senior TPM / Delivery Manager. AWS cloud migration, Angular modernization (v10→v18),
   5 Agile teams (60+ engineers), PI Planning, London Market Insurance.
2. Luxoft USA / Deutsche Bank (Cary, NC): Feb 2020 – May 2025
   VP, Product Owner. GCP migration, regulatory platforms (MiFID II, CAT, EMIR),
   50+ engineers across 15+ countries, $1.2M annual savings.
3. Luxoft Poland / UBS Bank: Apr 2016 – Jan 2020
   Senior Business Analyst → Product Owner. Trade reporting, EMIR compliance,
   cross-functional coordination across 8 teams.
4. Luxoft Europe / Barclays Capital: Oct 2013 – Mar 2016
   Business Analyst. Financial data systems, regulatory reporting.

CORE EXPERTISE:
- Program & Project Management: Full SDLC, roadmapping, scope/schedule/budget, risk mitigation
- Technical Delivery: Cloud migrations (AWS, GCP, Azure), distributed systems, platform modernization
- Large Complex Programs: Own business/technical vision, unblock teams, land business impact
- Methodologies: Agile (Scrum, Kanban, SAFe), Waterfall, Hybrid, PI Planning, CI/CD
- Stakeholder Management: Executive reporting, cross-functional collaboration, client-facing delivery
- Compliance & Governance: MiFID II, CAT, EMIR, SOX, GDPR, audit readiness

TECHNICAL SKILLS:
- Cloud: AWS (ECS, Lambda, Terraform), GCP (BigQuery, Vertex AI), Azure
- Tools: ServiceNow, Jira, MS Project, Confluence, Datadog, Splunk
- Programming: Java, Python, Angular, REST APIs
- Infrastructure: Terraform, CI/CD (Jenkins, GitHub Actions, ArgoCD)
- AI/ML: Amazon Q Developer, Vertex AI, BigQuery ML
- Data: SQL, Tableau, Power BI, BigQuery

CERTIFICATIONS:
- SAFe 5 POPM (Product Owner/Product Manager)
- PSM I (Professional Scrum Master)
- Google Cloud Architect
- MBA: International Institute of Management LINK / The Open University Business School (UK)

LOCATION & AUTHORIZATION:
- Based in Raleigh, NC (Research Triangle)
- Open to: Charlotte, RTP, Durham, Remote USA, relocation within US
- US Work Authorized (Green Card) - NO sponsorship required, ITAR/EAR eligible
- Available: Immediately

TARGET ROLES (what Anton IS looking for):
- Technical Program Manager (TPM) — STRONGEST FIT
- Senior Product Manager / Product Manager
- Program Manager / Delivery Manager / Portfolio Manager (projects, not investments)
- Product Owner (technical platforms)
- Scrum Master / Agile Coach / Delivery Lead
- VP-level or Senior IC positions
- Consulting: Tech Delivery Senior Manager, Client Engagement Manager

TARGET INDUSTRIES:
- Financial Services (FinTech, Banking, Insurance) — STRONGEST
- Enterprise Technology / SaaS
- Aerospace & Defense (program/portfolio management, ITAR eligible)
- Healthcare Technology (compliance expertise transfers)
- Consulting (Deloitte, PwC, Accenture — delivery/engagement management)

NOT A MATCH (important for analysis):
- Financial Portfolio Manager (investments/asset management) — DIFFERENT ROLE
- Sales, BD, Account Executive roles
- Pure software engineering (coding-focused, SWE, SDE)
- Data Scientist / ML Engineer (hands-on modeling, not managing)
- Contract/temporary positions (prefers full-time W2)
- Roles requiring relocation outside US
- Junior/mid-level roles (needs Senior+ or Manager+)
- Roles in India, EMEA-only, or non-US locations

ANALYSIS RULES:
1. "Portfolio Manager" in Tech/Aerospace = Project Portfolio Management = GOOD MATCH
2. "Portfolio Manager" in Bank/Investment = Financial Asset Management = NOT A MATCH
3. Always check company industry context before scoring
4. Regulatory compliance (MiFID II, CAT, EMIR) is a STRONG differentiator
5. Cloud transformation = infrastructure delivery expertise
6. London Market Insurance = commercial insurance domain knowledge
7. If role says "India", "Bengaluru", "Gurugram" etc — location mismatch, score LOW
8. If role says "Remote" with no US restriction — check if US-eligible
9. Senior/Staff/Principal/VP = appropriate level. Associate/Junior = too junior
10. "Agile Delivery Lead" matches well — Anton has SAFe POPM + PSM I + delivery experience
"""


def analyze_job_with_ai(job_title: str, company: str, jd: str, role_family: str) -> Dict:
    """Analyze job match using RAG - loads actual CV for comparison"""
    
    # Try to load actual CV for this role type (pass job_title for better detection)
    cv_path, cv_text = get_cv_for_role(role_family, job_title)
    
    # Use CV text if available, otherwise use static profile
    if cv_text and len(cv_text) > 500:
        candidate_info = f"""ANTON'S CV ({role_family.upper()} version):
{cv_text[:4000]}

ADDITIONAL CONTEXT:
- Location: Raleigh, NC (Open to Charlotte, RTP, Remote USA)
- Work Authorization: US Authorized, no sponsorship required
- Certifications: SAFe 5 POPM, PSM I, Google Cloud Digital Leader
- Looking for: Senior PM, TPM, Program Manager, Product Owner, Scrum Master roles
"""
        print(f"[PrepareApp] Using CV from: {cv_path.name} ({len(cv_text)} chars)")
    else:
        candidate_info = CANDIDATE_PROFILE
        print(f"[PrepareApp] Using static profile (CV not found or too short)")
    
    prompt = f"""{candidate_info}

=== JOB TO ANALYZE ===
Company: {company}
Title: {job_title}
Role Category: {role_family}

Job Description:
{jd[:6000]}

=== ANALYSIS INSTRUCTIONS ===

You are analyzing this job for Anton Kondakov. Be BRUTALLY HONEST — inflated scores waste the candidate's time on bad-fit jobs.

CRITICAL RULES:
1. CHECK COMPANY CONTEXT: "Portfolio Manager" at Frontgrade (Aerospace) = Program Management (GOOD MATCH)
   vs "Portfolio Manager" at Goldman Sachs = Investment Management (NOT A MATCH)
2. BE BRUTALLY HONEST about gaps. If JD asks for Tableau and Anton has zero Tableau → that's "no", not "partial"
3. "Transferable skills" alone should NEVER push score above 55%. Cloud migration ≠ Tableau. Enterprise PM ≠ Mobile UX.
4. Location: "Remote USA" or "NC-based" = match. Specific city only (SF, NYC, Seattle) without remote option = DEDUCT 15 pts
5. Domain-specific tools/platforms (Tableau, Salesforce, Slack APIs, etc.) with zero experience = DEDUCT 20 pts cumulative
6. Do NOT inflate score to be encouraging. A low score helps Anton skip bad matches and focus on real opportunities.

SCORING METHOD — FOLLOW THIS EXACTLY:
Step 1: List the top 5-8 HARD REQUIREMENTS from the JD (must-haves, not nice-to-haves).
Step 2: For each, honestly assess: does Anton have DIRECT experience with THIS SPECIFIC THING?
        - "yes" = he has done this exact thing (not something vaguely similar in a different domain)
        - "partial" = related but different domain/technology/product area
        - "no" = no relevant experience at all
Step 3: Calculate base_ratio = (yes_count * 1.0 + partial_count * 0.3) / total_requirements
Step 4: base_score = base_ratio * 100
Step 5: Apply adjustments:
  - Domain-specific platform/product Anton never used (e.g., Tableau, Salesforce, ServiceNow dev): -15 to -25 pts
  - Location mismatch (city-only, not remote-friendly): -15 pts
  - If mostly "partial" matches (transferable only): cap at 55% MAX
  - Nice-to-have matches from preferred section: +3 to +5 pts max total
  - Industry alignment (FinServ, Insurance, Enterprise Tech): +5 pts
  - Seniority level match: +5 pts

FINAL SCORE RANGES:
- 80-100%: DIRECT experience with 80%+ hard requirements AND matching domain. Rare — reserve for TPM/PM at FinServ/enterprise tech.
- 65-79%: Direct experience with 60%+ requirements, minor adaptation needed.
- 50-64%: Some direct matches but significant domain/technology gaps.
- 35-49%: Mostly transferable skills, wrong domain or product area.
- Below 35%: Wrong role type, wrong location, or fundamental mismatch.

CALIBRATION EXAMPLES:
- TPM at FinServ, cloud migration, distributed teams, regulatory → 85-95% (direct match)
- PM requiring Tableau/Salesforce/specific platform Anton never used → 35-50% (domain mismatch)
- Delivery Manager, Agile, enterprise tech, no specific platform → 75-85% (strong match)
- Role requiring mobile UX + consumer product experience → 30-45% (wrong domain)
- PM at enterprise SaaS, general requirements, no niche platform → 65-75% (good transferable)

Return ONLY valid JSON:
{{
  "match_score": <0-100>,
  "fit_level": "<excellent|good|moderate|low>",
  "analysis_summary": "<2-3 sentences explaining the match, be specific>",
  "role_type": "<what this role actually is: TPM, Product Manager, Financial Analyst, etc>",
  "industry_context": "<company industry and how it relates to Anton's background>",
  "location_info": "<location, remote status, salary if mentioned>",
  "key_requirements": [
    {{"requirement": "<from JD>", "anton_has": "<yes|partial|no>", "evidence": "<specific experience>"}}
  ],
  "matching_experience": ["<specific relevant experience>"],
  "gaps": ["<what's genuinely missing>"],
  "red_flags": ["<serious concerns if any>"],
  "pros": ["<reasons to apply>"],
  "cons": ["<reasons to hesitate>"],
  "recommendation": "<STRONG APPLY|APPLY|MAYBE|SKIP>",
  "recommendation_reason": "<clear explanation>",
  "cv_decision": "<base|optimize>",
  "cv_reason": "<why>",
  "keywords_to_add": ["<relevant keywords from JD>"],
  "cover_letter_focus": ["<key points to emphasize>"]
}}"""

    response = call_claude_api(prompt, 1500)
    if not response:
        return {"error": "AI failed"}
    try:
        match = re.search(r'\{[\s\S]*\}', response)
        if match:
            return json.loads(match.group())
    except:
        pass
    return {"error": "Parse failed"}


def generate_cover_letter(job_title: str, company: str, jd: str, analysis: Dict) -> str:
    focus = analysis.get("cover_letter_focus", [])
    prompt = f"""{CANDIDATE_PROFILE}

JOB: {company} - {job_title}
FOCUS: {', '.join(focus)}
JD: {jd[:2000]}

Write a professional 3-4 paragraph cover letter. No generic phrases. Be specific."""
    return call_claude_api(prompt, 1000) or ""


def create_optimized_cv(job_title: str, company: str, role_family: str, keywords: List[str]) -> Optional[Path]:
    try:
        from docx import Document
        cv_filename = ROLE_CV_MAP.get(role_family, "CV_Anton_Kondakov_Product Manager.docx")
        cv_path = GOLD_CV_PATH / cv_filename
        if not cv_path.exists():
            return None
        
        safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
        safe_position = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')[:40]
        folder_name = f"{safe_company}_{safe_position}_{datetime.now().strftime('%Y%m%d')}"
        app_folder = APPLICATIONS_PATH / folder_name
        app_folder.mkdir(parents=True, exist_ok=True)
        
        doc = Document(cv_path)
        if keywords:
            for i, para in enumerate(doc.paragraphs):
                if "COMPETENCIES" in para.text.upper():
                    for j in range(i+1, min(i+15, len(doc.paragraphs))):
                        if "technical" in doc.paragraphs[j].text.lower():
                            current = doc.paragraphs[j].text.rstrip('.')
                            doc.paragraphs[j].clear()
                            doc.paragraphs[j].add_run(f"{current} [+Added: {', '.join(keywords[:5])}]")
                            break
                    break
        
        output_path = app_folder / f"CV_Anton_Kondakov_{safe_company}.docx"
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"[PrepareApp] CV error: {e}")
        return None


def save_cover_letter(company: str, content: str, app_folder: Path) -> Optional[Path]:
    try:
        safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
        output_path = app_folder / f"Cover_Letter_{safe_company}.txt"
        output_path.write_text(content)
        return output_path
    except Exception as e:
        print(f"[PrepareApp] CL error: {e}")
        return None


def prepare_application(job_title: str, company: str, job_url: str, jd: str, role_family: str = "product") -> PrepareResult:
    result = PrepareResult()
    print(f"[PrepareApp] Starting: {company} - {job_title}")
    
    if not jd or len(jd) < 100:
        result.ok = False
        result.error = "Job description too short"
        result.application_url = job_url
        return result
    
    # Step 1: AI Analysis
    print("[PrepareApp] Step 1: AI Analysis...")
    analysis = analyze_job_with_ai(job_title, company, jd, role_family)
    
    if "error" in analysis:
        result.match_score = 70
        result.fit_level = "moderate"
        result.analysis_summary = "AI unavailable, using defaults"
        result.cv_decision = "base"
        analysis = {}
    else:
        result.match_score = analysis.get("match_score", 70)
        result.fit_level = analysis.get("fit_level", "moderate")
        result.analysis_summary = analysis.get("analysis_summary", "")
        result.key_requirements = analysis.get("key_requirements", [])
        result.matching_experience = analysis.get("matching_experience", [])
        result.gaps = analysis.get("gaps", [])
        result.red_flags = analysis.get("red_flags", [])
        result.cv_decision = analysis.get("cv_decision", "base")
        result.cv_reason = analysis.get("cv_reason", "")
    
    # Step 2: CV
    print(f"[PrepareApp] Step 2: CV ({result.cv_decision})...")
    keywords = analysis.get("keywords_to_add", [])
    
    if result.cv_decision == "optimize" and keywords:
        cv_path = create_optimized_cv(job_title, company, role_family, keywords)
        if cv_path:
            result.cv_path = str(cv_path)
            result.cv_filename = cv_path.name
            result.keywords_added = keywords[:5]
            result.application_folder = str(cv_path.parent)
    
    if not result.cv_path:
        cv_filename = ROLE_CV_MAP.get(role_family, "CV_Anton_Kondakov_Product Manager.docx")
        base_cv = GOLD_CV_PATH / cv_filename
        if base_cv.exists():
            result.cv_path = str(base_cv)
            result.cv_filename = cv_filename
    
    # Step 3: Cover Letter
    print("[PrepareApp] Step 3: Cover Letter...")
    cl_text = generate_cover_letter(job_title, company, jd, analysis)
    
    if cl_text:
        if not result.application_folder:
            safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
            safe_position = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')[:40]
            folder_name = f"{safe_company}_{safe_position}_{datetime.now().strftime('%Y%m%d')}"
            app_folder = APPLICATIONS_PATH / folder_name
            app_folder.mkdir(parents=True, exist_ok=True)
            result.application_folder = str(app_folder)
        
        cl_path = save_cover_letter(company, cl_text, Path(result.application_folder))
        if cl_path:
            result.cover_letter_path = str(cl_path)
            result.cover_letter_filename = cl_path.name
            result.cover_letter_preview = cl_text[:500] + "..." if len(cl_text) > 500 else cl_text
    
    result.application_url = job_url
    
    # Save metadata.json with keywords and analysis for future reference
    if result.application_folder:
        try:
            metadata = {
                "job_title": job_title,
                "company": company,
                "job_url": job_url,
                "role_family": role_family,
                "match_score": result.match_score,
                "fit_level": result.fit_level,
                "cv_decision": result.cv_decision,
                "cv_reason": result.cv_reason,
                "keywords_added": result.keywords_added or [],
                "gaps": result.gaps or [],
                "red_flags": result.red_flags or [],
                "created_at": datetime.now().isoformat()
            }
            metadata_path = Path(result.application_folder) / "metadata.json"
            with open(metadata_path, "w") as f:
                json.dump(metadata, f, indent=2)
        except Exception as e:
            print(f"[PrepareApp] Metadata save error: {e}")
    
    print(f"[PrepareApp] Done: score={result.match_score}")
    return result
